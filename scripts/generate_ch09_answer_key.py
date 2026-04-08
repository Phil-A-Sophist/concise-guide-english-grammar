#!/usr/bin/env python3
"""
Generate Chapter 9 Answer Key and Overhead Answer Key .docx files.
Revised homework: 5 parts, 15 exercises. Part 4 uses new sentences (10-12).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, add_sub_sentence, setup_document, add_title_page,
    add_part_heading, get_font_config, add_diagram_image,
    add_bracket_line,
    add_multilevel_from_bracket, load_chapter_roles,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch09'


PART1_BRACKETS = {
    1: {
        'bracket': '[S [NP [NP [DET The] [ADJ exhausted] [N marathon] [N runner] [PP [PREP from] [NP [N Kenya]]]] [CONJ and] [NP [DET her] [ADJ experienced] [N coach]]] [VP [V celebrated] [PP [PREP after] [NP [DET the] [N race]]]]]',
        'diagram': 'ch09_hw_ex01_marathon_coach',
    },
    2: {
        'bracket': '[S [DC [SUB Because] [NP [DET the] [N deadline]] [VP [AUX was] [V extended]]] [IC [NP [PRON I]] [VP [V had] [NP [N time] [VP [PART to] [V revise] [NP [DET my] [N paper]] [ADVP [ADV thoroughly]]]]]]]',
        'diagram': 'ch09_hw_ex02_deadline_revise',
    },
    3: {
        'bracket': '[S [DC [SUB Although] [NP [DET the] [N professor]] [VP [AUX has] [V retired]]] [IC [NP [PRON she]] [VP [ADV still] [ADV occasionally] [V gives] [NP [N guest] [N lectures]]]] [CONJ and] [IC [NP [DET her] [ADJ former] [N students]] [VP [V attend] [DC [SUB whenever] [NP [PRON they]] [VP [V can]]]]]]',
        'diagram': 'ch09_hw_ex03_professor_students',
    },
}

DIAGRAM_EXERCISES = [
    {
        'num': 10,
        'label': 'Compound noun phrase (simple sentence)',
        'sentence': 'The teacher and the principal met after school.',
        'bracket': '[S [NP [NP [DET The] [N teacher]] [CONJ and] [NP [DET the] [N principal]]] [VP [V met] [PP [PREP after] [NP [N school]]]]]',
        'diagram': 'ch09_hw_ex10_teacher_principal',
    },
    {
        'num': 11,
        'label': 'Compound sentence',
        'sentence': 'The train arrived late, but the passengers remained calm.',
        'bracket': '[S [IC [NP [DET The] [N train]] [VP [V arrived] [ADVP [ADV late]]]] [CONJ but] [IC [NP [DET the] [N passengers]] [VP [V remained] [AdjP [ADJ calm]]]]]',
        'diagram': 'ch09_hw_ex11_train_passengers',
    },
    {
        'num': 12,
        'label': 'Complex sentence',
        'sentence': 'Although the library was quiet, she remained distracted.',
        'bracket': '[S [DC [SUB Although] [NP [DET the] [N library]] [VP [V was] [AdjP [ADJ quiet]]]] [IC [NP [PRON she]] [VP [V remained] [AdjP [ADJ distracted]]]]]',
        'diagram': 'ch09_hw_ex12_library_distracted',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']
    table_font = 'Arial Narrow'
    table_size = 8

    add_title_page(doc, 'Chapter 9: Conjunctions and Clauses', cfg, overhead)

    # =============================================
    # Part 1: Sentence Type Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Sentence Type Identification', cfg, overhead)

    ch_roles = load_chapter_roles(9)
    mode = 'overhead' if overhead else 'answer_key'

    # Exercise 1 (simple — was old Ex 3)
    add_exercise(doc, 1,
        'The exhausted marathon runner from Kenya and her experienced coach celebrated after the race.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_answer_line(doc, 'Sentence type:', 'Simple', body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 One independent clause with a compound NP subject',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Compound NP: "The exhausted marathon runner from Kenya" + "her experienced coach"',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 "after the race" is a PP, not a dependent clause (no subject-verb pair)',
        body_size, indent=0.7, font_name=body_font)

    ex1 = PART1_BRACKETS[1]
    add_multilevel_from_bracket(doc, ex1['bracket'],
                                 roles_dict=ch_roles.get(' '.join(ex1['bracket'].split())),
                                 mode=mode, font_size=table_size, font_name=table_font)
    add_bracket_line(doc, ex1['bracket'], body_size, indent=0.7, font_name=body_font)
    add_diagram_image(doc, DIAGRAM_DIR, ex1['diagram'], width_inches=cfg['diagram_width'])

    question_page_break(doc, overhead)

    # Exercise 2 (complex — same as old Ex 2)
    add_exercise(doc, 2,
        'Because the deadline was extended, I had time to revise my paper thoroughly.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_answer_line(doc, 'Sentence type:', 'Complex', body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 DC: "Because the deadline was extended" (reason)',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 IC: "I had time to revise my paper thoroughly"',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Subordinating conjunction: "because"',
        body_size, indent=0.7, font_name=body_font)

    ex2 = PART1_BRACKETS[2]
    add_multilevel_from_bracket(doc, ex2['bracket'],
                                 roles_dict=ch_roles.get(' '.join(ex2['bracket'].split())),
                                 mode=mode, font_size=table_size, font_name=table_font)
    add_bracket_line(doc, ex2['bracket'], body_size, indent=0.7, font_name=body_font)
    add_diagram_image(doc, DIAGRAM_DIR, ex2['diagram'], width_inches=cfg['diagram_width'])

    question_page_break(doc, overhead)

    # Exercise 3 (compound-complex — was old Ex 1)
    add_exercise(doc, 3,
        'Although the professor has retired, she still occasionally gives guest lectures, and her former students attend whenever they can.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_answer_line(doc, 'Sentence type:', 'Compound-complex', body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 DC: "Although the professor has retired"',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 IC\u2081: "she still occasionally gives guest lectures"',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 IC\u2082: "her former students attend whenever they can"',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '\u2022 DC (nested): "whenever they can" (inside IC\u2082)',
        body_size, indent=0.7, font_name=body_font)

    ex3 = PART1_BRACKETS[3]
    add_multilevel_from_bracket(doc, ex3['bracket'],
                                 roles_dict=ch_roles.get(' '.join(ex3['bracket'].split())),
                                 mode=mode, font_size=table_size, font_name=table_font)
    add_bracket_line(doc, ex3['bracket'], body_size, indent=0.7, font_name=body_font)
    add_diagram_image(doc, DIAGRAM_DIR, ex3['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 2: Sentence Writing
    # =============================================
    add_part_heading(doc, 'Part 2: Sentence Writing', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 4\u20136 are open-ended. Accept any grammatically correct sentence that matches the requested structure.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    # Exercise 4: connect two clauses with semicolon + conjunctive adverb
    add_exercise(doc, 4, 'Connect two clauses using a semicolon and a conjunctive adverb.', body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Structure: Two independent clauses joined by semicolon + conjunctive adverb + comma', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc, 'Sample: "The test was difficult; however, most students passed."', body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc, 'Also acceptable: "The project was late; nevertheless, the client was satisfied."', body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 5: complex sentence with cause/reason
    add_exercise(doc, 5, 'Write a complex sentence with a dependent clause showing cause or reason.', body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Structure: Complex sentence with dependent clause showing cause/reason', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc, 'Sample: "Because the roads were icy, school was canceled."', body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc, 'Also acceptable: "She left early since she had an appointment."', body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 6: compound-complex sentence
    add_exercise(doc, 6, 'Write a compound-complex sentence (two ICs joined by FANBOYS + at least one DC).', body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Structure: Compound-complex (two ICs joined by FANBOYS + at least one DC)', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc, 'Sample: "Although the weather was terrible, the game continued, and the fans cheered."', body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc, 'Sample: "She studied all night because the exam was important, and she passed."', body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc, 'Check: two ICs connected by a coordinating conjunction + at least one DC with a subordinating conjunction.', body_size, indent=0.7, font_name=body_font)

    # =============================================
    # Part 3: Error Correction
    # =============================================
    add_part_heading(doc, 'Part 3: Error Correction', cfg, overhead)

    # Exercise 7 (run-on — was old Ex 9, simplest error type first)
    add_exercise(doc, 7,
        'She enjoys hiking he prefers swimming.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Error type: Run-on (fused sentence)', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc,
        'Correction 1: "She enjoys hiking, but he prefers swimming." (add comma + coordinating conjunction)',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        'Correction 2: "She enjoys hiking; he prefers swimming." (add semicolon)',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 8 (comma splice — was old Ex 8)
    add_exercise(doc, 8,
        'The assignment was challenging, many students struggled to finish it on time.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Error type: Comma splice', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc,
        'Correction 1: "The assignment was challenging, and many students struggled to finish it on time." (add coordinating conjunction)',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        'Correction 2: "The assignment was challenging; many students struggled to finish it on time." (replace comma with semicolon)',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 9 (comma splice — was old Ex 10)
    add_exercise(doc, 9,
        'The restaurant was crowded, we decided to order takeout instead.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Error type: Comma splice', body_size, bold_prefix='', font_name=body_font)
    add_plain_line(doc,
        'Correction 1: "The restaurant was crowded, so we decided to order takeout instead." (add coordinating conjunction)',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        'Correction 2: "Because the restaurant was crowded, we decided to order takeout instead." (subordinate one clause)',
        body_size, indent=0.7, font_name=body_font)

    # =============================================
    # Part 4: Sentence Tables and Diagrams
    # =============================================
    add_part_heading(doc, 'Part 4: Sentence Tables and Diagrams', cfg, overhead)

    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'],
            f'{ex["label"]}. Complete the labeling table and draw a tree diagram for:',
            body_size, font_name=body_font)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=2, space_after=4)
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(' '.join(ex['bracket'].split())),
                                     mode='student', font_size=table_size, font_name=table_font)
        answer_page_break(doc, overhead)
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(' '.join(ex['bracket'].split())),
                                     mode=mode, font_size=table_size, font_name=table_font)
        add_bracket_line(doc, ex['bracket'], body_size, indent=0.7, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Emphasis, End-Weight, and Clause Revision
    # =============================================
    add_part_heading(doc, 'Part 5: Emphasis, End-Weight, and Clause Revision', cfg, overhead)

    # Exercise 13: separate into simple sentences
    add_exercise(doc, 13,
        'Separate the following passage into individual simple sentences. Then explain what relationships between the ideas are lost.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Original: ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run('The lecture was long and the material was difficult and students were confused and they asked many questions and the professor stayed late to help.')
    run.italic = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Simple sentences:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        '1. The lecture was long. 2. The material was difficult. 3. Students were confused. '
        '4. They asked many questions. 5. The professor stayed late to help.',
        body_size, indent=0.7, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('What is lost:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        'The cause-effect relationships disappear. The original (however awkward) implies that '
        'the lecture length and difficulty caused the confusion, which caused the questions, which '
        'caused the professor to stay late. The simple sentences present these as five unrelated '
        'facts. Coordination and subordination are what make those relationships visible.',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 14: end-weight
    add_exercise(doc, 14, 'Revise the following front-loaded sentence using end-weight, then explain why the revision is easier to read.', body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Original (front-loaded): ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run('After the committee reviewed every proposal and discussed the budget in detail, they approved the new plan.')
    run.italic = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('a) End-weighted revision:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        '"The committee approved the new plan after they reviewed every proposal and discussed the budget in detail."',
        body_size, indent=0.7, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('b) Why is the revised version easier to read?')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        'End-weight: placing the heavy dependent clause ("after they reviewed every proposal and discussed '
        'the budget in detail") at the end allows readers to process the main point first ("The committee '
        'approved the new plan"), then receive the background detail. In the original, readers must hold '
        'the long clause in memory before they know what the sentence is about.',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 15: emphasis (most complex)
    add_exercise(doc, 15,
        'The experiment failed, and the researchers were disappointed.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('a) Emphasize disappointment (make "the researchers were disappointed" the main clause):')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        '"Because the experiment failed, the researchers were disappointed."',
        body_size, indent=0.7, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('b) Emphasize the failure (make "the experiment failed" the main clause):')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    add_plain_line(doc,
        '"Although the researchers were disappointed, the experiment had failed."',
        body_size, indent=0.7, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('c) ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run(
        'The coordinated version (original) presents both ideas as equally important. '
        'Coordination is the best choice when neither idea should be pushed into the background. '
        'Both facts carry equal weight in the original \u2014 using "and" signals this equality.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def _add_student_exercise(doc, num, text, fs=12, font='Garamond'):
    """Add a numbered exercise with blank answer lines."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run(f'Exercise {num}. ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = font
    run = p.add_run(text)
    run.font.size = Pt(fs)
    run.font.name = font


def _add_blank_lines(doc, labels, fs=12, font='Garamond'):
    """Add labeled blank answer lines."""
    for label in labels:
        p = doc.add_paragraph()
        run = p.add_run(f'{label} _______')
        run.font.size = Pt(fs)
        run.font.name = font
        set_paragraph_spacing(p, space_before=2, space_after=2)


def create_student_homework(output_path):
    """Create the Chapter 9 Student Homework with all 5 parts."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12
    font = 'Garamond'

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 9 Homework: Conjunctions and Clauses')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = font
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # =============================================
    # Part 1: Sentence Type Identification
    # =============================================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 1: Sentence Type Identification')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = font

    p = doc.add_paragraph()
    run = p.add_run('For each sentence, identify the sentence type (simple, compound, complex, or compound-complex) and identify all clauses. Label each clause as independent (IC) or dependent (DC).')
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=6)

    _add_student_exercise(doc, 1,
        'The exhausted marathon runner from Kenya and her experienced coach celebrated after the race.',
        fs, font)
    _add_blank_lines(doc, ['Sentence type:', 'Clauses:'], fs, font)

    _add_student_exercise(doc, 2,
        'Because the deadline was extended, I had time to revise my paper thoroughly.',
        fs, font)
    _add_blank_lines(doc, ['Sentence type:', 'Clauses:'], fs, font)

    _add_student_exercise(doc, 3,
        'Although the professor has retired, she still occasionally gives guest lectures, and her former students attend whenever they can.',
        fs, font)
    _add_blank_lines(doc, ['Sentence type:', 'Clauses (identify each as IC or DC):'], fs, font)

    # =============================================
    # Part 2: Sentence Writing
    # =============================================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 2: Sentence Writing')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = font

    p = doc.add_paragraph()
    run = p.add_run('Write original sentences following each prompt.')
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=6)

    _add_student_exercise(doc, 4,
        'Connect two clauses using a semicolon and a conjunctive adverb (such as however, therefore, moreover, consequently):',
        fs, font)
    _add_blank_lines(doc, [''], fs, font)

    _add_student_exercise(doc, 5,
        'Write a complex sentence with a dependent clause showing cause or reason (use because, since, or as):',
        fs, font)
    _add_blank_lines(doc, [''], fs, font)

    _add_student_exercise(doc, 6,
        'Write a compound-complex sentence. Use a coordinating conjunction (FANBOYS) to join two independent clauses, and add at least one dependent clause using a subordinating conjunction:',
        fs, font)
    _add_blank_lines(doc, [''], fs, font)

    # =============================================
    # Part 3: Error Correction
    # =============================================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 3: Error Correction')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = font

    p = doc.add_paragraph()
    run = p.add_run('Each sentence below contains a comma splice or run-on error. Provide two different corrections for each.')
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=6)

    _add_student_exercise(doc, 7, 'She enjoys hiking he prefers swimming.', fs, font)
    _add_blank_lines(doc, ['Correction 1:', 'Correction 2:'], fs, font)

    _add_student_exercise(doc, 8,
        'The assignment was challenging, many students struggled to finish it on time.', fs, font)
    _add_blank_lines(doc, ['Correction 1:', 'Correction 2:'], fs, font)

    _add_student_exercise(doc, 9,
        'The restaurant was crowded, we decided to order takeout instead.', fs, font)
    _add_blank_lines(doc, ['Correction 1:', 'Correction 2:'], fs, font)

    # =============================================
    # Part 4: Sentence Tables and Diagrams
    # =============================================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Sentence Tables and Diagrams')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = font

    p = doc.add_paragraph()
    run = p.add_run('For each sentence, complete the labeling table and draw a tree diagram. At the clause level, assign the role Main to independent clauses and Adverbial to dependent clauses.')
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=6)

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = font
        run = p.add_run(f'{ex["label"]}. Complete the labeling table and draw a tree diagram for: ')
        run.font.size = Pt(fs)
        run.font.name = font
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = font

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = font

    # =============================================
    # Part 5: Emphasis, End-Weight, and Clause Revision
    # =============================================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 5: Emphasis, End-Weight, and Clause Revision')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = font

    p = doc.add_paragraph()
    run = p.add_run('Apply what you have learned about emphasis, end-weight, and clause management to revise and analyze sentences.')
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=6)

    _add_student_exercise(doc, 13,
        'The following passage strings together several ideas using only coordination. Separate it into individual simple sentences. Then explain: what relationships between the ideas are lost?',
        fs, font)
    p = doc.add_paragraph()
    run = p.add_run('The lecture was long and the material was difficult and students were confused and they asked many questions and the professor stayed late to help.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=4)
    _add_blank_lines(doc, ['Simple sentences:', 'What is lost:'], fs, font)

    _add_student_exercise(doc, 14,
        'The following sentence is front-loaded\u2014the heavy element appears too early:',
        fs, font)
    p = doc.add_paragraph()
    run = p.add_run('After the committee reviewed every proposal and discussed the budget in detail, they approved the new plan.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=4)
    _add_blank_lines(doc, ['a) End-weighted revision:', 'b) Why is the revised version easier to read?'], fs, font)

    _add_student_exercise(doc, 15,
        'The following sentence uses coordination:',
        fs, font)
    p = doc.add_paragraph()
    run = p.add_run('The experiment failed, and the researchers were disappointed.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = font
    set_paragraph_spacing(p, space_before=2, space_after=4)
    _add_blank_lines(doc, [
        'a) Rewrite to emphasize disappointment (make "the researchers were disappointed" the main clause):',
        'b) Rewrite to emphasize the failure (make "the experiment failed" the main clause):',
        'c) Which version would you use if both ideas should feel equally important? Explain:',
    ], fs, font)

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 09 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 09 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 09 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

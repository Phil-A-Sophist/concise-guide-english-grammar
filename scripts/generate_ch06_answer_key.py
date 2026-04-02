#!/usr/bin/env python3
"""
Generate Chapter 6 Answer Key and Overhead Answer Key .docx files.
Chapter 6: Closed Classes — covers determiners, pronouns, prepositions, conjunctions.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page,
    add_part_heading, get_font_config, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles, add_bracket_line, blank_labels,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch06'


DIAGRAM_EXERCISES = [
    {
        'num': 11, 'sentence': 'She walked to the store.',
        'words':   ['She', 'walked', 'to', 'the', 'store'],
        'roles':   ['Subj', 'Pred', '', '', ''],
        'phrases': ['NP', 'VP', 'PP', 'NP', ''],
        'pos':     ['PRON', 'V', 'PREP', 'DET', 'N'],
        'bracket': '[S [NP [PRON She]] [VP [V walked] [PP [PREP to] [NP [DET the] [N store]]]]]',
        'diagram': 'ch06_hw_ex11_she_walked',
    },
    {
        'num': 12, 'sentence': 'They gave it to her.',
        'words':   ['They', 'gave', 'it', 'to', 'her'],
        'roles':   ['Subj', 'Pred', 'DO', '', ''],
        'phrases': ['NP', 'VP', 'NP', 'PP', ''],
        'pos':     ['PRON', 'V', 'PRON', 'PREP', 'PRON'],
        'bracket': '[S [NP [PRON They]] [VP [V gave] [NP [PRON it]] [PP [PREP to] [NP [PRON her]]]]]',
        'diagram': 'ch06_hw_ex12_they_gave',
    },
    {
        'num': 13, 'sentence': 'The book on the shelf belongs to him.',
        'words':   ['The', 'book', 'on', 'the', 'shelf', 'belongs', 'to', 'him'],
        'roles':   ['Subj', '', '', '', '', 'Pred', '', ''],
        'phrases': ['NP', '', 'PP', 'NP', '', 'VP', 'PP', ''],
        'pos':     ['DET', 'N', 'PREP', 'DET', 'N', 'V', 'PREP', 'PRON'],
        'bracket': '[S [NP [DET The] [N book] [PP [PREP on] [NP [DET the] [N shelf]]]] [VP [V belongs] [PP [PREP to] [NP [PRON him]]]]]',
        'diagram': 'ch06_hw_ex13_book_belongs',
    },
    {
        'num': 14, 'sentence': 'Everyone in the room listened carefully.',
        'words':   ['Everyone', 'in', 'the', 'room', 'listened', 'carefully'],
        'roles':   ['Subj', '', '', '', 'Pred', ''],
        'phrases': ['NP', 'PP', 'NP', '', 'VP', 'ADVP'],
        'pos':     ['PRON', 'PREP', 'DET', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [PRON Everyone] [PP [PREP in] [NP [DET the] [N room]]]] [VP [V listened] [ADVP [ADV carefully]]]]',
        'diagram': 'ch06_hw_ex14_everyone_listened',
    },
    {
        'num': 15, 'sentence': 'My sister and I drove to the park.',
        'words':   ['My', 'sister', 'and', 'I', 'drove', 'to', 'the', 'park'],
        'roles':   ['Subj', '', '', '', 'Pred', '', '', ''],
        'phrases': ['NP', '', 'CONJ', '', 'VP', 'PP', 'NP', ''],
        'pos':     ['DET', 'N', 'CONJ', 'PRON', 'V', 'PREP', 'DET', 'N'],
        'bracket': '[S [NP [DET My] [N sister] [CONJ and] [PRON I]] [VP [V drove] [PP [PREP to] [NP [DET the] [N park]]]]]',
        'diagram': 'ch06_hw_ex15_sister_drove',
    },
]


PASSAGE_PARA1 = (
    "Last year, the bank launched a cryptocurrency trading desk and began "
    "aggressively buying Bitcoin. The executives at the firm instructed "
    "the traders to promote it on social media and in client newsletters. "
    "Once the cryptocurrency had risen sharply in value, the bank quietly "
    "sold their holdings at a massive profit. Many investors who had "
    "followed the bank\u2019s public recommendations lost money when the "
    "scheme collapsed."
)

PASSAGE_PARA2 = (
    "The SEC is now looking into it. They told regulators that the bank had "
    "full compliance to all trading laws. That was confirmed by an internal "
    "review, but outside auditors disagreed. A few of the executives "
    "admitted that this put them in a difficult position. This eventually led "
    "to a class-action lawsuit filed on the bank. The bank has tried to "
    "distance itself from that, but the damage to its reputation may be "
    "permanent."
)


def add_bullet(doc, text, font_size, bold_prefix=None, indent=0.7, font_name=None):
    """Add a bullet-list item with optional bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=1)
    return p


def add_label_line(doc, bold_prefix, text, font_size, indent=0.35, font_name=None, italic_text=False):
    """Add a line with a bold label followed by body text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(bold_prefix)
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if italic_text:
        run.italic = True
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 6 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']
    bracket_size = cfg['bracket_size']

    if overhead:
        table_size = 16
    else:
        table_size = font_size - 1

    add_title_page(doc, 'Chapter 6: Closed Classes', cfg, overhead)

    # =====================================================
    # Part 1: Determiner and Pronoun Identification (Q1-Q3)
    # =====================================================
    add_part_heading(doc, 'Part 1: Determiner and Pronoun Identification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1,
        'The ambitious student submitted her application to several universities.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Determiners: ', '', body_size, font_name=body_font)
    add_bullet(doc, 'article (definite)', body_size, bold_prefix='The \u2014 ', font_name=body_font)
    add_bullet(doc, 'possessive determiner', body_size, bold_prefix='her \u2014 ', font_name=body_font)
    add_bullet(doc, 'quantifier', body_size, bold_prefix='several \u2014 ', font_name=body_font)

    add_label_line(doc, 'Pronouns: ', 'None in this sentence. Note: her is a possessive determiner here, '
        'not a pronoun, because it precedes and modifies a noun.',
        body_size, font_name=body_font)

    # Exercise 2
    question_page_break(doc, overhead)
    add_exercise(doc, 2,
        'Everyone who attended the conference received their materials before the first session.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Pronouns: ', '', body_size, font_name=body_font)
    add_bullet(doc, 'indefinite pronoun', body_size, bold_prefix='Everyone \u2014 ', font_name=body_font)
    add_bullet(doc, 'relative pronoun (introduces relative clause, refers to \u201Ceveryone\u201D)',
        body_size, bold_prefix='who \u2014 ', font_name=body_font)
    add_bullet(doc, 'possessive determiner (could also be classified as possessive pronoun '
        'depending on grammar framework)',
        body_size, bold_prefix='their \u2014 ', font_name=body_font)

    # Exercise 3
    question_page_break(doc, overhead)
    add_exercise(doc, 3,
        'Those books on the shelf belong to someone in this department.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Determiners: ', '', body_size, font_name=body_font)
    add_bullet(doc, 'demonstrative (far)', body_size, bold_prefix='Those \u2014 ', font_name=body_font)
    add_bullet(doc, 'article (definite)', body_size, bold_prefix='the \u2014 ', font_name=body_font)
    add_bullet(doc, 'demonstrative (near)', body_size, bold_prefix='this \u2014 ', font_name=body_font)

    add_label_line(doc, 'Pronouns: ', '', body_size, font_name=body_font)
    add_bullet(doc, 'indefinite pronoun', body_size, bold_prefix='someone \u2014 ', font_name=body_font)

    # ==========================================
    # Part 2: Prepositional Phrase Analysis (Q4-Q6)
    # ==========================================
    add_part_heading(doc, 'Part 2: Prepositional Phrase Analysis', cfg, overhead)

    # Exercise 4
    add_exercise(doc, 4,
        'The student with the red backpack studied in the library until midnight.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for pp_label, pp_phrase, modifies, question in [
        ('PP 1: ', 'with the red backpack', 'modifies \u201Cstudent\u201D (noun)', 'Which student?'),
        ('PP 2: ', 'in the library', 'modifies \u201Cstudied\u201D (verb)', 'Where?'),
        ('PP 3: ', 'until midnight', 'modifies \u201Cstudied\u201D (verb)', 'Until when? / How long?'),
    ]:
        add_label_line(doc, pp_label, pp_phrase, body_size, italic_text=True, font_name=body_font)
        add_bullet(doc, modifies, body_size, bold_prefix='Modifies: ', font_name=body_font)
        add_bullet(doc, question, body_size, bold_prefix='Question answered: ', font_name=body_font)

    # Exercise 5
    question_page_break(doc, overhead)
    add_exercise(doc, 5,
        'The child with the blue hat ran to the store for some milk.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for pp_label, pp_phrase, modifies, question in [
        ('PP 1: ', 'with the blue hat', 'modifies \u201Cchild\u201D (noun)', 'Which child?'),
        ('PP 2: ', 'to the store', 'modifies \u201Cran\u201D (verb)', 'Where? / To where?'),
        ('PP 3: ', 'for some milk', 'modifies \u201Cran\u201D (verb)', 'Why? / For what purpose?'),
    ]:
        add_label_line(doc, pp_label, pp_phrase, body_size, italic_text=True, font_name=body_font)
        add_bullet(doc, modifies, body_size, bold_prefix='Modifies: ', font_name=body_font)
        add_bullet(doc, question, body_size, bold_prefix='Question answered: ', font_name=body_font)

    # Exercise 6
    question_page_break(doc, overhead)
    add_exercise(doc, 6,
        'The author of the bestselling novel spoke to reporters about her new book.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for pp_label, pp_phrase, modifies, question in [
        ('PP 1: ', 'of the bestselling novel', 'modifies \u201Cauthor\u201D (noun)', 'Which author?'),
        ('PP 2: ', 'to reporters', 'modifies \u201Cspoke\u201D (verb)', 'To whom?'),
        ('PP 3: ', 'about her new book', 'modifies \u201Cspoke\u201D (verb)', 'About what?'),
    ]:
        add_label_line(doc, pp_label, pp_phrase, body_size, italic_text=True, font_name=body_font)
        add_bullet(doc, modifies, body_size, bold_prefix='Modifies: ', font_name=body_font)
        add_bullet(doc, question, body_size, bold_prefix='Question answered: ', font_name=body_font)

    # =====================================
    # Part 3: Sentence Completion (Q7-Q10)
    # =====================================
    add_part_heading(doc, 'Part 3: Sentence Completion', cfg, overhead)

    # Exercise 7
    add_exercise(doc, 7,
        'Add a prepositional phrase that modifies the noun book: The book won an award.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Sample revision: ',
        '\u201CThe book about climate change won an award.\u201D',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Other acceptable answers include any PP modifying \u201Cbook\u201D '
        '(e.g., on the bestseller list, by the famous author, with the blue cover).',
        body_size, font_name=body_font)

    # Exercise 8
    question_page_break(doc, overhead)
    add_exercise(doc, 8,
        'Add a prepositional phrase that modifies the verb and indicates when: '
        'She completed the project.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Sample revision: ',
        '\u201CShe completed the project before the deadline.\u201D',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Other acceptable answers include any adverbial PP indicating time '
        '(e.g., during the weekend, in the afternoon, after the meeting).',
        body_size, font_name=body_font)

    # Exercise 9
    question_page_break(doc, overhead)
    add_exercise(doc, 9,
        'Replace the noun phrases with appropriate pronouns: '
        'Maria told John that Maria would return John\u2019s laptop to John tomorrow.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Revised: ',
        '\u201CMaria told him that she would return his laptop to him tomorrow.\u201D',
        body_size, font_name=body_font)

    # Exercise 10
    question_page_break(doc, overhead)
    add_exercise(doc, 10,
        'Write a sentence with a determiner, a conjunction, and a prepositional phrase.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Sample: ',
        '\u201CThe cat and the dog played in the yard.\u201D',
        body_size, font_name=body_font)
    add_bullet(doc, 'The (article, definite) \u2014 appears three times',
        body_size, bold_prefix='Determiner: ', font_name=body_font)
    add_bullet(doc, 'and (coordinating conjunction)',
        body_size, bold_prefix='Conjunction: ', font_name=body_font)
    add_bullet(doc, 'in the yard \u2014 modifies \u201Cplayed\u201D (verb)',
        body_size, bold_prefix='Prepositional phrase: ', font_name=body_font)
    add_plain_line(doc,
        'Accept any sentence containing all three element types, correctly identified.',
        body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Closed-Class Words (Q11-Q15)
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Closed-Class Words', cfg, overhead)

    ch_roles = load_chapter_roles(6)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], bracket_size)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =========================================
    # Part 5: Analysis and Reflection (Q16-Q19)
    # =========================================
    add_part_heading(doc, 'Part 5: Analysis and Reflection', cfg, overhead)

    # Passage
    p = doc.add_paragraph()
    run = p.add_run('Passage:')
    run.bold = True
    run.font.size = Pt(body_size)
    if body_font:
        run.font.name = body_font
    set_paragraph_spacing(p, space_before=4, space_after=2)

    for para_text in [PASSAGE_PARA1, PASSAGE_PARA2]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(para_text)
        run.font.size = Pt(body_size)
        if body_font:
            run.font.name = body_font
        set_paragraph_spacing(p, space_before=2, space_after=4)

    # Exercise 16
    add_exercise(doc, 16,
        'Identify three pronouns with vague or ambiguous referents.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Answers will vary. Acceptable answers include any three of the following:',
        body_size, font_name=body_font)

    for pronoun_ref, problem, revision in [
        ('\u201Cpromote it\u201D (paragraph 1)',
         'Unclear referent. Does \u201Cit\u201D refer to Bitcoin, the trading desk, or cryptocurrency in general?',
         '\u201C\u2026instructed the traders to promote Bitcoin on social media\u2026\u201D'),
        ('\u201Clooking into it\u201D (paragraph 2)',
         'Vague \u201Cit.\u201D Could refer to the scheme, the trading, or the losses.',
         '\u201C\u2026looking into the trading scheme.\u201D'),
        ('\u201CThey told regulators\u201D (paragraph 2)',
         'Ambiguous \u201CThey.\u201D Could refer to the SEC, the executives, or the bank.',
         '\u201CBank representatives told regulators\u2026\u201D'),
        ('\u201CThat was confirmed\u201D (paragraph 2)',
         'Vague \u201CThat.\u201D The claim of compliance? The investigation?',
         '\u201CThe compliance claim was confirmed\u2026\u201D'),
        ('\u201Cthis put them\u201D (paragraph 2)',
         'Vague \u201Cthis\u201D and ambiguous \u201Cthem.\u201D',
         '\u201CThe auditors\u2019 disagreement put the executives in a difficult position.\u201D'),
        ('\u201CThis eventually led\u201D (paragraph 2)',
         'Vague \u201CThis.\u201D Multiple events precede it.',
         '\u201CThe resulting scandal eventually led to a class-action lawsuit\u2026\u201D'),
        ('\u201Cdistance itself from that\u201D (paragraph 2)',
         'Vague \u201Cthat.\u201D The lawsuit? The scheme? The negative publicity?',
         '\u201C\u2026distance itself from the scandal\u2026\u201D'),
        ('\u201Csold their holdings\u201D (paragraph 1)',
         'Singular/plural mismatch: \u201Cthe bank\u201D is singular, but \u201Ctheir\u201D is plural.',
         '\u201C\u2026the bank quietly sold its holdings\u2026\u201D'),
    ]:
        add_label_line(doc, 'Pronoun: ', pronoun_ref, body_size, font_name=body_font)
        add_bullet(doc, problem, body_size, bold_prefix='Problem: ', font_name=body_font)
        add_bullet(doc, revision, body_size, bold_prefix='Revision: ', font_name=body_font)

    # Exercise 17
    question_page_break(doc, overhead)
    add_exercise(doc, 17, 'Identify two determiner problems in the passage.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Problem 1 (premature \u201Cthe\u201D): ',
        'The passage opens with \u201Cthe bank\u201D as if the reader already knows which bank is being '
        'discussed. Similarly, \u201CThe executives at the firm,\u201D \u201Cthe traders,\u201D '
        '\u201Cthe cryptocurrency,\u201D and \u201Cthe scheme\u201D all use the definite article '
        'before these referents have been introduced.',
        body_size, font_name=body_font)
    add_bullet(doc,
        'First introduce the referent with an indefinite article or proper name, then use \u201Cthe\u201D '
        'in subsequent references. For example: \u201CLast year, a major investment bank launched\u2026\u201D '
        'or name the bank specifically.',
        body_size, bold_prefix='Fix: ', font_name=body_font)

    add_label_line(doc, 'Problem 2 (contradictory quantifiers): ',
        'Paragraph 1 says \u201CMany investors\u2026 lost money,\u201D while paragraph 2 says \u201CA few of the '
        'executives admitted\u2026\u201D The shift from \u201Cmany\u201D (suggesting large numbers) to '
        '\u201Ca few\u201D (suggesting small numbers) is potentially contradictory or confusing when describing '
        'related groups within the same event.',
        body_size, font_name=body_font)
    add_bullet(doc,
        'Use consistent quantifiers or clarify the different groups. For example: \u201CSeveral '
        'of the executives admitted\u2026\u201D or specify which executives.',
        body_size, bold_prefix='Fix: ', font_name=body_font)

    # Exercise 18
    question_page_break(doc, overhead)
    add_exercise(doc, 18, 'Identify two incorrect prepositions in the passage.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_label_line(doc, 'Error 1: ', '\u201Ccompliance to all trading laws\u201D',
        body_size, font_name=body_font)
    add_bullet(doc,
        '\u201Ccompliance with all trading laws.\u201D The standard collocation is \u201Ccompliance with,\u201D '
        'not \u201Ccompliance to.\u201D',
        body_size, bold_prefix='Correction: ', font_name=body_font)

    add_label_line(doc, 'Error 2: ', '\u201Ca class-action lawsuit filed on the bank\u201D',
        body_size, font_name=body_font)
    add_bullet(doc,
        '\u201Ca class-action lawsuit filed against the bank.\u201D Lawsuits are filed \u201Cagainst\u201D '
        'a party, not \u201Con\u201D a party.',
        body_size, bold_prefix='Correction: ', font_name=body_font)

    # Exercise 19
    question_page_break(doc, overhead)
    add_exercise(doc, 19,
        'Rewrite one paragraph correcting all closed-class word issues.',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_plain_line(doc, 'Answers will vary. Sample rewrite of paragraph 1:',
        body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(
        '\u201CLast year, a major investment bank launched a cryptocurrency trading desk and began '
        'aggressively buying Bitcoin. Executives at the bank instructed its traders to promote '
        'Bitcoin on social media and in client newsletters. Once the cryptocurrency had risen '
        'sharply in value, the bank quietly sold its holdings at a massive profit. Many investors '
        'who had followed the bank\u2019s public recommendations lost money when the price collapsed.\u201D'
    )
    run.font.size = Pt(body_size)
    if body_font:
        run.font.name = body_font
    set_paragraph_spacing(p, space_before=2, space_after=4)

    add_plain_line(doc, 'Sample rewrite of paragraph 2:', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(
        '\u201CThe SEC is now looking into the trading scheme. Bank representatives told regulators '
        'that the bank had maintained full compliance with all trading laws. The compliance claim '
        'was confirmed by an internal review, but outside auditors disagreed. Several executives '
        'admitted that the auditors\u2019 findings put the leadership team in a difficult position. '
        'The resulting scandal eventually led to a class-action lawsuit filed against the bank. '
        'The bank has tried to distance itself from the controversy, but the damage to its '
        'reputation may be permanent.\u201D'
    )
    run.font.size = Pt(body_size)
    if body_font:
        run.font.name = body_font
    set_paragraph_spacing(p, space_before=2, space_after=4)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 6 Student Homework with blank multi-level tables."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 6 Homework: Closed Classes')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Part 4 with blank multi-level tables
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Diagramming Closed-Class Words')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 06 Homework.docx'
    )

    # Create Answer Key (standard size)
    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 06 Answer Key.docx',
        font_size=12
    )

    # Create Overhead Answer Key
    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 06 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

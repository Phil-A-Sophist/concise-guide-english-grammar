"""
Reformat Chapter 05 Answer Key: Convert paragraph-form lists to bullet points.

Replaces "Test used:" paragraphs (which list multiple tests in a single paragraph)
with "Tests used:" followed by individual bullet items for each test.
Also converts the sentences list in section 5.9.5 to proper bullets.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "Homework/Chapter 05 Answer Key.docx"
OUTPUT = "Homework/Chapter 05 Answer Key.docx"


def make_numPr(numId="1001", ilvl="0"):
    """Create a numbering properties element (for bullet formatting)."""
    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), ilvl)
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), numId)
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    return numPr


def add_bullet_paragraph(doc, paragraph_before, text, bold_prefix=None, numId="1001"):
    """Insert a bullet paragraph after the given paragraph, with actual bullet formatting."""
    new_p = OxmlElement("w:p")
    paragraph_before._element.addnext(new_p)

    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph_before._parent)

    # Set style to Compact (same as existing bullets in the doc)
    new_para.style = doc.styles["Compact"]

    # Ensure paragraph properties exist and add numbering
    pPr = new_para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        new_para._element.insert(0, pPr)
    # Add numbering properties for bullet
    pPr.append(make_numPr(numId))

    if bold_prefix:
        run_b = new_para.add_run(bold_prefix)
        run_b.bold = True
        run = new_para.add_run(text)
    else:
        run = new_para.add_run(text)
    return new_para


def clear_paragraph_and_set(paragraph, bold_text):
    """Clear a paragraph and set it to just bold label text."""
    # Remove all existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    # Clear any remaining text
    for r in paragraph._element.findall(qn("w:r")):
        paragraph._element.remove(r)
    run = paragraph.add_run(bold_text)
    run.bold = True


# Define the bullet content for each "Test used" paragraph.
# Key: paragraph index (in original doc), Value: list of (bold_prefix, text) tuples
test_bullets = {
    # Q1: development
    7: [
        ("Morphological test: ", "The suffix -ment typically derives nouns from verbs (develop \u2192 development)."),
        ("Syntactic test: ", "\u201Cdevelopment\u201D follows the determiner \u201CThe\u201D and functions as the subject of the sentence."),
        ("Pronoun replacement: ", "It can be replaced by a pronoun: \u201CIt takes time.\u201D"),
    ],
    # Q2: quickly
    11: [
        ("Morphological test: ", "The suffix -ly attached to the adjective \u201Cquick\u201D forms an adverb."),
        ("Syntactic test: ", "\u201CQuickly\u201D modifies the verb \u201Csolved,\u201D telling us how she solved the problem."),
        ("Movability test: ", "It can be moved in the sentence: \u201CShe solved the problem quickly.\u201D"),
    ],
    # Q3: beautiful
    15: [
        ("Position test: ", "\u201CBeautiful\u201D appears between a determiner (\u201CThe\u201D) and a noun (\u201Cgarden\u201D), a characteristic position for adjectives."),
        ("Comparison test: ", "It can be used in comparative forms (more beautiful, most beautiful)."),
        ("Linking verb test: ", "It can appear after a linking verb: \u201CThe garden is beautiful.\u201D"),
    ],
    # Q4: investigate
    19: [
        ("Modal test: ", "\u201CInvestigate\u201D follows the modal verb \u201Cwill,\u201D which only combines with verbs."),
        ("Conjugation test: ", "It can be conjugated for tense (investigated, investigates, investigating)."),
        ("Object test: ", "It takes a direct object (\u201Cthe matter\u201D)."),
    ],
    # Q5: surprisingly
    23: [
        ("Morphological test: ", "The suffix -ly attached to \u201Csurprising\u201D forms an adverb."),
        ("Modification test: ", "\u201CSurprisingly\u201D modifies the adjective \u201Cconfident,\u201D indicating the degree or manner of confidence."),
        ("Pattern: ", "Adverbs commonly modify adjectives in this way."),
    ],
    # Q21: creation
    109: [
        ("Morphological test: ", "The suffix -tion derives nouns from verbs (create \u2192 creation)."),
        ("Syntactic test: ", "\u201CCreation\u201D follows the possessive \u201CThe artist\u2019s\u201D and functions as the subject of the sentence."),
        ("Pronoun replacement: ", "It can be replaced by a pronoun: \u201CIt amazed the critics.\u201D"),
    ],
    # Q22: created
    113: [
        ("Morphological test: ", "\u201CCreated\u201D shows past tense marking (-ed), a morphological feature of verbs."),
        ("Syntactic test: ", "It has a subject (\u201CThe artist\u201D) and takes an object (\u201Csomething amazing\u201D)."),
        ("Conjugation test: ", "It can be conjugated: creates, creating, will create."),
    ],
    # Q23: creative
    117: [
        ("Morphological test: ", "The suffix -ive typically forms adjectives (create \u2192 creative)."),
        ("Syntactic test: ", "\u201CCreative\u201D follows the linking verb \u201Cis\u201D and can be modified by the intensifier \u201Chighly.\u201D"),
        ("Comparison test: ", "It can be compared: more creative, most creative."),
    ],
    # Q24: creatively
    121: [
        ("Morphological test: ", "The suffix -ly attached to the adjective \u201Ccreative\u201D forms an adverb."),
        ("Modification test: ", "\u201CCreatively\u201D modifies the verb \u201Cworks,\u201D describing how the artist works."),
    ],
}

# Sentences list in section 5.9.5 (paragraph 105 in original)
sentences_bullets = [
    (None, "A. The artist\u2019s creation amazed the critics."),
    (None, "B. The artist created something amazing."),
    (None, "C. The artist is highly creative."),
    (None, "D. The artist works creatively."),
]


def main():
    doc = Document(INPUT)

    # Find the numId used by existing Compact/bullet paragraphs
    numId = "1001"  # default
    for p in doc.paragraphs:
        if p.style.name == "Compact":
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId_elem = numPr.find(qn("w:numId"))
                    if numId_elem is not None:
                        numId = numId_elem.get(qn("w:val"))
                        break
    print(f"Using numId={numId} for bullets")

    # Process "Test used" paragraphs - work in REVERSE order so indices stay valid
    for idx in sorted(test_bullets.keys(), reverse=True):
        para = doc.paragraphs[idx]
        bullets = test_bullets[idx]

        # Verify this is actually a "Test used:" paragraph
        if not para.text.startswith("Test used:"):
            print(f"WARNING: Paragraph {idx} doesn't start with 'Test used:': {para.text[:50]}")
            continue

        # Change paragraph text to "Tests used:"
        clear_paragraph_and_set(para, "Tests used:")

        # Insert bullets after the paragraph (in forward order, each after the last inserted)
        last = para
        for bold_prefix, text in bullets:
            last = add_bullet_paragraph(doc, last, text, bold_prefix, numId)

    # Now find and convert the sentences list paragraph (search by content since indices shifted)
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Sentences:") and "creation amazed" in p.text:
            clear_paragraph_and_set(p, "Sentences:")
            last = p
            for bold_prefix, text in sentences_bullets:
                last = add_bullet_paragraph(doc, last, text, bold_prefix, numId)
            break

    doc.save(OUTPUT)
    print(f"Saved reformatted answer key to {OUTPUT}")


if __name__ == "__main__":
    main()

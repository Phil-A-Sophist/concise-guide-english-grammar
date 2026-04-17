#!/usr/bin/env python3
"""
Validate labeling tables in generated homework .docx files.

Checks:
  - Answer keys: no empty role cells where merged spans exist (roles should be filled)
  - Student docs: no non-blank content in POS, Phrase, or Role rows (should be blank)

Exit code: 0 if all pass, 1 if violations found.

Usage:
    python scripts/validate_homework_tables.py [--verbose]
"""

import argparse
import sys
from pathlib import Path
from docx import Document


HOMEWORK_DIR = Path(__file__).resolve().parent.parent / 'Homework'

# Ch07 tables exempt from student validation (0-indexed):
# Part 3 examples and scaffolded exercises with intentionally pre-filled rows.
_CH07_EXEMPT_TABLES = {
    ('Chapter 07 Homework.docx', 0),   # Part 3 example "She arrived"
    ('Chapter 07 Homework.docx', 1),   # Ex 7 "Thunder rumbled" — POS pre-filled
    ('Chapter 07 Homework.docx', 3),   # Ex 9 "The cat chased the mouse" — Role pre-filled
    ('Chapter 07 Homework.docx', 4),   # Part 4 example "The cat slept"
}

# Row header labels that identify a labeling table
LABELING_HEADERS = {'Role', 'Phrase', 'POS', 'Word',
                    'Role 1', 'Role 2', 'Phrase 1', 'Phrase 2',
                    'Clause', 'Clause 1', 'Clause 2'}
ANSWER_ROWS = {'Role', 'Phrase', 'POS', 'Clause',
               'Role 1', 'Role 2', 'Phrase 1', 'Phrase 2',
               'Clause 1', 'Clause 2'}
WORD_ROWS = {'Word'}


def get_cell_text(cell):
    """Get trimmed text from a table cell."""
    return cell.text.strip()


def is_labeling_table(table):
    """Check if a table is a labeling table by examining first-column headers."""
    if len(table.rows) < 3:
        return False
    headers = set()
    for row in table.rows:
        if row.cells:
            text = get_cell_text(row.cells[0])
            if text in LABELING_HEADERS:
                headers.add(text)
    # Must have at least Word + one other labeling header
    return 'Word' in headers and len(headers) >= 2


def validate_answer_key(filepath, verbose=False):
    """Check answer key tables for empty role cells."""
    violations = []
    try:
        doc = Document(str(filepath))
    except Exception as e:
        return [f"  ERROR: Could not open {filepath.name}: {e}"]

    for t_idx, table in enumerate(doc.tables):
        if not is_labeling_table(table):
            continue

        # Skip student-mode tables embedded in answer key docs (question pages).
        # A table with ALL answer rows empty is a student table — skip it.
        has_any_content = False
        for row in table.rows:
            if not row.cells:
                continue
            header = get_cell_text(row.cells[0])
            if header not in ANSWER_ROWS:
                continue
            for c_idx in range(1, len(row.cells)):
                if get_cell_text(row.cells[c_idx]):
                    has_any_content = True
                    break
            if has_any_content:
                break
        if not has_any_content:
            continue  # Student table embedded in answer key — expected to be blank

        for row in table.rows:
            if not row.cells:
                continue
            header = get_cell_text(row.cells[0])
            if header not in ANSWER_ROWS:
                continue
            # Track already-seen XML elements to avoid duplicate reports on merged cells
            seen_tc_ids = set()
            for c_idx in range(1, len(row.cells)):
                cell = row.cells[c_idx]
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue  # Same merged cell, already reported
                seen_tc_ids.add(tc_id)
                text = get_cell_text(cell)
                # Check if cell has background shading (grey = no-info cell, skip)
                tc = cell._tc
                tcPr = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if tcPr is not None:
                    fill = tcPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '')
                    if fill.upper() in ('CCCCCC', 'C0C0C0', 'D9D9D9'):
                        continue  # Grey cell = intentionally blank
                if not text:
                    violations.append(
                        f"  Table {t_idx + 1}, row '{header}', col {c_idx}: "
                        f"EMPTY (should have content in answer key)")
                    if verbose:
                        # Show context: what's in the Word row?
                        for wr in table.rows:
                            if get_cell_text(wr.cells[0]) == 'Word':
                                words = [get_cell_text(wr.cells[i])
                                         for i in range(1, len(wr.cells))]
                                violations.append(f"    Words: {' '.join(words)}")
                                break
    return violations


def _is_example_table(doc, table):
    """Check if a table is preceded by an 'Example (completed)' paragraph.

    Scans the document body elements to find the paragraph immediately before
    this table. If it (or the one before it) contains 'example' and 'completed',
    the table is a worked example — intentionally filled in student docs.
    """
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    table_elem = table._tbl
    prev_paragraphs = []
    for child in body:
        if child is table_elem:
            break
        if child.tag == _qn('w:p'):
            text = child.text or ''
            # Also check runs inside the paragraph
            for run in child.findall('.//' + _qn('w:t')):
                if run.text:
                    text += run.text
            prev_paragraphs.append(text.lower())
    # Check last 3 paragraphs before the table for example indicators
    for p_text in prev_paragraphs[-3:]:
        # "Example (completed)" pattern in chapter homework
        if 'example' in p_text and 'completed' in p_text:
            return True
        # "Example: <sentence>" pattern in exam study guides
        if p_text.strip().startswith('example'):
            return True
    return False


def validate_student(filepath, verbose=False):
    """Check student tables for leaked answer content."""
    violations = []
    try:
        doc = Document(str(filepath))
    except Exception as e:
        return [f"  ERROR: Could not open {filepath.name}: {e}"]

    for t_idx, table in enumerate(doc.tables):
        if not is_labeling_table(table):
            continue
        # Skip worked example tables (intentionally filled in student docs)
        if _is_example_table(doc, table):
            continue
        if (filepath.name, t_idx) in _CH07_EXEMPT_TABLES:
            continue
        for row in table.rows:
            if not row.cells:
                continue
            header = get_cell_text(row.cells[0])
            if header not in ANSWER_ROWS:
                continue
            for c_idx in range(1, len(row.cells)):
                cell = row.cells[c_idx]
                text = get_cell_text(cell)
                # Allow blank, underscore placeholders, and single space (pre-merge)
                if text and text != '_____' and text.strip():
                    violations.append(
                        f"  Table {t_idx + 1}, row '{header}', col {c_idx}: "
                        f"LEAKED '{text}' (should be blank in student version)")
    return violations


def main():
    parser = argparse.ArgumentParser(description='Validate homework labeling tables')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Show additional context for violations')
    args = parser.parse_args()

    total_violations = 0
    files_checked = 0

    # Check Answer Keys
    ak_dir = HOMEWORK_DIR / 'Answer Keys'
    oh_dir = HOMEWORK_DIR / 'Overheads'
    st_dir = HOMEWORK_DIR / 'Student'

    print("=== Validating Answer Keys ===")
    for d in [ak_dir, oh_dir]:
        if not d.exists():
            print(f"  SKIP: {d} not found")
            continue
        for f in sorted(d.glob('*.docx')):
            if f.name.startswith('~'):
                continue
            files_checked += 1
            violations = validate_answer_key(f, verbose=args.verbose)
            if violations:
                print(f"\n  FAIL: {f.relative_to(HOMEWORK_DIR)}")
                for v in violations:
                    print(v)
                total_violations += len(violations)
            elif args.verbose:
                print(f"  OK: {f.relative_to(HOMEWORK_DIR)}")

    print("\n=== Validating Student Homework ===")
    if st_dir.exists():
        for f in sorted(st_dir.glob('*.docx')):
            if f.name.startswith('~'):
                continue
            files_checked += 1
            violations = validate_student(f, verbose=args.verbose)
            if violations:
                print(f"\n  FAIL: {f.relative_to(HOMEWORK_DIR)}")
                for v in violations:
                    print(v)
                total_violations += len(violations)
            elif args.verbose:
                print(f"  OK: {f.relative_to(HOMEWORK_DIR)}")
    else:
        print(f"  SKIP: {st_dir} not found")

    # Also check exam student versions
    exam_dir = HOMEWORK_DIR / 'Exams-Guides-Handouts'
    if exam_dir.exists():
        for f in sorted(exam_dir.glob('*.docx')):
            if f.name.startswith('~'):
                continue
            name_lower = f.name.lower()
            if 'answer key' in name_lower or 'overhead' in name_lower:
                files_checked += 1
                violations = validate_answer_key(f, verbose=args.verbose)
                if violations:
                    print(f"\n  FAIL: {f.relative_to(HOMEWORK_DIR)}")
                    for v in violations:
                        print(v)
                    total_violations += len(violations)
            elif 'study guide' in name_lower and 'answer' not in name_lower:
                files_checked += 1
                violations = validate_student(f, verbose=args.verbose)
                if violations:
                    print(f"\n  FAIL: {f.relative_to(HOMEWORK_DIR)}")
                    for v in violations:
                        print(v)
                    total_violations += len(violations)

    print(f"\n=== Summary ===")
    print(f"Files checked: {files_checked}")
    print(f"Violations: {total_violations}")

    if total_violations > 0:
        print("VALIDATION FAILED")
        return 1
    else:
        print("ALL TABLES VALID")
        return 0


if __name__ == '__main__':
    sys.exit(main())

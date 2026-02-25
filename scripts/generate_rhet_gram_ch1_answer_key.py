#!/usr/bin/env python3
"""
Generate Answer Key and Overhead Answer Key for
Homework Rhet Gram Ch 1 - S26.docx

Overhead reflects user-edited formatting:
  - Bulleted (one-per-line) headword/determiner lists
  - Expanded Ex 3.1 ambiguity note
  - Page break after each Ex 1 sub-exercise and between Ex 3.2 / 3.3
  - Spacers between PP items in Ex 3.3
  - Bracket notation + tree diagrams in Ex 3 (all three) and Ex 4
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch01_rg'


# ─── helpers ────────────────────────────────────────────────────────────────

def set_spacing(p, before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before * 20)))
    sp.set(qn('w:after'),  str(int(after  * 20)))
    pPr.append(sp)


def spacer(doc, font_name, size):
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(size)
    set_spacing(p, 0, 0)
    return p


def heading(doc, text, level, font_name, size, before=4, after=6):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name  = font_name
        run.font.size  = Pt(size)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    set_spacing(p, before, after)
    return p


def exercise_line(doc, number, sentence, font_name, size, before=6):
    p = doc.add_paragraph()
    r = p.add_run(f'Exercise {number}.  ')
    r.bold = True; r.font.name = font_name; r.font.size = Pt(size)
    if sentence:
        r2 = p.add_run(sentence)
        r2.italic = True; r2.font.name = font_name; r2.font.size = Pt(size)
    set_spacing(p, before, 2)
    return p


def labeled_line(doc, label, answer, font_name, size, indent=0.35):
    """Bold label + plain answer on one line, indented."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(label + ' ')
    r.bold = True; r.font.name = font_name; r.font.size = Pt(size)
    r2 = p.add_run(answer)
    r2.font.name = font_name; r2.font.size = Pt(size)
    set_spacing(p, 0, 2)
    return p


def label_only(doc, label, font_name, size, indent=0.35):
    """Bold label on its own line (for overhead bullet lists)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(label)
    r.bold = True; r.font.name = font_name; r.font.size = Pt(size)
    set_spacing(p, 0, 2)
    return p


def bullet_item(doc, text, font_name, size):
    """Plain item at left margin (for overhead bullet lists)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = font_name; r.font.size = Pt(size)
    set_spacing(p, 0, 2)
    return p


def plain_line(doc, text, font_name, size, indent=0, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.bold = bold; r.font.name = font_name; r.font.size = Pt(size)
    set_spacing(p, 0, 2)
    return p


def mono_line(doc, text, mono_size, indent=0.35, before=0, after=0):
    """Single Courier New line (for bracket notation)."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(mono_size)
    set_spacing(p, before, after)
    return p


def add_diagram(doc, image_name, font_name, body_size, width_inches=5.5):
    """Add a bold 'Tree diagram:' label then the PNG image, centred."""
    lbl = doc.add_paragraph()
    lbl.paragraph_format.left_indent = Inches(0.35)
    r = lbl.add_run('Tree diagram:')
    r.bold = True; r.font.name = font_name; r.font.size = Pt(body_size)
    set_spacing(lbl, 4, 2)

    img_path = DIAGRAM_DIR / f"{image_name}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_path.exists():
        p.add_run().add_picture(str(img_path), width=Inches(width_inches))
    else:
        p.add_run(f'[Diagram not found: {image_name}.png]')
    set_spacing(p, 2, 6)


def add_brackets(doc, text, mono_size, font_name, body_size):
    """Add a bold 'Bracket notation:' label then the bracket string."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    r = p.add_run('Bracket notation:')
    r.bold = True; r.font.name = font_name; r.font.size = Pt(body_size)
    set_spacing(p, 6, 2)
    mono_line(doc, text, mono_size, indent=0.5, before=0, after=4)


def add_table(doc, words, phrases, pos_list, font_name, size):
    n = len(words)
    table = doc.add_table(rows=3, cols=n + 1)
    table.style = 'Table Grid'
    for r_idx, (hdr, data) in enumerate(zip(
            ['Phrase', 'Word', 'Part of Speech'],
            [phrases,  words,  pos_list])):
        row = table.rows[r_idx]
        cell = row.cells[0]
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True; run.font.name = font_name; run.font.size = Pt(size)
        for c_idx, val in enumerate(data):
            run = row.cells[c_idx + 1].paragraphs[0].add_run(val)
            run.font.name = font_name; run.font.size = Pt(size)
            if r_idx == 0:
                run.bold = True
    p_after = doc.add_paragraph()
    set_spacing(p_after, 2, 2)
    return table


# ─── content data ───────────────────────────────────────────────────────────

EX1 = [
    {
        'number': '1a',
        'sentence': (
            'CGI helped produce the lush, cinematic rain forests of James Cameron\u2019s '
            'Avatar as well as the teeming orc armies of Peter Jackson\u2019s Hobbit series.'
        ),
        'headwords': ['CGI', 'rain forests', 'Avatar', 'armies', 'series'],
        'headwords_note': 'also Cameron and Jackson inside possessive DPs',
        'determiners': [
            'the (before \u201clush, cinematic rain forests\u201d)',
            'James Cameron\u2019s (possessive Det for \u201cAvatar\u201d)',
            'the (before \u201cteeming orc armies\u201d)',
            'Peter Jackson\u2019s (possessive Det for \u201cHobbit series\u201d)',
        ],
    },
    {
        'number': '1b',
        'sentence': (
            'In the operating room, surgeons may rely on three-dimensional images '
            'to improve their patients\u2019 prognoses.'
        ),
        'headwords': ['room', 'surgeons', 'images', 'prognoses'],
        'headwords_note': None,
        'determiners': [
            'the (before \u201coperating room\u201d)',
            'their (Determiner for patients\u2019)',
            'patients\u2019 (possessive Det for \u201cprognoses\u201d)',
        ],
    },
    {
        'number': '1c',
        'sentence': 'Clearly, this innovation has the potential to shape our future.',
        'headwords': ['innovation', 'potential', 'future'],
        'headwords_note': None,
        'determiners': [
            'this (before \u201cinnovation\u201d)',
            'the (before \u201cpotential\u201d)',
            'our (before \u201cfuture\u201d)',
        ],
    },
]

PRACTICE = [
    {
        'label': 'Practice One',
        'sentence': 'Bears seldomly attack without a very good reason.',
        'words':   ['Bears', 'seldomly', 'attack', 'without', 'a',   'very', 'good', 'reason'],
        'phrases': ['NP',    'VP',       'VP',     'PP',      'PP',  'PP',   'PP',   'PP'],
        'pos':     ['N',     'Adv',      'V',      'Prep',    'Det', 'Adv',  'Adj',  'N'],
        'brackets': '[NP Bears] [VP [ADVP seldomly] attack [PP without [NP a [ADJP [ADVP very] good] reason]]]',
        'image':   'ch01_rg_p1_bears',
        'width':    6.0,
    },
    {
        'label': 'Practice Two',
        'sentence': 'Stephen usually sits alone at home.',
        'words':   ['Stephen', 'usually', 'sits', 'alone', 'at',   'home'],
        'phrases': ['NP',      'VP',      'VP',   'VP',    'PP',   'PP'],
        'pos':     ['N',       'Adv',     'V',    'Adv',   'Prep', 'N'],
        'brackets': '[NP Stephen] [VP [ADVP usually] sits [ADVP alone] [PP at [NP home]]]',
        'image':   'ch01_rg_p2_stephen',
        'width':    4.5,
    },
    {
        'label': 'Practice Three',
        'sentence': 'My younger brother works for the city.',
        'words':   ['My',  'younger', 'brother', 'works', 'for',  'the', 'city'],
        'phrases': ['NP',  'NP',      'NP',      'VP',    'PP',   'PP',  'PP'],
        'pos':     ['Det', 'Adj',     'N',       'V',     'Prep', 'Det', 'N'],
        'brackets': '[NP My younger brother] [VP works [PP for [NP the city]]]',
        'image':   'ch01_rg_p3_brother',
        'width':    5.0,
    },
    {
        'label': 'Practice Four',
        'sentence': 'The painfully long discussion continued incessantly until noon.',
        'words':   ['The', 'painfully', 'long', 'discussion', 'continued', 'incessantly', 'until', 'noon'],
        'phrases': ['NP',  'NP',        'NP',   'NP',         'VP',        'VP',          'PP',    'PP'],
        'pos':     ['Det', 'Adv',       'Adj',  'N',          'V',         'Adv',         'Prep',  'N'],
        'brackets': '[NP The [ADJP [ADVP painfully] long] discussion] [VP continued [ADVP incessantly] [PP until [NP noon]]]',
        'image':   'ch01_rg_p4_discussion',
        'width':    5.5,
    },
    {
        'label': 'Example Five',
        'sentence': 'All my dearest friends from highschool suddenly left.',
        'words':   ['All', 'my',  'dearest', 'friends', 'from', 'highschool', 'suddenly', 'left'],
        'phrases': ['NP',  'NP',  'NP',      'NP',      'PP',   'PP',         'VP',       'VP'],
        'pos':     ['Det', 'Det', 'Adj',     'N',       'Prep', 'N',          'Adv',      'V'],
        'brackets': '[NP All my dearest friends [PP from [NP highschool]]] [VP [ADVP suddenly] left]',
        'image':   'ch01_rg_p5_friends',
        'width':    5.5,
    },
]


# ─── document builder ───────────────────────────────────────────────────────

def build(output_path, overhead=False):
    if overhead:
        body_font  = 'Arial Narrow'
        body_size  = 18
        h1_size    = 22
        h2_size    = 20
        h3_size    = 18
        tbl_size   = 14
        mono_size  = 14
    else:
        body_font  = 'Garamond'
        body_size  = 12
        h1_size    = 16
        h2_size    = 14
        h3_size    = 12
        tbl_size   = 11
        mono_size  = 10

    doc = Document()
    doc.styles['Normal'].font.name = body_font
    doc.styles['Normal'].font.size = Pt(body_size)

    # ── Title ──────────────────────────────────────────────────────────────
    heading(doc, 'Rhetorical Grammar: Chapter 1', 1, body_font, h1_size, before=0, after=4)
    heading(doc, 'Answer Key', 2, body_font, h2_size, before=0, after=10)
    if overhead:
        spacer(doc, body_font, body_size)

    # ══════════════════════════════════════════════════════════════════════
    # Exercise 1: Headwords and Determiners
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, 'Exercise 1: Headwords and Determiners', 3, body_font, h3_size, before=6, after=4)

    for ex in EX1:
        exercise_line(doc, ex['number'], ex['sentence'], body_font, body_size)

        if overhead:
            # ── Overhead: one item per line ────────────────────────────
            label_only(doc, 'Headwords:', body_font, body_size)
            for hw in ex['headwords']:
                bullet_item(doc, hw, body_font, body_size)
            if ex['headwords_note']:
                bullet_item(doc, f"({ex['headwords_note']})", body_font, body_size)

            label_only(doc, 'Determiners:', body_font, body_size)
            for det in ex['determiners']:
                bullet_item(doc, det, body_font, body_size)

            spacer(doc, body_font, body_size)
            doc.add_page_break()  # page break after each sub-exercise (overhead)

        else:
            # ── Answer key: inline bullet-separated list ───────────────
            hw_str = ' \u2022 '.join(ex['headwords'])
            if ex['headwords_note']:
                hw_str += f"  ({ex['headwords_note']})"
            labeled_line(doc, 'Headwords:', hw_str, body_font, body_size)
            labeled_line(doc, 'Determiners:', ' \u2022 '.join(ex['determiners']),
                         body_font, body_size)

    # ══════════════════════════════════════════════════════════════════════
    # Exercise 2: Subject / Predicate
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, 'Exercise 2: Subject / Predicate', 3, body_font, h3_size, before=6, after=4)

    ex2_data = [
        ('The mayor\u2019s husband / spoke against the ordinance.',
         'The mayor\u2019s husband', 'spoke against the ordinance'),
        ('The merchants in town / are unhappy.',
         'The merchants in town', 'are unhappy'),
        ('The three white monkeys / carelessly climbed on their tall tree.',
         'The three white monkeys', 'carelessly climbed on their tall tree'),
    ]
    for i, (marked, subj, pred) in enumerate(ex2_data, 1):
        exercise_line(doc, f'2.{i}', marked, body_font, body_size)
        labeled_line(doc, 'Subject:', subj, body_font, body_size)
        labeled_line(doc, 'Predicate:', pred, body_font, body_size)
        if overhead:
            spacer(doc, body_font, body_size)

    if overhead:
        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # Exercise 3: Prepositional Phrases
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, 'Exercise 3: Prepositional Phrases', 3, body_font, h3_size, before=6, after=4)

    # 3.1
    exercise_line(doc, '3.1', 'Birthday cakes are common in many Western cultures.',
                  body_font, body_size)
    labeled_line(doc, 'PP:', '\u201cin many Western cultures\u201d', body_font, body_size)
    if overhead:
        labeled_line(doc, 'Function:', 'Adverbial', body_font, body_size)
        plain_line(doc, 'Ambiguous what it modifies', body_font, body_size)
        plain_line(doc, 'Maybe \u201ccommon\u201d (adjective)', body_font, body_size)
        plain_line(doc, 'Maybe \u201care\u201d (verb)', body_font, body_size)
        plain_line(doc, 'Maybe the whole sentence', body_font, body_size)
    else:
        labeled_line(doc, 'Function:', 'Adverbial \u2014 modifies \u201care common,\u201d '
                     'indicating context/location (the exact attachment is somewhat ambiguous: '
                     'it could modify \u201ccommon,\u201d \u201care,\u201d or the whole clause)',
                     body_font, body_size)
    add_brackets(doc, '[NP Birthday cakes] [VP are [ADJP common] [PP in [NP many Western cultures]]]',
                 mono_size, body_font, body_size)
    add_diagram(doc, 'ch01_rg_ex31_birthday', body_font, body_size, width_inches=5.0)
    if overhead:
        spacer(doc, body_font, body_size)
        doc.add_page_break()  # page break after 3.1 (overhead)

    # 3.2
    exercise_line(doc, '3.2', 'Cupcakes are a popular alternative to birthday cakes.',
                  body_font, body_size)
    labeled_line(doc, 'PP:', '\u201cto birthday cakes\u201d', body_font, body_size)
    labeled_line(doc, 'Function:', 'Adjectival \u2014 modifies the noun \u201calternative\u201d',
                 body_font, body_size)
    add_brackets(doc, '[NP Cupcakes] [VP are [NP a popular alternative [PP to [NP birthday cakes]]]]',
                 mono_size, body_font, body_size)
    add_diagram(doc, 'ch01_rg_ex32_cupcakes', body_font, body_size, width_inches=5.0)
    if overhead:
        spacer(doc, body_font, body_size)
        doc.add_page_break()  # page break between 3.2 and 3.3 (overhead)

    # 3.3
    exercise_line(doc, '3.3', 'The man in the big red hat spoke eloquently about his vacation to Morocco.',
                  body_font, body_size)
    labeled_line(doc, 'PP 1:', '\u201cin the big red hat\u201d', body_font, body_size)
    labeled_line(doc, 'Function:', 'Adjectival \u2014 modifies the noun \u201cman\u201d',
                 body_font, body_size)
    if overhead:
        spacer(doc, body_font, body_size)
    labeled_line(doc, 'PP 2:', '\u201cabout his vacation\u201d', body_font, body_size)
    labeled_line(doc, 'Function:', 'Adverbial \u2014 modifies the verb \u201cspoke\u201d',
                 body_font, body_size)
    if overhead:
        spacer(doc, body_font, body_size)
    labeled_line(doc, 'PP 3:', '\u201cto Morocco\u201d', body_font, body_size)
    labeled_line(doc, 'Function:', 'Adjectival \u2014 modifies the noun \u201cvacation\u201d',
                 body_font, body_size)
    add_brackets(doc,
                 '[NP The man [PP in [NP the big red hat]]] '
                 '[VP spoke [ADVP eloquently] [PP about [NP his vacation [PP to [NP Morocco]]]]]',
                 mono_size, body_font, body_size)
    add_diagram(doc, 'ch01_rg_ex33_man', body_font, body_size, width_inches=5.5)
    if overhead:
        spacer(doc, body_font, body_size)
        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # Exercise 4: Sentence Labeling Tables
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, 'Exercise 4: Sentence Labeling Tables', 3, body_font, h3_size, before=6, after=4)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0)
    r = note.add_run('Note: ')
    r.bold = True; r.font.name = body_font; r.font.size = Pt(body_size)
    r2 = note.add_run('Bracket notation and tree diagram are provided below each table. '
                      'Student diagrams drawn in the diagramming app should match the tree structure.')
    r2.font.name = body_font; r2.font.size = Pt(body_size)
    set_spacing(note, 0, 8)

    for item in PRACTICE:
        # Sentence label
        lbl = doc.add_paragraph()
        r_l = lbl.add_run(item['label'] + ':  ')
        r_l.bold = True; r_l.font.name = body_font; r_l.font.size = Pt(body_size)
        r_s = lbl.add_run(item['sentence'])
        r_s.italic = True; r_s.font.name = body_font; r_s.font.size = Pt(body_size)
        set_spacing(lbl, 6, 3)

        # Table
        add_table(doc, item['words'], item['phrases'], item['pos'], body_font, tbl_size)

        # Bracket notation
        add_brackets(doc, item['brackets'], mono_size, body_font, body_size)

        # Tree diagram
        add_diagram(doc, item['image'], body_font, body_size, width_inches=item['width'])

        if overhead:
            spacer(doc, body_font, body_size)
            doc.add_page_break()

    # ── save ───────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'Saved: {output_path}')


# ─── main ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base = Path(__file__).parent.parent / 'Homework'
    build(base / 'Homework Rhet Gram Ch 1 - S26 Answer Key.docx', overhead=False)
    build(base / 'Homework Rhet Gram Ch 1 - S26 Overhead.docx',   overhead=True)

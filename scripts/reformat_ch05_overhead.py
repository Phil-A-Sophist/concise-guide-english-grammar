"""
Reformat Chapter 05 Overhead for classroom projection:
- Minimum 22pt font throughout (headings larger for hierarchy)
- Each question on its own page (page break before each Heading 2)
- Question at top of page, answer starts halfway down
"""
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

INPUT = "Homework/Homework 05 Overhead.docx"
OUTPUT = "Homework/Homework 05 Overhead.docx"

# Font sizes
SIZE_H1 = Pt(28)
SIZE_H2 = Pt(26)
SIZE_H3 = Pt(24)
SIZE_BODY = Pt(22)

# Space after question heading to push answer to ~halfway down page
# US Letter: 11" page, ~1" margins = 9" usable = 648pt
# Half = 324pt. Question takes ~50pt. Gap needed: ~270pt.
QUESTION_GAP = Pt(270)


def set_page_break_before(paragraph):
    """Set the 'page break before' property on a paragraph."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)
    pageBreak = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreak)


def set_font_size_on_runs(paragraph, size):
    """Set font size on all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.size = size


def set_space_after(paragraph, space):
    """Set space after a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_after = space


def main():
    doc = Document(INPUT)

    # Also resize fonts in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_font_size_on_runs(p, SIZE_BODY)

    # Track which H2s immediately follow an H1 (so we put the page break on the H1 instead)
    prev_style = None

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name

        # --- Font sizes ---
        if style_name == "Heading 1":
            set_font_size_on_runs(para, SIZE_H1)
        elif style_name == "Heading 2":
            set_font_size_on_runs(para, SIZE_H2)
        elif style_name == "Heading 3":
            set_font_size_on_runs(para, SIZE_H3)
        else:
            # Body Text, First Paragraph, Compact, Normal, etc.
            set_font_size_on_runs(para, SIZE_BODY)

        # --- Page breaks and spacing ---
        if style_name == "Heading 1":
            # Skip page break for the very first title paragraphs (indices 0 and 1)
            if i >= 2:
                # Section heading: add page break only if NOT followed by H2
                # (if followed by H2, the H2 will handle the page break and
                #  we'll keep the H1 on the same page as the first question)
                # We'll handle this by adding the page break here,
                # but the H2 below won't add one if it follows H1.
                set_page_break_before(para)

        elif style_name == "Heading 2":
            if prev_style != "Heading 1":
                # This question doesn't follow a section heading,
                # so it needs its own page break
                set_page_break_before(para)
            # Add gap after question heading to push answer to mid-page
            set_space_after(para, QUESTION_GAP)

        elif style_name == "Heading 3":
            # Sub-questions (25a, 25b, 25c) - each on its own page
            set_page_break_before(para)
            set_space_after(para, QUESTION_GAP)

        prev_style = style_name

    doc.save(OUTPUT)
    print(f"Saved reformatted overhead to {OUTPUT}")


if __name__ == "__main__":
    main()

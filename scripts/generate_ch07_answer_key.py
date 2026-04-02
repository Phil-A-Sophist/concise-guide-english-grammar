#!/usr/bin/env python3
"""
Generate Chapter 7 Answer Key and Overhead Answer Key .docx files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, add_diagram_image, setup_document, add_title_page,
    add_part_heading, get_font_config,
    add_labeling_table, add_bracket_line, compute_spans, blank_labels,
    add_multilevel_from_bracket, load_chapter_roles,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch07'


DIAGRAM_EXERCISES = [
    {
        'num': 10, 'sentence': 'The dog barked loudly.',
        'words':   ['The', 'dog', 'barked', 'loudly'],
        'roles':   ['Subj', '', 'Pred', ''],
        'phrases': ['NP', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [N dog]] [VP [V barked] [ADVP [ADV loudly]]]]',
        'image':   'ch07_hw_ex10_dog_barked',
    },
    {
        'num': 11, 'sentence': 'The talented student from Ohio won the award.',
        'words':   ['The', 'talented', 'student', 'from', 'Ohio', 'won', 'the', 'award'],
        'roles':   ['Subj', '', '', '', '', 'Pred', '', ''],
        'phrases': ['NP', '', '', 'PP', '', 'VP', 'NP', ''],
        'pos':     ['DET', 'ADJ', 'N', 'PREP', 'N', 'V', 'DET', 'N'],
        'bracket': '[S [NP [DET The] [ADJP [ADJ talented]] [N student] [PP [PREP from] [NP [N Ohio]]]] [VP [V won] [NP [DET the] [N award]]]]',
        'image':   'ch07_hw_ex11_student_won',
    },
    {
        'num': 12, 'sentence': 'She carefully read the interesting book.',
        'words':   ['She', 'carefully', 'read', 'the', 'interesting', 'book'],
        'roles':   ['Subj', 'Pred', '', '', '', ''],
        'phrases': ['NP', 'ADVP', 'VP', 'NP', '', ''],
        'pos':     ['PRON', 'ADV', 'V', 'DET', 'ADJ', 'N'],
        'bracket': '[S [NP [PRON She]] [VP [ADVP [ADV carefully]] [V read] [NP [DET the] [ADJP [ADJ interesting]] [N book]]]]',
        'image':   'ch07_hw_ex12_she_read',
    },
]

TABLE_EXERCISES = [
    {
        'num': 7, 'sentence': 'Thunder rumbled.',
        'words':   ['Thunder', 'rumbled'],
        'roles':   ['Subj', 'Pred'],
        'phrases': ['NP', 'VP'],
        'pos':     ['N', 'V'],
    },
    {
        'num': 8, 'sentence': 'The old man sat quietly.',
        'words':   ['The', 'old', 'man', 'sat', 'quietly'],
        'roles':   ['Subj', '', '', 'Pred', ''],
        'phrases': ['NP', '', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'ADJ', 'N', 'V', 'ADV'],
    },
    {
        'num': 9, 'sentence': 'The cat chased the mouse.',
        'words':   ['The', 'cat', 'chased', 'the', 'mouse'],
        'roles':   ['Subj', '', 'Pred', '', ''],
        'phrases': ['NP', '', 'VP', 'NP', ''],
        'pos':     ['DET', 'N', 'V', 'DET', 'N'],
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 7 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']
    bracket_size = cfg['bracket_size']
    diagram_width = cfg['diagram_width']

    # CH07 uses specific table and diagram sizes for overhead
    if overhead:
        table_size = 16
        diagram_width = 4.5
    else:
        table_size = font_size - 1

    add_title_page(doc, 'Chapter 7: Introduction to Sentence Diagramming', cfg, overhead)

    # =============================================
    # Part 1: Subject and Predicate Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Subject and Predicate Identification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'The curious students from the advanced chemistry class carefully examined the unusual compound.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for label, answer in [
        ('Subject NP:', 'The curious students from the advanced chemistry class'),
        ('Head of subject NP:', 'students'),
        ('Predicate VP:', 'carefully examined the unusual compound'),
        ('Head of predicate VP:', 'examined'),
    ]:
        add_answer_line(doc, label, answer, body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 2
    add_exercise(doc, 2, 'My extremely talented older sister from Portland won the national competition.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for label, answer in [
        ('Subject NP:', 'My extremely talented older sister from Portland'),
        ('Head of subject NP:', 'sister'),
        ('Predicate VP:', 'won the national competition'),
        ('Head of predicate VP:', 'won'),
    ]:
        add_answer_line(doc, label, answer, body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 3
    add_exercise(doc, 3, 'Several angry protesters outside the courthouse demanded immediate action.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for label, answer in [
        ('Subject NP:', 'Several angry protesters outside the courthouse'),
        ('Head of subject NP:', 'protesters'),
        ('Predicate VP:', 'demanded immediate action'),
        ('Head of predicate VP:', 'demanded'),
    ]:
        add_answer_line(doc, label, answer, body_size, font_name=body_font)

    # ==============================
    # Part 2: Heads and Modifiers
    # ==============================
    add_part_heading(doc, 'Part 2: Heads and Modifiers', cfg, overhead)

    # Exercise 4
    add_exercise(doc, 4, 'my grandmother\'s beautiful antique wooden jewelry box', body_size, font_name=body_font)

    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Head:', 'box', body_size, font_name=body_font)

    add_plain_line(doc, '', body_size, bold_prefix='Modifiers:', font_name=body_font)

    modifiers_4 = [
        "my grandmother's \u2014 possessive determiner",
        "beautiful \u2014 adjective",
        "antique \u2014 adjective",
        "wooden \u2014 adjective",
        "jewelry \u2014 noun (functioning adjectivally)",
    ]
    for mod in modifiers_4:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(mod)
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=0, space_after=1)

    question_page_break(doc, overhead)

    # Exercise 5
    add_exercise(doc, 5, 'extremely carefully', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_answer_line(doc, 'Head:', 'carefully', body_size, font_name=body_font)

    add_plain_line(doc, '', body_size, bold_prefix='Modifiers:', font_name=body_font)

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('extremely \u2014 adverb (degree modifier)')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=0, space_after=1)

    question_page_break(doc, overhead)

    # Exercise 6
    add_exercise(doc, 6, 'quite proud of her remarkable achievement', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_answer_line(doc, 'Head:', 'proud', body_size, font_name=body_font)

    add_plain_line(doc, '', body_size, bold_prefix='Modifiers:', font_name=body_font)

    modifiers_6 = [
        "quite \u2014 adverb (degree modifier)",
        "of her remarkable achievement \u2014 prepositional phrase (complement of 'proud')",
    ]
    for mod in modifiers_6:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(mod)
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=0, space_after=1)

    # ======================================
    # Part 3: Completing Sentence Tables
    # ======================================
    add_part_heading(doc, 'Part 3: Completing Sentence Tables', cfg, overhead)

    for i, ex in enumerate(TABLE_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_labeling_table(doc, ex['words'],
                           pos_labels=ex['pos'],
                           phrase_labels=ex['phrases'],
                           role_labels=ex['roles'],
                           font_size=table_size)

    # ===========================================
    # Part 4: Completing Diagrams and Tables
    # ===========================================
    add_part_heading(doc, 'Part 4: Completing Diagrams and Tables', cfg, overhead)

    ch_roles = load_chapter_roles(7)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], bracket_size)
        add_diagram_image(doc, DIAGRAM_DIR, ex['image'], width_inches=diagram_width)

    # ==========================================
    # Part 5: Structural Ambiguity Analysis
    # ==========================================
    add_part_heading(doc, 'Part 5: Structural Ambiguity Analysis', cfg, overhead)

    # Exercise 13
    add_exercise(doc, 13, 'I shot an elephant in my pajamas. How he got in my pajamas, I will never know.', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('13A) Two possible meanings:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Meaning 1: I was wearing my pajamas when I shot an elephant. (PP "in my pajamas" modifies VP \u2014 describes the circumstances of the shooting)')
    run.font.size = Pt(body_size)
    run.font.name = body_font

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Meaning 2: I shot an elephant that was wearing my pajamas. (PP "in my pajamas" modifies NP "an elephant" \u2014 describes which elephant)')
    run.font.size = Pt(body_size)
    run.font.name = body_font

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('13B) Diagrams and bracket notation for each reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Meaning 1 (VP attachment):')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=2, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('[S [NP [PRON I]] [VP [V shot] [NP [DET an] [N elephant]] [PP [PREP in] [NP [DET my] [N pajamas]]]]]')
    run.font.name = 'Consolas'
    run.font.size = Pt(bracket_size)

    add_diagram_image(doc, DIAGRAM_DIR, 'ch07_hw_ex13_elephant_vp', width_inches=diagram_width)

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Meaning 2 (NP attachment):')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=6, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('[S [NP [PRON I]] [VP [V shot] [NP [DET an] [N elephant] [PP [PREP in] [NP [DET my] [N pajamas]]]]]]')
    run.font.name = 'Consolas'
    run.font.size = Pt(bracket_size)

    add_diagram_image(doc, DIAGRAM_DIR, 'ch07_hw_ex13_elephant_np', width_inches=diagram_width)

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('13C) Model response:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'This sentence is funny because of structural ambiguity involving PP attachment. '
        'The audience initially interprets "in my pajamas" as modifying the VP \u2014 '
        'describing the shooter\'s attire, which is a plausible (if eccentric) reading. '
        'Groucho then reveals the absurd alternative: the elephant was wearing his pajamas. '
        'This reading comes from attaching the PP to the NP "an elephant" instead. '
        'The humor arises because both structures are grammatically valid, but one produces '
        'an absurd mental image. The joke exploits the fact that listeners commit to one '
        'structural analysis before realizing the other was intended.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    # Exercise 14
    question_page_break(doc, overhead)

    add_exercise(doc, 14, 'The horse raced past the barn fell.', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('14A) Initial reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'Most readers initially parse "The horse" as the subject NP and "raced past the barn" '
        'as the main VP \u2014 the horse is running past a barn. When "fell" appears, the sentence '
        'seems to "break" because the reader has already assigned "raced" as the main verb, '
        'and there appears to be no grammatical role for "fell" to play.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('14B) Correct reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'The correct reading is: "The horse [that was] raced past the barn fell." '
        'Here, "raced past the barn" is a reduced relative clause modifying "horse" \u2014 '
        'it tells us which horse (the one that was raced past the barn). '
        'The main verb of the sentence is "fell." The full subject NP is '
        '"The horse raced past the barn," and the VP is simply "fell."'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('14C) Diagrams and bracket notation for each reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Diagram 1 \u2014 Garden-path (incorrect) reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=2, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('[S [NP [DET The] [N horse]] [VP [V raced] [PP [PREP past] [NP [DET the] [N barn]]]]] + fell ???')
    run.font.name = 'Consolas'
    run.font.size = Pt(bracket_size)

    add_diagram_image(doc, DIAGRAM_DIR, 'ch07_hw_ex14_garden_path', width_inches=diagram_width)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'In the garden-path reading, "raced" is parsed as the main verb with "past the barn" '
        'as a PP inside the VP. This leaves "fell" with no grammatical role, which is why the '
        'sentence seems to break.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('Diagram 2 \u2014 Correct reading:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=6, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run('[S [NP [DET The] [N horse] [VP [V raced] [PP [PREP past] [NP [DET the] [N barn]]]]] [VP [V fell]]]')
    run.font.name = 'Consolas'
    run.font.size = Pt(bracket_size)

    add_diagram_image(doc, DIAGRAM_DIR, 'ch07_hw_ex14_correct', width_inches=diagram_width)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'In the correct reading, "raced past the barn" is a VP inside the '
        'subject NP (modifying "horse"), and "fell" is the main verb in the predicate.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('14D) Model response:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.7)
    run = p.add_run(
        'Garden-path sentences cause confusion because our brains process language incrementally \u2014 '
        'we build structural interpretations word by word as we read. When we encounter "The horse raced," '
        'the simplest analysis is that "raced" is the main verb, and we commit to that structure. '
        'When "fell" appears, it forces us to revise: "raced" was actually part of a reduced relative '
        'clause, not the main verb. This revision is cognitively costly, which is why the sentence '
        'feels confusing. Garden-path sentences demonstrate that sentence comprehension is not just '
        'about knowing the words \u2014 it requires actively building and sometimes revising hierarchical '
        'structure in real time.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 7 Student Homework with blank multi-level tables for Part 4."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 7 Homework: Introduction to Sentence Diagramming')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Part 4 with blank multi-level tables
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Completing Diagrams and Tables')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=8, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    # Create Student Homework
    create_student_homework(
        homework_dir / 'Student' / 'Chapter 07 Homework.docx'
    )

    # Create Answer Key (standard size)
    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 07 Answer Key.docx',
        font_size=12
    )

    # Create Overhead Answer Key (Arial Narrow, reduced sizes, spacer rows)
    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 07 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

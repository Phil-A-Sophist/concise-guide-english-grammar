#!/usr/bin/env python3
"""
Generate Chapter 14 homework files: Student Homework, Answer Key, and Overhead.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page, add_part_heading,
    exercise_separator, get_font_config, add_bracket_line, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles,
    parse_bracket_to_multilevel, add_multilevel_labeling_table,
    question_page_break, answer_page_break,
    load_canonical_trees,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch14'


def _build_diagram_exercises():
    """Load homework exercises from canonical data/trees/ch14/."""
    canonical = load_canonical_trees(14, purpose='homework')
    exercises = []
    for entry in canonical:
        if entry.get('exercise_num') is None:
            continue  # canonical homework items without exercise number — skip
        exercises.append({
            'num': entry['exercise_num'],
            'sentence': entry['sentence'],
            'bracket': entry['bracket'],
            'diagram': entry.get('diagram_filename'),
        })
    exercises.sort(key=lambda x: x['num'])
    return exercises


DIAGRAM_EXERCISES = _build_diagram_exercises()


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 14 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 14: Nominals', cfg, overhead)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part_heading(doc, 'Part 1: Identification and Classification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'I don\u2019t know whether she received my message.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'wh-clause (whether-clause)', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'direct object (of "know")', body_size, font_name=body_font)

    # Exercise 2
    question_page_break(doc, overhead)
    add_exercise(doc, 2, 'The problem is that we lack sufficient funding.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'that-clause', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'subject complement', body_size, font_name=body_font)

    # Exercise 3
    question_page_break(doc, overhead)
    add_exercise(doc, 3, 'To learn a new language requires dedication and practice.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'infinitive phrase', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'subject', body_size, font_name=body_font)

    # Exercise 4
    question_page_break(doc, overhead)
    add_exercise(doc, 4, 'What the scientist discovered changed the field of biology.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'wh-clause', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'subject', body_size, font_name=body_font)

    # Exercise 5
    question_page_break(doc, overhead)
    add_exercise(doc, 5, 'She enjoys reading mystery novels on rainy afternoons.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'gerund phrase', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'direct object (of "enjoys")', body_size, font_name=body_font)

    # Exercise 6
    question_page_break(doc, overhead)
    add_exercise(doc, 6, 'He asked who would be attending the conference.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'wh-clause', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'direct object (of "asked")', body_size, font_name=body_font)

    # Exercise 7
    question_page_break(doc, overhead)
    add_exercise(doc, 7, 'Her greatest fear is making a mistake in public.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'gerund phrase', body_size, font_name=body_font)
    add_answer_line(doc, 'Function:', 'subject complement', body_size, font_name=body_font)

    # =============================================
    # Part 2: Functional Analysis
    # =============================================
    add_part_heading(doc, 'Part 2: Functional Analysis', cfg, overhead)

    functions = [
        (8, 'That the project failed disappointed everyone.',
         'subject',
         'The that-clause is the subject of "disappointed."'),
        (9, 'The committee discussed how they would proceed.',
         'direct object',
         'The wh-clause is the direct object of "discussed."'),
        (10, 'She\u2019s interested in learning more about linguistics.',
         'object of preposition',
         'The gerund phrase is the object of the preposition "in."'),
        (11, 'The main issue is whether we should continue.',
         'subject complement',
         'The wh-clause follows the linking verb "is" and renames "the main issue."'),
        (12, 'I appreciate your helping us with the move.',
         'direct object',
         'The gerund phrase (with possessive) is the direct object of "appreciate."'),
    ]

    for i, (num, sentence, function, explanation) in enumerate(functions):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Function:', function, body_size, font_name=body_font)
        add_plain_line(doc, explanation, body_size, font_name=body_font)

    # =============================================
    # Part 3: Sentence Completion
    # =============================================
    add_part_heading(doc, 'Part 3: Sentence Completion', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 13\u201317 are open-ended. Accept any grammatically correct nominal of the requested type.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    completions = [
        (13, 'Gerund phrase as subject: __________ can be challenging for new employees.',
         '"Learning new software can be challenging for new employees."'),
        (14, 'Wh-clause as direct object: The detective investigated __________.',
         '"The detective investigated who had access to the building."'),
        (15, 'Infinitive phrase as subject complement: Her goal this year is __________.',
         '"Her goal this year is to complete her dissertation."'),
        (16, 'That-clause as subject: __________ surprised everyone at the meeting.',
         '"That the CEO resigned surprised everyone at the meeting."'),
        (17, 'Gerund phrase as object of preposition: She succeeded by __________.',
         '"She succeeded by studying consistently throughout the semester."'),
    ]

    for i, (num, prompt, sample) in enumerate(completions):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, prompt, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_plain_line(doc, f'Sample: {sample}', body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Nominals
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Nominals', cfg, overhead)

    ch_roles = load_chapter_roles(14)
    mode = 'overhead' if overhead else 'answer_key'

    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'], roles_dict=ch_roles.get(bracket_key), mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], body_size, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part_heading(doc, 'Part 5: Analysis and Application', cfg, overhead)

    # Exercise 23
    add_exercise(doc, 23, 'Explain the grammatical and meaning differences between these pairs.', body_size, font_name=body_font)
    add_plain_line(doc, '23A) "She stopped smoking." vs. 23B) "She stopped to smoke."', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_plain_line(doc,
        'Grammatical difference: In (a), "smoking" is a gerund \u2014 it functions as '
        'the direct object of "stopped." In (b), "to smoke" is an infinitive phrase \u2014 '
        'it functions as an adverbial of purpose.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Meaning difference: (a) means she quit the habit of smoking. '
        '(b) means she paused what she was doing in order to have a smoke.',
        body_size, font_name=body_font)

    # Exercise 24 \u2014 Prose Impact (paragraph swap)
    question_page_break(doc, overhead)
    add_exercise(doc, 24,
        'Read Paragraphs A and B (same events, different nominal forms) and '
        'answer questions a\u2013e.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Paragraph A: "The decision divided the staff. The chairperson called a vote. '
        'The outcome surprised the team. The minority report contained the strongest '
        'arguments. The press release was a careful attempt at unity."',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Paragraph B: "That the staff was divided became clear at once. Calling a vote '
        'was the chairperson\u2019s only option. What happened next surprised the team. '
        'Reading the minority report revealed the strongest arguments. To project unity '
        'required a careful press release."',
        body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_plain_line(doc,
        '(a) Sample identifications from Paragraph A: '
        '"The decision" \u2014 NP, subject of "divided"; '
        '"The staff" \u2014 NP, direct object of "divided"; '
        '"a vote" \u2014 NP, direct object of "called"; '
        '"The outcome" \u2014 NP, subject of "surprised". '
        'Accept any two correctly labeled nominals.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '(b) Sample identifications from Paragraph B: '
        '"That the staff was divided" \u2014 that-clause, subject of "became"; '
        '"Calling a vote" \u2014 gerund phrase, subject of "was"; '
        '"What happened next" \u2014 wh-clause, subject of "surprised"; '
        '"Reading the minority report" \u2014 gerund phrase, subject of "revealed"; '
        '"To project unity" \u2014 infinitive phrase, subject of "required". '
        'Accept any two correctly labeled nominals.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '(c) Paragraph A reads in short, even bursts: noun-verb-noun, noun-verb-noun. '
        'Paragraph B is heavier and slower in its subjects \u2014 each opens with a '
        'phrase or clause that has to be parsed before the main verb arrives. '
        'Reading aloud, B has more pause and more cognitive load up front.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '(d) Paragraph A foregrounds entities and outcomes \u2014 the decision, the '
        'staff, the vote, the outcome, the report, the press release. The actions are '
        'compressed into single verbs. Paragraph B foregrounds actions and processes '
        '\u2014 being divided, calling, what happened, reading, projecting unity. '
        'The verbs of the original sentences become the heads of the new subjects. '
        'Example contrast: "The chairperson called a vote" (entity-foregrounded) vs. '
        '"Calling a vote was the chairperson\u2019s only option" (process-foregrounded).',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '(e) Paragraph A would feel at home in a news report: short sentences, concrete '
        'entities, fast pace. Paragraph B would feel at home in an analytical essay or '
        'formal article: clausal and gerundive subjects let the writer treat actions and '
        'states as topics for analysis, and the heavier rhythm signals reflection rather '
        'than reporting. Either is acceptable for either genre when used deliberately, '
        'but the default expectations of those genres lean A and B respectively.',
        body_size, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 14 Student Homework with all 5 parts, mirroring the HTML."""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12
    fn = 'Garamond'

    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 14 Homework: Nominals')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = fn
    set_paragraph_spacing(p, space_before=0, space_after=4)

    def add_part(title):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=10, space_after=4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = fn

    def add_text(text, bold=False, italic=False, indent=0, space_before=3, space_after=2):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn
        run.bold = bold
        run.italic = italic

    def add_ex(num, sentence, italic=True):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {num}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = fn
        run = p.add_run(sentence)
        run.italic = italic
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_blank(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=1, space_after=1)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_ref_table(headers, rows):
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(fs - 1)
            run.font.name = fn
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(fs - 1)
                run.font.name = fn
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=0, space_after=2)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part('Part 1: Identification and Classification (approx. 10 minutes)')
    add_text('For each sentence, identify the underlined nominal, classify its form, '
             'and identify its function in the sentence. Use the reference tables below.')

    add_ref_table(['Nominal Forms', 'Example'], [
        ['Noun Phrase (NP)', 'the interesting book'],
        ['Pronoun', 'she, him, them'],
        ['Gerund Phrase', 'swimming every morning'],
        ['Infinitive Phrase', 'to win the race'],
        ['That-clause', 'that she resigned'],
        ['Wh-clause', 'what she said, whether he comes'],
    ])

    add_ref_table(['Nominal Functions', 'Diagnostic'], [
        ['Subject', 'Sits before the main verb; "who/what is doing or being something?"'],
        ['Direct Object', 'Follows a transitive verb; "verb what?" / "verb whom?"'],
        ['Indirect Object', 'Sits between the verb and the direct object; "to whom?"'],
        ['Object of Preposition', 'Follows a preposition; completes the prepositional phrase'],
        ['Subject Complement', 'Follows a linking verb (be, seem, become); renames the subject'],
    ])

    add_text('Example (completed)', bold=True)
    add_text('Swimming every morning has improved my health.', italic=True)
    add_blank('   Form: gerund phrase')
    add_blank('   Function: subject')

    add_text('Exercises', bold=True)

    part1_exercises = [
        (1, 'I don’t know whether she received my message.'),
        (2, 'The problem is that we lack sufficient funding.'),
        (3, 'To learn a new language requires dedication and practice.'),
        (4, 'What the scientist discovered changed the field of biology.'),
        (5, 'She enjoys reading mystery novels on rainy afternoons.'),
        (6, 'He asked who would be attending the conference.'),
        (7, 'Her greatest fear is making a mistake in public.'),
    ]
    for num, sentence in part1_exercises:
        add_ex(num, sentence)
        add_blank('   Form: __________')
        add_blank('   Function: __________')

    # =============================================
    # Part 2: Functional Analysis
    # =============================================
    add_part('Part 2: Functional Analysis (approx. 5 minutes)')
    add_text('Identify the function of each underlined nominal. Use the reference table below.')

    add_ref_table(['Function', 'Position in the Clause'], [
        ['Subject', 'Before the main verb'],
        ['Direct Object', 'After a transitive verb (the thing acted upon)'],
        ['Indirect Object', 'Between the verb and the direct object (the recipient)'],
        ['Object of Preposition', 'After a preposition (about, in, by, for, etc.)'],
        ['Subject Complement', 'After a linking verb (be, seem, become)'],
    ])

    add_text('Example (completed)', bold=True)
    add_text('She believes that honesty is important.', italic=True)
    add_blank('   Function: direct object (of the verb "believes")')

    add_text('Exercises', bold=True)

    part2_exercises = [
        (8, 'That the project failed disappointed everyone.'),
        (9, 'The committee discussed how they would proceed.'),
        (10, 'She’s interested in learning more about linguistics.'),
        (11, 'The main issue is whether we should continue.'),
        (12, 'I appreciate your helping us with the move.'),
    ]
    for num, sentence in part2_exercises:
        add_ex(num, sentence)
        add_blank('   Function: __________')

    # =============================================
    # Part 3: Sentence Completion
    # =============================================
    add_part('Part 3: Sentence Completion (approx. 5 minutes)')
    add_text('Complete each sentence with the requested nominal structure.')

    add_text('Example (completed)', bold=True)
    add_text('Add a that-clause as direct object: The scientists discovered __________.')
    add_text('Answer: The scientists discovered that the cells could regenerate.', italic=True)

    add_text('Exercises', bold=True)

    part3_exercises = [
        (13, 'Add a gerund phrase as subject: __________ can be challenging for new employees.'),
        (14, 'Add a wh-clause as direct object: The detective investigated __________.'),
        (15, 'Add an infinitive phrase as subject complement: Her goal this year is __________.'),
        (16, 'Add a that-clause as subject: __________ surprised everyone at the meeting.'),
        (17, 'Add a gerund phrase as object of a preposition: She succeeded by __________.'),
    ]
    for num, prompt in part3_exercises:
        add_ex(num, prompt, italic=False)
        add_blank('   Answer: __________')

    # =============================================
    # Part 4: Diagramming Nominals
    # =============================================
    add_part('Part 4: Diagramming Nominals')
    add_text('For each sentence, complete the labeling table (Role, Phrase, Word, POS), '
             'write the bracket notation, and draw a tree diagram.')

    for ex in DIAGRAM_EXERCISES:
        add_ex(ex['num'], ex['sentence'])

        table_data = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, table_data, mode='student', font_size=10)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = fn

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part('Part 5: Analysis and Application (approx. 10 minutes)')
    add_text('Read the sentences and answer the questions.')

    # Exercise 23 — gerund vs infinitive (stop)
    add_text('Gerund vs. Infinitive', bold=True)
    add_ex(23, 'Consider the verb "stop" in these two sentences:', italic=False)
    add_blank('   a) She stopped smoking.')
    add_blank('   b) She stopped to smoke.')
    add_blank('   Explain the grammatical difference (what form follows "stop" in each?) '
              'and the meaning difference:')
    add_blank('   Grammatical difference: __________')
    add_blank('   Meaning difference: __________')

    # Exercise 24 — Prose Impact (paragraph swap)
    add_text('Prose Impact: Choosing Among Nominal Forms', bold=True)
    add_ex(24,
           'The two paragraphs below describe the same sequence of events. '
           'Paragraph A uses mostly noun phrases and pronouns in its nominal slots. '
           'Paragraph B fills the same slots with that-clauses, gerund phrases, '
           'infinitive phrases, and wh-clauses. Read both, then answer the questions '
           'that follow.', italic=False)

    add_text('Paragraph A.', bold=True)
    add_text(
        'The decision divided the staff. The chairperson called a vote. The outcome '
        'surprised the team. The minority report contained the strongest arguments. '
        'The press release was a careful attempt at unity.',
        italic=True)

    add_text('Paragraph B.', bold=True)
    add_text(
        'That the staff was divided became clear at once. Calling a vote was the '
        'chairperson’s only option. What happened next surprised the team. '
        'Reading the minority report revealed the strongest arguments. To project '
        'unity required a careful press release.',
        italic=True)

    add_blank('   a) Identify two nominals from Paragraph A and label each one’s '
              'form (NP, pronoun, etc.) and function (subject, direct object, etc.).')
    add_blank('   b) Identify two nominals from Paragraph B and label each one’s '
              'form (that-clause, gerund phrase, infinitive phrase, wh-clause) and function.')
    add_blank('   c) How does the rhythm or pacing of the two paragraphs differ? '
              'Read each one aloud if it helps you decide.')
    add_blank('   d) One paragraph foregrounds entities and outcomes; the other '
              'foregrounds actions and processes. Which is which? Point to a specific '
              'sentence in each that supports your answer.')
    add_blank('   e) Which version would feel more at home in a quickly written news '
              'report? Which in an analytical essay or formal article? Explain why— '
              'appeal to specific features of the nominals used.')

    add_text('Total estimated time: 40 minutes', italic=True, space_before=12)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 14 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 14 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 14 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Generate Chapter 10 Answer Key and Overhead Answer Key .docx files.
Updated to match revised homework: Part 4 is now Diagramming Verb Phrases,
Part 5 exercises renumbered 21-23.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, add_bracket_line, add_diagram_image, setup_document,
    add_title_page, add_part_heading, exercise_separator, get_font_config,
    add_multilevel_from_bracket, load_chapter_roles,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch10'

PASSAGE_21 = (
    'Maria moved to Boston in 2018. She has lived there ever since. When I visited her last summer, '
    'she was working on her dissertation. She has been writing it for two years now. By next June, '
    'she will have finished the entire project. After that, she will be looking for a teaching position.'
)


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 10 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']
    bracket_size = cfg['bracket_size']
    diagram_width = cfg['diagram_width']

    add_title_page(doc, 'Chapter 10: Verbs Part One \u2014 Tense and Aspect', cfg, overhead)

    # =============================================
    # Part 1: Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Identification', cfg, overhead)

    exercises_p1 = [
        (1, 'The researchers have analyzed the experimental data.',
         [('Auxiliary verb(s):', 'have'),
          ('Main verb:', 'analyzed'),
          ('Tense:', 'present'),
          ('Aspect:', 'perfect')]),
        (2, 'Yesterday, she was working in the library when I called.',
         [('Auxiliary verb(s):', 'was'),
          ('Main verb:', 'working'),
          ('Tense:', 'past'),
          ('Aspect:', 'progressive')]),
        (3, 'By next month, they will have completed the entire project.',
         [('Auxiliary verb(s):', 'will, have'),
          ('Main verb:', 'completed'),
          ('Tense:', 'future'),
          ('Aspect:', 'perfect')]),
        (4, 'The professor teaches linguistics every semester.',
         [('Auxiliary verb(s):', 'none'),
          ('Main verb:', 'teaches'),
          ('Tense:', 'present'),
          ('Aspect:', 'simple')]),
        (5, 'The students had been studying for three hours before the test began.',
         [('Auxiliary verb(s):', 'had, been'),
          ('Main verb:', 'studying'),
          ('Tense:', 'past'),
          ('Aspect:', 'perfect progressive')]),
    ]

    for i, (num, sentence, answers) in enumerate(exercises_p1):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        for label, answer in answers:
            add_answer_line(doc, label, answer, body_size, font_name=body_font)

    # =============================================
    # Part 2: Sentence Completion
    # =============================================
    add_part_heading(doc, 'Part 2: Sentence Completion', cfg, overhead)

    completions = [
        (6, 'Present progressive: Right now, the children ________ (play) in the park.',
         'are playing'),
        (7, 'Past perfect: By the time I arrived, they ________ (already / leave).',
         'had already left'),
        (8, 'Present perfect progressive: She ________ (work) on this project for six months.',
         'has been working'),
        (9, 'Future progressive: At noon tomorrow, I ________ (meet) with the committee.',
         'will be meeting'),
        (10, 'Past simple: The team ________ (finish) the assignment last night.',
         'finished'),
    ]

    for i, (num, prompt, answer) in enumerate(completions):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, prompt, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Answer:', answer, body_size, font_name=body_font)

    # =============================================
    # Part 3: Sentence Writing
    # =============================================
    add_part_heading(doc, 'Part 3: Sentence Writing', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 11\u201315 are open-ended. Accept any grammatically correct sentence that demonstrates the requested tense-aspect combination.')
    run.font.size = body_size
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    writing = [
        (11, 'Write a sentence in present perfect that shows an experience up to now.',
         'Present perfect (experience up to now)',
         '"She has traveled to Japan three times."'),
        (12, 'Write a sentence in past progressive that describes a background action interrupted by another event.',
         'Past progressive (background action interrupted)',
         '"I was reading when the doorbell rang."'),
        (13, 'Write a sentence in future perfect that describes an action completed before a future point.',
         'Future perfect (action completed before future point)',
         '"By December, we will have finished the renovation."'),
        (14, 'Write a sentence in present simple that expresses a general truth.',
         'Present simple (general truth)',
         '"The Earth revolves around the Sun."'),
        (15, 'Write a sentence in past perfect that positions one past event before another.',
         'Past perfect (one past event before another)',
         '"By the time the ambulance arrived, the patient had already recovered."'),
    ]

    for i, (num, question, structure, sample) in enumerate(writing):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, question, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_plain_line(doc, f'{structure}:', body_size, bold_prefix='Structure: ', font_name=body_font)
        add_plain_line(doc, f'Sample: {sample}', body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Verb Phrases
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Verb Phrases', cfg, overhead)

    diagram_exercises = [
        {'num': 16, 'sentence': 'The students are studying for the exam.',
         'bracket': '[S [NP [DET The] [N students]] [VP [AUX are] [V studying] [PP [PREP for] [NP [DET the] [N exam]]]]]',
         'tense': 'present', 'aspect': 'progressive',
         'diagram': 'ch10_hw_ex16_students_studying'},
        {'num': 17, 'sentence': 'He had finished the assignment before class.',
         'bracket': '[S [NP [PRON He]] [VP [AUX had] [V finished] [NP [DET the] [N assignment]] [PP [PREP before] [NP [N class]]]]]',
         'tense': 'past', 'aspect': 'perfect',
         'diagram': 'ch10_hw_ex17_had_finished'},
        {'num': 18, 'sentence': 'Does the professor teach on Fridays?',
         'bracket': '[S [AUX Does] [NP [DET the] [N professor]] [VP [V teach] [PP [PREP on] [NP [N Fridays]]]]]',
         'tense': 'present', 'aspect': 'simple (do-support)',
         'diagram': 'ch10_hw_ex18_does_teach'},
        {'num': 19, 'sentence': 'The report was written by the committee.',
         'bracket': '[S [NP [DET The] [N report]] [VP [AUX was] [V written] [PP [PREP by] [NP [DET the] [N committee]]]]]',
         'tense': 'past', 'aspect': 'passive voice',
         'diagram': 'ch10_hw_ex19_was_written'},
        {'num': 20, 'sentence': 'They have been waiting at the station for an hour.',
         'bracket': '[S [NP [PRON They]] [VP [AUX have] [AUX been] [V waiting] [PP [PREP at] [NP [DET the] [N station]]] [PP [PREP for] [NP [DET an] [N hour]]]]]',
         'tense': 'present', 'aspect': 'perfect progressive',
         'diagram': 'ch10_hw_ex20_have_been_waiting'},
    ]

    ch_roles = load_chapter_roles(10)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(diagram_exercises):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], bracket_size)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=diagram_width)
        add_answer_line(doc, 'Tense:', ex['tense'], body_size, font_name=body_font)
        add_answer_line(doc, 'Aspect:', ex['aspect'], body_size, font_name=body_font)

    # =============================================
    # Part 5: Contextual Analysis
    # =============================================
    add_part_heading(doc, 'Part 5: Contextual Analysis', cfg, overhead)

    # Exercise 21
    add_exercise(doc, 21, 'Identify the tense-aspect of each verb phrase in the passage.', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Passage: ')
    run.bold = True
    run.font.size = body_size
    run.font.name = body_font
    run = p.add_run(PASSAGE_21)
    run.italic = True
    run.font.size = body_size
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=4)

    answer_page_break(doc, overhead)

    verb_phrases = [
        ('moved:', 'past simple'),
        ('has lived:', 'present perfect'),
        ('visited:', 'past simple'),
        ('was working:', 'past progressive'),
        ('has been writing:', 'present perfect progressive'),
        ('will have finished:', 'future perfect'),
        ('will be looking:', 'future progressive'),
    ]

    for verb, tense_aspect in verb_phrases:
        add_answer_line(doc, verb, tense_aspect, body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 22
    add_exercise(doc, 22,
        'The passage uses both "moved" (past simple) and "has lived" (present perfect). '
        'Both refer to events that began in 2018. Explain why the writer chose different tense-aspects for these two verbs.',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    add_plain_line(doc,
        '"Moved" (past simple) presents the action as a completed event in the past \u2014 the move happened '
        'and is over. "Has lived" (present perfect) connects the past event to the present \u2014 she moved '
        'in 2018 and STILL lives there now. The writer uses past simple for the completed action of moving '
        'and present perfect for the ongoing state of living there, because the living continues into the present.',
        body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 23
    add_exercise(doc, 23,
        'Rewrite the following sentence in three different tense-aspect combinations and explain how the meaning changes with each: "She studies linguistics."',
        body_size, font_name=body_font)

    answer_page_break(doc, overhead)

    for sub, rewrite, explanation in [
        ('23A) Past progressive:',
         '"She was studying linguistics."',
         'Changes from a habitual/general statement to a temporary, ongoing activity at a specific past moment.'),
        ('23B) Present perfect:',
         '"She has studied linguistics."',
         'Changes from a current habit to a completed experience with present relevance (she has this knowledge now).'),
        ('23C) Future perfect:',
         '"She will have studied linguistics (by graduation)."',
         'Projects the activity into the future as something that will be completed before a reference point.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(sub)
        run.bold = True
        run.font.size = body_size
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=3, space_after=2)

        add_plain_line(doc, f'Rewrite: {rewrite}', body_size, indent=0.7, font_name=body_font)
        add_plain_line(doc, f'Meaning change: {explanation}', body_size, indent=0.7, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 10 Student Homework with blank multi-level tables for Part 4."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 10 Homework: Verbs Part One \u2014 Tense and Aspect')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Part 4 with blank multi-level tables
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Diagramming Verb Phrases')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    diagram_exercises = [
        {'num': 16, 'sentence': 'The students are studying for the exam.',
         'bracket': '[S [NP [DET The] [N students]] [VP [AUX are] [V studying] [PP [PREP for] [NP [DET the] [N exam]]]]]'},
        {'num': 17, 'sentence': 'He had finished the assignment before class.',
         'bracket': '[S [NP [PRON He]] [VP [AUX had] [V finished] [NP [DET the] [N assignment]] [PP [PREP before] [NP [N class]]]]]'},
        {'num': 18, 'sentence': 'Does the professor teach on Fridays?',
         'bracket': '[S [AUX Does] [NP [DET the] [N professor]] [VP [V teach] [PP [PREP on] [NP [N Fridays]]]]]'},
        {'num': 19, 'sentence': 'The report was written by the committee.',
         'bracket': '[S [NP [DET The] [N report]] [VP [AUX was] [V written] [PP [PREP by] [NP [DET the] [N committee]]]]]'},
        {'num': 20, 'sentence': 'They have been waiting at the station for an hour.',
         'bracket': '[S [NP [PRON They]] [VP [AUX have] [AUX been] [V waiting] [PP [PREP at] [NP [DET the] [N station]]] [PP [PREP for] [NP [DET an] [N hour]]]]]'},
    ]

    for ex in diagram_exercises:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 10 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 10 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 10 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

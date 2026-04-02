#!/usr/bin/env python3
"""
Generate Chapter 11 homework files: Student Homework, Answer Key, and Overhead.

All Part 4 labeling table data is defined once (DIAGRAM_EXERCISES) and reused
across all three outputs with proper cell merging.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page, add_part_heading,
    exercise_separator, get_font_config, add_bracket_line, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles,
    parse_bracket_to_multilevel, add_multilevel_labeling_table,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch11'


# =============================================================================
# SHARED EXERCISE DATA — single source of truth for all outputs
# =============================================================================

PART1_EXERCISES = [
    {'num': 1, 'sentence': 'The researchers carefully analyzed the data.',
     'voice': 'active', 'actor': None, 'active_version': None},
    {'num': 2, 'sentence': 'Three errors were discovered in the code.',
     'voice': 'passive', 'actor': 'none stated',
     'active_version': 'Someone discovered three errors in the code.'},
    {'num': 3, 'sentence': 'The new policy will be announced tomorrow.',
     'voice': 'passive', 'actor': 'none stated',
     'active_version': 'Someone/They will announce the new policy tomorrow.'},
    {'num': 4, 'sentence': 'Someone stole my bicycle last night.',
     'voice': 'active', 'actor': None, 'active_version': None},
    {'num': 5, 'sentence': 'The building was constructed in 1920.',
     'voice': 'passive', 'actor': 'none stated',
     'active_version': 'Someone/They constructed the building in 1920.'},
]

PART2_EXERCISES = [
    (6, 'Active to passive: The team is preparing the presentation.',
     'The presentation is being prepared by the team.'),
    (7, 'Active to passive: Someone had stolen the documents before the investigation began.',
     'The documents had been stolen before the investigation began.'),
    (8, 'Passive to active: The experiment was conducted by the research team.',
     'The research team conducted the experiment.'),
    (9, 'Passive to active: The proposal will be reviewed by the committee next week.',
     'The committee will review the proposal next week.'),
    (10, 'Active to passive: The company will hire fifty new employees.',
     'Fifty new employees will be hired by the company.'),
]

PART3_MODALS = [
    (11, 'She can speak three languages fluently.',
     [('Modal:', 'can'), ('Meaning:', 'ability')]),
    (12, 'That might be the correct answer, but I\u2019m not certain.',
     [('Modal:', 'might'), ('Meaning:', 'possibility')]),
    (13, 'You should apologize for your mistake.',
     [('Modal:', 'should'), ('Meaning:', 'advice')]),
    (14, 'He must be exhausted after running the marathon.',
     [('Modal:', 'must'), ('Meaning:', 'deduction')]),
    (15, 'May I leave the room early?',
     [('Modal:', 'may'), ('Meaning:', 'permission')]),
    (16, 'They could have won the game if they had practiced more.',
     [('Modal:', 'could (have)'), ('Meaning:', 'past possibility (unrealized)')]),
]

# Part 4: Diagramming — labeling table data defined once, used everywhere
DIAGRAM_EXERCISES = [
    {
        'num': 18,
        'sentence': 'The letter was delivered yesterday.',
        'words':   ['The', 'letter', 'was', 'delivered', 'yesterday'],
        'roles':   ['Subj', '', 'Pred', '', 'Advl'],
        'phrases': ['NP', '', 'VP', '', 'ADVP'],
        'pos':     ['DET', 'N', 'AUX', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [N letter]] [VP [AUX was] [V delivered] [ADVP [ADV yesterday]]]]',
        'diagram': 'ch11_hw_ex18_was_delivered',
    },
    {
        'num': 19,
        'sentence': 'She must leave before noon.',
        'words':   ['She', 'must', 'leave', 'before', 'noon'],
        'roles':   ['Subj', 'Pred', '', 'Advl', ''],
        'phrases': ['NP', 'VP', '', 'PP', ''],
        'pos':     ['PRON', 'MOD', 'V', 'PREP', 'N'],
        'bracket': '[S [NP [PRON She]] [VP [MOD must] [V leave] [PP [PREP before] [NP [N noon]]]]]',
        'diagram': 'ch11_hw_ex19_must_leave',
    },
    {
        'num': 20,
        'sentence': 'The report can be finished tomorrow.',
        'words':   ['The', 'report', 'can', 'be', 'finished', 'tomorrow'],
        'roles':   ['Subj', '', 'Pred', '', '', 'Advl'],
        'phrases': ['NP', '', 'VP', '', '', 'ADVP'],
        'pos':     ['DET', 'N', 'MOD', 'AUX', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [N report]] [VP [MOD can] [AUX be] [V finished] [ADVP [ADV tomorrow]]]]',
        'diagram': 'ch11_hw_ex20_can_be_finished',
    },
    {
        'num': 21,
        'sentence': 'He should have called earlier.',
        'words':   ['He', 'should', 'have', 'called', 'earlier'],
        'roles':   ['Subj', 'Pred', '', '', 'Advl'],
        'phrases': ['NP', 'VP', '', '', 'ADVP'],
        'pos':     ['PRON', 'MOD', 'AUX', 'V', 'ADV'],
        'bracket': '[S [NP [PRON He]] [VP [MOD should] [AUX have] [V called] [ADVP [ADV earlier]]]]',
        'diagram': 'ch11_hw_ex21_should_have_called',
    },
]

EXAMPLE_DIAGRAM = {
    'sentence': 'She can swim.',
    'words':   ['She', 'can', 'swim'],
    'roles':   ['Subj', 'Pred', ''],
    'phrases': ['NP', 'VP', ''],
    'pos':     ['PRON', 'MOD', 'V'],
    'bracket': '[S [NP [PRON She]] [VP [MOD can] [V swim]]]',
    'diagram': 'ch11_hw_example_can_swim',
}

PART5_PASSIVES = [
    ('was announced (yesterday by the CEO)',
     'Focuses on the policy (the topic) rather than the CEO; maintains topic continuity.'),
    ('will be made (after all responses have been reviewed)',
     'Actor is unspecified, emphasizing the process and the decision rather than who will make it.'),
    ('have been reviewed',
     'Embedded passive in the subordinate clause; keeps "responses" as the focus.'),
]

PART5_MODALS = [
    ('must (submit)', 'obligation/permission \u2014 employees are required to submit feedback'),
    ('should (improve)', 'certainty/possibility \u2014 management expects changes will improve efficiency'),
    ('might (create)', 'certainty/possibility \u2014 workers think it is possible the policy will create challenges'),
    ('will (be made)', 'certainty/possibility \u2014 prediction about the future'),
]

PASSAGE = ('The new policy was announced yesterday by the CEO. All employees must submit '
           'their feedback by next Friday. According to management, the changes should improve '
           'workplace efficiency. Some workers believe the policy might create additional '
           'challenges. However, the final decision will be made after all responses have '
           'been reviewed.')


# =============================================================================
# STUDENT HOMEWORK GENERATOR
# =============================================================================

def create_student_homework(output_path):
    """Create the complete Chapter 11 Student Homework with proper merged tables."""
    doc = Document()

    # Basic styling
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12  # body font size

    # Title
    p = doc.add_paragraph()
    run = p.add_run('Chapter 11 Homework: Voice and Modals')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    p = doc.add_paragraph()
    run = p.add_run('Total estimated time: 45 minutes')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    # --- Part 1 ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 1: Voice Identification')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('For each sentence, identify whether it is in active or passive voice. '
                     'If passive, identify the actor (if present) and the original active construction.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    # Example
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Example (completed)')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('The report was written by the committee.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for line in ['Voice: passive', 'Actor: the committee',
                 'Active version: The committee wrote the report.']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(line)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # Exercises
    for ex in PART1_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        blanks = ['Voice: _____']
        if ex['voice'] == 'passive':
            blanks.append('Actor (if present): _____')
        for line in blanks:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(line)
            run.font.size = Pt(fs)
            run.font.name = 'Garamond'

    # --- Part 2 ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 2: Voice Transformation')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Convert each sentence to the opposite voice (active to passive or passive to active). '
                     'Maintain the same tense and aspect.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    # Example
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Example (completed)')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Active to passive: ')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('The architect has designed the new building.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Answer: The new building has been designed by the architect.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for num, prompt, _ in PART2_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {num}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(prompt)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run('Answer: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # --- Part 3 ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 3: Modal Meaning')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('For each sentence, identify the modal and classify its meaning '
                     'using one of the following categories:')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for cat in ['Ability', 'Possibility', 'Permission', 'Obligation/Necessity', 'Deduction', 'Advice']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(f'\u2022 {cat}')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # Example
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Example (completed)')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('You must submit the application by Friday.')
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for line in ['Modal: must', 'Meaning: obligation/necessity']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(line)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    for num, sentence, _ in PART3_MODALS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {num}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(sentence)
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        for line in ['Modal: _____', 'Meaning: _____']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(line)
            run.font.size = Pt(fs)
            run.font.name = 'Garamond'

    # Exercise 17 (special)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Exercise 17. ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Explain the difference between the two uses of must:')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for line in ['a) You must wear a seatbelt. (Meaning type: _____)',
                 'b) She\u2019s not answering the phone. She must be asleep. (Meaning type: _____)']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(line)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # --- Part 4: Diagramming (with proper merged tables) ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Diagramming Voice and Modals')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('For each sentence, complete the labeling table, write the bracket notation, '
                     'and draw a tree diagram. Use the abbreviations from the Diagram Examples section '
                     '(MOD for modals, AUX for auxiliaries, V for main verbs).')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    # Completed example
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Example (completed)')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run(EXAMPLE_DIAGRAM['sentence'])
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    ex_td = parse_bracket_to_multilevel(EXAMPLE_DIAGRAM['bracket'])
    add_multilevel_labeling_table(doc, ex_td, mode='answer_key', font_size=10)

    p = doc.add_paragraph()
    run = p.add_run('Bracket notation: ')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run(EXAMPLE_DIAGRAM['bracket'])
    run.font.name = 'Consolas'
    run.font.size = Pt(11)

    add_diagram_image(doc, DIAGRAM_DIR, EXAMPLE_DIAGRAM['diagram'], width_inches=5.5)

    # Exercises with blank pre-merged tables
    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=8, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=10)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        p = doc.add_paragraph()
        run = p.add_run('Diagram:')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # --- Part 5: Analysis and Application ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 5: Analysis and Application')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Read the passage below and answer the questions.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=4, space_after=2)
    run = p.add_run('Passage')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run(PASSAGE)
    run.italic = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    # Exercise 22
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Exercise 22. ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Identify all passive voice constructions in the passage. '
                     'For each, explain why the writer might have chosen passive over active voice.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for line in ['Passive construction 1: _____', 'Reason: _____',
                 'Passive construction 2: _____', 'Reason: _____']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(line)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    # Exercise 23
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    run = p.add_run('Exercise 23. ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Identify the modals in the passage and classify each as expressing '
                     'certainty/possibility or obligation/permission:')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for line in ['Modal 1: _____ \u2014 Type: _____',
                 'Modal 2: _____ \u2014 Type: _____',
                 'Modal 3: _____ \u2014 Type: _____']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(line)
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f"Created: {output_path}")


# =============================================================================
# ANSWER KEY / OVERHEAD GENERATOR
# =============================================================================

def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 11 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 11: Verbs Part Two \u2014 Voice and Modals', cfg, overhead)

    # =============================================
    # Part 1: Voice Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Voice Identification', cfg, overhead)

    for i, ex in enumerate(PART1_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Voice:', ex['voice'], body_size, font_name=body_font)
        if ex['actor']:
            add_answer_line(doc, 'Actor:', ex['actor'], body_size, font_name=body_font)
        if ex['active_version']:
            add_plain_line(doc, f'Active version: "{ex["active_version"]}"',
                           body_size, font_name=body_font)

    # =============================================
    # Part 2: Voice Transformation
    # =============================================
    add_part_heading(doc, 'Part 2: Voice Transformation', cfg, overhead)

    for i, (num, prompt, answer) in enumerate(PART2_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, prompt, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Answer:', answer, body_size, font_name=body_font)

    # =============================================
    # Part 3: Modal Meaning
    # =============================================
    add_part_heading(doc, 'Part 3: Modal Meaning', cfg, overhead)

    for i, (num, sentence, answers) in enumerate(PART3_MODALS):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        for label, answer in answers:
            add_answer_line(doc, label, answer, body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 17
    add_exercise(doc, 17, 'Explain the difference between the two uses of must:', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('17A) ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run('You must wear a seatbelt.')
    run.italic = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    answer_page_break(doc, overhead)

    add_plain_line(doc,
        'Meaning type: obligation. The speaker is stating a rule or requirement '
        'that the listener is obligated to follow.',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('17B) ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run('She\u2019s not answering the phone. She must be asleep.')
    run.italic = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)

    answer_page_break(doc, overhead)

    add_plain_line(doc,
        'Meaning type: deduction. The speaker is drawing a logical conclusion '
        'based on evidence (she\u2019s not answering), not imposing an obligation.',
        body_size, indent=0.7, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Voice and Modals
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Voice and Modals', cfg, overhead)

    ch_roles = load_chapter_roles(11)
    mode = 'overhead' if overhead else 'answer_key'

    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)

        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                    roles_dict=ch_roles.get(bracket_key),
                                    mode=mode,
                                    font_size=body_size - 2)

        add_bracket_line(doc, ex['bracket'], body_size, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part_heading(doc, 'Part 5: Analysis and Application', cfg, overhead)

    # Exercise 22
    add_exercise(doc, 22, 'Identify passive voice constructions in the passage:', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for i, (construction, reason) in enumerate(PART5_PASSIVES, 1):
        add_plain_line(doc, f'Passive {i}: "{construction}"', body_size, indent=0.35, font_name=body_font)
        add_plain_line(doc, f'Reason: {reason}', body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 23
    add_exercise(doc, 23, 'Identify modals and classify as certainty/possibility or obligation/permission:', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for modal, classification in PART5_MODALS:
        add_answer_line(doc, f'{modal}:', classification, body_size, indent=0.35, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 11 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 11 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 11 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

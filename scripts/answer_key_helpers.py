#!/usr/bin/env python3
"""
Shared helper functions for generating answer key and overhead .docx files.
All chapter answer key scripts should import from this module.
"""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set paragraph spacing in points."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def add_spacer_row(doc):
    """Add a blank spacer paragraph in Times New Roman 20 (no text, for instructor notes)."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '40')  # 20pt = 40 half-points
    rPr.append(sz)
    pPr.append(rPr)
    set_paragraph_spacing(p, space_before=0, space_after=0)
    return p


def add_exercise(doc, number, sentence, font_size, font_name=None):
    """Add an exercise header with sentence."""
    p = doc.add_paragraph()
    run = p.add_run(f'Exercise {number}. ')
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if sentence:
        run = p.add_run(sentence)
        run.italic = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    set_paragraph_spacing(p, space_before=6, space_after=3)
    return p


def add_answer_line(doc, label, answer, font_size, indent=0.35, font_name=None):
    """Add a label: answer line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    run = p.add_run(answer)
    run.italic = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def add_plain_line(doc, text, font_size, indent=0.35, bold_prefix=None, font_name=None):
    """Add a plain text line with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def add_sub_sentence(doc, sub, sentence, font_size, font_name=None):
    """Add a sub-item like a), b), c) with sentence."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(f'{sub} ')
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if sentence:
        run = p.add_run(sentence)
        run.italic = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    set_paragraph_spacing(p, space_before=3, space_after=2)
    return p


def compute_spans(labels):
    """Convert flat label list to (label, span_count, start_index) tuples.

    Non-empty labels start a new span; consecutive empty strings extend it.
    """
    spans = []
    i = 0
    while i < len(labels):
        label = labels[i]
        if label:
            span = 1
            while i + span < len(labels) and labels[i + span] == "":
                span += 1
            spans.append((label, span, i))
            i += span
        else:
            i += 1
    return spans


def blank_labels(labels):
    """Convert answer labels to blank labels preserving span structure for pre-merging.

    Non-empty labels become a space (visually blank but truthy for compute_spans),
    empty strings stay empty to continue the previous span.
    """
    return ["" if label == "" else " " for label in labels]


def add_labeling_table(doc, words, pos_labels=None, phrase_labels=None,
                       role_labels=None, font_size=10):
    """Add a sentence labeling table with merged cells for Role and Phrase rows.

    When labels are provided for Role/Phrase rows, cells are merged to show
    how words group into phrases and roles. When labels are None (student
    version), cells are left unmerged so students can write in each cell.
    Pass blank_labels(answer_labels) for pre-merged blank cells.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    num_cols = len(words) + 1  # +1 for row headers
    table = doc.add_table(rows=4, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_headers = ["Role", "Phrase", "Word", "POS"]
    row_data = [role_labels, phrase_labels, words, pos_labels]

    for i, (header, data) in enumerate(zip(row_headers, row_data)):
        # Row header cell
        cell = table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        # Shade header cell
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)

        if i <= 1 and data:
            # Role or Phrase row with data — merge cells to show groupings
            spans = compute_spans(data)
            for label, span, start_idx in spans:
                col_start = start_idx + 1  # +1 for row header column
                col_end = col_start + span - 1
                if span > 1:
                    table.rows[i].cells[col_start].merge(table.rows[i].cells[col_end])
                merged_cell = table.rows[i].cells[col_start]
                for paragraph in merged_cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                p = merged_cell.paragraphs[0]
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(label.strip())
                run.font.size = Pt(font_size)
        else:
            # Normal unmerged row (Word, POS, or blank student rows)
            for j, val in enumerate(data if data else [""] * len(words)):
                cell = table.rows[i].cells[j + 1]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 2:  # Word row always filled
                    run = p.add_run(val)
                    run.font.size = Pt(font_size)
                elif data:
                    run = p.add_run(val)
                    run.font.size = Pt(font_size)
                # else leave blank for student version

    return table


def parse_bracket_to_multilevel(bracket):
    """Parse bracket notation into multi-level table data.

    Returns dict with:
        words: list of terminal words
        pos_labels: list of POS labels (one per word)
        levels: list of dicts, each with entries [{label, start_col, end_col}, ...]
                ordered by depth (level 0 = top phrase/clause level)

    Example:
        bracket = "[S [NP [DET The] [N cat]] [VP [V chased] [NP [DET the] [N mouse]]]]"
        Result has 2 levels:
            level 0: NP(0-2), VP(2-5)
            level 1: NP(3-5)
    """
    # Tokenize
    tokens = []
    i = 0
    while i < len(bracket):
        if bracket[i] in '[]':
            tokens.append(bracket[i])
            i += 1
        elif bracket[i].isspace():
            i += 1
        else:
            j = i
            while j < len(bracket) and bracket[j] not in '[] ':
                j += 1
            tokens.append(bracket[j - (j - i):j] if j > i else bracket[i:j])
            tokens.append(bracket[i:j])
            i = j

    # Re-tokenize cleanly
    tokens = []
    i = 0
    while i < len(bracket):
        if bracket[i] in '[]':
            tokens.append(bracket[i])
            i += 1
        elif bracket[i].isspace():
            i += 1
        else:
            j = i
            while j < len(bracket) and bracket[j] not in '[] ':
                j += 1
            tokens.append(bracket[i:j])
            i = j

    # Build tree as nested lists: [label, child1, child2, ...]
    stack = []
    current = None
    for tok in tokens:
        if tok == '[':
            node = []
            if current is not None:
                stack.append(current)
            current = node
        elif tok == ']':
            completed = current
            if stack:
                current = stack.pop()
                current.append(completed)
            else:
                current = completed
        else:
            if current is not None:
                current.append(tok)

    root = current

    # Known POS tags
    POS_TAGS = {'N', 'NOUN', 'V', 'VERB', 'ADJ', 'ADV', 'DET', 'PRON', 'PREP',
                'MOD', 'AUX', 'REL', 'COMP', 'CONJ', 'SUB', 'PART', 'MODAL'}

    # Collect terminals and POS in order, and build phrase/clause node info
    words = []
    pos_labels = []
    nodes = []  # (label, depth, start_col, end_col)

    def is_pos_node(node):
        """A POS node is [TAG word] where TAG is in POS_TAGS and word is a string."""
        return (isinstance(node, list) and len(node) == 2
                and isinstance(node[0], str) and node[0] in POS_TAGS
                and isinstance(node[1], str))

    def walk(node, depth):
        if isinstance(node, str):
            return
        if not isinstance(node, list) or len(node) == 0:
            return

        label = node[0] if isinstance(node[0], str) else None
        children = node[1:] if label else node

        if is_pos_node(node):
            # Terminal: POS + word
            pos_labels.append(label)
            words.append(node[1])
            return

        start_col = len(words)
        for child in children:
            walk(child, depth + 1)
        end_col = len(words)

        if label and end_col > start_col:
            # Skip the root S node (depth 0) — use its children as top level
            CLAUSE_TAGS = {'S', 'IC', 'DC', 'RC', 'CC'}
            if depth == 0 and label in CLAUSE_TAGS:
                return  # Skip root clause, children already processed at depth+1
            nodes.append((label, depth, start_col, end_col))

    walk(root, 0)

    # Normalize depths to start at 0
    if nodes:
        min_depth = min(d for _, d, _, _ in nodes)
        nodes = [(l, d - min_depth, s, e) for l, d, s, e in nodes]

    # Group by depth level
    max_depth = max((d for _, d, _, _ in nodes), default=-1)
    levels = []
    for d in range(max_depth + 1):
        entries = sorted(
            [{'label': l, 'start_col': s, 'end_col': e}
             for l, dep, s, e in nodes if dep == d],
            key=lambda x: x['start_col']
        )
        if entries:
            levels.append(entries)

    return {
        'words': words,
        'pos_labels': pos_labels,
        'levels': levels,
    }


def add_multilevel_labeling_table(doc, table_data, roles=None, font_size=10,
                                   mode='answer_key'):
    """Add a multi-level labeling table matching SyntaxTreeHybrid styling.

    Args:
        doc: python-docx Document
        table_data: dict from parse_bracket_to_multilevel() with words, pos_labels, levels
        roles: dict like {"0": {"0": "Subject", "2": "Predicate"}, "1": {"3": "DO"}}
               Keys are level index (str), values are dicts mapping start_col (str) to label.
        font_size: font size in points
        mode: 'answer_key' (all filled), 'student' (words only, rest blank),
              'overhead' (same as answer_key)

    Returns:
        The table object.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    words = table_data['words']
    pos_labels = table_data['pos_labels']
    levels = table_data['levels']
    num_word_cols = len(words)
    num_cols = num_word_cols + 1  # +1 for row headers
    show_numbers = len(levels) > 1
    is_student = (mode == 'student')

    # Calculate rows: 2 per level (Role + Phrase) + POS + Word
    num_rows = len(levels) * 2 + 2
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # SyntaxTreeHybrid color scheme
    GREY_FILL = 'CCCCCC'
    HEADER_FILL = 'E0E0E0'
    ROLE_FILL = 'FFF3E0'
    PHRASE_FILL = 'E3F2FD'
    CLAUSE_FILL = 'F3E5F5'
    POS_FILL = 'E8F5E9'
    CLAUSE_TAGS = {'S', 'IC', 'DC', 'RC', 'CC'}

    # Border constants (eighth-points): thick=18 (~2.25pt), thin=4 (0.5pt)
    THICK = '18'
    THIN = '4'
    BORDER_COLOR = '222222'
    INNER_COLOR = '333333'

    def shade_cell(cell, color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)

    def set_cell_text(cell, text, bold=False, center=True, font_name='Calibri'):
        cell.text = ""
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Tight cell spacing
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        pPr.append(spacing)
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run.bold = bold

    def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
        """Set individual cell borders. Each param is (size, color) or None."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge, val in [('top', top), ('bottom', bottom),
                          ('start', left), ('end', right)]:
            if val:
                sz, clr = val
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), sz)
                el.set(qn('w:color'), clr)
                el.set(qn('w:space'), '0')
                borders.append(el)
        tcPr.append(borders)

    def apply_borders(row_idx, col_idx, is_level_start=False):
        """Apply SyntaxTreeHybrid border scheme to a cell."""
        cell = table.rows[row_idx].cells[col_idx]
        top = (THICK, BORDER_COLOR) if (row_idx == 0 or is_level_start) else (THIN, INNER_COLOR)
        bottom = (THICK, BORDER_COLOR) if row_idx == num_rows - 1 else (THIN, INNER_COLOR)
        left = (THICK, BORDER_COLOR) if col_idx == 0 else (THIN, INNER_COLOR)
        right = (THICK, BORDER_COLOR) if col_idx == num_cols - 1 else (THIN, INNER_COLOR)
        set_cell_border(cell, top=top, bottom=bottom, left=left, right=right)

    row_idx = 0

    # Level rows (Role + Phrase/Clause pairs)
    for level_idx, entries in enumerate(levels):
        suffix = f' {level_idx + 1}' if show_numbers else ''
        has_clauses = any(e['label'] in CLAUSE_TAGS for e in entries)
        type_label = 'Clause' if has_clauses else 'Phrase'
        is_level_start = (level_idx > 0)  # thick top border on level 2+

        # Role row
        role_cell = table.rows[row_idx].cells[0]
        set_cell_text(role_cell, f'Role{suffix}', bold=True, center=False)
        shade_cell(role_cell, HEADER_FILL)

        # Fill role cells — merge spans, grey gaps
        col = 0
        for entry in entries:
            # Grey gap before
            if entry['start_col'] > col:
                gap_start = col + 1
                gap_end = entry['start_col']
                if gap_end > gap_start:
                    table.rows[row_idx].cells[gap_start].merge(
                        table.rows[row_idx].cells[gap_end])
                merged = table.rows[row_idx].cells[gap_start]
                shade_cell(merged, GREY_FILL)
                merged.text = ""

            # Role cell with span
            col_start = entry['start_col'] + 1
            col_end = entry['end_col']  # exclusive in data, but +1 offset makes it inclusive
            if col_end > col_start:
                table.rows[row_idx].cells[col_start].merge(
                    table.rows[row_idx].cells[col_end])
            merged = table.rows[row_idx].cells[col_start]
            shade_cell(merged, ROLE_FILL)
            # Fill role if provided and not student mode
            role_text = ''
            if not is_student and roles and str(level_idx) in roles:
                role_text = roles[str(level_idx)].get(str(entry['start_col']), '')
            set_cell_text(merged, role_text)
            col = entry['end_col']

        # Grey gap after last entry
        if col < num_word_cols:
            gap_start = col + 1
            gap_end = num_word_cols
            if gap_end > gap_start:
                table.rows[row_idx].cells[gap_start].merge(
                    table.rows[row_idx].cells[gap_end])
            merged = table.rows[row_idx].cells[gap_start]
            shade_cell(merged, GREY_FILL)
            merged.text = ""

        # Apply borders to all cells in this row
        for ci in range(num_cols):
            apply_borders(row_idx, ci, is_level_start=is_level_start)
        row_idx += 1

        # Phrase/Clause row
        phrase_cell = table.rows[row_idx].cells[0]
        set_cell_text(phrase_cell, f'{type_label}{suffix}', bold=True, center=False)
        shade_cell(phrase_cell, HEADER_FILL)

        col = 0
        for entry in entries:
            if entry['start_col'] > col:
                gap_start = col + 1
                gap_end = entry['start_col']
                if gap_end > gap_start:
                    table.rows[row_idx].cells[gap_start].merge(
                        table.rows[row_idx].cells[gap_end])
                merged = table.rows[row_idx].cells[gap_start]
                shade_cell(merged, GREY_FILL)
                merged.text = ""

            col_start = entry['start_col'] + 1
            col_end = entry['end_col']
            if col_end > col_start:
                table.rows[row_idx].cells[col_start].merge(
                    table.rows[row_idx].cells[col_end])
            merged = table.rows[row_idx].cells[col_start]
            fill = CLAUSE_FILL if entry['label'] in CLAUSE_TAGS else PHRASE_FILL
            shade_cell(merged, fill)
            # In student mode, phrase/clause labels are blank
            label_text = '' if is_student else entry['label']
            set_cell_text(merged, label_text)
            col = entry['end_col']

        if col < num_word_cols:
            gap_start = col + 1
            gap_end = num_word_cols
            if gap_end > gap_start:
                table.rows[row_idx].cells[gap_start].merge(
                    table.rows[row_idx].cells[gap_end])
            merged = table.rows[row_idx].cells[gap_start]
            shade_cell(merged, GREY_FILL)
            merged.text = ""

        for ci in range(num_cols):
            apply_borders(row_idx, ci)
        row_idx += 1

    # POS row
    pos_header = table.rows[row_idx].cells[0]
    set_cell_text(pos_header, 'POS', bold=True, center=False)
    shade_cell(pos_header, HEADER_FILL)
    for j, pos in enumerate(pos_labels):
        cell = table.rows[row_idx].cells[j + 1]
        pos_text = '' if is_student else pos
        set_cell_text(cell, pos_text)
        shade_cell(cell, POS_FILL)
    for ci in range(num_cols):
        apply_borders(row_idx, ci, is_level_start=True)
    row_idx += 1

    # Word row
    word_header = table.rows[row_idx].cells[0]
    set_cell_text(word_header, 'Word', bold=True, center=False)
    shade_cell(word_header, HEADER_FILL)
    for j, word in enumerate(words):
        cell = table.rows[row_idx].cells[j + 1]
        set_cell_text(cell, word, bold=True)
    for ci in range(num_cols):
        apply_borders(row_idx, ci)

    return table


def add_diagram_image(doc, diagram_dir, image_name, width_inches=5.5):
    """Add a diagram PNG image to the document."""
    img_path = diagram_dir / f"{image_name}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        set_paragraph_spacing(p, space_before=4, space_after=4)
        return p
    else:
        p = doc.add_paragraph(f"[Diagram not found: {image_name}.png]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p


def add_bracket_line(doc, bracket, font_size, indent=0.35, font_name=None):
    """Add a bracket notation line in monospace."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(bracket)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size - 1)
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def add_multilevel_from_bracket(doc, bracket, roles_json_path=None, roles_dict=None,
                                mode='answer_key', font_size=10):
    """Parse a bracket notation and add a multi-level labeling table.

    Convenience wrapper: parses the bracket, looks up roles from JSON or dict,
    and calls add_multilevel_labeling_table().

    Args:
        doc: python-docx Document
        bracket: bracket notation string
        roles_json_path: Path to chapter's table-roles JSON file. Roles are
                         looked up by matching bracket text.
        roles_dict: direct roles dict (overrides JSON lookup)
        mode: 'answer_key', 'student', or 'overhead'
        font_size: font size in points

    Returns:
        The table object.
    """
    import json
    from pathlib import Path

    table_data = parse_bracket_to_multilevel(bracket)

    roles = roles_dict
    if roles is None and roles_json_path:
        jp = Path(roles_json_path)
        if jp.exists():
            data = json.loads(jp.read_text(encoding='utf-8'))
            normalized = ' '.join(bracket.split())
            for entry in data['sentences']:
                if ' '.join(entry['bracket'].split()) == normalized:
                    roles = entry.get('roles')
                    break

    return add_multilevel_labeling_table(doc, table_data, roles=roles,
                                         font_size=font_size, mode=mode)


def load_chapter_roles(chapter_num):
    """Load the roles JSON for a chapter. Returns dict mapping bracket -> roles."""
    import json
    from pathlib import Path
    json_dir = Path(__file__).resolve().parent.parent / 'data' / 'static' / 'table-roles'
    jp = json_dir / f'ch{chapter_num:02d}_tables.json'
    if not jp.exists():
        return {}
    data = json.loads(jp.read_text(encoding='utf-8'))
    result = {}
    for entry in data['sentences']:
        key = ' '.join(entry['bracket'].split())
        result[key] = entry.get('roles', {})
    return result


def generate_pretext_labeling_xml(words, filled=False, roles=None, phrases=None, pos_labels=None):
    """Generate PreTeXt XML for a 4-row labeling table.

    Args:
        words: list of words in the sentence
        filled: if True, include answer values (for completed examples);
                if False, leave Role/Phrase/POS rows blank (student version)
        roles: list of role labels (same length as words, '' for span continuation)
        phrases: list of phrase labels
        pos_labels: list of POS labels

    Returns:
        String of PreTeXt XML for a <tabular> element.
    """
    num_cols = len(words) + 1  # +1 for row header
    lines = []
    lines.append('      <tabular top="minor" left="minor">')
    for _ in range(num_cols):
        halign = 'left' if _ == 0 else 'center'
        lines.append(f'        <col halign="{halign}" right="minor"/>')

    # Row 1: Role
    lines.append('        <row bottom="minor">')
    lines.append('          <cell>Role</cell>')
    for i in range(len(words)):
        val = roles[i] if (filled and roles) else ''
        lines.append(f'          <cell>{val}</cell>')
    lines.append('        </row>')

    # Row 2: Phrase
    lines.append('        <row bottom="minor">')
    lines.append('          <cell>Phrase</cell>')
    for i in range(len(words)):
        val = phrases[i] if (filled and phrases) else ''
        lines.append(f'          <cell>{val}</cell>')
    lines.append('        </row>')

    # Row 3: Word (always filled)
    lines.append('        <row bottom="minor">')
    lines.append('          <cell>Word</cell>')
    for w in words:
        lines.append(f'          <cell><foreign>{w}</foreign></cell>')
    lines.append('        </row>')

    # Row 4: POS
    lines.append('        <row bottom="minor">')
    lines.append('          <cell>POS</cell>')
    for i in range(len(words)):
        val = pos_labels[i] if (filled and pos_labels) else ''
        lines.append(f'          <cell>{val}</cell>')
    lines.append('        </row>')

    lines.append('      </tabular>')
    return '\n'.join(lines)


def generate_pretext_exercise_block(exercise_data, filled=False):
    """Generate a complete PreTeXt exercise block with labeling table.

    Args:
        exercise_data: dict with keys: num, sentence, words, roles, phrases,
                       pos, bracket (optional), image_source (optional)
        filled: if True, show answers (for completed examples)

    Returns:
        String of PreTeXt XML for the exercise.
    """
    ex = exercise_data
    lines = []
    lines.append(f'      <p><em>Exercise {ex["num"]}.</em> <foreign>{ex["sentence"]}</foreign></p>')
    lines.append('')
    lines.append(generate_pretext_labeling_xml(
        ex['words'], filled=filled,
        roles=ex.get('roles'), phrases=ex.get('phrases'), pos_labels=ex.get('pos')
    ))
    lines.append('')
    lines.append('      <ul>')
    lines.append('        <li><p>Bracket notation: _____</p></li>')
    lines.append('        <li><p>Diagram:</p></li>')
    lines.append('      </ul>')
    return '\n'.join(lines)


def overhead_page_break(doc):
    """Add a page break for overhead mode (each question starts on a new page)."""
    doc.add_page_break()


def get_font_config(overhead=False, font_size=12):
    """Return font configuration dict for answer key or overhead."""
    if overhead:
        return {
            'body_font': 'Arial Narrow',
            'body_size': 18,
            'heading1_size': 22,
            'heading2_size': 20,
            'heading3_size': 16,
            'bracket_size': 15,
            'diagram_width': 9.0,
        }
    else:
        return {
            'body_font': 'Garamond',
            'body_size': font_size,
            'heading1_size': 16,
            'heading2_size': 14,
            'heading3_size': 12,
            'bracket_size': font_size - 1,
            'diagram_width': 9.0,
        }


def setup_document(doc, overhead=False):
    """Configure document styles for answer key or overhead."""
    cfg = get_font_config(overhead)

    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    style = doc.styles['Normal']
    style.font.name = cfg['body_font']
    style.font.size = Pt(cfg['body_size'])

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Open Sans' if not overhead else 'Arial Narrow'
        heading_style.font.bold = True

    return cfg


def add_title_page(doc, chapter_title, cfg, overhead=False):
    """Add the title page for an answer key."""
    title = doc.add_heading(chapter_title, level=1)
    title.runs[0].font.size = Pt(cfg['heading1_size'])
    set_paragraph_spacing(title, space_before=0, space_after=6)

    subtitle = doc.add_heading('Answer Key', level=2)
    subtitle.runs[0].font.size = Pt(cfg['heading2_size'])
    set_paragraph_spacing(subtitle, space_before=0, space_after=12)

    if overhead:
        add_spacer_row(doc)


def add_part_heading(doc, part_title, cfg, overhead=False):
    """Add a part heading with page break."""
    doc.add_page_break()
    part = doc.add_heading(part_title, level=3)
    part.runs[0].font.size = Pt(cfg['heading3_size'])
    return part


def exercise_separator(doc, overhead=False):
    """Add appropriate separator between exercises.
    For overhead: page break so each question starts on a new page.
    For answer key: just a spacer for readability."""
    if overhead:
        doc.add_page_break()

#!/usr/bin/env python3
"""
Generate Chapter 5 Answer Key and Overhead Answer Key .docx files.
Chapter 5: Open Classes (5 parts, 25 exercises).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, add_sub_sentence, setup_document, add_title_page,
    add_part_heading, exercise_separator, get_font_config,
    add_bracket_line, blank_labels, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch05'


DIAGRAM_EXERCISES = [
    {
        'num': 15, 'sentence': 'The tall girl runs quickly.',
        'words':   ['The', 'tall', 'girl', 'runs', 'quickly'],
        'roles':   ['Subj', '', '', 'Pred', ''],
        'phrases': ['NP', '', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'ADJ', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [ADJP [ADJ tall]] [N girl]] [VP [V runs] [ADVP [ADV quickly]]]]',
        'diagram': 'ch05_hw_ex15_girl_runs',
    },
    {
        'num': 16, 'sentence': 'Heavy rain fell suddenly.',
        'words':   ['Heavy', 'rain', 'fell', 'suddenly'],
        'roles':   ['Subj', '', 'Pred', ''],
        'phrases': ['NP', '', 'VP', 'ADVP'],
        'pos':     ['ADJ', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [ADJP [ADJ Heavy]] [N rain]] [VP [V fell] [ADVP [ADV suddenly]]]]',
        'diagram': 'ch05_hw_ex16_rain_fell',
    },
    {
        'num': 17, 'sentence': 'The young artist painted beautifully.',
        'words':   ['The', 'young', 'artist', 'painted', 'beautifully'],
        'roles':   ['Subj', '', '', 'Pred', ''],
        'phrases': ['NP', '', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'ADJ', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [ADJP [ADJ young]] [N artist]] [VP [V painted] [ADVP [ADV beautifully]]]]',
        'diagram': 'ch05_hw_ex17_artist_painted',
    },
    {
        'num': 18, 'sentence': 'Small birds sing loudly.',
        'words':   ['Small', 'birds', 'sing', 'loudly'],
        'roles':   ['Subj', '', 'Pred', ''],
        'phrases': ['NP', '', 'VP', 'ADVP'],
        'pos':     ['ADJ', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [ADJP [ADJ Small]] [N birds]] [VP [V sing] [ADVP [ADV loudly]]]]',
        'diagram': 'ch05_hw_ex18_birds_sing',
    },
    {
        'num': 19, 'sentence': 'The clever student solved the difficult problem easily.',
        'words':   ['The', 'clever', 'student', 'solved', 'the', 'difficult', 'problem', 'easily'],
        'roles':   ['Subj', '', '', 'Pred', '', '', '', ''],
        'phrases': ['NP', '', '', 'VP', 'NP', '', '', 'ADVP'],
        'pos':     ['DET', 'ADJ', 'N', 'V', 'DET', 'ADJ', 'N', 'ADV'],
        'bracket': '[S [NP [DET The] [ADJP [ADJ clever]] [N student]] [VP [V solved] [NP [DET the] [ADJP [ADJ difficult]] [N problem]] [ADVP [ADV easily]]]]',
        'diagram': 'ch05_hw_ex19_student_solved',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 5 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']
    bracket_size = cfg['bracket_size']

    if overhead:
        table_size = 16
    else:
        table_size = font_size - 1

    add_title_page(doc, 'Chapter 5: Open Classes', cfg, overhead)

    # Helper to add bullet list of tests
    def add_test_bullets(tests):
        for bold_prefix, text in tests:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.7)
            run = p.add_run(bold_prefix)
            run.bold = True
            run.font.size = Pt(body_size)
            run.font.name = body_font
            run = p.add_run(text)
            run.font.size = Pt(body_size)
            run.font.name = body_font
            set_paragraph_spacing(p, space_before=0, space_after=1)

    # =============================================
    # Part 1: Word Class Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Word Class Identification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'The development of new technology takes time.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Noun', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -ment typically derives nouns from verbs (develop \u2192 development).'),
        ('Syntactic test: ', '\u201CDevelopment\u201D follows the determiner \u201CThe\u201D and functions as the subject of the sentence.'),
        ('Pronoun replacement: ', 'It can be replaced by a pronoun: \u201CIt takes time.\u201D'),
    ])

    # Exercise 2
    question_page_break(doc, overhead)
    add_exercise(doc, 2, 'She quickly solved the problem.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Adverb', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -ly attached to the adjective \u201Cquick\u201D forms an adverb.'),
        ('Syntactic test: ', '\u201CQuickly\u201D modifies the verb \u201Csolved,\u201D telling us how she solved the problem.'),
        ('Movability test: ', 'It can be moved in the sentence: \u201CShe solved the problem quickly.\u201D'),
    ])

    # Exercise 3
    question_page_break(doc, overhead)
    add_exercise(doc, 3, 'The beautiful garden attracted visitors.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Adjective', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Position test: ', '\u201CBeautiful\u201D appears between a determiner (\u201CThe\u201D) and a noun (\u201Cgarden\u201D), a characteristic position for adjectives.'),
        ('Comparison test: ', 'It can be used in comparative forms (more beautiful, most beautiful).'),
        ('Linking verb test: ', 'It can appear after a linking verb: \u201CThe garden is beautiful.\u201D'),
    ])

    # Exercise 4
    question_page_break(doc, overhead)
    add_exercise(doc, 4, 'The committee will investigate the matter.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Verb', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Modal test: ', '\u201CInvestigate\u201D follows the modal verb \u201Cwill,\u201D which only combines with verbs.'),
        ('Conjugation test: ', 'It can be conjugated for tense (investigated, investigates, investigating).'),
        ('Object test: ', 'It takes a direct object (\u201Cthe matter\u201D).'),
    ])

    # Exercise 5
    question_page_break(doc, overhead)
    add_exercise(doc, 5, 'His response was surprisingly confident.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Adverb', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -ly attached to \u201Csurprising\u201D forms an adverb.'),
        ('Modification test: ', '\u201CSurprisingly\u201D modifies the adjective \u201Cconfident,\u201D indicating the degree or manner of confidence.'),
        ('Pattern: ', 'Adverbs commonly modify adjectives in this way.'),
    ])

    # =============================================
    # Part 2: Phrase Identification
    # =============================================
    add_part_heading(doc, 'Part 2: Phrase Identification', cfg, overhead)

    phrase_exercises = [
        (6,  'very carefully',                 'carefully',        'AdvP (Adverb Phrase)'),
        (7,  'read the entire chapter',         'read',             'VP (Verb Phrase)'),
        (8,  'extremely proud',                 'proud',            'AdjP (Adjective Phrase)'),
        (9,  "my sister's new apartment",       'apartment',        'NP (Noun Phrase)'),
        (10, 'gave her the news',               'gave',             'VP (Verb Phrase)'),
    ]

    for i, (num, phrase, headword, phrase_type) in enumerate(phrase_exercises):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, phrase, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Headword:', headword, body_size, font_name=body_font)
        add_answer_line(doc, 'Phrase type:', phrase_type, body_size, font_name=body_font)

    # =============================================
    # Part 3: Phrase Replacement
    # =============================================
    add_part_heading(doc, 'Part 3: Phrase Replacement', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 11\u201314 are open-ended. Accept any grammatically correct replacement phrase of the specified type that changes the meaning while keeping the sentence grammatical.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    # Exercise 11
    add_exercise(doc, 11, 'The committee reviewed the lengthy proposal.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Find and replace:', 'The NP that serves as object (after the verb)', body_size, font_name=body_font)
    add_answer_line(doc, 'Original NP:', 'the lengthy proposal', body_size, font_name=body_font)
    add_plain_line(doc,
        'Sample replacement: "The committee reviewed the new budget." (Any NP that fits grammatically is acceptable.)',
        body_size, indent=0.7, font_name=body_font)

    # Exercise 12
    question_page_break(doc, overhead)
    add_exercise(doc, 12, 'She seemed extremely confident during the interview.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Find and replace:', 'The AdjP (adjective phrase)', body_size, font_name=body_font)
    add_answer_line(doc, 'Original AdjP:', 'extremely confident', body_size, font_name=body_font)
    add_plain_line(doc,
        'Sample replacement: "She seemed very nervous during the interview." (Any AdjP after the linking verb is acceptable.)',
        body_size, indent=0.7, font_name=body_font)

    # Exercise 13
    question_page_break(doc, overhead)
    add_exercise(doc, 13, 'The dog barked quite loudly at the mail carrier.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Find and replace:', 'The AdvP (adverb phrase)', body_size, font_name=body_font)
    add_answer_line(doc, 'Original AdvP:', 'quite loudly', body_size, font_name=body_font)
    add_plain_line(doc,
        'Sample replacement: "The dog barked surprisingly softly at the mail carrier." (Any AdvP modifying the verb is acceptable.)',
        body_size, indent=0.7, font_name=body_font)

    # Exercise 14
    question_page_break(doc, overhead)
    add_exercise(doc, 14, 'The talented musician plays the guitar.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Find and replace:', 'The VP (verb phrase \u2014 includes the verb and everything that follows it)', body_size, font_name=body_font)
    add_answer_line(doc, 'Original VP:', 'plays the guitar', body_size, font_name=body_font)
    add_plain_line(doc,
        'Sample replacement: "The talented musician composed a symphony." (Any VP with a verb and appropriate complements is acceptable.)',
        body_size, indent=0.7, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Open-Class Words
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Open-Class Words', cfg, overhead)

    ch_roles = load_chapter_roles(5)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], bracket_size)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Word Class Transformations
    # =============================================
    add_part_heading(doc, 'Part 5: Word Class Transformations', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Sentences:')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=3)

    for sent in [
        'A. The artist\u2019s creation amazed the critics.',
        'B. The artist created something amazing.',
        'C. The artist is highly creative.',
        'D. The artist works creatively.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(sent)
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=0, space_after=1)

    # Exercise 21
    add_exercise(doc, 21, 'Sentence A: creation', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Noun', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -tion derives nouns from verbs (create \u2192 creation).'),
        ('Syntactic test: ', '\u201CCreation\u201D follows the possessive \u201CThe artist\u2019s\u201D and functions as the subject of the sentence.'),
        ('Pronoun replacement: ', 'It can be replaced by a pronoun: \u201CIt amazed the critics.\u201D'),
    ])

    # Exercise 22
    question_page_break(doc, overhead)
    add_exercise(doc, 22, 'Sentence B: created', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Verb', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', '\u201CCreated\u201D shows past tense marking (-ed), a morphological feature of verbs.'),
        ('Syntactic test: ', 'It has a subject (\u201CThe artist\u201D) and takes an object (\u201Csomething amazing\u201D).'),
        ('Conjugation test: ', 'It can be conjugated: creates, creating, will create.'),
    ])

    # Exercise 23
    question_page_break(doc, overhead)
    add_exercise(doc, 23, 'Sentence C: creative', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Adjective', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -ive typically forms adjectives (create \u2192 creative).'),
        ('Syntactic test: ', '\u201CCreative\u201D follows the linking verb \u201Cis\u201D and can be modified by the intensifier \u201Chighly.\u201D'),
        ('Comparison test: ', 'It can be compared: more creative, most creative.'),
    ])

    # Exercise 24
    question_page_break(doc, overhead)
    add_exercise(doc, 24, 'Sentence D: creatively', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Part of speech:', 'Adverb', body_size, font_name=body_font)
    add_plain_line(doc, 'Tests used:', body_size, bold_prefix='', font_name=body_font)
    add_test_bullets([
        ('Morphological test: ', 'The suffix -ly attached to the adjective \u201Ccreative\u201D forms an adverb.'),
        ('Modification test: ', '\u201CCreatively\u201D modifies the verb \u201Cworks,\u201D describing how the artist works.'),
    ])

    # Exercise 25: Reflection (sub-items 25A, 25B, 25C)
    question_page_break(doc, overhead)
    add_exercise(doc, 25, 'All four sentences communicate something about an artist making things. Compare them:', body_size, font_name=body_font)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('25A) Which sentence feels clearest or most direct to you? Why?')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        'Model response: Sentence B (\u201CThe artist created something amazing.\u201D) tends to feel most direct '
        'because it uses a simple active-voice construction with verb, subject, and object clearly in their '
        'expected positions. However, any sentence can feel clearest depending on context \u2014 accept any '
        'answer with a sound explanation.',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('25B) What does each version emphasize?')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    answer_page_break(doc, overhead)
    for text in [
        'A (\u201Ccreation\u201D): emphasizes the product \u2014 the noun form makes the output the subject and main topic.',
        'B (\u201Ccreated\u201D): emphasizes the action/process \u2014 the verb form foregrounds what the artist did.',
        'C (\u201Ccreative\u201D): emphasizes the quality/trait of the artist as a person.',
        'D (\u201Ccreatively\u201D): emphasizes the manner of the work \u2014 how the artist does what they do.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(text)
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=0, space_after=1)

    question_page_break(doc, overhead)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('25C) In what writing situation might you choose one version over another?')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=2)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        'Model response: Use version A (\u201Ccreation\u201D) in an art review focusing on the work itself. '
        'Use version B (\u201Ccreated\u201D) in a narrative or biography focusing on events. '
        'Use version C (\u201Ccreative\u201D) in a character description or recommendation letter. '
        'Use version D (\u201Ccreatively\u201D) to describe working style or process. '
        'Accept any answer that connects word class to communicative purpose.',
        body_size, indent=0.7, font_name=body_font)

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def create_student_homework(output_path):
    """Create the Chapter 5 Student Homework with blank multi-level tables."""
    from answer_key_helpers import set_paragraph_spacing, parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 5 Homework: Open Classes')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Part 4 with blank multi-level tables
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Diagramming Open-Class Words')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 05 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 05 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 05 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

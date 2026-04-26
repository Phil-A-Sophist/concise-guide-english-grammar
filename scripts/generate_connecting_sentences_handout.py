#!/usr/bin/env python3
"""
Generate Connecting Sentences Handout with SyntaxTreeHybrid diagrams
and completed multi-level labeling tables.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from answer_key_helpers import (
    set_paragraph_spacing, setup_document,
    add_multilevel_labeling_table, parse_bracket_to_multilevel,
    add_diagram_image,
)

DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'connecting_sentences_diagrams'
OUTPUT_PATH = Path(__file__).parent.parent / 'Homework' / 'Connecting_Sentences_Diagrams.docx'

# Each entry: section header, sentence, bracket, roles dict
EXAMPLES = [
    {
        'header': 'Coordination of Phrases',
        'sentence': 'Frodo and Sam travel.',
        'bracket': '[S [NP [NP [N Frodo]] [CONJ and] [NP [N Sam]]] [VP [V travel]]]',
        'image': 'conn_01_coord_phrases_np',
        'roles': {
            "0": {"0": "Subject", "3": "Predicate"},
        },
    },
    {
        'header': None,  # continuation of same section
        'sentence': 'Gandalf thinks and acts.',
        'bracket': '[S [NP [N Gandalf]] [VP [VP [V thinks]] [CONJ and] [VP [V acts]]]]',
        'image': 'conn_02_coord_phrases_vp',
        'roles': {
            "0": {"0": "Subject", "1": "Predicate"},
        },
    },
    {
        'header': 'Coordination of Clauses',
        'sentence': 'Suzie yelled and Tommy cried.',
        'bracket': '[S [IC [NP [N Suzie]] [VP [V yelled]]] [CONJ and] [IC [NP [N Tommy]] [VP [V cried]]]]',
        'image': 'conn_03_coord_clauses',
        'roles': {
            "0": {"0": "Main", "3": "Main"},
            "1": {"0": "Subject", "1": "Predicate", "3": "Subject", "4": "Predicate"},
        },
    },
    {
        'header': 'Subordination',
        'sentence': 'Suzie yelled when Tommy cried.',
        'bracket': '[S [IC [NP [N Suzie]] [VP [V yelled]]] [DC [SUB when] [NP [N Tommy]] [VP [V cried]]]]',
        'image': 'conn_04_subord_dc_second',
        'roles': {
            "0": {"0": "Main", "2": "Adverbial"},
            "1": {"0": "Subject", "1": "Predicate", "3": "Subject", "4": "Predicate"},
        },
    },
    {
        'header': None,
        'sentence': 'When Tommy cried, Suzie yelled.',
        'bracket': '[S [DC [SUB When] [NP [N Tommy]] [VP [V cried]]] [IC [NP [N Suzie]] [VP [V yelled]]]]',
        'image': 'conn_05_subord_dc_first',
        'roles': {
            "0": {"0": "Adverbial", "3": "Main"},
            "1": {"1": "Subject", "2": "Predicate", "3": "Subject", "4": "Predicate"},
        },
    },
    {
        'header': 'Conjunctive Adverbs and Connecting Phrases',
        'sentence': 'Suzie yelled.',
        'bracket': '[S [NP [N Suzie]] [VP [V yelled]]]',
        'image': 'conn_06_simple',
        'roles': {
            "0": {"0": "Subject", "1": "Predicate"},
        },
    },
    {
        'header': None,
        'sentence': 'Also, Tommy cried.',
        'bracket': '[S [ADVP [ADV Also]] [NP [N Tommy]] [VP [V cried]]]',
        'image': 'conn_07_conj_adverb',
        'roles': {
            "0": {"0": "Adverbial", "1": "Subject", "2": "Predicate"},
        },
    },
]


def add_section_header(doc, text):
    """Add a bold section header."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Open Sans'
    set_paragraph_spacing(p, space_before=12, space_after=4)
    return p


def add_sentence_line(doc, sentence):
    """Add an italic sentence line."""
    p = doc.add_paragraph()
    run = p.add_run(sentence)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=4, space_after=4)
    return p


def main():
    doc = Document()
    cfg = setup_document(doc, overhead=False)

    # Title
    title = doc.add_heading('Connecting Sentences — Diagrams and Tables', level=1)
    title.runs[0].font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=12)

    for ex in EXAMPLES:
        if ex['header']:
            add_section_header(doc, ex['header'])

        add_sentence_line(doc, ex['sentence'])

        # Diagram image
        add_diagram_image(doc, DIAGRAM_DIR, ex['image'], width_inches=4.0)

        # Labeling table
        table_data = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(
            doc, table_data, roles=ex['roles'],
            font_size=11, mode='answer_key',
        )

        # Small spacer
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=0, space_after=6)

    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == '__main__':
    main()

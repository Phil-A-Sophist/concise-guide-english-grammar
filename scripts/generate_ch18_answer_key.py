#!/usr/bin/env python3
"""
Generate Chapter 18 Answer Key and Overhead Answer Key .docx files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set paragraph spacing in points."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def add_spacer_row(doc):
    """Add a blank spacer paragraph in Times New Roman 20 (no text, for instructor notes)."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '40')  # 20pt = 40 half-points
    rPr.append(sz)
    pPr.append(rPr)
    set_paragraph_spacing(p, space_before=0, space_after=0)
    return p


def add_exercise(doc, number, sentence, font_size, font_name=None):
    """Add an exercise header with sentence."""
    p = doc.add_paragraph()
    run = p.add_run(f'Exercise {number}. ')
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if sentence:
        run = p.add_run(sentence)
        run.italic = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    set_paragraph_spacing(p, space_before=6, space_after=3)
    return p


def add_answer_line(doc, label, answer, font_size, indent=0.35, font_name=None):
    """Add a label: answer line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    run = p.add_run(answer)
    run.italic = True
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def add_plain_line(doc, text, font_size, indent=0.35, bold_prefix=None, font_name=None):
    """Add a plain text line with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    set_paragraph_spacing(p, space_before=0, space_after=2)
    return p


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 18 Answer Key document."""
    if overhead:
        body_font = 'Arial Narrow'
        body_size = 18
        heading1_size = 22
        heading2_size = 20
        heading3_size = 16
        table_size = 16
        bracket_size = 15
    else:
        body_font = 'Garamond'
        body_size = font_size
        heading1_size = 16
        heading2_size = 14
        heading3_size = 12
        table_size = font_size - 1
        bracket_size = font_size - 1

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(body_size)

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Open Sans' if not overhead else 'Arial Narrow'
        heading_style.font.bold = True

    # Title
    title = doc.add_heading('Chapter 18: Clarity and Readability', level=1)
    title.runs[0].font.size = Pt(heading1_size)
    set_paragraph_spacing(title, space_before=0, space_after=6)

    subtitle = doc.add_heading('Answer Key', level=2)
    subtitle.runs[0].font.size = Pt(heading2_size)
    set_paragraph_spacing(subtitle, space_before=0, space_after=12)

    if overhead:
        add_spacer_row(doc)

    # =============================================
    # Part 1: Pronoun Reference
    # =============================================
    part = doc.add_heading('Part 1: Pronoun Reference', level=3)
    part.runs[0].font.size = Pt(heading3_size)

    # Exercise 1
    add_exercise(doc, 1, 'When John met Mark, he was surprised.', body_size, font_name=body_font)
    add_answer_line(doc, 'Problem:', 'Ambiguous reference \u2014 "he" could refer to John or Mark.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'When John met Mark, John was surprised.', body_size, font_name=body_font)
    add_plain_line(doc, '(Or: "When John met Mark, Mark was surprised.")', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 2
    add_exercise(doc, 2, 'They say the economy is improving.', body_size, font_name=body_font)
    add_answer_line(doc, 'Problem:', 'Vague reference \u2014 "they" has no identifiable antecedent.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'Economists say the economy is improving.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 3
    add_exercise(doc, 3, 'She failed the test, which disappointed her parents.', body_size, font_name=body_font)
    add_answer_line(doc, 'Problem:',
        'Broad reference \u2014 "which" refers to the whole clause, not a specific noun.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'Her test failure disappointed her parents.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 4
    add_exercise(doc, 4,
        'The committee reviewed the proposal and rejected it. This caused problems.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Problem:',
        'Broad reference \u2014 "this" could refer to the review, the rejection, or both.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'The committee\'s rejection of the proposal caused problems.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 5
    add_exercise(doc, 5,
        'The teacher told the student that her presentation needed work.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Problem:',
        'Ambiguous reference \u2014 "her" could refer to the teacher or the student.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'The teacher told the student that the student\'s presentation needed work.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # =============================================
    # Part 2: Modifier Placement
    # =============================================
    if overhead:
        doc.add_page_break()
    part = doc.add_heading('Part 2: Modifier Placement', level=3)
    part.runs[0].font.size = Pt(heading3_size)

    # Exercise 6
    add_exercise(doc, 6, 'Having finished dinner, the movie was started.', body_size, font_name=body_font)
    add_answer_line(doc, 'Error type:', 'Dangling modifier \u2014 the movie did not finish dinner.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'Having finished dinner, we started the movie.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 7
    add_exercise(doc, 7, 'She almost failed every exam.', body_size, font_name=body_font)
    add_answer_line(doc, 'Error type:',
        'Misplaced modifier \u2014 "almost" modifies "failed," but the intended meaning is '
        '"almost every."',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'She failed almost every exam.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 8
    add_exercise(doc, 8, 'Students who cheat often get caught.', body_size, font_name=body_font)
    add_answer_line(doc, 'Error type:',
        'Squinting modifier \u2014 "often" could modify "cheat" or "get caught."',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 1:', 'Students who often cheat get caught.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 2:', 'Students who cheat get caught often.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 9
    add_exercise(doc, 9,
        'To earn a good grade, the assignment must be completed carefully.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Error type:',
        'Dangling modifier \u2014 the assignment cannot earn a grade.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'To earn a good grade, you must complete the assignment carefully.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 10
    add_exercise(doc, 10, 'He only eats organic food on weekdays.', body_size, font_name=body_font)
    add_answer_line(doc, 'Error type:',
        'Misplaced modifier \u2014 "only" modifies "eats," but the intended meaning '
        'is "only on weekdays" or "only organic food."',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised (if "only organic"):', 'He eats only organic food on weekdays.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised (if "only weekdays"):', 'He eats organic food only on weekdays.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # =============================================
    # Part 3: Structural Ambiguity
    # =============================================
    if overhead:
        doc.add_page_break()
    part = doc.add_heading('Part 3: Structural Ambiguity', level=3)
    part.runs[0].font.size = Pt(heading3_size)

    # Exercise 11
    add_exercise(doc, 11, 'I photographed the elephant with a camera.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 1:',
        'I used a camera to photograph the elephant. (PP modifies VP)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'Using a camera, I photographed the elephant.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 2:',
        'The elephant had a camera. (PP modifies NP)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'I photographed the elephant that had a camera.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 12
    add_exercise(doc, 12, 'Bright students and teachers attended the workshop.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 1:',
        'Only the students are bright. (ADJ modifies first conjunct only)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'Teachers and bright students attended the workshop.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 2:',
        'Both the students and teachers are bright. (ADJ modifies entire coordination)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'Bright students and bright teachers attended the workshop.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 13
    add_exercise(doc, 13, 'The professor\'s assistant who was sick left early.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 1:',
        'The assistant was sick. (Relative clause modifies "assistant")',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'The professor\'s assistant, who was sick, left early.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 2:',
        'The professor was sick. (Relative clause modifies "professor")',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'The assistant of the professor who was sick left early.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 14
    add_exercise(doc, 14, 'She watched the children playing in the park.', body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 1:',
        'She was in the park when she watched. (PP modifies VP)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'In the park, she watched the children playing.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Meaning 2:',
        'The children were playing in the park. (PP modifies "playing" or NP)',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'She watched the children who were playing in the park.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # =============================================
    # Part 4: Parallel Structure and Sentence Complexity
    # =============================================
    if overhead:
        doc.add_page_break()
    part = doc.add_heading('Part 4: Parallel Structure and Sentence Complexity', level=3)
    part.runs[0].font.size = Pt(heading3_size)

    # Exercise 15
    add_exercise(doc, 15, 'She enjoys reading, writing, and to paint.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'She enjoys reading, writing, and painting.', body_size, font_name=body_font)
    add_plain_line(doc, 'All three items are now gerunds, creating parallel structure.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 16
    add_exercise(doc, 16, 'The job requires experience, dedication, and being creative.', body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:', 'The job requires experience, dedication, and creativity.', body_size, font_name=body_font)
    add_plain_line(doc, 'All three items are now nouns, creating parallel structure.', body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 17
    add_exercise(doc, 17,
        'He not only finished the report but also he proofread the entire document.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'He not only finished the report but also proofread the entire document.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'The correlative conjunction "not only...but also" now connects parallel verb '
        'phrases ("finished the report" and "proofread the entire document").',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 18
    add_exercise(doc, 18,
        'The report, which was commissioned by the board that was established last year '
        'to oversee operations, contains recommendations that, if implemented, would '
        'significantly improve efficiency.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'The board established a committee last year to oversee operations. The committee\'s '
        'report contains recommendations that would significantly improve efficiency if implemented.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # Exercise 19
    add_exercise(doc, 19,
        'The student, who had already completed the assignment that the professor assigned '
        'last week during the lecture that was held in the auditorium, submitted it early.',
        body_size, font_name=body_font)
    add_answer_line(doc, 'Revised:',
        'Last week, the professor assigned an assignment during a lecture in the auditorium. '
        'One student had already completed it and submitted it early.',
        body_size, font_name=body_font)
    if overhead:
        add_spacer_row(doc)

    # =============================================
    # Part 5: Comprehensive Revision
    # =============================================
    if overhead:
        doc.add_page_break()
    part = doc.add_heading('Part 5: Comprehensive Revision', level=3)
    part.runs[0].font.size = Pt(heading3_size)

    # Exercise 20
    add_exercise(doc, 20, None, body_size, font_name=body_font)
    add_plain_line(doc, 'Sample revised paragraph:', body_size, indent=0, bold_prefix='', font_name=body_font)
    add_plain_line(doc,
        'When the team walked into the meeting, the tension was immediately apparent. '
        'The manager told the employees that their performance needed to improve. '
        'The employees\' frustration grew. The proposal was not only expensive but also '
        'time-consuming, requiring years to implement. After reviewing all options carefully, '
        'the leadership decided to wait. Everyone understood that patience would be '
        'necessary.',
        body_size, font_name=body_font)

    add_plain_line(doc, '', body_size, indent=0, font_name=body_font)
    add_plain_line(doc, 'Issues to identify (any four):', body_size, indent=0, bold_prefix='', font_name=body_font)
    add_plain_line(doc,
        '\u2022 Dangling modifier: "Walking into the meeting, the tension..." '
        '\u2014 tension was not walking',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Ambiguous pronoun: "they needed to improve" '
        '\u2014 "they" is unclear (manager or employees?)',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Broad reference: "This led to frustration" '
        '\u2014 "this" has no specific antecedent',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Faulty parallelism: "not only expensive but also it would take years" '
        '\u2014 not parallel',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Dangling modifier: "Having reviewed all options carefully, the decision was made" '
        '\u2014 the decision did not review options',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Vague reference: "They say that patience is a virtue" '
        '\u2014 "they" has no antecedent',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '\u2022 Broad reference: "which everyone understood" '
        '\u2014 "which" refers to an entire clause',
        body_size, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_answer_key(
        homework_dir / 'Chapter 18 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Homework 18 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

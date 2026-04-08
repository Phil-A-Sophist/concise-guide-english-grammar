#!/usr/bin/env python3
"""
Generate Chapter 15 homework files: Student Homework, Answer Key, and Overhead.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page, add_part_heading,
    exercise_separator, get_font_config, add_bracket_line, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles,
    parse_bracket_to_multilevel, add_multilevel_labeling_table,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch15'


DIAGRAM_EXERCISES = [
    {
        'num': 16, 'sentence': 'The storm ended, and the sun came out.',
        'words':   ['The', 'storm', 'ended', 'and', 'the', 'sun', 'came', 'out'],
        'roles':   ['IC', '', '', '', 'IC', '', '', ''],
        'phrases': ['NP', '', 'VP', 'CC', 'NP', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'N', 'V', 'CONJ', 'DET', 'N', 'V', 'ADV'],
        'bracket': '[S [IC [NP [DET The] [N storm]] [VP [V ended]]] [CC [CONJ and]] [IC [NP [DET the] [N sun]] [VP [V came] [ADVP [ADV out]]]]]',
        'diagram': 'ch15_hw_ex16_storm_ended',
    },
    {
        'num': 17, 'sentence': 'Although she was tired, she finished the report.',
        'words':   ['Although', 'she', 'was', 'tired', 'she', 'finished', 'the', 'report'],
        'roles':   ['DC', '', '', '', 'IC', '', '', ''],
        'phrases': ['COMP', 'NP', 'VP', 'ADJP', 'NP', 'VP', 'NP', ''],
        'pos':     ['COMP', 'PRON', 'V', 'ADJ', 'PRON', 'V', 'DET', 'N'],
        'bracket': '[S [DC [COMP Although] [NP [PRON she]] [VP [V was] [ADJP [ADJ tired]]]] [IC [NP [PRON she]] [VP [V finished] [NP [DET the] [N report]]]]]',
        'diagram': 'ch15_hw_ex17_although_tired',
    },
    {
        'num': 18, 'sentence': 'The professor, however, disagreed completely.',
        'words':   ['The', 'professor', 'however', 'disagreed', 'completely'],
        'roles':   ['Subj', '', 'Conjunct', 'Pred', ''],
        'phrases': ['NP', '', 'ADVP', 'VP', 'ADVP'],
        'pos':     ['DET', 'N', 'ADV', 'V', 'ADV'],
        'bracket': '[S [NP [DET The] [N professor]] [ADVP [ADV however]] [VP [V disagreed] [ADVP [ADV completely]]]]',
        'diagram': 'ch15_hw_ex18_however',
    },
    {
        'num': 19, 'sentence': 'My sister, who lives in Boston, visits often.',
        'words':   ['My', 'sister', 'who', 'lives', 'in', 'Boston', 'visits', 'often'],
        'roles':   ['Subj', '', '', '', '', '', 'Pred', ''],
        'phrases': ['NP', '', 'RC', 'VP', 'PP', '', 'VP', 'ADVP'],
        'pos':     ['DET', 'N', 'REL', 'V', 'PREP', 'N', 'V', 'ADV'],
        'bracket': '[S [NP [DET My] [N sister] [RC [REL who] [VP [V lives] [PP [PREP in] [NP [N Boston]]]]]] [VP [V visits] [ADVP [ADV often]]]]',
        'diagram': 'ch15_hw_ex19_sister_boston',
    },
    {
        'num': 20, 'sentence': 'After the lecture, students asked many questions.',
        'words':   ['After', 'the', 'lecture', 'students', 'asked', 'many', 'questions'],
        'roles':   ['Advl', '', '', 'Subj', 'Pred', '', ''],
        'phrases': ['PP', 'NP', '', 'NP', 'VP', 'NP', ''],
        'pos':     ['PREP', 'DET', 'N', 'N', 'V', 'DET', 'N'],
        'bracket': '[S [PP [PREP After] [NP [DET the] [N lecture]]] [NP [N students]] [VP [V asked] [NP [DET many] [N questions]]]]',
        'diagram': 'ch15_hw_ex20_after_lecture',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 15 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 15: Punctuation', cfg, overhead)

    # =============================================
    # Part 1: Comma Usage
    # =============================================
    add_part_heading(doc, 'Part 1: Comma Usage', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'When the storm passed we surveyed the damage and began cleanup efforts.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Corrected:', 'When the storm passed, we surveyed the damage and began cleanup efforts.', body_size, font_name=body_font)
    add_plain_line(doc, 'Function: comma after introductory adverb clause', body_size, font_name=body_font)

    # Exercise 2
    question_page_break(doc, overhead)
    add_exercise(doc, 2, 'She is talented hardworking and creative.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Corrected:', 'She is talented, hardworking, and creative.', body_size, font_name=body_font)
    add_plain_line(doc, 'Function: commas separating items in a series (Oxford comma before "and")', body_size, font_name=body_font)

    # Exercise 3
    question_page_break(doc, overhead)
    add_exercise(doc, 3, 'My brother who lives in Seattle is visiting next week.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Corrected:', 'My brother, who lives in Seattle, is visiting next week.', body_size, font_name=body_font)
    add_plain_line(doc,
        'Function: commas setting off nonrestrictive relative clause (assumes the speaker has only one brother; '
        'if the speaker has multiple brothers, no commas would be needed \u2014 the clause would be restrictive)',
        body_size, font_name=body_font)

    # Exercise 4
    question_page_break(doc, overhead)
    add_exercise(doc, 4, 'The meeting was productive but it ran overtime.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Corrected:', 'The meeting was productive, but it ran overtime.', body_size, font_name=body_font)
    add_plain_line(doc, 'Function: comma before coordinating conjunction joining two independent clauses', body_size, font_name=body_font)

    # Exercise 5
    question_page_break(doc, overhead)
    add_exercise(doc, 5, 'The tall distinguished professor gave an inspiring lecture.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Corrected:', 'The tall, distinguished professor gave an inspiring lecture.', body_size, font_name=body_font)
    add_plain_line(doc,
        'Function: comma between coordinate adjectives (you can say "tall and distinguished," '
        'so a comma is appropriate)',
        body_size, font_name=body_font)

    # Exercise 6
    question_page_break(doc, overhead)
    add_exercise(doc, 6, 'The students who completed the assignment received extra credit.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        'No comma is needed because "who completed the assignment" is a restrictive relative clause \u2014 '
        'it identifies which students received extra credit (only those who completed the assignment, '
        'not all students). Removing the clause would change the meaning.',
        body_size, font_name=body_font)

    # =============================================
    # Part 2: Semicolons and Colons
    # =============================================
    add_part_heading(doc, 'Part 2: Semicolons and Colons', cfg, overhead)

    punctuation = [
        (7, 'She had one goal ( : / ; ) to finish the project on time.',
         'colon',
         'A colon introduces an explanation or elaboration after a complete sentence.'),
        (8, 'The rain stopped ( : / ; ) we went outside immediately.',
         'semicolon',
         'A semicolon joins two related independent clauses without a conjunction.'),
        (9, 'The committee includes three officers ( : / ; ) Dr. Lee, president; Ms. Park, secretary; and Mr. Kim, treasurer.',
         'colon',
         'A colon introduces the list. Semicolons are already used within the list items to separate names from titles.'),
        (10, 'He was exhausted ( : / ; ) however, he continued working.',
         'semicolon',
         'A semicolon is needed before a conjunctive adverb ("however") joining two independent clauses.'),
    ]

    for i, (num, sentence, choice, reasoning) in enumerate(punctuation):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Choice:', choice, body_size, font_name=body_font)
        add_plain_line(doc, reasoning, body_size, font_name=body_font)

    # =============================================
    # Part 3: Apostrophes
    # =============================================
    add_part_heading(doc, 'Part 3: Apostrophes', cfg, overhead)

    apostrophes = [
        (11, 'Its important to understand its function in the sentence.',
         'It\u2019s important to understand its function in the sentence.',
         'The first "its" should be "it\u2019s" (contraction of "it is"). The second "its" is correct (possessive).'),
        (12, 'The students books were left in the classroom.',
         'The students\u2019 books were left in the classroom.',
         'Plural possessive: "students\u2019" (the books belonging to the students).'),
        (13, 'The Joneses car is parked in the driveway.',
         'The Joneses\u2019 car is parked in the driveway.',
         'Plural possessive of a name ending in -s: "Joneses\u2019" (the car belonging to the Joneses).'),
        (14, 'Theyre going to their house over there.',
         'They\u2019re going to their house over there.',
         '"Theyre" should be "They\u2019re" (contraction of "they are"). "Their" and "there" are correct as used.'),
        (15, 'The womens team won the championship.',
         'The women\u2019s team won the championship.',
         'Irregular plural possessive: "women\u2019s" (the team belonging to the women). '
         'Since "women" doesn\u2019t end in -s, add \u2019s.'),
    ]

    for i, (num, original, corrected, explanation) in enumerate(apostrophes):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, original, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Corrected:', corrected, body_size, font_name=body_font)
        add_plain_line(doc, explanation, body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Punctuation-Relevant Structures
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Punctuation-Relevant Structures', cfg, overhead)

    ch_roles = load_chapter_roles(15)
    mode = 'overhead' if overhead else 'answer_key'

    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'], roles_dict=ch_roles.get(bracket_key), mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], body_size, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part_heading(doc, 'Part 5: Analysis and Application', cfg, overhead)

    # Exercise 21
    add_exercise(doc, 21, 'Explain how punctuation changes meaning in these two sentences.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        '(21A) "The students who studied passed the exam." \u2014 Restrictive: only those students '
        'who studied passed. Implies some students didn\u2019t study and didn\u2019t pass.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        '(21B) "The students, who studied, passed the exam." \u2014 Non-restrictive: all the students '
        'studied, and all of them passed. The clause adds extra information about what the students did.',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'The commas change the meaning from identifying a subset (restrictive) to describing the '
        'whole group (non-restrictive). This is a key example of how punctuation affects meaning.',
        body_size, font_name=body_font)

    # Exercise 22
    question_page_break(doc, overhead)
    add_exercise(doc, 22, 'Find a paragraph from a newspaper, textbook, or article and identify at least four punctuation marks with grammatical explanations.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        'Open-ended. Accept any paragraph that correctly identifies at least four punctuation marks '
        'with accurate grammatical explanations for each.',
        body_size, font_name=body_font)

    # Exercise 23
    question_page_break(doc, overhead)
    add_exercise(doc, 23, 'Reflect on which punctuation rules you find most challenging and why.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        'Open-ended reflection. Accept thoughtful answers that demonstrate awareness of '
        'punctuation rules and self-assessment of challenges.',
        body_size, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 15 Student Homework with blank multi-level tables."""
    doc = Document()

    # Basic styling — Garamond 12pt, landscape
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Title
    p = doc.add_paragraph()
    run = p.add_run('Chapter 15 Homework: Punctuation')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # --- Part 4: Diagramming Punctuation-Relevant Structures ---
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Diagramming Punctuation-Relevant Structures')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('For each sentence, complete the labeling table and write the bracket notation.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=8, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        table_data = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, table_data, mode='student', font_size=10)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 15 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 15 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 15 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

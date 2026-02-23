"""
Generate Chapter 06 homework documents:
  1. Chapter 06 Homework.docx       (student worksheet)
  2. Chapter 06 Answer Key.docx     (detailed answer key)
  3. Homework 06 Overhead.docx      (classroom projection)

Uses Chapter 5 files as formatting models and PreTeXt ch-06.ptx as the source of truth.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = "Homework"

# ---------------------------------------------------------------------------
#  Passage for Section 6.7.5 — Bank pump-and-dump Bitcoin scheme
# ---------------------------------------------------------------------------

PASSAGE_PARA1 = (
    "Last year, the bank launched a cryptocurrency trading desk and began "
    "aggressively buying Bitcoin. The executives at the firm instructed "
    "the traders to promote it on social media and in client newsletters. "
    "Once the cryptocurrency had risen sharply in value, the bank quietly "
    "sold their holdings at a massive profit. Many investors who had "
    "followed the bank\u2019s public recommendations lost money when the "
    "scheme collapsed."
)

PASSAGE_PARA2 = (
    "The SEC is now looking into it. They told regulators that the bank had "
    "full compliance to all trading laws. That was confirmed by an internal "
    "review, but outside auditors disagreed. A few of the executives "
    "admitted that this put them in a difficult position. This eventually led "
    "to a class-action lawsuit filed on the bank. The bank has tried to "
    "distance itself from that, but the damage to its reputation may be "
    "permanent."
)


# ---------------------------------------------------------------------------
#  Helper utilities
# ---------------------------------------------------------------------------

def make_numPr(numId, ilvl="0"):
    numPr = OxmlElement("w:numPr")
    ilvl_e = OxmlElement("w:ilvl")
    ilvl_e.set(qn("w:val"), ilvl)
    numId_e = OxmlElement("w:numId")
    numId_e.set(qn("w:val"), numId)
    numPr.append(ilvl_e)
    numPr.append(numId_e)
    return numPr


def set_page_break_before(para):
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    pPr.append(OxmlElement("w:pageBreakBefore"))


def add_run(para, text, bold=False, italic=False, size=None):
    r = para.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = size
    return r


# ---------------------------------------------------------------------------
#  1.  HOMEWORK  (student worksheet)
# ---------------------------------------------------------------------------

def build_homework():
    # Use Chapter 05 Homework as a template for styles
    doc = Document(f"{OUTPUT_DIR}/Chapter 05 Homework.docx")
    # Clear all content
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    def h1(text):
        p = doc.add_paragraph(style="Heading 1")
        add_run(p, text, bold=True)
        return p

    def h2(text):
        p = doc.add_paragraph(style="Heading 2")
        add_run(p, text, bold=True)
        return p

    def first(text="", bold_prefix=None):
        p = doc.add_paragraph(style="First Paragraph")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
            add_run(p, text)
        else:
            add_run(p, text)
        return p

    def body_text(text="", bold_prefix=None, italic_text=None):
        p = doc.add_paragraph(style="Body Text")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
        if italic_text:
            add_run(p, italic_text, italic=True)
        elif text:
            add_run(p, text)
        return p

    def blank():
        doc.add_paragraph(style="Normal")

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="Body Text")
        add_run(p, "\u2022 ", bold=False)
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
            add_run(p, text)
        else:
            add_run(p, text)
        return p

    # --- Title ---
    h1("Chapter 6: Closed Classes")
    h2("Homework: Closed Classes")
    first("Estimated completion time: 30\u201340 minutes")
    blank()

    # --- Section 1: Determiner and Pronoun Identification (Q1-Q3) ---
    h2("6.7.1 Determiner and Pronoun Identification")
    first("Instructions: ", "For each sentence, identify and classify the requested closed-class words.")
    body_text("Example (completed):")
    bullet("", bold_prefix="Sentence: ")
    p = doc.paragraphs[-1]
    add_run(p, "My brother gave his old car to me.", italic=True)
    bullet("", bold_prefix="Determiners: ")
    p = doc.paragraphs[-1]
    add_run(p, "my", italic=True)
    add_run(p, " (possessive), ")
    add_run(p, "his", italic=True)
    add_run(p, " (possessive). Note: ")
    add_run(p, "old", italic=True)
    add_run(p, " is NOT a determiner\u2014it\u2019s an adjective.")
    bullet("", bold_prefix="Pronouns: ")
    p = doc.paragraphs[-1]
    add_run(p, "me", italic=True)
    add_run(p, " (personal pronoun, 1st person singular, object form)")
    blank()

    # Exercise 1
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "1. ", bold=True)
    add_run(p, "The ambitious student submitted her application to several universities.", italic=True)
    body_text("Determiners (identify each and classify by type\u2014article, demonstrative, possessive, quantifier, number):")
    body_text("Pronouns:")
    blank()

    # Exercise 2
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "2. ", bold=True)
    add_run(p, "Everyone who attended the conference received their materials before the first session.", italic=True)
    body_text("Pronouns (identify each and classify by type\u2014personal, relative, indefinite, possessive):")
    blank()

    # Exercise 3
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "3. ", bold=True)
    add_run(p, "Those books on the shelf belong to someone in this department.", italic=True)
    body_text("Determiners:")
    body_text("Pronouns:")
    blank()

    # --- Section 2: Prepositional Phrase Analysis (Q4-Q6) ---
    h2("6.7.2 Prepositional Phrase Analysis")
    first("Instructions: ", "Identify all prepositional phrases in each sentence. "
          "For each PP, state the specific word it modifies, that word\u2019s part of speech, "
          "and what question the PP answers.")
    body_text("Example (completed):")
    bullet("", bold_prefix="Sentence: ")
    p = doc.paragraphs[-1]
    add_run(p, "The professor from Harvard lectured about linguistics in the auditorium.", italic=True)
    bullet("", bold_prefix="PP 1: ")
    p = doc.paragraphs[-1]
    add_run(p, "from Harvard", italic=True)
    body_text("Modifies: \u201Cprofessor\u201D (noun)")
    body_text("Question answered: Which professor?")

    bullet("", bold_prefix="PP 2: ")
    p = doc.paragraphs[-1]
    add_run(p, "about linguistics", italic=True)
    body_text("Modifies: \u201Clectured\u201D (verb)")
    body_text("Question answered: About what?")

    bullet("", bold_prefix="PP 3: ")
    p = doc.paragraphs[-1]
    add_run(p, "in the auditorium", italic=True)
    body_text("Modifies: \u201Clectured\u201D (verb)")
    body_text("Question answered: Where?")
    blank()

    # Exercise 4 (was Q5)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "4. ", bold=True)
    add_run(p, "The student with the red backpack studied in the library until midnight.", italic=True)
    for label in ["PP 1:", "Modifies:", "Question answered:",
                  "PP 2:", "Modifies:", "Question answered:",
                  "PP 3:", "Modifies:", "Question answered:"]:
        body_text(label)
    blank()

    # Exercise 5 (new sentence replacing old Q6)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "5. ", bold=True)
    add_run(p, "The child with the blue hat ran to the store for some milk.", italic=True)
    for label in ["PP 1:", "Modifies:", "Question answered:",
                  "PP 2:", "Modifies:", "Question answered:",
                  "PP 3:", "Modifies:", "Question answered:"]:
        body_text(label)
    blank()

    # Exercise 6 (new sentence replacing old Q7)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "6. ", bold=True)
    add_run(p, "The author of the bestselling novel spoke to reporters about her new book.", italic=True)
    for label in ["PP 1:", "Modifies:", "Question answered:",
                  "PP 2:", "Modifies:", "Question answered:",
                  "PP 3:", "Modifies:", "Question answered:"]:
        body_text(label)
    blank()

    # --- Section 3: Sentence Completion (Q7-Q9) ---
    h2("6.7.3 Sentence Completion")
    first("Instructions: ", "Complete each sentence by adding the requested element.")
    blank()

    # Exercise 7 (was Q8)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "7. ", bold=True)
    add_run(p, "Add a prepositional phrase that modifies the noun ")
    add_run(p, "book", italic=True)
    add_run(p, ":")
    body_text("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "The book won an award.", italic=True)
    body_text("Your revision:")
    blank()

    # Exercise 8 (was Q9)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "8. ", bold=True)
    add_run(p, "Add a prepositional phrase that modifies the verb and indicates when:")
    body_text("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "She completed the project.", italic=True)
    body_text("Your revision:")
    blank()

    # Exercise 9 (was Q10)
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "9. ", bold=True)
    add_run(p, "Replace the noun phrases with appropriate pronouns:")
    body_text("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "Maria told John that Maria would return John\u2019s laptop to John tomorrow.", italic=True)
    body_text("Your revision:")
    blank()

    # --- Section 4: Sentence Writing (Q10-Q12) ---
    h2("6.7.4 Sentence Writing")
    first("Instructions: ", "Write original sentences following each prompt. Identify the requested elements in your sentence.")
    blank()

    # Exercise 10
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "10. ", bold=True)
    add_run(p, "Write a sentence that includes both a prepositional phrase and a pronoun. Identify each.")
    body_text("Your sentence:")
    body_text("Prepositional phrase:")
    body_text("Pronoun:")
    blank()

    # Exercise 11
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "11. ", bold=True)
    add_run(p, "Write a sentence that includes at least two prepositional phrases. Identify each.")
    body_text("Your sentence:")
    body_text("PP 1:")
    body_text("PP 2:")
    blank()

    # Exercise 12
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "12. ", bold=True)
    add_run(p, "Write a sentence that includes a determiner, a conjunction, and a prepositional phrase. Identify each.")
    body_text("Your sentence:")
    body_text("Determiner:")
    body_text("Conjunction:")
    body_text("Prepositional phrase:")
    blank()

    # --- Section 5: Analysis and Reflection (Q13-Q16) ---
    h2("6.7.5 Analysis and Reflection")
    first("Instructions: ", "Read the following passage carefully. It contains intentional errors "
          "involving closed-class words (determiners, pronouns, and prepositions). "
          "Answer the questions that follow.")
    blank()

    first("", bold_prefix="Passage:")
    blank()
    body_text(PASSAGE_PARA1)
    blank()
    body_text(PASSAGE_PARA2)
    blank()

    # Exercise 13
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "13. ", bold=True)
    add_run(p, "Identify three pronouns in the passage that have vague or ambiguous referents. "
            "For each, explain the problem and suggest a revision.")
    body_text("Pronoun 1:")
    body_text("Problem:")
    body_text("Revision:")
    blank()
    body_text("Pronoun 2:")
    body_text("Problem:")
    body_text("Revision:")
    blank()
    body_text("Pronoun 3:")
    body_text("Problem:")
    body_text("Revision:")
    blank()

    # Exercise 14
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "14. ", bold=True)
    add_run(p, "Find two determiner problems in the passage: one where ")
    add_run(p, "the", italic=True)
    add_run(p, " is used before the reader knows which specific thing is meant, "
            "and one where quantifiers contradict each other. Explain each problem and suggest a fix.")
    body_text("Problem 1 (premature \u201Cthe\u201D):")
    body_text("Explanation:")
    body_text("Fix:")
    blank()
    body_text("Problem 2 (contradictory quantifiers):")
    body_text("Explanation:")
    body_text("Fix:")
    blank()

    # Exercise 15
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "15. ", bold=True)
    add_run(p, "Find two incorrectly used prepositions in the passage. "
            "State the error and provide the correct preposition.")
    body_text("Error 1:")
    body_text("Correction:")
    blank()
    body_text("Error 2:")
    body_text("Correction:")
    blank()

    # Exercise 16
    first("")
    p = doc.paragraphs[-1]
    p.clear()
    add_run(p, "16. ", bold=True)
    add_run(p, "Choose one of the two paragraphs and rewrite it, correcting all closed-class word issues "
            "(pronoun reference, determiner usage, and preposition choice).")
    body_text("Your rewritten paragraph:")

    doc.save(f"{OUTPUT_DIR}/Chapter 06 Homework.docx")
    print("Created Chapter 06 Homework.docx")


# ---------------------------------------------------------------------------
#  2.  ANSWER KEY  (detailed, following README format)
# ---------------------------------------------------------------------------

def build_answer_key():
    doc = Document(f"{OUTPUT_DIR}/Chapter 05 Answer Key.docx")
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Discover numId for bullets from styles
    numId = "1001"

    def h1(text):
        p = doc.add_paragraph(style="Heading 1")
        add_run(p, text, bold=True)
        return p

    def h2(text):
        p = doc.add_paragraph(style="Heading 2")
        add_run(p, text, bold=True)
        return p

    def first(text="", bold_prefix=None, italic_suffix=None):
        p = doc.add_paragraph(style="First Paragraph")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
        if italic_suffix:
            add_run(p, italic_suffix, italic=True)
        elif text:
            add_run(p, text)
        return p

    def body_text(text="", bold_prefix=None, italic_prefix=None):
        p = doc.add_paragraph(style="Body Text")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
            add_run(p, text)
        elif italic_prefix:
            add_run(p, italic_prefix, italic=True)
            add_run(p, text)
        else:
            add_run(p, text)
        return p

    def bullet_item(text, bold_prefix=None):
        p = doc.add_paragraph(style="Compact")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        pPr.append(make_numPr(numId))
        if bold_prefix:
            add_run(p, bold_prefix, bold=True)
            add_run(p, text)
        else:
            add_run(p, text)
        return p

    def blank():
        doc.add_paragraph(style="Normal")

    # -----------------------------------------------------------------------
    h1("Chapter 6: Closed Classes \u2014 Answer Key")
    blank()

    # --- Section 1: Q1-Q3 ---
    h2("6.7.1 Determiner and Pronoun Identification")
    first("Instructions: ", "Identify and classify the requested closed-class words in each sentence.")
    blank()

    # Q1
    first("", bold_prefix="1. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe ambitious student submitted her application to several universities.\u201D", italic=True)
    body_text("", bold_prefix="Determiners:")
    bullet_item("article (definite)", bold_prefix="The \u2014 ")
    bullet_item("possessive determiner", bold_prefix="her \u2014 ")
    bullet_item("quantifier", bold_prefix="several \u2014 ")
    body_text("", bold_prefix="Pronouns: ")
    p = doc.paragraphs[-1]
    add_run(p, "None in this sentence. Note: ")
    add_run(p, "her", italic=True)
    add_run(p, " is a possessive determiner here, not a pronoun, because it precedes and modifies a noun.")
    blank()

    # Q2
    first("", bold_prefix="2. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CEveryone who attended the conference received their materials before the first session.\u201D", italic=True)
    body_text("", bold_prefix="Pronouns:")
    bullet_item("indefinite pronoun", bold_prefix="Everyone \u2014 ")
    bullet_item("relative pronoun (introduces relative clause, refers to \u201Ceveryone\u201D)", bold_prefix="who \u2014 ")
    bullet_item("possessive determiner (could also be classified as possessive pronoun depending on grammar framework)", bold_prefix="their \u2014 ")
    blank()

    # Q3
    first("", bold_prefix="3. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThose books on the shelf belong to someone in this department.\u201D", italic=True)
    body_text("", bold_prefix="Determiners:")
    bullet_item("demonstrative (far)", bold_prefix="Those \u2014 ")
    bullet_item("article (definite)", bold_prefix="the \u2014 ")
    bullet_item("demonstrative (near)", bold_prefix="this \u2014 ")
    body_text("", bold_prefix="Pronouns:")
    bullet_item("indefinite pronoun", bold_prefix="someone \u2014 ")
    blank()

    # --- Section 2: Q4-Q6 ---
    h2("6.7.2 Prepositional Phrase Analysis")
    first("Instructions: ", "Identify all prepositional phrases. For each, state what specific word "
          "it modifies, that word\u2019s part of speech, and what question the PP answers.")
    blank()

    # Q4 (was Q5)
    first("", bold_prefix="4. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe student with the red backpack studied in the library until midnight.\u201D", italic=True)
    body_text("", bold_prefix="PP 1: ")
    p = doc.paragraphs[-1]
    add_run(p, "with the red backpack", italic=True)
    bullet_item("modifies \u201Cstudent\u201D (noun)", bold_prefix="Modifies: ")
    bullet_item("Which student?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 2: ")
    p = doc.paragraphs[-1]
    add_run(p, "in the library", italic=True)
    bullet_item("modifies \u201Cstudied\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("Where?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 3: ")
    p = doc.paragraphs[-1]
    add_run(p, "until midnight", italic=True)
    bullet_item("modifies \u201Cstudied\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("Until when? / How long?", bold_prefix="Question answered: ")
    blank()

    # Q5 (new sentence)
    first("", bold_prefix="5. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe child with the blue hat ran to the store for some milk.\u201D", italic=True)
    body_text("", bold_prefix="PP 1: ")
    p = doc.paragraphs[-1]
    add_run(p, "with the blue hat", italic=True)
    bullet_item("modifies \u201Cchild\u201D (noun)", bold_prefix="Modifies: ")
    bullet_item("Which child?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 2: ")
    p = doc.paragraphs[-1]
    add_run(p, "to the store", italic=True)
    bullet_item("modifies \u201Cran\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("Where? / To where?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 3: ")
    p = doc.paragraphs[-1]
    add_run(p, "for some milk", italic=True)
    bullet_item("modifies \u201Cran\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("Why? / For what purpose?", bold_prefix="Question answered: ")
    blank()

    # Q6 (new sentence)
    first("", bold_prefix="6. ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe author of the bestselling novel spoke to reporters about her new book.\u201D", italic=True)
    body_text("", bold_prefix="PP 1: ")
    p = doc.paragraphs[-1]
    add_run(p, "of the bestselling novel", italic=True)
    bullet_item("modifies \u201Cauthor\u201D (noun)", bold_prefix="Modifies: ")
    bullet_item("Which author?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 2: ")
    p = doc.paragraphs[-1]
    add_run(p, "to reporters", italic=True)
    bullet_item("modifies \u201Cspoke\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("To whom?", bold_prefix="Question answered: ")
    body_text("", bold_prefix="PP 3: ")
    p = doc.paragraphs[-1]
    add_run(p, "about her new book", italic=True)
    bullet_item("modifies \u201Cspoke\u201D (verb)", bold_prefix="Modifies: ")
    bullet_item("About what?", bold_prefix="Question answered: ")
    blank()

    # --- Section 3: Q7-Q9 ---
    h2("6.7.3 Sentence Completion")
    first("Instructions: ", "Complete each sentence by adding the requested element.")
    blank()

    # Q7
    first("", bold_prefix="7. ")
    p = doc.paragraphs[-1]
    add_run(p, "Add a prepositional phrase that modifies the noun ")
    add_run(p, "book", italic=True)
    add_run(p, ":")
    body_text("The book won an award.", bold_prefix="Original: ")
    body_text("\u201CThe book about climate change won an award.\u201D", bold_prefix="Sample revision: ")
    body_text("", italic_prefix="Other acceptable answers include any PP modifying \u201Cbook\u201D (e.g., on the bestseller list, by the famous author, with the blue cover).")
    blank()

    # Q8
    first("", bold_prefix="8. ")
    p = doc.paragraphs[-1]
    add_run(p, "Add a prepositional phrase that modifies the verb and indicates when:")
    body_text("She completed the project.", bold_prefix="Original: ")
    body_text("\u201CShe completed the project before the deadline.\u201D", bold_prefix="Sample revision: ")
    body_text("", italic_prefix="Other acceptable answers include any adverbial PP indicating time (e.g., during the weekend, in the afternoon, after the meeting).")
    blank()

    # Q9
    first("", bold_prefix="9. ")
    p = doc.paragraphs[-1]
    add_run(p, "Replace the noun phrases with appropriate pronouns:")
    body_text("Maria told John that Maria would return John\u2019s laptop to John tomorrow.", bold_prefix="Original: ")
    body_text("\u201CMaria told him that she would return his laptop to him tomorrow.\u201D", bold_prefix="Revised: ")
    blank()

    # --- Section 4: Q10-Q12 ---
    h2("6.7.4 Sentence Writing")
    first("Instructions: ", "Write original sentences following each prompt.")
    blank()

    # Q10
    first("", bold_prefix="10. ")
    p = doc.paragraphs[-1]
    add_run(p, "Sentence with a prepositional phrase and a pronoun.")
    body_text("\u201CShe walked to the park after lunch.\u201D", bold_prefix="Sample: ")
    bullet_item("to the park (PP), after lunch (PP)", bold_prefix="Prepositional phrases: ")
    bullet_item("She (personal pronoun, 3rd person singular, subject form)", bold_prefix="Pronoun: ")
    body_text("", italic_prefix="Accept any sentence with at least one PP and one pronoun, correctly identified.")
    blank()

    # Q11
    first("", bold_prefix="11. ")
    p = doc.paragraphs[-1]
    add_run(p, "Sentence with at least two prepositional phrases.")
    body_text("\u201CThe student from Canada worked in the library.\u201D", bold_prefix="Sample: ")
    bullet_item("from Canada \u2014 modifies \u201Cstudent\u201D (noun)", bold_prefix="PP 1: ")
    bullet_item("in the library \u2014 modifies \u201Cworked\u201D (verb)", bold_prefix="PP 2: ")
    body_text("", italic_prefix="Accept any sentence with two or more PPs, correctly identified.")
    blank()

    # Q12
    first("", bold_prefix="12. ")
    p = doc.paragraphs[-1]
    add_run(p, "Sentence with a determiner, a conjunction, and a prepositional phrase.")
    body_text("\u201CThe cat and the dog played in the yard.\u201D", bold_prefix="Sample: ")
    bullet_item("The (article, definite) \u2014 appears three times", bold_prefix="Determiner: ")
    bullet_item("and (coordinating conjunction)", bold_prefix="Conjunction: ")
    bullet_item("in the yard \u2014 modifies \u201Cplayed\u201D (verb)", bold_prefix="Prepositional phrase: ")
    body_text("", italic_prefix="Accept any sentence containing all three element types, correctly identified.")
    blank()

    # --- Section 5: Q13-Q16 ---
    h2("6.7.5 Analysis and Reflection")
    first("", bold_prefix="Passage:")
    blank()
    body_text(PASSAGE_PARA1)
    blank()
    body_text(PASSAGE_PARA2)
    blank()

    # Q13
    first("", bold_prefix="13. ")
    p = doc.paragraphs[-1]
    add_run(p, "Identify three pronouns with vague or ambiguous referents.")
    blank()
    body_text("", italic_prefix="Answers will vary. Acceptable answers include any three of the following:")
    blank()

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Cpromote it\u201D (paragraph 1)")
    bullet_item("Unclear referent. Does \u201Cit\u201D refer to Bitcoin, the trading desk, or cryptocurrency in general?", bold_prefix="Problem: ")
    bullet_item("\u201C\u2026instructed the traders to promote Bitcoin on social media\u2026\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Clooking into it\u201D (paragraph 2)")
    bullet_item("Vague \u201Cit.\u201D Could refer to the scheme, the trading, or the losses.", bold_prefix="Problem: ")
    bullet_item("\u201C\u2026looking into the trading scheme.\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThey told regulators\u201D (paragraph 2)")
    bullet_item("Ambiguous \u201CThey.\u201D Could refer to the SEC, the executives, or the bank.", bold_prefix="Problem: ")
    bullet_item("\u201CBank representatives told regulators\u2026\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThat was confirmed\u201D (paragraph 2)")
    bullet_item("Vague \u201CThat.\u201D The claim of compliance? The investigation?", bold_prefix="Problem: ")
    bullet_item("\u201CThe compliance claim was confirmed\u2026\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Cthis put them\u201D (paragraph 2)")
    bullet_item("Vague \u201Cthis\u201D and ambiguous \u201Cthem.\u201D", bold_prefix="Problem: ")
    bullet_item("\u201CThe auditors\u2019 disagreement put the executives in a difficult position.\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThis eventually led\u201D (paragraph 2)")
    bullet_item("Vague \u201CThis.\u201D Multiple events precede it.", bold_prefix="Problem: ")
    bullet_item("\u201CThe resulting scandal eventually led to a class-action lawsuit\u2026\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Cdistance itself from that\u201D (paragraph 2)")
    bullet_item("Vague \u201Cthat.\u201D The lawsuit? The scheme? The negative publicity?", bold_prefix="Problem: ")
    bullet_item("\u201C\u2026distance itself from the scandal\u2026\u201D", bold_prefix="Revision: ")

    body_text("", bold_prefix="Pronoun: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Csold their holdings\u201D (paragraph 1)")
    bullet_item("Singular/plural mismatch: \u201Cthe bank\u201D is singular, but \u201Ctheir\u201D is plural.", bold_prefix="Problem: ")
    bullet_item("\u201C\u2026the bank quietly sold its holdings\u2026\u201D", bold_prefix="Revision: ")
    blank()

    # Q14
    first("", bold_prefix="14. ")
    p = doc.paragraphs[-1]
    add_run(p, "Determiner problems.")
    blank()

    body_text("", bold_prefix="Problem 1 (premature \u201Cthe\u201D): ")
    p = doc.paragraphs[-1]
    add_run(p, "The passage opens with \u201Cthe bank\u201D as if the reader already knows which bank is being "
            "discussed. Similarly, \u201CThe executives at the firm,\u201D \u201Cthe traders,\u201D \u201Cthe cryptocurrency,\u201D "
            "and \u201Cthe scheme\u201D all use the definite article before these referents have been introduced.")
    bullet_item("First introduce the referent with an indefinite article or proper name, then use \u201Cthe\u201D "
                "in subsequent references. For example: \u201CLast year, a major investment bank launched\u2026\u201D "
                "or name the bank specifically.", bold_prefix="Fix: ")
    blank()

    body_text("", bold_prefix="Problem 2 (contradictory quantifiers): ")
    p = doc.paragraphs[-1]
    add_run(p, "Paragraph 1 says \u201CMany investors\u2026 lost money,\u201D while paragraph 2 says \u201CA few of the "
            "executives admitted\u2026\u201D The shift from \u201Cmany\u201D (suggesting large numbers) to \u201Ca few\u201D "
            "(suggesting small numbers) is potentially contradictory or confusing when describing "
            "related groups within the same event.")
    bullet_item("Use consistent quantifiers or clarify the different groups. For example: \u201CSeveral "
                "of the executives admitted\u2026\u201D or specify which executives.", bold_prefix="Fix: ")
    blank()

    # Q15
    first("", bold_prefix="15. ")
    p = doc.paragraphs[-1]
    add_run(p, "Incorrect prepositions.")
    blank()

    body_text("", bold_prefix="Error 1: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Ccompliance to all trading laws\u201D")
    bullet_item("\u201Ccompliance with all trading laws.\u201D The standard collocation is \u201Ccompliance with,\u201D "
                "not \u201Ccompliance to.\u201D", bold_prefix="Correction: ")
    blank()

    body_text("", bold_prefix="Error 2: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201Ca class-action lawsuit filed on the bank\u201D")
    bullet_item("\u201Ca class-action lawsuit filed against the bank.\u201D Lawsuits are filed \u201Cagainst\u201D "
                "a party, not \u201Con\u201D a party.", bold_prefix="Correction: ")
    blank()

    # Q16
    first("", bold_prefix="16. ")
    p = doc.paragraphs[-1]
    add_run(p, "Rewrite one paragraph correcting all closed-class word issues.")
    blank()
    body_text("", italic_prefix="Answers will vary. Sample rewrite of paragraph 1:")
    blank()
    body_text("\u201CLast year, a major investment bank launched a cryptocurrency trading desk and began "
              "aggressively buying Bitcoin. Executives at the bank instructed its traders to promote "
              "Bitcoin on social media and in client newsletters. Once the cryptocurrency had risen "
              "sharply in value, the bank quietly sold its holdings at a massive profit. Many investors "
              "who had followed the bank\u2019s public recommendations lost money when the price collapsed.\u201D")
    blank()

    body_text("", italic_prefix="Sample rewrite of paragraph 2:")
    blank()
    body_text("\u201CThe SEC is now looking into the trading scheme. Bank representatives told regulators "
              "that the bank had maintained full compliance with all trading laws. The compliance claim "
              "was confirmed by an internal review, but outside auditors disagreed. Several executives "
              "admitted that the auditors\u2019 findings put the leadership team in a difficult position. "
              "The resulting scandal eventually led to a class-action lawsuit filed against the bank. "
              "The bank has tried to distance itself from the controversy, but the damage to its "
              "reputation may be permanent.\u201D")
    blank()

    first("", italic_suffix="End of Answer Key")

    doc.save(f"{OUTPUT_DIR}/Chapter 06 Answer Key.docx")
    print("Created Chapter 06 Answer Key.docx")


# ---------------------------------------------------------------------------
#  3.  OVERHEAD  (classroom projection)
# ---------------------------------------------------------------------------

def build_overhead():
    doc = Document(f"{OUTPUT_DIR}/Homework 05 Overhead.docx")
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    SIZE_H1 = Pt(28)
    SIZE_H2 = Pt(26)
    SIZE_H3 = Pt(24)
    SIZE_BODY = Pt(22)
    QUESTION_GAP = Pt(270)

    prev_was_h1 = False

    def h1(text, page_break=True):
        nonlocal prev_was_h1
        p = doc.add_paragraph(style="Heading 1")
        add_run(p, text, bold=True, size=SIZE_H1)
        if page_break:
            set_page_break_before(p)
        prev_was_h1 = True
        return p

    def h2(text):
        nonlocal prev_was_h1
        p = doc.add_paragraph(style="Heading 2")
        add_run(p, text, bold=True, size=SIZE_H2)
        if not prev_was_h1:
            set_page_break_before(p)
        p.paragraph_format.space_after = QUESTION_GAP
        prev_was_h1 = False
        return p

    def h3(text):
        nonlocal prev_was_h1
        p = doc.add_paragraph(style="Heading 3")
        add_run(p, text, bold=True, size=SIZE_H3)
        set_page_break_before(p)
        p.paragraph_format.space_after = QUESTION_GAP
        prev_was_h1 = False
        return p

    def first(text="", bold_prefix=None):
        nonlocal prev_was_h1
        prev_was_h1 = False
        p = doc.add_paragraph(style="First Paragraph")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True, size=SIZE_BODY)
            add_run(p, text, size=SIZE_BODY)
        else:
            add_run(p, text, size=SIZE_BODY)
        return p

    def body_t(text="", bold_prefix=None):
        nonlocal prev_was_h1
        prev_was_h1 = False
        p = doc.add_paragraph(style="Body Text")
        if bold_prefix:
            add_run(p, bold_prefix, bold=True, size=SIZE_BODY)
            add_run(p, text, size=SIZE_BODY)
        else:
            add_run(p, text, size=SIZE_BODY)
        return p

    # Use existing numId from original doc's Compact paragraphs
    numId = "1001"

    def bullet(text, bold_prefix=None):
        nonlocal prev_was_h1
        prev_was_h1 = False
        p = doc.add_paragraph(style="Compact")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        pPr.append(make_numPr(numId))
        if bold_prefix:
            add_run(p, bold_prefix, bold=True, size=SIZE_BODY)
            add_run(p, text, size=SIZE_BODY)
        else:
            add_run(p, text, size=SIZE_BODY)
        return p

    # -----------------------------------------------------------------------
    # Title page
    h1("Homework 06: Overhead", page_break=False)
    h1("Closed Classes Answer Key", page_break=False)

    # --- Section 1: Q1-Q3 ---
    h1("6.7.1 Determiner and Pronoun Identification")

    h2('1. "The ambitious student submitted her application to several universities."')
    first("", bold_prefix="Determiners:")
    bullet("article (definite)", bold_prefix="The \u2014 ")
    bullet("possessive determiner", bold_prefix="her \u2014 ")
    bullet("quantifier", bold_prefix="several \u2014 ")
    first("", bold_prefix="Pronouns: ")
    p = doc.paragraphs[-1]
    add_run(p, "None (her is a determiner here)", size=SIZE_BODY)

    h2('2. "Everyone who attended the conference received their materials before the first session."')
    first("", bold_prefix="Pronouns:")
    bullet("indefinite pronoun", bold_prefix="Everyone \u2014 ")
    bullet("relative pronoun (refers to \u201Ceveryone\u201D)", bold_prefix="who \u2014 ")
    bullet("possessive determiner", bold_prefix="their \u2014 ")

    h2('3. "Those books on the shelf belong to someone in this department."')
    first("", bold_prefix="Determiners:")
    bullet("demonstrative (far)", bold_prefix="Those \u2014 ")
    bullet("article (definite)", bold_prefix="the \u2014 ")
    bullet("demonstrative (near)", bold_prefix="this \u2014 ")
    first("", bold_prefix="Pronouns:")
    bullet("indefinite pronoun", bold_prefix="someone \u2014 ")

    # --- Section 2: Q4-Q6 ---
    h1("6.7.2 Prepositional Phrase Analysis")

    h2('4. "The student with the red backpack studied in the library until midnight."')
    bullet("modifies \u201Cstudent\u201D (noun) \u2014 Which student?", bold_prefix="with the red backpack: ")
    bullet("modifies \u201Cstudied\u201D (verb) \u2014 Where?", bold_prefix="in the library: ")
    bullet("modifies \u201Cstudied\u201D (verb) \u2014 Until when?", bold_prefix="until midnight: ")

    h2('5. "The child with the blue hat ran to the store for some milk."')
    bullet("modifies \u201Cchild\u201D (noun) \u2014 Which child?", bold_prefix="with the blue hat: ")
    bullet("modifies \u201Cran\u201D (verb) \u2014 Where?", bold_prefix="to the store: ")
    bullet("modifies \u201Cran\u201D (verb) \u2014 For what purpose?", bold_prefix="for some milk: ")

    h2('6. "The author of the bestselling novel spoke to reporters about her new book."')
    bullet("modifies \u201Cauthor\u201D (noun) \u2014 Which author?", bold_prefix="of the bestselling novel: ")
    bullet("modifies \u201Cspoke\u201D (verb) \u2014 To whom?", bold_prefix="to reporters: ")
    bullet("modifies \u201Cspoke\u201D (verb) \u2014 About what?", bold_prefix="about her new book: ")

    # --- Section 3: Q7-Q9 ---
    h1("6.7.3 Sentence Completion")

    h2("7. Add PP modifying the noun book")
    first("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "The book won an award.", size=SIZE_BODY)
    first("", bold_prefix="Sample: ")
    p = doc.paragraphs[-1]
    add_run(p, "The book ", size=SIZE_BODY)
    add_run(p, "about climate change", bold=True, size=SIZE_BODY)
    add_run(p, " won an award.", size=SIZE_BODY)

    h2("8. Add PP indicating when")
    first("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "She completed the project.", size=SIZE_BODY)
    first("", bold_prefix="Sample: ")
    p = doc.paragraphs[-1]
    add_run(p, "She completed the project ", size=SIZE_BODY)
    add_run(p, "before the deadline", bold=True, size=SIZE_BODY)
    add_run(p, ".", size=SIZE_BODY)

    h2("9. Replace noun phrases with pronouns")
    first("", bold_prefix="Original: ")
    p = doc.paragraphs[-1]
    add_run(p, "Maria told John that Maria would return John\u2019s laptop to John tomorrow.", size=SIZE_BODY)
    first("", bold_prefix="Revised: ")
    p = doc.paragraphs[-1]
    add_run(p, "Maria told ", size=SIZE_BODY)
    add_run(p, "him", bold=True, size=SIZE_BODY)
    add_run(p, " that ", size=SIZE_BODY)
    add_run(p, "she", bold=True, size=SIZE_BODY)
    add_run(p, " would return ", size=SIZE_BODY)
    add_run(p, "his", bold=True, size=SIZE_BODY)
    add_run(p, " laptop to ", size=SIZE_BODY)
    add_run(p, "him", bold=True, size=SIZE_BODY)
    add_run(p, " tomorrow.", size=SIZE_BODY)

    # --- Section 4: Q10-Q12 ---
    h1("6.7.4 Sentence Writing")

    h2("10. PP + pronoun")
    first("", bold_prefix="Sample: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CShe walked to the park after lunch.\u201D", size=SIZE_BODY)
    bullet("to the park, after lunch", bold_prefix="PPs: ")
    bullet("She (personal pronoun)", bold_prefix="Pronoun: ")

    h2("11. Two prepositional phrases")
    first("", bold_prefix="Sample: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe student from Canada worked in the library.\u201D", size=SIZE_BODY)
    bullet("modifies \u201Cstudent\u201D (noun)", bold_prefix="from Canada: ")
    bullet("modifies \u201Cworked\u201D (verb)", bold_prefix="in the library: ")

    h2("12. Determiner + conjunction + PP")
    first("", bold_prefix="Sample: ")
    p = doc.paragraphs[-1]
    add_run(p, "\u201CThe cat and the dog played in the yard.\u201D", size=SIZE_BODY)
    bullet("The (article)", bold_prefix="Determiner: ")
    bullet("and (coordinating)", bold_prefix="Conjunction: ")
    bullet("in the yard \u2014 modifies \u201Cplayed\u201D", bold_prefix="PP: ")

    # --- Section 5: Q13-Q16 ---
    h1("6.7.5 Analysis and Reflection")

    h2("13. Three vague/ambiguous pronouns")
    bullet("\u201Cpromote it\u201D \u2014 unclear referent (Bitcoin? trading desk?)", bold_prefix="Example: ")
    bullet("\u201Clooking into it\u201D \u2014 vague \u201Cit\u201D (scheme? trading? losses?)", bold_prefix="Example: ")
    bullet("\u201CThey told regulators\u201D \u2014 ambiguous \u201CThey\u201D (SEC? executives? bank?)", bold_prefix="Example: ")
    bullet("Also: \u201CThat was confirmed,\u201D \u201Cthis put them,\u201D \u201CThis eventually led,\u201D \u201Cfrom that,\u201D \u201Ctheir holdings\u201D")

    h2("14. Two determiner problems")
    bullet("Passage opens with \u201Cthe bank\u201D before identifying which bank", bold_prefix="Premature \u201Cthe\u201D: ")
    bullet("Fix: Use \u201Ca major investment bank\u201D or name the bank first")
    bullet("\u201CMany investors\u201D (para 1) vs. \u201CA few executives\u201D (para 2) \u2014 unclear contrast", bold_prefix="Contradictory quantifiers: ")
    bullet("Fix: Clarify the different groups or use consistent quantifiers")

    h2("15. Two incorrect prepositions")
    bullet("\u201Ccompliance to\u201D \u2192 should be \u201Ccompliance with\u201D", bold_prefix="Error 1: ")
    bullet("\u201Cfiled on the bank\u201D \u2192 should be \u201Cfiled against the bank\u201D", bold_prefix="Error 2: ")

    h2("16. Rewrite one paragraph")
    first("See answer key for full sample rewrites.", bold_prefix="")
    bullet("Fix all premature \u201Cthe\u201D references")
    bullet("Replace vague pronouns with specific nouns")
    bullet("Correct \u201Ccompliance to\u201D \u2192 \u201Ccompliance with\u201D")
    bullet("Correct \u201Cfiled on\u201D \u2192 \u201Cfiled against\u201D")

    h1("End of Answer Key")

    doc.save(f"{OUTPUT_DIR}/Homework 06 Overhead.docx")
    print("Created Homework 06 Overhead.docx")


# ---------------------------------------------------------------------------
if __name__ == "__main__":
    build_homework()
    build_answer_key()
    build_overhead()

#!/usr/bin/env python3
"""
Generate Chapter 13 Answer Key and Overhead Answer Key .docx files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page, add_part_heading,
    get_font_config, add_bracket_line, add_diagram_image,
    add_multilevel_from_bracket, load_chapter_roles,
    parse_bracket_to_multilevel, add_multilevel_labeling_table,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch13'


DIAGRAM_EXERCISES = [
    {
        'num': 16, 'sentence': 'The student who won the award celebrated.',
        'words':   ['The', 'student', 'who', 'won', 'the', 'award', 'celebrated'],
        'roles':   ['Subj', '', '', '', '', '', 'Pred'],
        'phrases': ['NP', '', 'RC', 'VP', 'NP', '', 'VP'],
        'pos':     ['DET', 'N', 'REL', 'V', 'DET', 'N', 'V'],
        'bracket': '[S [NP [DET The] [N student] [RC [NP [REL who]] [VP [V won] [NP [DET the] [N award]]]]] [VP [V celebrated]]]',
        'diagram': 'ch13_hw_ex16_student_award',
    },
    {
        'num': 17, 'sentence': 'The students selected for the team celebrated.',
        'words':   ['The', 'students', 'selected', 'for', 'the', 'team', 'celebrated'],
        'roles':   ['Subj', '', '', '', '', '', 'Pred'],
        'phrases': ['NP', '', 'VP', 'PP', 'NP', '', 'VP'],
        'pos':     ['DET', 'N', 'V', 'PREP', 'DET', 'N', 'V'],
        'bracket': '[S [NP [DET The] [N students] [VP [V selected] [PP [PREP for] [NP [DET the] [N team]]]]] [VP [V celebrated]]]',
        'diagram': 'ch13_hw_ex17_selected_team',
    },
    {
        'num': 18, 'sentence': 'The team has a plan to win the tournament.',
        'words':   ['The', 'team', 'has', 'a', 'plan', 'to win', 'the', 'tournament'],
        'roles':   ['Subj', '', 'Pred', 'DO', '', '', '', ''],
        'phrases': ['NP', '', 'VP', 'NP', '', 'VP', 'NP', ''],
        'pos':     ['DET', 'N', 'V', 'DET', 'N', 'V', 'DET', 'N'],
        'bracket': '[S [NP [DET The] [N team]] [VP [V has] [NP [DET a] [N plan] [VP [V to win] [NP [DET the] [N tournament]]]]]]',
        'diagram': 'ch13_hw_ex18_plan_win',
    },
    {
        'num': 19, 'sentence': 'Running water flowed through the pipe.',
        'words':   ['Running', 'water', 'flowed', 'through', 'the', 'pipe'],
        'roles':   ['Subj', '', 'Pred', '', '', ''],
        'phrases': ['VP', '', 'VP', 'PP', 'NP', ''],
        'pos':     ['V', 'N', 'V', 'PREP', 'DET', 'N'],
        'bracket': '[S [NP [VP [V Running]] [N water]] [VP [V flowed] [PP [PREP through] [NP [DET the] [N pipe]]]]]',
        'diagram': 'ch13_hw_ex19_running_water',
    },
    {
        'num': 20, 'sentence': 'The woman wearing the red coat smiled.',
        'words':   ['The', 'woman', 'wearing', 'the', 'red', 'coat', 'smiled'],
        'roles':   ['Subj', '', '', '', '', '', 'Pred'],
        'phrases': ['NP', '', 'VP', 'NP', '', '', 'VP'],
        'pos':     ['DET', 'N', 'V', 'DET', 'ADJ', 'N', 'V'],
        'bracket': '[S [NP [DET The] [N woman] [VP [V wearing] [NP [DET the] [ADJP [ADJ red]] [N coat]]]] [VP [V smiled]]]',
        'diagram': 'ch13_hw_ex20_woman_coat',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 13 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 13: Adjectivals', cfg, overhead)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part_heading(doc, 'Part 1: Identification and Classification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'The book on the top shelf belongs to my professor.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'prepositional phrase', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "book" \u2014 tells which book', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 2
    add_exercise(doc, 2, 'The woman who won the award gave an inspiring speech.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'relative clause', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "woman" \u2014 identifies which woman', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 3
    add_exercise(doc, 3, 'The broken window needs to be repaired immediately.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'past participle (single-word adjectival)', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "window" \u2014 describes the window\u2019s state', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 4
    add_exercise(doc, 4, 'I need something to eat before the meeting.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'infinitive phrase', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "something" \u2014 specifies what kind of something', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 5
    add_exercise(doc, 5, 'The government report was released yesterday.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'noun (used as adjectival)', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "report" \u2014 classifies the type of report', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 6
    add_exercise(doc, 6, 'The students waiting in line seemed impatient.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'present participial phrase', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "students" \u2014 identifies which students', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 7
    add_exercise(doc, 7, 'We found a very comfortable chair at the antique store.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Form:', 'adjective phrase', body_size, font_name=body_font)
    add_plain_line(doc, 'Modifies "chair" \u2014 describes the chair', body_size, font_name=body_font)

    # =============================================
    # Part 2: Restrictive vs. Non-Restrictive
    # =============================================
    add_part_heading(doc, 'Part 2: Restrictive vs. Non-Restrictive', cfg, overhead)

    classifications = [
        (8, 'The students who completed the extra assignment received bonus points.',
         'Restrictive (R)',
         'The clause identifies which students received bonus points \u2014 '
         'only those who completed the extra assignment, not all students.',
         'No change needed \u2014 the sentence is correctly punctuated without commas.'),
        (9, 'The Eiffel Tower which was built in 1889 attracts millions of visitors.',
         'Non-restrictive (NR)',
         'The Eiffel Tower is already uniquely identified; '
         'the clause adds supplementary information about when it was built.',
         'The Eiffel Tower, which was built in 1889, attracts millions of visitors.'),
        (10, 'The car that I bought last year already needs repairs.',
         'Restrictive (R)',
         '"That" signals a restrictive clause identifying which car \u2014 '
         'specifically the one bought last year.',
         'No change needed \u2014 the sentence is correctly punctuated without commas.'),
        (11, 'Professor Adams who teaches linguistics won a research award.',
         'Non-restrictive (NR)',
         'Professor Adams is already uniquely identified by name; '
         'the clause adds extra information about what he teaches.',
         'Professor Adams, who teaches linguistics, won a research award.'),
    ]

    for i, (num, sentence, classification, explanation, rewrite) in enumerate(classifications):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Type:', classification, body_size, font_name=body_font)
        add_plain_line(doc, explanation, body_size, font_name=body_font)
        add_answer_line(doc, 'Rewrite:', rewrite, body_size, font_name=body_font)

    # =============================================
    # Part 3: Sentence Combining
    # =============================================
    add_part_heading(doc, 'Part 3: Sentence Combining', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 12\u201315 are open-ended. Accept any grammatically correct combination using the requested structure.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    combinations = [
        (12, 'Relative clause: This is the book. I told you about it.',
         '"This is the book that I told you about."'),
        (13, 'Relative clause: The scientist won a Nobel Prize. Her research changed medicine.',
         '"The scientist whose research changed medicine won a Nobel Prize."'),
        (14, 'Participial phrase: The students were exhausted from the exam. They went home early.',
         '"Exhausted from the exam, the students went home early."'),
        (15, 'Participial phrase: The letter was written in 1945. The letter was found in the attic.',
         '"Written in 1945, the letter was found in the attic." '
         'OR "The letter, written in 1945, was found in the attic."'),
    ]

    for i, (num, prompt, sample) in enumerate(combinations):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, f'Combine using the specified structure: {prompt}', body_size, font_name=body_font)
        add_plain_line(doc, prompt, body_size, bold_prefix='Prompt: ', font_name=body_font)
        answer_page_break(doc, overhead)
        add_plain_line(doc, f'Sample: {sample}', body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Adjectivals
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Adjectivals', cfg, overhead)

    ch_roles = load_chapter_roles(13)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], body_size, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Error Correction and Analysis
    # =============================================
    add_part_heading(doc, 'Part 5: Error Correction and Analysis', cfg, overhead)

    # Exercise 21: Dangling Participle Correction
    add_exercise(doc, 21, 'Correct each dangling participle:', body_size, font_name=body_font)

    danglers = [
        ('21A)', 'Walking through the park, the flowers were beautiful.',
         '"Walking through the park, I thought the flowers were beautiful." '
         'OR "As I walked through the park, the flowers were beautiful."',
         'The original implies the flowers were walking.'),
        ('21B)', 'Having finished the report, the computer was shut down.',
         '"Having finished the report, she shut down the computer."',
         'The original implies the computer finished the report.'),
        ('21C)', 'Exhausted from the journey, the bed looked inviting.',
         '"Exhausted from the journey, I thought the bed looked inviting."',
         'The original implies the bed was exhausted.'),
    ]

    answer_page_break(doc, overhead)
    for label, original, corrected, explanation in danglers:
        add_plain_line(doc, f'{label} {original}', body_size, indent=0.35, font_name=body_font)
        add_plain_line(doc, f'Corrected: {corrected}', body_size, indent=0.7, font_name=body_font)
        add_plain_line(doc, f'Explanation: {explanation}', body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 22: Meaning Analysis
    add_exercise(doc, 22, 'Explain the meaning difference between restrictive and non-restrictive versions:', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_plain_line(doc,
        '22A) "My brother who lives in Chicago is a doctor."',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Restrictive: implies the speaker has more than one brother. The clause '
        'identifies which brother \u2014 the one in Chicago (as opposed to brothers elsewhere).',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '22B) "My brother, who lives in Chicago, is a doctor."',
        body_size, font_name=body_font)
    add_plain_line(doc,
        'Non-restrictive: implies the speaker has only one brother. The clause '
        'adds supplementary information about where he lives; it doesn\u2019t serve '
        'to distinguish him from other brothers.',
        body_size, indent=0.7, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 13 Student Homework mirroring the HTML homework."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12
    fn = 'Garamond'

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 13 Homework: Adjectivals')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = fn
    set_paragraph_spacing(p, space_before=0, space_after=4)

    def add_part(title):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=10, space_after=4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = fn

    def add_text(text, bold=False, italic=False, indent=0, space_before=3, space_after=2):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn
        run.bold = bold
        run.italic = italic
        return p

    def add_ex(num, sentence, italic=True):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {num}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = fn
        run = p.add_run(sentence)
        run.italic = italic
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_blank(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=1, space_after=1)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_ref_table(headers, rows):
        """Add a simple reference table matching the HTML style."""
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(fs - 1)
            run.font.name = fn
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(fs - 1)
                run.font.name = fn
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=0, space_after=2)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part('Part 1: Identification and Classification')
    add_text('For each sentence, identify the underlined adjectival and classify its form '
             '(adjective phrase, prepositional phrase, relative clause, participial phrase, '
             'infinitive phrase, or noun).')

    add_ref_table(['Adjectival Forms', 'Example'], [
        ['Adjective Phrase (ADJP)', 'the tall man, very interesting'],
        ['Noun Adjectival', 'the coffee table, government report'],
        ['Prepositional Phrase (PP)', 'the book on the shelf'],
        ['Relative Clause (RC)', 'the student who won'],
        ['Participial Phrase (VP)', 'the woman singing, the report written'],
        ['Infinitive Phrase (VP)', 'a book to read'],
    ])

    add_text('Example (completed)', bold=True)
    add_text('The extremely talented musician performed last night.', italic=True)
    add_blank('   Form: adjective phrase')
    add_blank('   Function: modifies "musician"')

    add_text('Exercises', bold=True)

    add_ex(1, 'The book on the top shelf belongs to my professor.')
    add_blank('   Form:')

    add_ex(2, 'The woman who won the award gave an inspiring speech.')
    add_blank('   Form:')

    add_ex(3, 'The broken window needs to be repaired immediately.')
    add_blank('   Form:')

    add_ex(4, 'I need something to eat before the meeting.')
    add_blank('   Form:')

    add_ex(5, 'The government report was released yesterday.')
    add_blank('   Form:')

    add_ex(6, 'The students waiting in line seemed impatient.')
    add_blank('   Form:')

    add_ex(7, 'We found a very comfortable chair at the antique store.')
    add_blank('   Form:')

    # =============================================
    # Part 2: Restrictive vs. Non-Restrictive
    # =============================================
    add_part('Part 2: Restrictive vs. Non-Restrictive')
    add_text('The commas have been removed from the following sentences. For each, '
             'determine whether the modifier is restrictive (R) or non-restrictive (NR), '
             'explain your reasoning, and rewrite the sentence with correct punctuation if needed.')

    add_text('Example (completed)', bold=True)
    add_text('My sister who lives in Portland is visiting next week.', italic=True)
    add_blank('   Type: Non-restrictive (NR)')
    add_blank('   Reasoning: The speaker has only one sister, so the clause adds extra '
              'information rather than identifying which sister.')
    add_blank('   Rewrite: My sister, who lives in Portland, is visiting next week.')

    add_text('Exercises', bold=True)

    add_ex(8, 'The students who completed the extra assignment received bonus points.')
    add_blank('   Type:')
    add_blank('   Reasoning:')
    add_blank('   Rewrite:')

    add_ex(9, 'The Eiffel Tower which was built in 1889 attracts millions of visitors.')
    add_blank('   Type:')
    add_blank('   Reasoning:')
    add_blank('   Rewrite:')

    add_ex(10, 'The car that I bought last year already needs repairs.')
    add_blank('   Type:')
    add_blank('   Reasoning:')
    add_blank('   Rewrite:')

    add_ex(11, 'Professor Adams who teaches linguistics won a research award.')
    add_blank('   Type:')
    add_blank('   Reasoning:')
    add_blank('   Rewrite:')

    # =============================================
    # Part 3: Sentence Combining
    # =============================================
    add_part('Part 3: Sentence Combining')
    add_text('Combine each pair of sentences using a relative clause or participial phrase.')

    add_text('Example (completed)', bold=True)
    add_text('Sentences: I met a professor. She specializes in linguistics.', italic=True)
    add_blank('   Combined (relative clause): I met a professor who specializes in linguistics.')

    add_text('Exercises', bold=True)

    add_ex(12, 'Combine with a relative clause: This is the book. I told you about it.', italic=False)
    add_blank('')

    add_ex(13, 'Combine with a relative clause: The scientist won a Nobel Prize. Her research changed medicine.', italic=False)
    add_blank('')

    add_ex(14, 'Combine with a participial phrase: The students were exhausted from the exam. They went home early.', italic=False)
    add_blank('')

    add_ex(15, 'Combine with a participial phrase: The letter was written in 1945. The letter was found in the attic.', italic=False)
    add_blank('')

    # =============================================
    # Part 4: Diagramming Adjectivals
    # =============================================
    add_part('Part 4: Diagramming Adjectivals')
    add_text('For each sentence, complete the labeling table (Role, Phrase, Word, POS), '
             'write the bracket notation, and draw a tree diagram. Pay special attention '
             'to adjectival elements: relative clauses (RC), adjective phrases (ADJP), '
             'and prepositional phrases (PP) modifying nouns.')

    add_text('Exercises', bold=True)

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {ex["num"]}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = fn
        run = p.add_run(ex['sentence'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = fn

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = fn

    # =============================================
    # Part 5: Error Correction and Analysis
    # =============================================
    add_part('Part 5: Error Correction and Analysis')

    add_text('Dangling Participle Correction', bold=True)

    add_ex(21, 'Correct each dangling participle by rewriting the sentence:', italic=False)

    add_text('a) Walking through the park, the flowers were beautiful.', italic=True)
    add_blank('   Corrected:')

    add_text('b) Having finished the report, the computer was shut down.', italic=True)
    add_blank('   Corrected:')

    add_text('c) Exhausted from the journey, the bed looked inviting.', italic=True)
    add_blank('   Corrected:')

    add_text('Meaning Analysis', bold=True)

    add_ex(22, 'Explain the difference in meaning between these two sentences. '
           'What does each sentence imply about how many brothers the speaker has?', italic=False)

    add_text('a) My brother who lives in Chicago is a doctor.', italic=True)
    add_blank('   Sentence (a) implies:')

    add_text('b) My brother, who lives in Chicago, is a doctor.', italic=True)
    add_blank('   Sentence (b) implies:')

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 13 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 13 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 13 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Generate Chapter 8 Answer Key and Overhead Answer Key .docx files.
Updated to match revised homework structure (5 parts, 14 exercises).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, add_sub_sentence, setup_document, add_title_page,
    add_part_heading, get_font_config, add_diagram_image,
    add_bracket_line, blank_labels,
    add_multilevel_from_bracket, load_chapter_roles,
    question_page_break, answer_page_break,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch08'


DIAGRAM_EXERCISES = [
    {
        'sub': '11A)',
        'label': 'Pattern 1 (Intransitive): Birds sing.',
        'words':   ['Birds', 'sing'],
        'roles':   ['Subj', 'Pred'],
        'phrases': ['NP', 'VP'],
        'pos':     ['N', 'V'],
        'bracket': '[S [NP [N Birds]] [VP [V sing]]]',
        'diagram': 'ch08_hw_ex11a_birds_sing',
    },
    {
        'sub': '11B)',
        'label': 'Pattern 2 (Copular Be): The solution was simple.',
        'words':   ['The', 'solution', 'was', 'simple'],
        'roles':   ['Subj', '', 'Pred', 'SC'],
        'phrases': ['NP', '', 'VP', 'ADJP'],
        'pos':     ['DET', 'N', 'V', 'ADJ'],
        'bracket': '[S [NP [DET The] [N solution]] [VP [V was] [ADJP [ADJ simple]]]]',
        'diagram': 'ch08_hw_ex11b_solution_simple',
    },
    {
        'sub': '11C)',
        'label': 'Pattern 3 (Linking Verb): The music sounded beautiful.',
        'words':   ['The', 'music', 'sounded', 'beautiful'],
        'roles':   ['Subj', '', 'Pred', 'SC'],
        'phrases': ['NP', '', 'VP', 'ADJP'],
        'pos':     ['DET', 'N', 'V', 'ADJ'],
        'bracket': '[S [NP [DET The] [N music]] [VP [V sounded] [ADJP [ADJ beautiful]]]]',
        'diagram': 'ch08_hw_ex11c_music_sounded',
    },
    {
        'sub': '11D)',
        'label': 'Pattern 4 (Transitive): The student finished the report.',
        'words':   ['The', 'student', 'finished', 'the', 'report'],
        'roles':   ['Subj', '', 'Pred', 'DO', ''],
        'phrases': ['NP', '', 'VP', 'NP', ''],
        'pos':     ['DET', 'N', 'V', 'DET', 'N'],
        'bracket': '[S [NP [DET The] [N student]] [VP [V finished] [NP [DET the] [N report]]]]',
        'diagram': 'ch08_hw_ex11d_student_finished',
    },
    {
        'sub': '11E)',
        'label': 'Pattern 5 (Ditransitive, IO + DO): The professor gave the class a deadline.',
        'words':   ['The', 'professor', 'gave', 'the', 'class', 'a', 'deadline'],
        'roles':   ['Subj', '', 'Pred', 'IO', '', 'DO', ''],
        'phrases': ['NP', '', 'VP', 'NP', '', 'NP', ''],
        'pos':     ['DET', 'N', 'V', 'DET', 'N', 'DET', 'N'],
        'bracket': '[S [NP [DET The] [N professor]] [VP [V gave] [NP [DET the] [N class]] [NP [DET a] [N deadline]]]]',
        'diagram': 'ch08_hw_ex11e_professor_gave',
    },
    {
        'sub': '11F)',
        'label': 'Pattern 6 (Ditransitive, DO + OC): The board declared the plan inadequate.',
        'words':   ['The', 'board', 'declared', 'the', 'plan', 'inadequate'],
        'roles':   ['Subj', '', 'Pred', 'DO', '', 'OC'],
        'phrases': ['NP', '', 'VP', 'NP', '', 'ADJP'],
        'pos':     ['DET', 'N', 'V', 'DET', 'N', 'ADJ'],
        'bracket': '[S [NP [DET The] [N board]] [VP [V declared] [NP [DET the] [N plan]] [ADJP [ADJ inadequate]]]]',
        'diagram': 'ch08_hw_ex11f_board_declared',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 8: Basic Sentence Elements and Sentence Patterns', cfg, overhead)

    # =============================================
    # Part 1: Sentence Element Identification
    # =============================================
    add_part_heading(doc, 'Part 1: Sentence Element Identification', cfg, overhead)

    # Exercise 1: Identify DO/IO/SC/OC
    add_exercise(doc, 1, 'Identify the direct object, indirect object, subject complement, or object complement in each sentence.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_sub_sentence(doc, '1A)', 'The committee awarded the outstanding student a prestigious scholarship.', body_size, font_name=body_font)
    add_answer_line(doc, 'Indirect Object (IO):', 'the outstanding student', body_size, indent=0.7, font_name=body_font)
    add_answer_line(doc, 'Direct Object (DO):', 'a prestigious scholarship', body_size, indent=0.7, font_name=body_font)

    add_sub_sentence(doc, '1B)', 'The homemade soup tasted absolutely delicious.', body_size, font_name=body_font)
    add_answer_line(doc, 'Subject Complement (SC):', 'absolutely delicious (AdjP)', body_size, indent=0.7, font_name=body_font)

    add_sub_sentence(doc, '1C)', 'The judges declared the young contestant the winner.', body_size, font_name=body_font)
    add_answer_line(doc, 'Direct Object (DO):', 'the young contestant', body_size, indent=0.7, font_name=body_font)
    add_answer_line(doc, 'Object Complement (OC):', 'the winner (NP)', body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 2: Argument vs. adverbial
    add_exercise(doc, 2, 'Determine whether the underlined element is an argument (required) or an adverbial (optional). Explain your reasoning.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for sub, sentence, verdict, explanation in [
        ('2A)', 'She placed the documents on the desk.',
         'Argument (required)',
         '"On the desk" is required by "placed." Remove it: *She placed the documents. \u2717 \u2014 ungrammatical without a location argument.'),
        ('2B)', 'She found the documents on the desk.',
         'Adverbial (optional)',
         '"On the desk" is an optional location modifier. Remove it: She found the documents. \u2713 \u2014 still grammatical.'),
        ('2C)', 'The professor is extremely knowledgeable about linguistics.',
         'Argument (required)',
         '"Extremely knowledgeable about linguistics" is the subject complement required by "is." Remove it: *The professor is. \u2717 \u2014 incomplete.'),
        ('2D)', 'The professor lectured extremely knowledgeably about linguistics.',
         'Adverbial (optional)',
         '"Extremely knowledgeably about linguistics" is an optional manner/topic modifier. Remove it: The professor lectured. \u2713 \u2014 still grammatical.'),
    ]:
        add_sub_sentence(doc, sub, sentence, body_size, font_name=body_font)
        add_answer_line(doc, 'Answer:', verdict, body_size, indent=0.7, font_name=body_font)
        add_plain_line(doc, explanation, body_size, indent=0.7, font_name=body_font)

    # =============================================
    # Part 2: Sentence Pattern Identification
    # =============================================
    add_part_heading(doc, 'Part 2: Sentence Pattern Identification', cfg, overhead)

    patterns = [
        (3, 'The exhausted marathon runner collapsed at the finish line yesterday.',
         'Pattern 1 (Intransitive)',
         'Main verb: "collapsed." "At the finish line" and "yesterday" are adverbials (optional \u2014 answer where? and when?). '
         'Without adverbials: "The exhausted marathon runner collapsed." \u2014 complete with subject + intransitive verb. '
         '"Collapsed" does not require an object or complement.'),
        (4, "My grandmother's secret recipe remains a family treasure.",
         'Pattern 3 (Linking verb)',
         'Main verb: "remains." "A family treasure" is a subject complement (NP identifying the subject). '
         'Be substitution test: "My grandmother\'s secret recipe is a family treasure" \u2713. '
         'Since the verb is not "be" itself but passes the be-substitution test, this is Pattern 3 (Linking), not Pattern 2 (Copular be).'),
        (5, 'The committee considered the proposal inadequate.',
         'Pattern 6 (Ditransitive: DO + OC)',
         'Main verb: "considered." Two elements follow the verb: "the proposal" (NP) + "inadequate" (AdjP). '
         'Do they refer to the same thing? Yes \u2014 the proposal is described as inadequate. '
         'Therefore "the proposal" = Direct Object, "inadequate" = Object Complement.'),
        (6, 'The chef prepared the guests an extraordinary seven-course meal.',
         'Pattern 5 (Ditransitive: IO + DO)',
         'Main verb: "prepared." Two NPs follow the verb: "the guests" and "an extraordinary seven-course meal." '
         'Do they refer to the same thing? No \u2014 the guests \u2260 the meal. '
         'Can rephrase with "for": "prepared an extraordinary seven-course meal for the guests." '
         '"The guests" = Indirect Object, "an extraordinary seven-course meal" = Direct Object.'),
        (7, 'The situation grew increasingly tense during the negotiations.',
         'Pattern 3 (Linking verb)',
         'Main verb: "grew." "During the negotiations" is an adverbial (time) \u2014 set aside. '
         '"Increasingly tense" is a subject complement (AdjP describing the subject). '
         'Be substitution test: "The situation was increasingly tense" \u2713. Pattern 3 (Linking).'),
    ]

    for i, (num, sentence, pattern, explanation) in enumerate(patterns):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Pattern:', pattern, body_size, font_name=body_font)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run('Explanation: ')
        run.bold = True
        run.font.size = Pt(body_size)
        run.font.name = body_font
        run = p.add_run(explanation)
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=0, space_after=2)

    # =============================================
    # Part 3: Sentence Writing
    # =============================================
    add_part_heading(doc, 'Part 3: Sentence Writing', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 8\u201310 are open-ended. Accept any grammatically correct sentence that follows the requested pattern with elements correctly labeled.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    for i, (num, pattern_label, sample) in enumerate([
        (8, 'Pattern 4 (S + V + DO)',
         'Sample: "[The dog]_S [chased]_V [the cat]_DO."'),
        (9, 'Pattern 5 (S + V + IO + DO)',
         'Sample: "[The teacher]_S [gave]_V [the students]_IO [a quiz]_DO."'),
        (10, 'Pattern 6 (S + V + DO + OC)',
         'Sample: "[The class]_S [elected]_V [Maria]_DO [president]_OC."'),
    ]):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, f'Write a sentence using {pattern_label} and label each element.', body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_plain_line(doc, f'{pattern_label}:', body_size, bold_prefix='Pattern: ', font_name=body_font)
        add_plain_line(doc, sample, body_size, font_name=body_font)

    # =============================================
    # Part 4: Sentence Tables and Diagrams
    # =============================================
    add_part_heading(doc, 'Part 4: Sentence Tables and Diagrams', cfg, overhead)

    add_exercise(doc, 11, 'Complete each table and draw a tree diagram.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    ch_roles = load_chapter_roles(8)
    mode = 'overhead' if overhead else 'answer_key'
    for ex in DIAGRAM_EXERCISES:
        add_sub_sentence(doc, ex['sub'], ex['label'], body_size, font_name=body_font)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], body_size, indent=0.7, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Analysis and Reflection
    # =============================================
    add_part_heading(doc, 'Part 5: Analysis and Reflection', cfg, overhead)

    # Exercise 12: "put" valency
    add_exercise(doc, 12, 'She put the book on the shelf.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    for sub, answer in [
        ('What happens if you remove "the book"?',
         '*She put on the shelf. \u2717 \u2014 ungrammatical. "The book" is a required argument (Direct Object).'),
        ('What happens if you remove "on the shelf"?',
         '*She put the book. \u2717 \u2014 ungrammatical/incomplete. "On the shelf" is a required locative argument.'),
        ('What does this tell you about the valency of "put"?',
         '"Put" requires THREE arguments: a subject, a direct object, and a locative phrase. '
         'It has valency 3 \u2014 unusual among English verbs, most of which require only two.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        run = p.add_run(sub)
        run.bold = True
        run.font.size = Pt(body_size)
        run.font.name = body_font
        set_paragraph_spacing(p, space_before=3, space_after=2)
        add_plain_line(doc, answer, body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 13: linking vs transitive
    add_exercise(doc, 13, 'Explain the difference between the two uses of "smells" in the following sentences.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    add_sub_sentence(doc, '13A)', 'The milk smells sour. vs. The detective smells trouble.', body_size, font_name=body_font)
    add_plain_line(doc,
        '"The milk smells sour" \u2014 linking verb (Pattern 3). "Sour" is a subject complement describing the milk.',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        '"The detective smells trouble" \u2014 transitive verb (Pattern 4). "Trouble" is a direct object.',
        body_size, indent=0.7, font_name=body_font)
    add_plain_line(doc,
        'Be substitution test: "The milk is sour" \u2713 (makes sense \u2192 linking). '
        '"The detective is trouble" \u2717 (doesn\'t make sense \u2192 not linking, therefore transitive).',
        body_size, indent=0.7, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 14: argument vs adverbial reflection
    add_exercise(doc, 14, 'In your own words, explain the difference between an argument and an adverbial. Why does this distinction matter for identifying sentence patterns? Give an example.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run('Model response: ')
    run.bold = True
    run.font.size = Pt(body_size)
    run.font.name = body_font
    run = p.add_run(
        'Arguments are elements required by the verb to form a grammatical sentence; '
        'removing them makes the sentence ungrammatical or changes its meaning dramatically. '
        'Adverbials provide optional information about time, place, manner, or reason; '
        'removing them leaves a grammatical sentence intact. This distinction matters because '
        'sentence patterns are defined by the required elements (arguments), not the optional ones (adverbials). '
        'For example, in "She put the book on the table," "on the table" is an argument '
        '(removing it yields *She put the book \u2014 ungrammatical). But in "She read the book on the table," '
        '"on the table" is an adverbial (removing it yields She read the book \u2014 fine). '
        'The first sentence requires a locative argument; the second does not.'
    )
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=0, space_after=2)

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def create_student_homework(output_path):
    """Create the Chapter 8 Student Homework with blank multi-level tables."""
    from answer_key_helpers import set_paragraph_spacing, parse_bracket_to_multilevel, add_multilevel_labeling_table
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 8 Homework: Basic Sentence Elements and Sentence Patterns')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Part 4 with blank multi-level tables
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=10, space_after=4)
    run = p.add_run('Part 4: Sentence Tables and Diagrams')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    p = doc.add_paragraph()
    run = p.add_run('Exercise 11. ')
    run.bold = True
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'
    run = p.add_run('Complete each table and draw a tree diagram.')
    run.font.size = Pt(fs)
    run.font.name = 'Garamond'

    for ex in DIAGRAM_EXERCISES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'{ex["sub"]} ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'
        run = p.add_run(ex['label'])
        run.italic = True
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = 'Garamond'

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 08 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 08 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 08 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

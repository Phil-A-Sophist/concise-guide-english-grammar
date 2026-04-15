#!/usr/bin/env python3
"""
Generate Chapter 12 Answer Key and Overhead Answer Key .docx files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

from answer_key_helpers import (
    set_paragraph_spacing, add_spacer_row, add_exercise, add_answer_line,
    add_plain_line, setup_document, add_title_page, add_part_heading,
    question_page_break, answer_page_break, get_font_config, add_bracket_line,
    add_multilevel_from_bracket, load_chapter_roles,
    parse_bracket_to_multilevel, add_multilevel_labeling_table, add_diagram_image,
)


DIAGRAM_DIR = Path(__file__).parent.parent / 'Homework' / 'diagrams' / 'ch12'


DIAGRAM_EXERCISES = [
    {
        'num': 14, 'sentence': 'She spoke very clearly.',
        'words':   ['She', 'spoke', 'very', 'clearly'],
        'roles':   ['Subj', 'Pred', '', ''],
        'phrases': ['NP', 'VP', 'ADVP', ''],
        'pos':     ['PRON', 'V', 'ADV', 'ADV'],
        'bracket': '[S [NP [PRON She]] [VP [V spoke] [ADVP [ADV very] [ADV clearly]]]]',
        'diagram': 'ch12_hw_ex14_spoke_clearly',
    },
    {
        'num': 15, 'sentence': 'The train arrived after midnight.',
        'words':   ['The', 'train', 'arrived', 'after', 'midnight'],
        'roles':   ['Subj', '', 'Pred', 'Advl', ''],
        'phrases': ['NP', '', 'VP', 'PP', ''],
        'pos':     ['DET', 'N', 'V', 'PREP', 'N'],
        'bracket': '[S [NP [DET The] [N train]] [VP [V arrived] [PP [PREP after] [NP [N midnight]]]]]',
        'diagram': 'ch12_hw_ex15_train_arrived',
    },
    {
        'num': 16, 'sentence': 'He walked slowly through the park.',
        'words':   ['He', 'walked', 'slowly', 'through', 'the', 'park'],
        'roles':   ['Subj', 'Pred', '', 'Advl', '', ''],
        'phrases': ['NP', 'VP', 'ADVP', 'PP', 'NP', ''],
        'pos':     ['PRON', 'V', 'ADV', 'PREP', 'DET', 'N'],
        'bracket': '[S [NP [PRON He]] [VP [V walked] [ADVP [ADV slowly]] [PP [PREP through] [NP [DET the] [N park]]]]]',
        'diagram': 'ch12_hw_ex16_walked_park',
    },
    {
        'num': 17, 'sentence': 'Unfortunately, the game was cancelled.',
        'words':   ['Unfortunately', 'the', 'game', 'was', 'cancelled'],
        'roles':   ['Advl', 'Subj', '', 'Pred', ''],
        'phrases': ['ADVP', 'NP', '', 'VP', ''],
        'pos':     ['ADV', 'DET', 'N', 'AUX', 'V'],
        'bracket': '[S [ADVP [ADV Unfortunately]] [NP [DET the] [N game]] [VP [AUX was] [V cancelled]]]',
        'diagram': 'ch12_hw_ex17_unfortunately',
    },
    {
        'num': 18, 'sentence': 'She left early because the roads were icy.',
        'words':   ['She', 'left', 'early', 'because', 'the', 'roads', 'were', 'icy'],
        'roles':   ['Subj', 'Pred', '', 'Advl', '', '', '', ''],
        'phrases': ['NP', 'VP', 'ADVP', 'SBAR', 'NP', '', 'VP', 'ADJP'],
        'pos':     ['PRON', 'V', 'ADV', 'COMP', 'DET', 'N', 'V', 'ADJ'],
        'bracket': '[S [NP [PRON She]] [VP [V left] [ADVP [ADV early]]] [SBAR [COMP because] [S [NP [DET the] [N roads]] [VP [V were] [ADJP [ADJ icy]]]]]]',
        'diagram': 'ch12_hw_ex18_left_early',
    },
]


def create_answer_key(output_path, font_size=12, overhead=False):
    """Create the Chapter 12 Answer Key document."""
    doc = Document()
    cfg = setup_document(doc, overhead)
    body_font = cfg['body_font']
    body_size = cfg['body_size']

    add_title_page(doc, 'Chapter 12: Adverbials', cfg, overhead)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part_heading(doc, 'Part 1: Identification and Classification', cfg, overhead)

    # Exercise 1
    add_exercise(doc, 1, 'Last week, the students studied diligently in the library.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Adverbial 1:', 'Last week \u2014 NP \u2014 time', body_size, font_name=body_font)
    add_answer_line(doc, 'Adverbial 2:', 'diligently \u2014 AdvP \u2014 manner', body_size, font_name=body_font)
    add_answer_line(doc, 'Adverbial 3:', 'in the library \u2014 PP \u2014 place', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 2
    add_exercise(doc, 2, 'If you need assistance, call the help desk immediately.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Adverbial 1:', 'If you need assistance \u2014 Dependent Clause \u2014 condition', body_size, font_name=body_font)
    add_answer_line(doc, 'Adverbial 2:', 'immediately \u2014 AdvP \u2014 time', body_size, font_name=body_font)

    question_page_break(doc, overhead)

    # Exercise 3
    add_exercise(doc, 3, 'She left early to catch her flight.', body_size, font_name=body_font)
    answer_page_break(doc, overhead)
    add_answer_line(doc, 'Adverbial 1:', 'early \u2014 AdvP \u2014 time', body_size, font_name=body_font)
    add_answer_line(doc, 'Adverbial 2:', 'to catch her flight \u2014 infinitive phrase \u2014 purpose', body_size, font_name=body_font)

    # =============================================
    # Part 2: Adverbial Scope and Form
    # =============================================
    add_part_heading(doc, 'Part 2: Adverbial Scope and Form', cfg, overhead)

    scope_exercises = [
        (4, 'She studied with great focus before the exam.',
         'the verb studied (tells how she studied \u2014 manner)',
         'PP'),
        (5, 'Smiling broadly, she accepted the award.',
         'the verb accepted (tells the circumstance of accepting)',
         'Present Participial Phrase'),
        (6, 'The results were surprisingly clear.',
         'the adjective clear (intensifies/evaluates the adjective)',
         'AdvP'),
        (7, 'He practiced every morning before the competition.',
         'the verb practiced (tells when he practiced \u2014 frequency)',
         'NP'),
        (8, 'She left early to catch the train.',
         'the verb left (tells why she left \u2014 purpose)',
         'Infinitive Phrase'),
    ]

    for i, (num, sentence, modifies, form) in enumerate(scope_exercises):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, sentence, body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_answer_line(doc, 'Modifies:', modifies, body_size, font_name=body_font)
        add_answer_line(doc, 'Form:', form, body_size, font_name=body_font)

    # =============================================
    # Part 3: Sentence Completion
    # =============================================
    add_part_heading(doc, 'Part 3: Sentence Completion', cfg, overhead)

    p = doc.add_paragraph()
    run = p.add_run('Exercises 9\u201313 are open-ended. Accept any grammatically correct adverbial of the requested type.')
    run.font.size = Pt(body_size)
    run.font.name = body_font
    set_paragraph_spacing(p, space_before=3, space_after=6)

    completions = [
        (9, 'Prepositional phrase: __________, the committee will announce its decision.',
         '"After the meeting, the committee will announce its decision."'),
        (10, 'Dependent clause: She stayed home __________.',
         '"She stayed home because she was feeling ill."'),
        (11, 'Infinitive phrase: He went to the store __________.',
         '"He went to the store to buy groceries."'),
        (12, 'Past participial phrase: __________, the runners collapsed at the finish line.',
         '"Exhausted from the race, the runners collapsed at the finish line."'),
        (13, 'Present participial phrase: __________, she answered all the questions correctly.',
         '"Having studied all night, she answered all the questions correctly."'),
    ]

    for i, (num, prompt, sample) in enumerate(completions):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, num, f'Complete the sentence with the specified adverbial type: {prompt}', body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        add_plain_line(doc, prompt, body_size, bold_prefix='Prompt: ', font_name=body_font)
        add_plain_line(doc, f'Sample: {sample}', body_size, font_name=body_font)

    # =============================================
    # Part 4: Diagramming Adverbials
    # =============================================
    add_part_heading(doc, 'Part 4: Diagramming Adverbials', cfg, overhead)

    ch_roles = load_chapter_roles(12)
    mode = 'overhead' if overhead else 'answer_key'
    for i, ex in enumerate(DIAGRAM_EXERCISES):
        if i > 0:
            question_page_break(doc, overhead)
        add_exercise(doc, ex['num'], ex['sentence'], body_size, font_name=body_font)
        answer_page_break(doc, overhead)
        bracket_key = ' '.join(ex['bracket'].split())
        add_multilevel_from_bracket(doc, ex['bracket'],
                                     roles_dict=ch_roles.get(bracket_key),
                                     mode=mode, font_size=body_size)
        add_bracket_line(doc, ex['bracket'], body_size, font_name=body_font)
        add_diagram_image(doc, DIAGRAM_DIR, ex['diagram'], width_inches=cfg['diagram_width'])

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part_heading(doc, 'Part 5: Analysis and Application', cfg, overhead)

    # Exercise 19
    add_exercise(doc, 19, 'Identify five adverbials in the passage:', body_size, font_name=body_font)
    answer_page_break(doc, overhead)

    adverbials = [
        ('Yesterday', 'NP', 'time'),
        ('finally', 'AdvP', 'time (completion)'),
        ('Surprisingly', 'AdvP (sentence-level)', 'speaker attitude'),
        ('diligently', 'AdvP', 'manner'),
        ('for three years', 'PP', 'time (duration)'),
        ('because funding was severely limited', 'Dependent Clause', 'reason'),
        ('in a prestigious journal', 'PP', 'place'),
        ('last month', 'NP', 'time'),
        ('If additional funding becomes available', 'Dependent Clause', 'condition'),
        ('next year', 'NP', 'time'),
        ('in a new laboratory', 'PP', 'place'),
    ]

    add_plain_line(doc, 'Any five of the following are acceptable:', body_size, font_name=body_font)

    for adv, form, role in adverbials:
        add_plain_line(doc, f'"{adv}" \u2014 {form} \u2014 {role}', body_size, indent=0.7, font_name=body_font)

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_student_homework(output_path):
    """Create the Chapter 12 Student Homework mirroring the HTML homework."""
    from answer_key_helpers import parse_bracket_to_multilevel, add_multilevel_labeling_table
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    fs = 12
    fn = 'Garamond'

    # Set landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    run = p.add_run('Chapter 12 Homework: Adverbials')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = fn
    set_paragraph_spacing(p, space_before=0, space_after=4)

    def add_part(title):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=10, space_after=4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = fn

    def add_text(text, bold=False, italic=False, indent=0, space_before=3, space_after=2):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn
        run.bold = bold
        run.italic = italic
        return p

    def add_ex(num, sentence, italic=True):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=6, space_after=2)
        run = p.add_run(f'Exercise {num}. ')
        run.bold = True
        run.font.size = Pt(fs)
        run.font.name = fn
        run = p.add_run(sentence)
        run.italic = italic
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_blank(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=1, space_after=1)
        run = p.add_run(text)
        run.font.size = Pt(fs)
        run.font.name = fn

    def add_ref_table(headers, rows):
        """Add a simple reference table matching the HTML style."""
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(fs - 1)
            run.font.name = fn
        # Data rows
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(fs - 1)
                run.font.name = fn
        # Add spacing after table
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=0, space_after=2)

    # =============================================
    # Part 1: Identification and Classification
    # =============================================
    add_part('Part 1: Identification and Classification')
    add_text('For each sentence, identify all adverbials, their form, and their semantic role. Use the reference tables below.')

    add_ref_table(['Adverbial Forms'], [
        ['AdvP (Adverb Phrase)'], ['PP (Prepositional Phrase)'], ['NP (Noun Phrase)'],
        ['Infinitive Phrase'], ['Present Participial Phrase'],
        ['Past Participial Phrase'], ['Dependent Clause'],
    ])

    add_ref_table(['Semantic Role', 'Question Answered'], [
        ['Time', 'When? How long? How often?'],
        ['Place', 'Where? In what direction?'],
        ['Manner', 'How? In what way?'],
        ['Reason', 'Why? For what reason?'],
        ['Purpose', 'What for? In order to?'],
        ['Condition', 'Under what circumstances?'],
        ['Concession', 'Despite what?'],
    ])

    add_text('Example (completed)', bold=True)
    add_text('Yesterday, she worked carefully at the office because she had a deadline.', italic=True)

    add_ref_table(['Adverbial', 'Form', 'Semantic Role'], [
        ['Yesterday', 'NP', 'time'],
        ['carefully', 'AdvP', 'manner'],
        ['at the office', 'PP', 'place'],
        ['because she had a deadline', 'Dependent Clause', 'reason'],
    ])

    add_text('Exercises', bold=True)

    add_ex(1, 'Last week, the students studied diligently in the library.')
    for i in range(1, 4):
        add_blank(f'   Adverbial {i}: __________ Form: __________ Semantic role: __________')

    add_ex(2, 'If you need assistance, call the help desk immediately.')
    for i in range(1, 3):
        add_blank(f'   Adverbial {i}: __________ Form: __________ Semantic role: __________')

    add_ex(3, 'She left early to catch her flight.')
    for i in range(1, 3):
        add_blank(f'   Adverbial {i}: __________ Form: __________ Semantic role: __________')

    # =============================================
    # Part 2: Adverbial Scope and Form
    # =============================================
    add_part('Part 2: Adverbial Scope and Form')
    add_text('For each underlined adverbial, state (a) what it modifies and (b) its form. Use the reference tables below.')

    add_ref_table(['What Adverbials Can Modify'], [
        ['a verb'], ['an adjective'], ['another adverb'], ['the whole sentence'],
    ])

    add_ref_table(['Adverbial Forms', 'Example'], [
        ['AdvP (Adverb Phrase)', 'quickly, very carefully'],
        ['PP (Prepositional Phrase)', 'in the library, with care'],
        ['NP (Noun Phrase)', 'last week, every morning'],
        ['Infinitive Phrase', 'to catch the train'],
        ['Present Participial Phrase', 'knowing the answer'],
        ['Past Participial Phrase', 'exhausted from the race'],
        ['Dependent Clause', 'because it rained'],
    ])

    add_text('Example (completed)', bold=True)
    add_text('Fortunately, the rain stopped before the game.', italic=True)
    add_text('Fortunately \u2014 modifies: the whole sentence; form: AdvP')

    add_text('Exercises', bold=True)

    part2_exercises = [
        (4, 'She studied with great focus before the exam.'),
        (5, 'Smiling broadly, she accepted the award.'),
        (6, 'The results were surprisingly clear.'),
        (7, 'He practiced every morning before the competition.'),
        (8, 'She left early to catch the train.'),
    ]
    for num, sentence in part2_exercises:
        add_ex(num, sentence)
        add_blank('   Modifies: __________')
        add_blank('   Form: __________')

    # =============================================
    # Part 3: Sentence Completion
    # =============================================
    add_part('Part 3: Sentence Completion')
    add_text('Add an adverbial of the requested form to complete each sentence.')

    add_text('Example (completed)', bold=True)
    add_text('Add a dependent clause: __________, you should bring an umbrella.')
    add_text('Answer: If it looks like rain, you should bring an umbrella.', italic=True)

    add_text('Exercises', bold=True)

    part3_exercises = [
        (9, 'Add a prepositional phrase: __________, the committee will announce its decision.'),
        (10, 'Add a dependent clause: She stayed home __________.'),
        (11, 'Add an infinitive phrase: He went to the store __________.'),
        (12, 'Add a past participial phrase: __________, the runners collapsed at the finish line.'),
        (13, 'Add a present participial phrase: __________, she answered all the questions correctly.'),
    ]
    for num, prompt in part3_exercises:
        add_ex(num, prompt, italic=False)
        add_blank('   Answer: __________')

    # =============================================
    # Part 4: Diagramming Adverbials
    # =============================================
    add_part('Part 4: Diagramming Adverbials')
    add_text('For each sentence, complete the labeling table, write the bracket notation, and draw a tree diagram.')

    for ex in DIAGRAM_EXERCISES:
        add_ex(ex['num'], ex['sentence'])

        td = parse_bracket_to_multilevel(ex['bracket'])
        add_multilevel_labeling_table(doc, td, mode='student', font_size=fs)

        p = doc.add_paragraph()
        run = p.add_run('Bracket notation: _____')
        run.font.size = Pt(fs)
        run.font.name = fn

    # =============================================
    # Part 5: Analysis and Application
    # =============================================
    add_part('Part 5: Analysis and Application')
    add_text('Read the passage and answer the question.')

    add_text('Passage', bold=True)
    add_text(
        'Yesterday, the researchers finally completed their groundbreaking study. '
        'Surprisingly, the results contradicted earlier findings. They had worked '
        'diligently for three years because funding was severely limited. Nevertheless, '
        'they published their findings in a prestigious journal last month. If additional '
        'funding becomes available, they will continue their research next year in a new laboratory.',
        italic=True
    )

    add_ex(19, 'Identify five adverbials in the passage. For each, state its form and semantic role.', italic=False)
    for i in range(1, 6):
        add_blank(f'   Adverbial {i}: __________ Form: __________ Semantic role: __________')

    doc.save(str(output_path))
    print(f'Created: {output_path}')


def main():
    script_dir = Path(__file__).parent
    homework_dir = script_dir.parent / 'Homework'

    create_student_homework(
        homework_dir / 'Student' / 'Chapter 12 Homework.docx'
    )

    create_answer_key(
        homework_dir / 'Answer Keys' / 'Chapter 12 Answer Key.docx',
        font_size=12
    )

    create_answer_key(
        homework_dir / 'Overheads' / 'Homework 12 Overhead.docx',
        overhead=True
    )


if __name__ == '__main__':
    main()

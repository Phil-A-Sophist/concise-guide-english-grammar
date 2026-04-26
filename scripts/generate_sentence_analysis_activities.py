#!/usr/bin/env python3
"""
Generate three Sentence Analysis Activities (student + answer key each).
Activity 4: Sentence Patterns
Activity 5: Tense and Aspect
Activity 6: Compound and Complex Sentences
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(__file__).parent.parent / 'Homework'


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def make_doc(title):
    """Create a portrait document with standard formatting."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)

    # Title
    h = doc.add_heading(title, level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].font.name = 'Open Sans'
    set_paragraph_spacing(h, space_before=0, space_after=4)

    return doc


def add_name_date(doc):
    p = doc.add_paragraph()
    run = p.add_run('Name: ___________________________    Date: _______________')
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=8)


def add_instructions(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('Instructions: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=6)
    return p


def add_reference_block(doc, lines):
    """Add a block of reference info (patterns, roles, etc.)."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        # Check for bold prefix
        if line.startswith('**') and '**' in line[2:]:
            end = line.index('**', 2)
            bold_text = line[2:end]
            rest = line[end+2:]
            run = p.add_run(bold_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Garamond'
            if rest:
                run = p.add_run(rest)
                run.font.size = Pt(11)
                run.font.name = 'Garamond'
        else:
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Garamond'
        set_paragraph_spacing(p, space_before=0, space_after=1)


def add_sentence(doc, num, sentence, font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(f'{num}.  ')
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = 'Garamond'
    run = p.add_run(sentence)
    run.italic = True
    run.font.size = Pt(font_size)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=6, space_after=2)
    return p


def add_blank_line(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    run = p.add_run('___________________________________________')
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=2)


def add_answer_line(doc, label, answer):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    run = p.add_run(answer)
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=2)


def add_blank_rewrite(doc, target_tense):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(f'Rewrite in the {target_tense}: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    run = p.add_run('___________________________________________')
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    set_paragraph_spacing(p, space_before=0, space_after=2)


# ============================================================
# DATA
# ============================================================

PATTERN_SENTENCES = [
    {
        'sentence': 'The old lighthouse still stands on the rocky coast.',
        'pattern': 'Pattern 1: Intransitive (S + V)',
        'roles': 'Subject: The old lighthouse | Predicate: stands | Adverbial: still, on the rocky coast',
    },
    {
        'sentence': 'The committee has reviewed all of the applications.',
        'pattern': 'Pattern 4: Transitive (S + V + DO)',
        'roles': 'Subject: The committee | Predicate: has reviewed | Direct Object: all of the applications',
    },
    {
        'sentence': 'My grandmother was a talented seamstress.',
        'pattern': 'Pattern 2: Copular Be (S + Be + SC)',
        'roles': 'Subject: My grandmother | Predicate: was | Subject Complement: a talented seamstress',
    },
    {
        'sentence': 'The bread smells wonderful this morning.',
        'pattern': 'Pattern 3: Linking Verb (S + LV + SC)',
        'roles': 'Subject: The bread | Predicate: smells | Subject Complement: wonderful | Adverbial: this morning',
    },
    {
        'sentence': 'The librarian handed me a dusty first edition.',
        'pattern': 'Pattern 5: Ditransitive (S + V + IO + DO)',
        'roles': 'Subject: The librarian | Predicate: handed | Indirect Object: me | Direct Object: a dusty first edition',
    },
    {
        'sentence': 'Rain has been falling steadily all afternoon.',
        'pattern': 'Pattern 1: Intransitive (S + V)',
        'roles': 'Subject: Rain | Predicate: has been falling | Adverbial: steadily, all afternoon',
    },
    {
        'sentence': 'The jury found the defendant guilty.',
        'pattern': 'Pattern 6: Ditransitive (S + V + DO + OC)',
        'roles': 'Subject: The jury | Predicate: found | Direct Object: the defendant | Object Complement: guilty',
    },
    {
        'sentence': 'She had already finished her dissertation before the deadline.',
        'pattern': 'Pattern 4: Transitive (S + V + DO)',
        'roles': 'Subject: She | Predicate: had finished | Direct Object: her dissertation | Adverbial: already, before the deadline',
    },
    {
        'sentence': 'The company will offer employees a new benefits package.',
        'pattern': 'Pattern 5: Ditransitive (S + V + IO + DO)',
        'roles': 'Subject: The company | Predicate: will offer | Indirect Object: employees | Direct Object: a new benefits package',
    },
    {
        'sentence': 'The children were building a sandcastle at the beach.',
        'pattern': 'Pattern 4: Transitive (S + V + DO)',
        'roles': 'Subject: The children | Predicate: were building | Direct Object: a sandcastle | Adverbial: at the beach',
    },
]

TENSE_SENTENCES = [
    {
        'sentence': 'The dog barks loudly.',
        'tense': 'Present Simple',
        'target': 'past simple',
        'rewrite': 'The dog barked loudly.',
    },
    {
        'sentence': 'She painted a picture.',
        'tense': 'Past Simple',
        'target': 'present perfect',
        'rewrite': 'She has painted a picture.',
    },
    {
        'sentence': 'They are running in the park.',
        'tense': 'Present Progressive',
        'target': 'past progressive',
        'rewrite': 'They were running in the park.',
    },
    {
        'sentence': 'He has read the book.',
        'tense': 'Present Perfect',
        'target': 'past simple',
        'rewrite': 'He read the book.',
    },
    {
        'sentence': 'The baby was crying.',
        'tense': 'Past Progressive',
        'target': 'present progressive',
        'rewrite': 'The baby is crying.',
    },
    {
        'sentence': 'We had left the building.',
        'tense': 'Past Perfect',
        'target': 'present perfect',
        'rewrite': 'We have left the building.',
    },
    {
        'sentence': 'The flowers bloom beautifully.',
        'tense': 'Present Simple',
        'target': 'present progressive',
        'rewrite': 'The flowers are blooming beautifully.',
    },
    {
        'sentence': 'She has been studying all day.',
        'tense': 'Present Perfect Progressive',
        'target': 'past simple',
        'rewrite': 'She studied all day.',
    },
    {
        'sentence': 'The ice melted quickly.',
        'tense': 'Past Simple',
        'target': 'past progressive',
        'rewrite': 'The ice was melting quickly.',
    },
    {
        'sentence': 'He had been waiting for hours.',
        'tense': 'Past Perfect Progressive',
        'target': 'present perfect progressive',
        'rewrite': 'He has been waiting for hours.',
    },
]

COMPOUND_COMPLEX_SENTENCES = [
    {
        'sentence': 'Marcus and Elena traveled to Spain last summer.',
        'classification': 'Simple sentence with compound subject',
        'keywords': 'Coordinating conjunction: and | Compound element: compound subject (Marcus and Elena)',
    },
    {
        'sentence': 'The critics panned the film, but audiences loved it.',
        'classification': 'Compound sentence',
        'keywords': 'Coordinating conjunction: but | Two independent clauses (Main + Main)',
    },
    {
        'sentence': 'The dog barked loudly and chased the squirrel across the yard.',
        'classification': 'Simple sentence with compound predicate',
        'keywords': 'Coordinating conjunction: and | Compound element: compound predicate (barked ... and chased ...)',
    },
    {
        'sentence': 'Although the task was difficult, the team persevered.',
        'classification': 'Complex sentence',
        'keywords': 'Subordinating conjunction: Although | Dependent clause: Although the task was difficult (Adverbial) | Independent clause: the team persevered (Main)',
    },
    {
        'sentence': 'The restaurant serves breakfast, lunch, and dinner.',
        'classification': 'Simple sentence with compound direct object',
        'keywords': 'Coordinating conjunction: and | Compound element: compound direct object (breakfast, lunch, and dinner)',
    },
    {
        'sentence': 'The experiment failed; however, we learned from the results.',
        'classification': 'Two simple sentences joined by a conjunctive adverb',
        'keywords': 'Conjunctive adverb: however | Sentence 1: The experiment failed (Main) | Sentence 2: we learned from the results (Main)',
    },
    {
        'sentence': 'The children played inside because it was raining.',
        'classification': 'Complex sentence',
        'keywords': 'Subordinating conjunction: because | Independent clause: The children played inside (Main) | Dependent clause: because it was raining (Adverbial)',
    },
    {
        'sentence': 'She studied all weekend, so she felt confident on the exam.',
        'classification': 'Compound sentence',
        'keywords': 'Coordinating conjunction: so | Two independent clauses (Main + Main)',
    },
    {
        'sentence': 'The hikers packed their supplies and checked the weather forecast before departure.',
        'classification': 'Simple sentence with compound predicate',
        'keywords': 'Coordinating conjunction: and | Compound element: compound predicate (packed ... and checked ...)',
    },
    {
        'sentence': 'She looked under the bed and behind the dresser.',
        'classification': 'Simple sentence with compound prepositional phrase',
        'keywords': 'Coordinating conjunction: and | Compound element: compound PP (under the bed and behind the dresser)',
    },
]


# ============================================================
# GENERATORS
# ============================================================

def generate_activity_4_student():
    doc = make_doc('Activity 4: Sentence Patterns')
    add_name_date(doc)
    add_instructions(doc,
        'For each sentence below, identify the sentence pattern (1\u20136). '
        'Then open the sentence in SyntaxTreeHybrid and assign a role to each '
        'phrase in the sentence.')

    # Spacer
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=2)

    add_reference_block(doc, [
        '**Sentence Patterns:**',
        'Pattern 1: Intransitive (S + V)',
        'Pattern 2: Copular Be (S + Be + SC)',
        'Pattern 3: Linking Verb (S + LV + SC)',
        'Pattern 4: Transitive (S + V + DO)',
        'Pattern 5: Ditransitive (S + V + IO + DO)',
        'Pattern 6: Ditransitive (S + V + DO + OC)',
    ])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=2, space_after=2)

    add_reference_block(doc, [
        '**Assignable Roles:** Subject, Predicate, Direct Object, Indirect Object, '
        'Subject Complement, Object Complement, Adverbial',
    ])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=4)

    for i, item in enumerate(PATTERN_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_blank_line(doc, 'Pattern')

    path = OUTPUT_DIR / 'Activity_4_Sentence_Patterns.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def generate_activity_4_key():
    doc = make_doc('Activity 4: Sentence Patterns \u2014 Answer Key')

    for i, item in enumerate(PATTERN_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_answer_line(doc, 'Pattern', item['pattern'])
        add_answer_line(doc, 'Roles', item['roles'])

    path = OUTPUT_DIR / 'Activity_4_Sentence_Patterns_Answer_Key.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def generate_activity_5_student():
    doc = make_doc('Activity 5: Tense and Aspect')
    add_name_date(doc)
    add_instructions(doc,
        'For each sentence below, (1) identify the tense and aspect of the main verb, '
        'then (2) rewrite the sentence in the target tense/aspect. Change only the '
        'verb form\u2014keep the rest of the sentence the same.')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=2)

    add_reference_block(doc, [
        '**Tenses:** Present, Past',
        '**Aspects:** Simple, Progressive, Perfect, Perfect Progressive',
    ])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=4)

    for i, item in enumerate(TENSE_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_blank_line(doc, 'Tense/Aspect')
        add_blank_rewrite(doc, item['target'])

    path = OUTPUT_DIR / 'Activity_5_Tense_and_Aspect.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def generate_activity_5_key():
    doc = make_doc('Activity 5: Tense and Aspect \u2014 Answer Key')

    for i, item in enumerate(TENSE_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_answer_line(doc, 'Tense/Aspect', item['tense'])
        add_answer_line(doc, f'Rewrite in the {item["target"]}', item['rewrite'])

    path = OUTPUT_DIR / 'Activity_5_Tense_and_Aspect_Answer_Key.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def generate_activity_6_student():
    doc = make_doc('Activity 6: Compound and Complex Sentences')
    add_name_date(doc)
    add_instructions(doc,
        'Classify each sentence as one of the following, then identify the '
        'coordinating conjunction, conjunctive adverb, or subordinating conjunction.')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=2)

    add_reference_block(doc, [
        '**Classifications:**',
        '(a) Simple sentence with a compound element '
        '(compound subject, compound predicate, compound object, or compound PP)',
        '(b) Compound sentence (two independent clauses joined by a coordinating conjunction)',
        '(c) Two simple sentences joined by a conjunctive adverb',
        '(d) Complex sentence (independent clause + dependent clause)',
    ])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=2, space_after=2)

    add_reference_block(doc, [
        '**Clause Roles:** Main (independent clause), Adverbial (dependent clause)',
    ])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=0, space_after=4)

    for i, item in enumerate(COMPOUND_COMPLEX_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_blank_line(doc, 'Classification')
        add_blank_line(doc, 'Key word(s) / element')

    path = OUTPUT_DIR / 'Activity_6_Compound_and_Complex.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def generate_activity_6_key():
    doc = make_doc('Activity 6: Compound and Complex Sentences \u2014 Answer Key')

    for i, item in enumerate(COMPOUND_COMPLEX_SENTENCES, 1):
        add_sentence(doc, i, item['sentence'])
        add_answer_line(doc, 'Classification', item['classification'])
        add_answer_line(doc, 'Key word(s) / element', item['keywords'])

    path = OUTPUT_DIR / 'Activity_6_Compound_and_Complex_Answer_Key.docx'
    doc.save(str(path))
    print(f'Saved: {path}')


def main():
    generate_activity_4_student()
    generate_activity_4_key()
    generate_activity_5_student()
    generate_activity_5_key()
    generate_activity_6_student()
    generate_activity_6_key()
    print('\nAll 6 documents generated.')


if __name__ == '__main__':
    main()

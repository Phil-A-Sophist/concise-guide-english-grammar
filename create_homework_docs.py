#!/usr/bin/env python3
"""
Extract homework sections from grammar book chapters and create Word documents.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Chapter number to word mapping
NUMBER_WORDS = {
    1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
    6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten",
    11: "Eleven", 12: "Twelve", 13: "Thirteen", 14: "Fourteen", 15: "Fifteen",
    16: "Sixteen", 17: "Seventeen", 18: "Eighteen", 19: "Nineteen", 20: "Twenty",
    21: "Twenty-One"
}


def extract_homework_section(content):
    """Extract the homework section from chapter markdown."""
    homework_match = re.search(r'^## Homework:.*?(?=^## Glossary|\Z)', content, re.MULTILINE | re.DOTALL)
    if homework_match:
        return homework_match.group(0).strip()
    return None


def parse_homework_content(homework_text):
    """Parse homework markdown into structured content."""
    lines = homework_text.split('\n')

    # Extract title
    title_match = re.match(r'^## Homework:\s*(.+)$', lines[0])
    title = title_match.group(1).strip() if title_match else "Homework"

    # Parse the rest
    sections = []
    current_section = None
    current_items = []

    for line in lines[1:]:
        line = line.rstrip()

        # Skip empty lines at the start
        if not line and not current_section:
            continue

        # Check for part headers (### Part A, ### Part B, etc.)
        part_match = re.match(r'^###\s+Part\s+([A-Z]):\s*(.+)$', line)
        if part_match:
            if current_section:
                sections.append({
                    'type': 'part',
                    'label': current_section['label'],
                    'title': current_section['title'],
                    'items': current_items
                })
            current_section = {
                'label': f"Part {part_match.group(1)}",
                'title': part_match.group(2).strip()
            }
            current_items = []
            continue

        # Check for numbered questions
        question_match = re.match(r'^\*\*(\d+)\.\*\*\s*(.+)$', line)
        if question_match:
            current_items.append({
                'type': 'question',
                'number': question_match.group(1),
                'text': question_match.group(2).strip()
            })
            continue

        # Check for sub-items with bullets
        bullet_match = re.match(r'^[-*]\s+\*\*(.+?)\*\*:\s*(.*)$', line)
        if bullet_match:
            current_items.append({
                'type': 'sub_item',
                'label': bullet_match.group(1),
                'text': bullet_match.group(2).strip() if bullet_match.group(2) else ""
            })
            continue

        # Regular bullet items
        bullet_match2 = re.match(r'^[-*]\s+(.+)$', line)
        if bullet_match2:
            current_items.append({
                'type': 'bullet',
                'text': bullet_match2.group(1).strip()
            })
            continue

        # Continuation of previous item or regular text
        if line.strip() and current_items:
            # Append to last item
            if current_items[-1].get('text'):
                current_items[-1]['text'] += ' ' + line.strip()
            else:
                current_items[-1]['text'] = line.strip()
        elif line.strip():
            # Standalone text (like instructions)
            current_items.append({
                'type': 'text',
                'text': line.strip()
            })

    # Add last section
    if current_section:
        sections.append({
            'type': 'part',
            'label': current_section['label'],
            'title': current_section['title'],
            'items': current_items
        })
    elif current_items:
        # No parts, just items
        sections.append({
            'type': 'content',
            'items': current_items
        })

    return {
        'title': title,
        'sections': sections
    }


def clean_markdown(text):
    """Remove markdown formatting from text."""
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*([^*]+?)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    para._p.get_or_add_pPr().append(pBdr)
    return para


def create_word_document(chapter_num, homework_data, output_path):
    """Create a Word document for the homework."""
    doc = Document()

    # Set default font and margins
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title
    chapter_word = NUMBER_WORDS.get(chapter_num, str(chapter_num))
    title = doc.add_paragraph()
    title_run = title.add_run(f"Chapter {chapter_word} Homework")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=6)

    # Add subtitle with homework topic
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(homework_data['title'])
    subtitle_run.font.size = Pt(13)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=0, after=12)

    # Add a thin horizontal line
    add_horizontal_line(doc)

    # Process sections
    for section in homework_data['sections']:
        if section['type'] == 'part':
            # Add part header
            part_para = doc.add_paragraph()
            part_run = part_para.add_run(f"{section['label']}: {section['title']}")
            part_run.bold = True
            part_run.font.size = Pt(12)
            set_paragraph_spacing(part_para, before=18, after=10)

            # Add items
            for item in section['items']:
                add_homework_item(doc, item)
        else:
            # Content without parts
            for item in section['items']:
                add_homework_item(doc, item)

    # Save the document
    doc.save(output_path)
    print(f"  -> Created: {output_path}")


def add_homework_item(doc, item):
    """Add a homework item to the document."""
    if item['type'] == 'question':
        # Add question with tight spacing
        para = doc.add_paragraph()
        para.add_run(f"{item['number']}. ").bold = True
        para.add_run(clean_markdown(item['text']))
        set_paragraph_spacing(para, before=6, after=3)

        # Add generous answer space (about 1.5 inches of blank space)
        for _ in range(6):
            blank = doc.add_paragraph()
            set_paragraph_spacing(blank, before=0, after=0)

    elif item['type'] == 'sub_item':
        # Add labeled sub-item with tight spacing
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.4)
        para.add_run(f"• {item['label']}: ").bold = True
        if item['text']:
            para.add_run(clean_markdown(item['text']))
        set_paragraph_spacing(para, before=3, after=2)

        # Add answer space for sub-items (about 0.75 inch)
        for _ in range(3):
            blank = doc.add_paragraph()
            blank.paragraph_format.left_indent = Inches(0.4)
            set_paragraph_spacing(blank, before=0, after=0)

    elif item['type'] == 'bullet':
        # Add bullet item with tight spacing
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.4)
        para.add_run(f"• {clean_markdown(item['text'])}")
        set_paragraph_spacing(para, before=2, after=2)

    elif item['type'] == 'text':
        # Add regular text
        text = clean_markdown(item['text'])
        # Skip certain lines
        if text.startswith('Name:') or text.startswith('Date:') or text == '---':
            return
        if text.startswith('Due ') or text.startswith('*Due'):
            # Skip due date lines
            return

        para = doc.add_paragraph(text)
        set_paragraph_spacing(para, before=3, after=3)


def main():
    chapters_dir = Path("Writing Projects/concise_guide_to_english_grammar/chapters")
    output_dir = Path("homework")

    # Create output directory
    output_dir.mkdir(exist_ok=True)

    # Chapter files in order
    chapter_files = [
        "chapter_01_introduction_to_linguistics.md",
        "chapter_02_prescriptive_vs_descriptive.md",
        "chapter_03_language_variation.md",
        "chapter_04_morphology.md",
        "chapter_05_open_classes.md",
        "chapter_06_closed_classes.md",
        "chapter_07_sentence_diagramming.md",
        "chapter_08_sentence_elements_patterns.md",
        "chapter_09_compound_complex_sentences.md",
        "chapter_10_verbs_tense_aspect.md",
        "chapter_11_verbs_voice_modals.md",
        "chapter_12_adverbials.md",
        "chapter_13_nominals.md",
        "chapter_14_adjectivals.md",
        "chapter_15_other_grammatical_forms.md",
        "chapter_16_stylistic_choices.md",
        "chapter_17_punctuation.md",
        "chapter_18_clarity_readability.md",
        "chapter_19_organization_concision.md",
        "chapter_20_genre_register.md",
        "chapter_21_teaching_grammar.md",
    ]

    print("Creating homework Word documents...\n")

    for i, filename in enumerate(chapter_files, 1):
        filepath = chapters_dir / filename
        print(f"Processing Chapter {i}: {filename}")

        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract homework section
        homework_text = extract_homework_section(content)

        if homework_text:
            # Parse homework content
            homework_data = parse_homework_content(homework_text)

            # Create Word document
            chapter_word = NUMBER_WORDS[i]
            output_file = output_dir / f"Chapter {chapter_word} Homework.docx"
            create_word_document(i, homework_data, output_file)
        else:
            print(f"  -> No homework section found")

    print(f"\nDone! Created homework documents in {output_dir}/")


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Generate Word homework documents from PreTeXt source files.
Extracts homework sections from ch-XX.ptx files and creates formatted .docx files.

Formatting requirements:
- No name/date fields (digital completion)
- Minimal space between headings and text
- Space between activities for student answers
- Section numbers (1, 2, 3) and problem letters (A, B, C)
- Overview with sections and time estimates
"""

import os
import re
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set paragraph spacing in points."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))  # Convert to twips
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def extract_text_from_element(elem):
    """Extract text content from an XML element, handling nested elements."""
    text_parts = []

    if elem.text:
        text_parts.append(elem.text)

    for child in elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'em':
            text_parts.append(extract_text_from_element(child))
        elif tag == 'term':
            text_parts.append(extract_text_from_element(child))
        elif tag == 'c':
            text_parts.append(extract_text_from_element(child))
        elif tag == 'q':
            inner = extract_text_from_element(child)
            text_parts.append(f'"{inner}"')
        else:
            text_parts.append(extract_text_from_element(child))

        if child.tail:
            text_parts.append(child.tail)

    return ''.join(text_parts)


def add_formatted_text(paragraph, elem, skip_question_prefix=False):
    """Add formatted text to a paragraph, preserving italic/bold."""
    text = elem.text or ''

    # Optionally skip "Question X." or "Exercise X." prefix
    if skip_question_prefix and text:
        text = re.sub(r'^(Question|Exercise)\s+\d+\.?\s*', '', text)

    if text:
        paragraph.add_run(text)

    for child in elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'em':
            # Check if this is "Question X." that should be skipped
            child_text = extract_text_from_element(child)
            if skip_question_prefix and re.match(r'^(Question|Exercise)\s+\d+\.?$', child_text.strip()):
                # Skip the question number, but keep the tail
                if child.tail:
                    paragraph.add_run(child.tail)
                continue
            run = paragraph.add_run(child_text)
            run.italic = True
        elif tag == 'term':
            run = paragraph.add_run(extract_text_from_element(child))
            run.bold = True
        elif tag == 'c':
            run = paragraph.add_run(extract_text_from_element(child))
            run.font.name = 'Consolas'
        elif tag == 'q':
            paragraph.add_run('"')
            add_formatted_text(paragraph, child)
            paragraph.add_run('"')
        else:
            add_formatted_text(paragraph, child)

        if child.tail:
            paragraph.add_run(child.tail)


def estimate_time_for_section(section_info):
    """Estimate time for a section based on problem count and types."""
    problem_count = section_info.get('problem_count', 0)
    has_writing = section_info.get('has_writing', False)
    has_analysis = section_info.get('has_analysis', False)

    # Base time per problem
    if has_writing:
        time_per_problem = 8  # Writing tasks take longer
    elif has_analysis:
        time_per_problem = 5  # Analysis tasks
    else:
        time_per_problem = 3  # Basic identification/matching

    total_minutes = problem_count * time_per_problem

    if total_minutes < 5:
        return "~5 min"
    elif total_minutes <= 10:
        return "~10 min"
    elif total_minutes <= 15:
        return "~15 min"
    elif total_minutes <= 20:
        return "~20 min"
    elif total_minutes <= 30:
        return "~30 min"
    else:
        return f"~{(total_minutes // 10) * 10} min"


def analyze_homework_structure(homework_section):
    """Analyze homework section to extract sections and problems."""
    sections = []
    current_section = None

    for elem in homework_section:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'title':
            continue

        elif tag == 'subsection':
            # Subsection becomes a section
            title_elem = elem.find('.//{*}title')
            title = extract_text_from_element(title_elem) if title_elem is not None else "Section"

            section_info = {
                'title': title,
                'element': elem,
                'type': 'subsection',
                'problems': [],
                'problem_count': 0,
                'has_writing': False,
                'has_analysis': False
            }

            # Count problems in subsection
            for sub_elem in elem:
                sub_tag = sub_elem.tag.split('}')[-1] if '}' in sub_elem.tag else sub_elem.tag
                if sub_tag == 'p':
                    text = extract_text_from_element(sub_elem)
                    if re.match(r'^\s*(Question|Exercise)\s+\d+', text):
                        section_info['problem_count'] += 1
                        section_info['problems'].append(sub_elem)
                        if 'paragraph' in text.lower() or 'write' in text.lower() or 'explain' in text.lower():
                            section_info['has_writing'] = True
                        if 'analyze' in text.lower() or 'identify' in text.lower():
                            section_info['has_analysis'] = True
                elif sub_tag == 'ol':
                    items = sub_elem.findall('.//{*}li')
                    section_info['problem_count'] += len(items)
                elif sub_tag == 'ul':
                    # Check if list items are answerable questions
                    for li in sub_elem.findall('.//{*}li'):
                        p_elem = li.find('.//{*}p')
                        if p_elem is not None:
                            text = extract_text_from_element(p_elem)
                            if '?' in text or text.endswith(':'):
                                section_info['problem_count'] += 1

            sections.append(section_info)

        elif tag == 'paragraphs':
            # Paragraphs with title becomes a section
            title_elem = elem.find('.//{*}title')
            if title_elem is not None:
                title = extract_text_from_element(title_elem)
                current_section = {
                    'title': title,
                    'element': elem,
                    'type': 'paragraphs',
                    'problems': [],
                    'problem_count': 0,
                    'has_writing': False,
                    'has_analysis': False
                }
                sections.append(current_section)

        elif tag == 'p' and current_section is not None:
            text = extract_text_from_element(elem)
            if re.match(r'^\s*(Question|Exercise)\s+\d+', text):
                current_section['problem_count'] += 1
                current_section['problems'].append(elem)
                if 'paragraph' in text.lower() or 'write' in text.lower() or 'explain' in text.lower():
                    current_section['has_writing'] = True
                if 'analyze' in text.lower() or 'identify' in text.lower():
                    current_section['has_analysis'] = True
        elif tag == 'p' and not sections:
            # Problem before any section - create default section
            text = extract_text_from_element(elem)
            if re.match(r'^\s*(Question|Exercise)\s+\d+', text):
                if not current_section:
                    current_section = {
                        'title': 'Questions',
                        'element': None,
                        'type': 'default',
                        'problems': [],
                        'problem_count': 0,
                        'has_writing': False,
                        'has_analysis': False
                    }
                    sections.append(current_section)
                current_section['problem_count'] += 1
                current_section['problems'].append(elem)

        elif tag == 'ul' and current_section is not None:
            # Check for answerable list items
            for li in elem.findall('.//{*}li'):
                p_elem = li.find('.//{*}p')
                if p_elem is not None:
                    text = extract_text_from_element(p_elem)
                    if '?' in text or text.endswith(':'):
                        current_section['problem_count'] += 1

    return sections


def add_overview(doc, sections):
    """Add overview section with sections and time estimates."""
    para = doc.add_paragraph()
    run = para.add_run("Overview")
    run.bold = True
    run.font.size = Pt(12)
    set_paragraph_spacing(para, space_before=6, space_after=6)

    total_time = 0

    for i, section in enumerate(sections, 1):
        title = section['title']
        # Clean up title (remove "Part A:", etc. if present)
        title = re.sub(r'^Part\s+[A-Z]:\s*', '', title)

        time_est = estimate_time_for_section(section)
        # Extract minutes for total
        match = re.search(r'(\d+)', time_est)
        if match:
            total_time += int(match.group(1))

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(f"Section {i}: ")
        run.bold = True
        para.add_run(f"{title} ({time_est})")
        set_paragraph_spacing(para, space_before=0, space_after=2)

    # Total time
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(f"Total estimated time: ")
    run.bold = True
    para.add_run(f"~{total_time} minutes")
    set_paragraph_spacing(para, space_before=6, space_after=12)

    # Separator
    para = doc.add_paragraph()
    para.add_run("─" * 50)
    set_paragraph_spacing(para, space_before=0, space_after=12)


def process_homework_with_numbering(doc, homework_section, sections):
    """Process homework with section numbers and problem letters."""

    section_num = 0
    problem_letter_idx = 0  # 0=A, 1=B, etc.

    for elem in homework_section:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'title':
            continue

        elif tag == 'subsection':
            section_num += 1
            problem_letter_idx = 0
            process_subsection_with_numbering(doc, elem, section_num)

        elif tag == 'paragraphs':
            title_elem = elem.find('.//{*}title')
            if title_elem is not None:
                section_num += 1
                problem_letter_idx = 0
                title_text = extract_text_from_element(title_elem)

                para = doc.add_paragraph()
                run = para.add_run(f"Section {section_num}: {title_text}")
                run.bold = True
                run.font.size = Pt(12)
                set_paragraph_spacing(para, space_before=12, space_after=6)

        elif tag == 'p':
            text = extract_text_from_element(elem)

            if re.match(r'^\s*(Question|Exercise)\s+\d+', text):
                # This is a numbered question - use letter
                letter = chr(ord('A') + problem_letter_idx)
                problem_letter_idx += 1

                para = doc.add_paragraph()
                run = para.add_run(f"{section_num}.{letter}  ")
                run.bold = True

                # Add the question text without the original "Question X." prefix
                add_formatted_text(para, elem, skip_question_prefix=True)
                set_paragraph_spacing(para, space_before=6, space_after=3)

                # Add answer space
                answer_para = doc.add_paragraph()
                answer_para.paragraph_format.left_indent = Inches(0.35)
                run = answer_para.add_run('[Your answer here]')
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                set_paragraph_spacing(answer_para, space_before=0, space_after=12)
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                add_formatted_text(para, elem)
                set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'ul':
            for li in elem.findall('.//{*}li'):
                p_elem = li.find('.//{*}p')
                if p_elem is not None:
                    para = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(para, p_elem)
                    set_paragraph_spacing(para, space_before=0, space_after=3)

                    text = extract_text_from_element(p_elem)
                    if '?' in text or text.endswith(':'):
                        answer_para = doc.add_paragraph()
                        answer_para.paragraph_format.left_indent = Inches(0.5)
                        run = answer_para.add_run('[Answer]')
                        run.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        set_paragraph_spacing(answer_para, space_before=0, space_after=6)

        elif tag == 'ol':
            for i, li in enumerate(elem.findall('.//{*}li'), 1):
                p_elem = li.find('.//{*}p')
                if p_elem is not None:
                    para = doc.add_paragraph(style='List Number')
                    add_formatted_text(para, p_elem)
                    set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'blockquote':
            for p_elem in elem.findall('.//{*}p'):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                add_formatted_text(para, p_elem)
                set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'tabular':
            add_table(doc, elem)


def process_subsection_with_numbering(doc, subsection, section_num):
    """Process a homework subsection with numbering."""
    problem_letter_idx = 0

    title_elem = subsection.find('.//{*}title')
    if title_elem is not None:
        title_text = extract_text_from_element(title_elem)
        para = doc.add_paragraph()
        run = para.add_run(f"Section {section_num}: {title_text}")
        run.bold = True
        run.font.size = Pt(12)
        set_paragraph_spacing(para, space_before=12, space_after=6)

    for elem in subsection:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'title':
            continue

        elif tag == 'paragraphs':
            title_elem = elem.find('.//{*}title')
            if title_elem is not None:
                title_text = extract_text_from_element(title_elem)
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
                run.italic = True
                set_paragraph_spacing(para, space_before=9, space_after=3)

        elif tag == 'p':
            text = extract_text_from_element(elem)

            if re.match(r'^\s*(Question|Exercise)\s+\d+', text):
                letter = chr(ord('A') + problem_letter_idx)
                problem_letter_idx += 1

                para = doc.add_paragraph()
                run = para.add_run(f"{section_num}.{letter}  ")
                run.bold = True
                add_formatted_text(para, elem, skip_question_prefix=True)
                set_paragraph_spacing(para, space_before=6, space_after=3)

                answer_para = doc.add_paragraph()
                answer_para.paragraph_format.left_indent = Inches(0.35)
                run = answer_para.add_run('[Your answer here]')
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                set_paragraph_spacing(answer_para, space_before=0, space_after=12)
            else:
                para = doc.add_paragraph()
                add_formatted_text(para, elem)
                set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'ul':
            for li in elem.findall('.//{*}li'):
                p_elem = li.find('.//{*}p')
                if p_elem is not None:
                    para = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(para, p_elem)
                    set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'ol':
            for i, li in enumerate(elem.findall('.//{*}li'), 1):
                p_elem = li.find('.//{*}p')
                if p_elem is not None:
                    para = doc.add_paragraph(style='List Number')
                    add_formatted_text(para, p_elem)
                    set_paragraph_spacing(para, space_before=0, space_after=3)

        elif tag == 'tabular':
            add_table(doc, elem)


def add_table(doc, tabular_elem):
    """Add a table to the document, supporting colspan for merged cells."""
    rows = tabular_elem.findall('.//{*}row')
    if not rows:
        return

    # Calculate actual number of columns by finding the max column span sum
    num_cols = 0
    for row in rows:
        cells = row.findall('.//{*}cell')
        row_cols = sum(int(cell.get('colspan', '1')) for cell in cells)
        num_cols = max(num_cols, row_cols)

    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        cells = row.findall('.//{*}cell')
        col_idx = 0
        for cell in cells:
            if col_idx >= num_cols:
                break
            colspan = int(cell.get('colspan', '1'))
            cell_text = extract_text_from_element(cell)

            # Merge cells if colspan > 1
            if colspan > 1 and col_idx + colspan - 1 < num_cols:
                table.rows[i].cells[col_idx].merge(
                    table.rows[i].cells[col_idx + colspan - 1]
                )

            merged_cell = table.rows[i].cells[col_idx]
            # Clear any residual text from merge
            for paragraph in merged_cell.paragraphs:
                paragraph.text = ""
            merged_cell.paragraphs[0].text = cell_text

            if row.get('header') == 'yes' or i == 0:
                for paragraph in merged_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            col_idx += colspan


def extract_homework_from_ptx(ptx_path):
    """Extract homework section from a PreTeXt file."""
    try:
        tree = etree.parse(str(ptx_path))
        root = tree.getroot()

        homework_sections = root.xpath('//*[contains(@xml:id, "homework")]',
                                       namespaces={'xml': 'http://www.w3.org/XML/1998/namespace'})

        if not homework_sections:
            for section in root.iter():
                xml_id = section.get('{http://www.w3.org/XML/1998/namespace}id', '')
                if 'homework' in xml_id:
                    return section
        else:
            return homework_sections[0]

    except Exception as e:
        print(f"Error parsing {ptx_path}: {e}")

    return None


def get_chapter_title(ptx_path):
    """Get chapter title from PreTeXt file."""
    try:
        tree = etree.parse(str(ptx_path))
        root = tree.getroot()

        title_elem = root.find('.//{*}chapter/{*}title')
        if title_elem is None:
            title_elem = root.find('.//{*}title')

        if title_elem is not None:
            return extract_text_from_element(title_elem)

    except Exception as e:
        print(f"Error getting title from {ptx_path}: {e}")

    return None


def get_homework_title(homework_section):
    """Get homework section title."""
    title_elem = homework_section.find('.//{*}title')
    if title_elem is not None:
        title = extract_text_from_element(title_elem)
        if title.startswith('Homework:'):
            title = title[9:].strip()
        return title
    return None


def create_homework_document(chapter_num, chapter_title, homework_title, homework_section, output_path):
    """Create a Word document for the homework."""
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)  # Increased to match book

    # Set up heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Open Sans'
        heading_style.font.bold = True

    # Add title
    title = doc.add_heading(f'Chapter {chapter_num}: {chapter_title}', level=1)
    title.runs[0].font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=6)

    # Add homework subtitle
    if homework_title:
        subtitle = doc.add_heading(f'Homework: {homework_title}', level=2)
        subtitle.runs[0].font.size = Pt(14)
        set_paragraph_spacing(subtitle, space_before=0, space_after=12)

    # Analyze homework structure
    sections = analyze_homework_structure(homework_section)

    # Add overview if there are sections
    if sections:
        add_overview(doc, sections)

    # Process homework content with numbering
    process_homework_with_numbering(doc, homework_section, sections)

    # Save document
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    script_dir = Path(__file__).parent
    pretext_source = script_dir.parent / 'pretext' / 'source'
    output_dir = script_dir

    for chapter_num in range(1, 22):
        ptx_file = pretext_source / f'ch-{chapter_num:02d}.ptx'

        if not ptx_file.exists():
            print(f"Chapter {chapter_num} file not found: {ptx_file}")
            continue

        print(f"\nProcessing Chapter {chapter_num}...")

        chapter_title = get_chapter_title(ptx_file)
        if not chapter_title:
            print(f"  Could not get chapter title")
            continue

        homework_section = extract_homework_from_ptx(ptx_file)
        if homework_section is None:
            print(f"  No homework section found")
            continue

        homework_title = get_homework_title(homework_section)

        output_file = output_dir / f'Chapter {chapter_num:02d} Homework.docx'

        create_homework_document(
            chapter_num,
            chapter_title,
            homework_title,
            homework_section,
            output_file
        )

    print("\nDone!")


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Split homework files into questions and answer keys, then convert to Word.
Designed for digital completion - no name/date fields, spaces for typed answers.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HOMEWORK_DIR = r'C:\Users\irphy\Documents\bookmaker\Writing Projects\concise_guide_to_english_grammar\Homework'

def get_chapter_title(content):
    """Extract chapter title from homework content."""
    match = re.search(r'^#\s*Homework:\s*(.+)$', content, re.MULTILINE)
    if match:
        return match.group(1).strip()
    return "Unknown"

def split_homework_and_answers(content):
    """Split content into homework questions and answer key."""
    patterns = [
        r'\n---\s*\n---\s*\n+#\s*ANSWER\s*KEY',
        r'\n#\s*ANSWER\s*KEY'
    ]

    for pattern in patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            homework = content[:match.start()].strip()
            answers = "# ANSWER KEY" + content[match.end():].strip()
            return homework, answers

    return content.strip(), None

def clean_markdown_for_display(text):
    """Convert markdown formatting for display."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    return text.strip()

def add_answer_space(doc, space_type='short'):
    """Add space for digital answers."""
    if space_type == 'short':
        # Add 2-3 empty paragraphs for short answers
        for _ in range(3):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
    elif space_type == 'medium':
        # Add more space for medium answers
        for _ in range(5):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
    elif space_type == 'long':
        # Add substantial space for long answers/paragraphs
        for _ in range(8):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
    elif space_type == 'diagram':
        # Add space for diagram insertion with placeholder text
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run('[Insert your diagram here]')
        run.italic = True
        run.font.color.rgb = None
        # Add extra space after
        for _ in range(6):
            doc.add_paragraph()

def add_markdown_to_doc(doc, content, is_answer_key=False):
    """Convert markdown to Word document elements for digital completion."""
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Main title
        if stripped.startswith('# ') and not 'ANSWER KEY' in stripped.upper():
            text = stripped[2:].strip()
            heading = doc.add_heading(text, level=0)
            doc.add_paragraph()  # Space after title

        # Answer key title
        elif 'ANSWER KEY' in stripped.upper():
            doc.add_heading('ANSWER KEY', level=0)
            doc.add_paragraph()

        # Part headers
        elif stripped.startswith('## Part') or stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_paragraph()  # Space before part
            doc.add_heading(text, level=1)

        # Subheaders
        elif stripped.startswith('### '):
            text = stripped[4:].strip()
            if 'Example' in text:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                run = para.add_run('Example (completed):')
                run.bold = True
                run.italic = True
            else:
                doc.add_heading(text, level=2)

        # Horizontal rules - add space
        elif stripped == '---':
            doc.add_paragraph()

        # Instructions
        elif stripped.startswith('**Instructions:**') or stripped.startswith('**Instructions**'):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            run = para.add_run('Instructions: ')
            run.bold = True
            text = re.sub(r'\*\*Instructions:?\*\*\s*', '', stripped)
            para.add_run(clean_markdown_for_display(text))

        # Numbered questions
        elif re.match(r'^\*\*\d+\.?\*\*', stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            text = clean_markdown_for_display(stripped)
            run = para.add_run(text)
            run.bold = True

            # Collect continuation lines
            i += 1
            has_subparts = False
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    continue
                if next_line.startswith('**') and re.match(r'^\*\*\d+\.?\*\*', next_line):
                    i -= 1
                    break
                if next_line.startswith('#') or next_line.startswith('---'):
                    i -= 1
                    break

                # Check if it's a sub-part prompt (like "Part of speech:", "Morphemes:", etc.)
                if ':' in next_line and not next_line.startswith('-'):
                    has_subparts = True
                    para = doc.add_paragraph()
                    para.add_run(clean_markdown_for_display(next_line))
                    # Add space for answer after each sub-part
                    if not is_answer_key:
                        add_answer_space(doc, 'short')
                elif next_line.startswith('- ') or next_line.startswith('* '):
                    text = clean_markdown_for_display(next_line[2:])
                    doc.add_paragraph(text, style='List Bullet')
                else:
                    para = doc.add_paragraph()
                    para.add_run(clean_markdown_for_display(next_line))
                i += 1

            # Add answer space if not answer key and no sub-parts already handled
            if not is_answer_key and not has_subparts:
                # Check if this seems like a diagram question
                question_text = text.lower()
                if 'diagram' in question_text or 'tree' in question_text:
                    add_answer_space(doc, 'diagram')
                elif 'paragraph' in question_text or 'passage' in question_text or 'explain' in question_text or 'reflection' in question_text:
                    add_answer_space(doc, 'long')
                else:
                    add_answer_space(doc, 'medium')
            continue

        # List items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = clean_markdown_for_display(stripped[2:])
            doc.add_paragraph(text, style='List Bullet')

        # Passage blocks (italicized paragraphs)
        elif stripped.startswith('*') and stripped.endswith('*') and len(stripped) > 50:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            text = stripped.strip('*')
            run = para.add_run(text)
            run.italic = True

        # Regular paragraphs
        else:
            para = doc.add_paragraph()
            para.add_run(clean_markdown_for_display(stripped))

        i += 1

def create_word_doc(chapter_num, title, content, doc_type, is_answer_key=False):
    """Create Word document for digital completion."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set default paragraph spacing
    para_format = style.paragraph_format
    para_format.space_after = Pt(6)
    para_format.line_spacing = 1.15

    # Add content
    add_markdown_to_doc(doc, content, is_answer_key=is_answer_key)

    # Save
    if doc_type == 'homework':
        filename = f'Chapter {chapter_num:02d} Homework.docx'
    else:
        filename = f'Chapter {chapter_num:02d} Answer Key.docx'

    filepath = os.path.join(HOMEWORK_DIR, filename)
    doc.save(filepath)
    print(f'Created: {filename}')

def create_answer_key_md(chapter_num, title, answer_content):
    """Create separate answer key markdown file."""
    filename = f'ch{chapter_num:02d}-answer-key.md'
    filepath = os.path.join(HOMEWORK_DIR, filename)

    content = f"# Answer Key: {title}\n\n{answer_content}"

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'Created: {filename}')

def process_homework_file(md_path):
    """Process a homework markdown file."""
    filename = os.path.basename(md_path)
    match = re.search(r'ch(\d+)', filename)
    if not match:
        print(f'Could not extract chapter number from {filename}')
        return

    chapter_num = int(match.group(1))

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    title = get_chapter_title(content)
    homework_content, answer_content = split_homework_and_answers(content)

    if not answer_content:
        print(f'No answer key found in {filename}')
        return

    # Create separate answer key markdown
    answer_text = re.sub(r'^#\s*ANSWER\s*KEY\s*\n*', '', answer_content, flags=re.IGNORECASE)
    create_answer_key_md(chapter_num, title, answer_text)

    # Create Word documents
    create_word_doc(chapter_num, title, homework_content, 'homework', is_answer_key=False)
    create_word_doc(chapter_num, title, answer_content, 'answer_key', is_answer_key=True)

def main():
    """Process all homework files for chapters 4-15."""
    print("Processing homework files for digital completion...")
    print("=" * 50)

    for chapter_num in range(4, 16):
        md_file = f'ch{chapter_num:02d}-homework.md'
        md_path = os.path.join(HOMEWORK_DIR, md_file)

        if os.path.exists(md_path):
            print(f'\nProcessing Chapter {chapter_num}...')
            process_homework_file(md_path)
        else:
            print(f'\nChapter {chapter_num}: {md_file} not found')

    print("\n" + "=" * 50)
    print("Done!")

if __name__ == '__main__':
    main()

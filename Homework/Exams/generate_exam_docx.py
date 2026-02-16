#!/usr/bin/env python3
"""
Generate Exam One .docx files (student exam and answer key) with embedded diagrams.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
DIAGRAM_DIR = SCRIPT_DIR / "diagrams"


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set paragraph spacing in points."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_run(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_question(doc, number, text, pts=2, font_size=11):
    """Add a multiple choice question header."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    add_run(p, f"Question {number}: ", bold=True, font_size=font_size)
    add_run(p, text, font_size=font_size)
    add_run(p, f" ({pts} pts)", font_size=font_size)
    return p


def add_mc_option(doc, letter, text, font_size=11, is_answer=False, explanation=""):
    """Add a multiple choice option."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=1, space_after=1)
    p.paragraph_format.left_indent = Inches(0.5)
    if is_answer:
        add_run(p, f"{letter}. {text}", bold=True, font_size=font_size)
        if explanation:
            add_run(p, f" — {explanation}", italic=True, font_size=font_size - 1)
    else:
        add_run(p, f"{letter}. {text}", font_size=font_size)
        if explanation:
            add_run(p, f" — {explanation}", italic=True, font_size=font_size - 1)
    return p


def add_answer_line(doc, answer_letter, font_size=11):
    """Add 'Answer: ___' line for student version or 'Answer: X' for answer key."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=2, space_after=2)
    p.paragraph_format.left_indent = Inches(0.3)
    if answer_letter:
        add_run(p, "Answer: ", bold=True, font_size=font_size)
        add_run(p, answer_letter, bold=True, font_size=font_size)
    else:
        add_run(p, "Answer: _____", font_size=font_size)
    return p


def add_section_header(doc, text, level=2, font_size=14):
    """Add a section heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(font_size)
    return h


def compute_spans(labels):
    """Convert flat label list to (label, span_count, start_index) tuples.

    Non-empty labels start a new span; consecutive empty strings extend it.
    """
    spans = []
    i = 0
    while i < len(labels):
        label = labels[i]
        if label:
            span = 1
            while i + span < len(labels) and labels[i + span] == "":
                span += 1
            spans.append((label, span, i))
            i += span
        else:
            i += 1
    return spans


def blank_labels(labels):
    """Convert answer labels to blank labels preserving span structure for pre-merging.

    Non-empty labels become a space (visually blank but truthy for compute_spans),
    empty strings stay empty to continue the previous span.
    """
    return ["" if label == "" else " " for label in labels]


def add_labeling_table(doc, words, pos_labels=None, phrase_labels=None, role_labels=None, font_size=10):
    """Add a sentence labeling table with merged cells for Role and Phrase rows.

    When labels are provided for Role/Phrase rows, cells are merged to show
    how words group into phrases and roles. When labels are None (student
    version), cells are left unmerged so students can write in each cell.
    """
    num_cols = len(words) + 1  # +1 for row headers
    table = doc.add_table(rows=4, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_headers = ["Role", "Phrase", "Word", "POS"]
    row_data = [role_labels, phrase_labels, words, pos_labels]

    for i, (header, data) in enumerate(zip(row_headers, row_data)):
        # Row header cell
        cell = table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, header, bold=True, font_size=font_size)
        set_cell_shading(cell, "E8E8E8")

        if i <= 1 and data:
            # Role or Phrase row with data — merge cells to show groupings
            spans = compute_spans(data)
            for label, span, start_idx in spans:
                col_start = start_idx + 1  # +1 for row header column
                col_end = col_start + span - 1
                if span > 1:
                    table.rows[i].cells[col_start].merge(table.rows[i].cells[col_end])
                merged_cell = table.rows[i].cells[col_start]
                # Clear residual paragraphs from merge
                for paragraph in merged_cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                p = merged_cell.paragraphs[0]
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, label, font_size=font_size)
        else:
            # Normal unmerged row (Word, POS, or blank student rows)
            for j, val in enumerate(data if data else [""] * len(words)):
                cell = table.rows[i].cells[j + 1]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 2:  # Word row always filled
                    add_run(p, val, font_size=font_size)
                elif data:
                    add_run(p, val, font_size=font_size)
                # else leave blank for student version

    return table


def add_diagram_image(doc, image_name, width_inches=5.5):
    """Add a diagram PNG image to the document."""
    img_path = DIAGRAM_DIR / f"{image_name}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        set_paragraph_spacing(p, space_before=4, space_after=4)
        return p
    else:
        p = doc.add_paragraph(f"[Diagram not found: {image_name}.png]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p


# =============================================================================
# QUESTION DATA
# =============================================================================

MC_QUESTIONS = {
    "A1": {
        "title": "Section A1: Morphology and Word Structure (Questions 1\u20138, 16 points)",
        "questions": [
            {
                "num": 1,
                "text": 'How many morphemes does the word "disagreements" contain?',
                "options": ["1", "2", "3", "4"],
                "answer": "d",
                "explanations": {
                    "d": 'dis- (bound) + agree (free) + -ment (bound) + -s (bound)'
                }
            },
            {
                "num": 2,
                "text": "Which of the following words contains exactly one morpheme?",
                "options": ["unfriendly", "describe", "replayable", "misplace"],
                "answer": "b",
                "explanations": {
                    "a": "3 morphemes: un- + friend + -ly",
                    "b": '1 morpheme. Although it appears to begin with de-, the remaining -scribe does not combine with de- to produce the word\u2019s meaning compositionally.',
                    "c": "3 morphemes: re- + play + -able",
                    "d": "2 morphemes: mis- + place"
                }
            },
            {
                "num": 3,
                "text": "Which of the following is a bound morpheme?",
                "options": ["book", "happy", "-ness", "walk"],
                "answer": "c",
                "explanations": {
                    "a": "free morpheme (can stand alone)",
                    "b": "free morpheme (can stand alone)",
                    "c": "bound morpheme (cannot stand alone; must attach to another morpheme)",
                    "d": "free morpheme (can stand alone)"
                }
            },
            {
                "num": 4,
                "text": "The suffix -tion typically creates which word class?",
                "options": ["Verb", "Adjective", "Noun", "Adverb"],
                "answer": "c",
                "explanations": {
                    "c": "-tion converts verbs into nouns: action, creation, decision"
                }
            },
            {
                "num": 5,
                "text": "Which word is an example of a compound (formed from two free morphemes)?",
                "options": ["understand", "carpet", "bookcase", "receive"],
                "answer": "c",
                "explanations": {
                    "a": '1 morpheme. The meaning does not derive from "under" + "stand."',
                    "b": '1 morpheme. "Car" + "pet" is coincidental resemblance.',
                    "c": '2 free morphemes: book + case. The meaning is compositional.',
                    "d": '1 morpheme. "Ceive" is not a word in modern English.'
                }
            },
            {
                "num": 6,
                "text": 'The word "unhelpful" contains how many morphemes?',
                "options": ["1", "2", "3", "4"],
                "answer": "c",
                "explanations": {
                    "c": "un- (bound) + help (free) + -ful (bound)"
                }
            },
            {
                "num": 7,
                "text": 'Which of the following correctly identifies the morphemes in "rebuilding"?',
                "options": [
                    "re- (bound) + build (free) + -ing (bound)",
                    "re- (free) + build (free) + -ing (bound)",
                    "rebuild (free) + -ing (bound)",
                    "re- (bound) + building (free)"
                ],
                "answer": "a",
                "explanations": {
                    "a": 'Correct. re- is a bound prefix meaning "again," build is the free base, -ing is bound.',
                    "b": "Incorrect. re- is bound, not free.",
                    "c": "Incorrect. rebuild can be further divided.",
                    "d": "Incorrect. building should be divided into build + -ing."
                }
            },
            {
                "num": 8,
                "text": "Which word best demonstrates morphological productivity \u2014 the active use of existing morpheme patterns to create new words?",
                "options": ["receive", "podcaster", "understand", "carpet"],
                "answer": "b",
                "explanations": {
                    "a": "Borrowed as a whole word from Latin.",
                    "b": 'A recent creation using the productive -er suffix added to podcast.',
                    "c": "A single morpheme; no productive process visible.",
                    "d": "A single morpheme; no productive process visible."
                }
            },
        ]
    },
    "A2": {
        "title": "Section A2: Open and Closed Classes (Questions 9\u201316, 16 points)",
        "questions": [
            {
                "num": 9,
                "text": "Which of the following word classes is an open class?",
                "options": ["Determiners", "Pronouns", "Adjectives", "Prepositions"],
                "answer": "c",
                "explanations": {
                    "a": "closed class", "b": "closed class",
                    "c": "open class (readily accepts new members: viral, binge-worthy, gluten-free)",
                    "d": "closed class"
                }
            },
            {
                "num": 10,
                "text": "Which feature best distinguishes determiners from adjectives?",
                "options": [
                    "Determiners appear after the noun",
                    'Determiners can be modified by "very"',
                    "Determiners must appear before any adjectives in a noun phrase",
                    "Determiners describe qualities of the noun"
                ],
                "answer": "c",
                "explanations": {
                    "a": "Incorrect. Determiners appear before the noun.",
                    "b": "Incorrect. Adjectives can be modified by \"very\"; determiners cannot.",
                    "c": "Correct. DET + ADJ + N (the tall man, not tall the man).",
                    "d": "Incorrect. Adjectives describe qualities; determiners specify reference."
                }
            },
            {
                "num": 11,
                "text": 'In the phrase "the tall student from Ohio," which word is the head of the noun phrase?',
                "options": ["the", "tall", "student", "Ohio"],
                "answer": "c",
                "explanations": {
                    "a": "determiner (modifier)", "b": "adjective (modifier)",
                    "c": "The head noun. Cannot be removed without destroying the phrase.",
                    "d": "head of a nested NP inside the PP, not the head of the main NP."
                }
            },
            {
                "num": 12,
                "text": "Which of the following words is a determiner?",
                "options": ["beautiful", "quickly", "every", "student"],
                "answer": "c",
                "explanations": {
                    "a": "adjective", "b": "adverb",
                    "c": 'determiner (appears before adjectives; cannot be modified by "very")',
                    "d": "noun"
                }
            },
            {
                "num": 13,
                "text": '"The professor entered the room and placed her notes on the desk." What is the antecedent of "her"?',
                "options": ["The room", "The professor", "The desk", "her notes"],
                "answer": "b",
                "explanations": {
                    "b": '"her" (third person, singular, feminine) refers back to "The professor."'
                }
            },
            {
                "num": 14,
                "text": "Pronouns replace which of the following elements?",
                "options": [
                    "Only the head noun",
                    "The entire noun phrase",
                    "Only the determiner and noun",
                    "The verb phrase"
                ],
                "answer": "b",
                "explanations": {
                    "b": 'Correct. Pronouns replace the full NP including all modifiers: "the tall student from Ohio" \u2192 "she."'
                }
            },
            {
                "num": 15,
                "text": "Which of the following words is a preposition?",
                "options": ["quickly", "beautiful", "through", "they"],
                "answer": "c",
                "explanations": {
                    "a": "adverb", "b": "adjective", "c": "preposition", "d": "pronoun"
                }
            },
            {
                "num": 16,
                "text": "Which of the following is NOT a determiner?",
                "options": ["the", "every", "very", "some"],
                "answer": "c",
                "explanations": {
                    "a": "determiner", "b": "determiner",
                    "c": "adverb (modifies adjectives and adverbs, not nouns)",
                    "d": "determiner"
                }
            },
        ]
    },
    "A3": {
        "title": "Section A3: Parts of Speech in Context (Questions 17\u201325, 18 points)",
        "preamble": 'Sentence for Questions 17\u201318: "Several enthusiastic students quietly studied the challenging material."',
        "questions": [
            {
                "num": 17,
                "text": "Which word in the sentence above is a determiner?",
                "options": ["enthusiastic", "Several", "quietly", "challenging"],
                "answer": "b",
                "explanations": {
                    "a": "adjective",
                    "b": 'determiner (indicates amount; cannot be modified by "very"; appears before adjectives)',
                    "c": "adverb", "d": "adjective"
                }
            },
            {
                "num": 18,
                "text": "Which word in the sentence above is an adverb?",
                "options": ["Several", "enthusiastic", "quietly", "material"],
                "answer": "c",
                "explanations": {
                    "a": "determiner", "b": "adjective",
                    "c": 'adverb (modifies the verb "studied"; formed with -ly from quiet)',
                    "d": "noun"
                }
            },
            {
                "num": 19,
                "text": '"The professor with the briefcase lectured brilliantly." What part of speech is "with"?',
                "options": ["Conjunction", "Determiner", "Preposition", "Adverb"],
                "answer": "c",
                "explanations": {
                    "c": '"With" introduces the prepositional phrase "with the briefcase."'
                }
            },
            {
                "num": 20,
                "text": '"My sister carefully examined their proposal." Which word is functioning as a verb?',
                "options": ["My", "sister", "carefully", "examined"],
                "answer": "d",
                "explanations": {
                    "a": "determiner", "b": "noun", "c": "adverb",
                    "d": "verb (the main action; takes past tense -ed)"
                }
            },
            {
                "num": 21,
                "text": '"Three playful kittens explored the mysterious basement." Which word is an adjective?',
                "options": ["Three", "playful", "explored", "basement"],
                "answer": "b",
                "explanations": {
                    "a": "determiner",
                    "b": 'adjective (describes a quality; can be modified by "very"; formed with -ful)',
                    "c": "verb", "d": "noun"
                }
            },
            {
                "num": 22,
                "text": '"She quickly ran to the store." What type of phrase is "to the store"?',
                "options": ["Noun phrase", "Verb phrase", "Adjective phrase", "Prepositional phrase"],
                "answer": "d",
                "explanations": {
                    "d": 'Headed by the preposition "to," with "the store" as its NP object.'
                }
            },
            {
                "num": 23,
                "text": 'In the phrase "the extremely tall building," what word class does "extremely" belong to?',
                "options": ["Adjective", "Adverb", "Determiner", "Noun"],
                "answer": "b",
                "explanations": {
                    "b": '"Extremely" modifies the adjective "tall," telling the degree of tallness.'
                }
            },
            {
                "num": 24,
                "text": "Which test reliably distinguishes adjectives from determiners?",
                "options": [
                    "Adjectives can be pluralized",
                    'Adjectives can be modified by "very"',
                    "Adjectives always appear after the noun",
                    "Adjectives cannot appear before nouns"
                ],
                "answer": "b",
                "explanations": {
                    "a": "Incorrect. Nouns are pluralized, not adjectives.",
                    "b": 'Correct. "Very tall" but not "very the" or "very my."',
                    "c": "Incorrect. Adjectives typically appear before the noun.",
                    "d": "Incorrect. Adjectives routinely appear before nouns."
                }
            },
            {
                "num": 25,
                "text": "Every sentence (S) divides into which two parts?",
                "options": [
                    "Determiner + Sentence",
                    "Noun + Verb",
                    "Subject NP + Predicate VP",
                    "Adjective Phrase + Verb Phrase"
                ],
                "answer": "c",
                "explanations": {
                    "b": "Too narrow; subjects and predicates are phrases.",
                    "c": "Correct. S \u2192 NP + VP is the fundamental structure."
                }
            },
        ]
    }
}

# Section B sentence data
SECTION_B = [
    {
        "num": 26,
        "sentence": "The cat slept.",
        "words": ["The", "cat", "slept"],
        "pos": ["DET", "N", "V"],
        "phrases": ["NP", "", "VP"],
        "roles": ["Subject", "", "Predicate"],
        "bracket": "[S [NP [DET The] [N cat]] [VP [V slept]]]",
        "diagram": "q26_cat_slept",
        "width": 3.5,
    },
    {
        "num": 27,
        "sentence": "The cheerful birds sang beautifully in the garden.",
        "words": ["The", "cheerful", "birds", "sang", "beautifully", "in", "the", "garden"],
        "pos": ["DET", "ADJ", "N", "V", "ADV", "PREP", "DET", "N"],
        "phrases": ["NP", "", "", "VP", "ADVP", "PP", "", ""],
        "roles": ["Subject", "", "", "Predicate", "", "", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ cheerful]] [N birds]] [VP [V sang] [ADVP [ADV beautifully]] [PP [PREP in] [NP [DET the] [N garden]]]]]",
        "diagram": "q27_birds_sang",
        "width": 5.5,
    },
    {
        "num": 28,
        "sentence": "The students from Ohio studied carefully.",
        "words": ["The", "students", "from", "Ohio", "studied", "carefully"],
        "pos": ["DET", "N", "PREP", "N", "V", "ADV"],
        "phrases": ["NP", "", "PP", "", "VP", "ADVP"],
        "roles": ["Subject", "", "", "", "Predicate", ""],
        "bracket": "[S [NP [DET The] [N students] [PP [PREP from] [NP [N Ohio]]]] [VP [V studied] [ADVP [ADV carefully]]]]",
        "diagram": "q28_students_studied",
        "width": 5.0,
    },
    {
        "num": 29,
        "sentence": "Several brave firefighters worked tirelessly during the dangerous storm.",
        "words": ["Several", "brave", "firefighters", "worked", "tirelessly", "during", "the", "dangerous", "storm"],
        "pos": ["DET", "ADJ", "N", "V", "ADV", "PREP", "DET", "ADJ", "N"],
        "phrases": ["NP", "", "", "VP", "ADVP", "PP", "", "", ""],
        "roles": ["Subject", "", "", "Predicate", "", "", "", "", ""],
        "bracket": "[S [NP [DET Several] [ADJP [ADJ brave]] [N firefighters]] [VP [V worked] [ADVP [ADV tirelessly]] [PP [PREP during] [NP [DET the] [ADJP [ADJ dangerous]] [N storm]]]]]",
        "diagram": "q29_firefighters_worked",
        "width": 6.0,
    },
    {
        "num": 30,
        "sentence": "The tall professor with gray hair spoke very eloquently.",
        "words": ["The", "tall", "professor", "with", "gray", "hair", "spoke", "very", "eloquently"],
        "pos": ["DET", "ADJ", "N", "PREP", "ADJ", "N", "V", "ADV", "ADV"],
        "phrases": ["NP", "", "", "PP", "", "", "VP", "ADVP", ""],
        "roles": ["Subject", "", "", "", "", "", "Predicate", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ tall]] [N professor] [PP [PREP with] [NP [ADJP [ADJ gray]] [N hair]]]] [VP [V spoke] [ADVP [ADVP [ADV very]] [ADV eloquently]]]]",
        "diagram": "q30_professor_spoke",
        "width": 6.0,
    },
]

EXAMPLE_SENTENCE = {
    "sentence": "The old man sat quietly.",
    "words": ["The", "old", "man", "sat", "quietly"],
    "pos": ["DET", "ADJ", "N", "V", "ADV"],
    "phrases": ["NP", "", "", "VP", "ADVP"],
    "roles": ["Subject", "", "", "Predicate", ""],
    "bracket": "[S [NP [DET The] [ADJP [ADJ old]] [N man]] [VP [V sat] [ADVP [ADV quietly]]]]",
    "diagram": "example_old_man_sat",
    "width": 5.5,
}

BONUS_QUESTIONS = [
    {
        "num": 1,
        "text": "Write a sentence containing a compound word (a word formed from two free morphemes). Underline or circle the compound word.",
        "example": 'I placed the flowers on the bookshelf.',
        "rubric": "Award 1 pt for a valid sentence and 1 pt for correctly identifying a compound word (book + shelf)."
    },
    {
        "num": 2,
        "text": "Write a sentence containing a prepositional phrase that modifies a noun. Underline or circle the prepositional phrase.",
        "example": 'The student from Colorado passed the exam.',
        "rubric": "Award 1 pt for a valid sentence and 1 pt for correctly identifying a PP that modifies a noun (not a verb)."
    },
    {
        "num": 3,
        "text": "Write a word that contains at least three morphemes. List each morpheme and label it as free or bound.",
        "example": '"unhappiness" \u2014 un- (bound) + happy (free) + -ness (bound) = 3 morphemes',
        "rubric": "Award 1 pt for a valid word with 3+ morphemes and 1 pt for correct labeling."
    },
    {
        "num": 4,
        "text": "Write a noun phrase that contains a determiner, an adjective, and a noun.",
        "example": '"the brave soldier"',
        "rubric": "Award 1 pt for including all three elements and 1 pt for correct word class identification (DET + ADJ + N)."
    },
    {
        "num": 5,
        "text": "Write a sentence that is structurally ambiguous (has two possible meanings due to different possible structures). Briefly explain the two meanings.",
        "example": '"I saw the man with binoculars."\nMeaning 1: I used binoculars to see the man. (PP modifies VP)\nMeaning 2: I saw the man who had binoculars. (PP modifies NP)',
        "rubric": "Award 1 pt for a genuinely ambiguous sentence and 1 pt for correctly explaining both structural readings."
    },
]


def create_document(is_answer_key=False):
    """Create either the student exam or the answer key."""
    doc = Document()

    # Set up default style
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Open Sans'
        h_style.font.bold = True

    # =========================================================================
    # HEADER
    # =========================================================================
    title_text = "ENGL 3110 - Exam One - Spring 2026"
    if is_answer_key:
        title_text += " \u2014 ANSWER KEY"

    title = doc.add_heading(title_text, level=1)
    for run in title.runs:
        run.font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=4)

    p = doc.add_paragraph()
    add_run(p, "Total Points: 100  \u2022  Time: 70 Minutes", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=2)

    p = doc.add_paragraph()
    add_run(p, "Resources: ", bold=True, font_size=11)
    add_run(p, "Anything except help from other people or AI/LLM (ChatGPT, etc.)", font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Section overview
    p = doc.add_paragraph()
    add_run(p, "Sections:", bold=True, font_size=11)
    for line in [
        "Section A: Multiple Choice \u2014 50 points (25 questions, 2 points each) \u2014 Suggested Time: 30 minutes",
        "Section B: Sentence Labeling and Diagramming \u2014 50 points (5 sentences, 10 points each) \u2014 Suggested Time: 30 minutes",
        "Bonus \u2014 10 points (5 questions, 2 points each)"
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, line, font_size=10)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "STUDENT NAME: ", bold=True, font_size=11)
        add_run(p, "_" * 60, font_size=11)

    # =========================================================================
    # SECTION A: MULTIPLE CHOICE
    # =========================================================================
    add_section_header(doc, "Section A: Multiple Choice (50 points)", level=2, font_size=14)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "Instructions: ", bold=True, font_size=11)
        add_run(p, "For each question, select the best answer and write the letter in the blank provided.", font_size=11)

    letters = ["a", "b", "c", "d"]

    for section_key in ["A1", "A2", "A3"]:
        section = MC_QUESTIONS[section_key]
        add_section_header(doc, section["title"], level=3, font_size=12)

        if "preamble" in section:
            p = doc.add_paragraph()
            add_run(p, section["preamble"], italic=True, font_size=11)
            set_paragraph_spacing(p, space_before=4, space_after=4)

        for q in section["questions"]:
            add_question(doc, q["num"], q["text"])

            if is_answer_key:
                add_answer_line(doc, q["answer"])
            else:
                add_answer_line(doc, None)

            for i, option_text in enumerate(q["options"]):
                letter = letters[i]
                is_correct = (letter == q["answer"])
                explanation = q.get("explanations", {}).get(letter, "")

                if is_answer_key:
                    add_mc_option(doc, letter, option_text, is_answer=is_correct, explanation=explanation)
                else:
                    add_mc_option(doc, letter, option_text)

    # =========================================================================
    # SECTION B: SENTENCE LABELING AND DIAGRAMMING
    # =========================================================================
    doc.add_page_break()
    add_section_header(doc, "Section B: Sentence Labeling and Diagramming (50 points)", level=2, font_size=14)

    # Instructions
    p = doc.add_paragraph()
    add_run(p, "Instructions: ", bold=True, font_size=11)
    add_run(p, "For each sentence below, take two steps.", font_size=11)

    p = doc.add_paragraph()
    add_run(p, "1. ", bold=True, font_size=11)
    add_run(p, "Complete the labeling table ", bold=True, font_size=11)
    add_run(p, "by filling in the Role, Phrase, and Part of Speech (POS) rows.", font_size=11)

    p = doc.add_paragraph()
    add_run(p, "2. ", bold=True, font_size=11)
    add_run(p, "Write the bracket notation ", bold=True, font_size=11)
    add_run(p, "for the sentence's tree diagram.", font_size=11)

    # Scoring
    p = doc.add_paragraph()
    add_run(p, "Scoring (per sentence):", bold=True, font_size=10)
    for score_line in [
        "4 points \u2014 Correct Parts of Speech",
        "3 points \u2014 Correct Phrase labels",
        "1 point \u2014 Correct Role labels (Subject, Predicate)",
        "2 points \u2014 Correct bracket notation"
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, score_line, font_size=10)

    # Abbreviation key table
    p = doc.add_paragraph()
    add_run(p, "Abbreviation Key", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=6, space_after=2)

    abbrev_table = doc.add_table(rows=8, cols=3)
    abbrev_table.style = 'Table Grid'
    headers = ["Parts of Speech", "Phrase Types", "Roles"]
    abbrev_data = [
        ["Noun = N", "Noun Phrase = NP", "Subject = Subj"],
        ["Verb = V", "Verb Phrase = VP", "Predicate = Pred"],
        ["Adjective = ADJ", "", ""],
        ["Adverb = ADV", "Adjective Phrase = ADJP", ""],
        ["Determiner = DET", "Adverb Phrase = ADVP", ""],
        ["Preposition = PREP", "Prepositional Phrase = PP", ""],
        ["", "Sentence = S", ""],
    ]
    for j, h in enumerate(headers):
        cell = abbrev_table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, font_size=9)
        set_cell_shading(cell, "E8E8E8")
    for i, row_data in enumerate(abbrev_data):
        for j, val in enumerate(row_data):
            cell = abbrev_table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_run(p, val, font_size=9)

    # Notes
    p = doc.add_paragraph()
    add_run(p, "Notes:", bold=True, font_size=10)
    for note in [
        "In the Phrase row, label the main verb as VP (Verb Phrase). The VP label in the table applies to just the main verb itself.",
        "In the Role row, label the subject noun phrase as Subject and the verb phrase as Predicate. Leave other role cells empty.",
        "Always include a phrase level for nouns, verbs, adjectives, adverbs, and prepositions."
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, note, font_size=10)

    # Worked example
    p = doc.add_paragraph()
    add_run(p, "Example: ", bold=True, font_size=11)
    add_run(p, EXAMPLE_SENTENCE["sentence"], italic=True, font_size=11)
    set_paragraph_spacing(p, space_before=8, space_after=4)

    add_labeling_table(
        doc,
        EXAMPLE_SENTENCE["words"],
        pos_labels=EXAMPLE_SENTENCE["pos"],
        phrase_labels=EXAMPLE_SENTENCE["phrases"],
        role_labels=EXAMPLE_SENTENCE["roles"],
        font_size=10,
    )

    p = doc.add_paragraph()
    add_run(p, "Bracket notation: ", bold=True, font_size=10)
    add_run(p, EXAMPLE_SENTENCE["bracket"], font_size=9, font_name="Consolas")
    set_paragraph_spacing(p, space_before=4, space_after=2)

    add_diagram_image(doc, EXAMPLE_SENTENCE["diagram"], width_inches=EXAMPLE_SENTENCE["width"])

    # Questions 26-30
    for q in SECTION_B:
        doc.add_page_break()

        p = doc.add_paragraph()
        add_run(p, f"Question {q['num']}: ", bold=True, font_size=12)
        add_run(p, f"(10 pts) ", font_size=11)
        add_run(p, q["sentence"], italic=True, font_size=11)
        set_paragraph_spacing(p, space_before=6, space_after=4)

        if is_answer_key:
            add_labeling_table(
                doc,
                q["words"],
                pos_labels=q["pos"],
                phrase_labels=q["phrases"],
                role_labels=q["roles"],
                font_size=10,
            )
        else:
            add_labeling_table(
                doc,
                q["words"],
                phrase_labels=blank_labels(q["phrases"]),
                role_labels=blank_labels(q["roles"]),
                font_size=10,
            )

        p = doc.add_paragraph()
        if is_answer_key:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
            add_run(p, q["bracket"], font_size=9, font_name="Consolas")
        else:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
            add_run(p, "_" * 70, font_size=10)
        set_paragraph_spacing(p, space_before=4, space_after=2)

        if is_answer_key:
            add_diagram_image(doc, q["diagram"], width_inches=q["width"])

    # =========================================================================
    # BONUS
    # =========================================================================
    doc.add_page_break()
    add_section_header(doc, "Bonus (10 points)", level=2, font_size=14)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "Instructions: ", bold=True, font_size=11)
        add_run(p, "Write original sentences or phrases meeting each requirement below.", font_size=11)
    else:
        p = doc.add_paragraph()
        add_run(p, "Answers will vary. Accept any response that correctly demonstrates the required feature.", italic=True, font_size=11)

    for bq in BONUS_QUESTIONS:
        p = doc.add_paragraph()
        add_run(p, f"Bonus {bq['num']}: ", bold=True, font_size=11)
        add_run(p, f"{bq['text']} (2 pts)", font_size=11)
        set_paragraph_spacing(p, space_before=6, space_after=2)

        if is_answer_key:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_run(p, "Example: ", bold=True, font_size=10)
            add_run(p, bq["example"], italic=True, font_size=10)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_run(p, bq["rubric"], font_size=10)
        else:
            if bq["num"] == 3:
                # Word + morphemes fields
                p = doc.add_paragraph()
                add_run(p, "Word: ", font_size=11)
                add_run(p, "_" * 30, font_size=11)
                p = doc.add_paragraph()
                add_run(p, "Morphemes: ", font_size=11)
                add_run(p, "_" * 60, font_size=11)
            elif bq["num"] == 5:
                # Sentence + two meanings
                p = doc.add_paragraph()
                add_run(p, "Sentence: ", font_size=11)
                add_run(p, "_" * 60, font_size=11)
                p = doc.add_paragraph()
                add_run(p, "Meaning 1: ", font_size=11)
                add_run(p, "_" * 55, font_size=11)
                p = doc.add_paragraph()
                add_run(p, "Meaning 2: ", font_size=11)
                add_run(p, "_" * 55, font_size=11)
            else:
                p = doc.add_paragraph()
                add_run(p, "_" * 75, font_size=11)

    # =========================================================================
    # QUICK ANSWER REFERENCE (answer key only)
    # =========================================================================
    if is_answer_key:
        doc.add_page_break()
        add_section_header(doc, "Quick Answer Reference", level=3, font_size=12)

        answers = {
            1: "d", 2: "b", 3: "c", 4: "c", 5: "c", 6: "c", 7: "a",
            8: "b", 9: "c", 10: "c", 11: "c", 12: "c", 13: "b",
            14: "b", 15: "c", 16: "c", 17: "b", 18: "c", 19: "c",
            20: "d", 21: "b", 22: "d", 23: "b", 24: "b", 25: "c"
        }

        ref_table = doc.add_table(rows=14, cols=4)
        ref_table.style = 'Table Grid'
        for j, h in enumerate(["Question", "Answer", "Question", "Answer"]):
            cell = ref_table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, h, bold=True, font_size=10)
            set_cell_shading(cell, "E8E8E8")

        for row_idx in range(13):
            q1 = row_idx + 1
            q2 = row_idx + 14

            cell = ref_table.rows[row_idx + 1].cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(q1), font_size=10)

            cell = ref_table.rows[row_idx + 1].cells[1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, answers[q1], bold=True, font_size=10)

            if q2 <= 25:
                cell = ref_table.rows[row_idx + 1].cells[2]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, str(q2), font_size=10)

                cell = ref_table.rows[row_idx + 1].cells[3]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, answers[q2], bold=True, font_size=10)

    return doc


def main():
    print("=" * 60)
    print("Generating Exam One .docx files")
    print("=" * 60)

    # Generate student exam
    exam_path = SCRIPT_DIR / "ENGL 3110 - S26 - Exam One.docx"
    print(f"\nGenerating student exam: {exam_path.name}")
    doc = create_document(is_answer_key=False)
    doc.save(str(exam_path))
    print(f"  Saved: {exam_path}")

    # Generate answer key
    key_path = SCRIPT_DIR / "ENGL 3110 - S26 - Exam One - Answer Key.docx"
    print(f"\nGenerating answer key: {key_path.name}")
    doc = create_document(is_answer_key=True)
    doc.save(str(key_path))
    print(f"  Saved: {key_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()

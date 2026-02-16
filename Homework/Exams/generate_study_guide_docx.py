#!/usr/bin/env python3
"""
Generate Exam One Study Guide .docx files (student version and answer key) with embedded diagrams.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
DIAGRAM_DIR = SCRIPT_DIR / "diagrams"


def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set paragraph spacing in points."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_run(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_question(doc, number, text, pts=None, font_size=11):
    """Add a multiple choice question header."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=2)
    add_run(p, f"Question {number}: ", bold=True, font_size=font_size)
    add_run(p, text, font_size=font_size)
    if pts:
        add_run(p, f" ({pts} pts)", font_size=font_size)
    return p


def add_mc_option(doc, letter, text, font_size=11, is_answer=False, explanation=""):
    """Add a multiple choice option."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=1, space_after=1)
    p.paragraph_format.left_indent = Inches(0.5)
    if is_answer:
        add_run(p, f"{letter}. {text}", bold=True, font_size=font_size)
        if explanation:
            add_run(p, f" — {explanation}", italic=True, font_size=font_size - 1)
    else:
        add_run(p, f"{letter}. {text}", font_size=font_size)
        if explanation:
            add_run(p, f" — {explanation}", italic=True, font_size=font_size - 1)
    return p


def add_answer_line(doc, answer_letter, font_size=11):
    """Add 'Answer: ___' line for student version or 'Answer: X' for answer key."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=2, space_after=2)
    p.paragraph_format.left_indent = Inches(0.3)
    if answer_letter:
        add_run(p, "Answer: ", bold=True, font_size=font_size)
        add_run(p, answer_letter, bold=True, font_size=font_size)
    else:
        add_run(p, "Answer: _____", font_size=font_size)
    return p


def add_section_header(doc, text, level=2, font_size=14):
    """Add a section heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(font_size)
    return h


def compute_spans(labels):
    """Convert flat label list to (label, span_count, start_index) tuples.

    Non-empty labels start a new span; consecutive empty strings extend it.
    """
    spans = []
    i = 0
    while i < len(labels):
        label = labels[i]
        if label:
            span = 1
            while i + span < len(labels) and labels[i + span] == "":
                span += 1
            spans.append((label, span, i))
            i += span
        else:
            i += 1
    return spans


def add_labeling_table(doc, words, pos_labels=None, phrase_labels=None, role_labels=None, font_size=10):
    """Add a sentence labeling table with merged cells for Role and Phrase rows.

    When labels are provided for Role/Phrase rows, cells are merged to show
    how words group into phrases and roles. When labels are None (student
    version), cells are left unmerged so students can write in each cell.
    """
    num_cols = len(words) + 1  # +1 for row headers
    table = doc.add_table(rows=4, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_headers = ["Role", "Phrase", "Word", "POS"]
    row_data = [role_labels, phrase_labels, words, pos_labels]

    for i, (header, data) in enumerate(zip(row_headers, row_data)):
        # Row header cell
        cell = table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, header, bold=True, font_size=font_size)
        set_cell_shading(cell, "E8E8E8")

        if i <= 1 and data:
            # Role or Phrase row with data — merge cells to show groupings
            spans = compute_spans(data)
            for label, span, start_idx in spans:
                col_start = start_idx + 1  # +1 for row header column
                col_end = col_start + span - 1
                if span > 1:
                    table.rows[i].cells[col_start].merge(table.rows[i].cells[col_end])
                merged_cell = table.rows[i].cells[col_start]
                # Clear residual paragraphs from merge
                for paragraph in merged_cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                p = merged_cell.paragraphs[0]
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, label, font_size=font_size)
        else:
            # Normal unmerged row (Word, POS, or blank student rows)
            for j, val in enumerate(data if data else [""] * len(words)):
                cell = table.rows[i].cells[j + 1]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 2:  # Word row always filled
                    add_run(p, val, font_size=font_size)
                elif data:
                    add_run(p, val, font_size=font_size)
                # else leave blank for student version

    return table


def add_diagram_image(doc, image_name, width_inches=5.5):
    """Add a diagram PNG image to the document."""
    img_path = DIAGRAM_DIR / f"{image_name}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        set_paragraph_spacing(p, space_before=4, space_after=4)
        return p
    else:
        p = doc.add_paragraph(f"[Diagram not found: {image_name}.png]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p


# =============================================================================
# QUESTION DATA
# =============================================================================

MC_QUESTIONS = {
    "A1": {
        "title": "Section A1: Morphology and Word Structure (Questions 1\u20134)",
        "questions": [
            {
                "num": 1,
                "text": 'How many morphemes does the word "unbreakable" contain?',
                "options": ["1", "2", "3", "4"],
                "answer": "c",
                "explanations": {
                    "c": "un- (bound) + break (free) + -able (bound)"
                }
            },
            {
                "num": 2,
                "text": "Which of the following words contains exactly one morpheme?",
                "options": ["replay", "kitchen", "untrue", "singer"],
                "answer": "b",
                "explanations": {
                    "a": "2 morphemes: re- + play",
                    "b": "1 morpheme. It cannot be meaningfully divided into smaller units.",
                    "c": "2 morphemes: un- + true",
                    "d": "2 morphemes: sing + -er"
                }
            },
            {
                "num": 3,
                "text": "Which of the following is a free morpheme?",
                "options": ["-ly", "-ment", "tree", "un-"],
                "answer": "c",
                "explanations": {
                    "a": "bound morpheme (cannot stand alone)",
                    "b": "bound morpheme (cannot stand alone)",
                    "c": "free morpheme (can stand alone as an independent word)",
                    "d": "bound morpheme (cannot stand alone)"
                }
            },
            {
                "num": 4,
                "text": 'Which of the following correctly identifies the morphemes in "prewashing"?',
                "options": [
                    "pre- (bound) + wash (free) + -ing (bound)",
                    "pre- (free) + wash (free) + -ing (bound)",
                    "prewash (free) + -ing (bound)",
                    "pre- (bound) + washing (free)"
                ],
                "answer": "a",
                "explanations": {
                    "a": 'Correct. pre- is a bound prefix meaning "before," wash is the free base, -ing is bound.',
                    "b": "Incorrect. pre- is bound, not free.",
                    "c": "Incorrect. prewash can be further divided.",
                    "d": "Incorrect. washing should be divided into wash + -ing."
                }
            },
        ]
    },
    "A2": {
        "title": "Section A2: Open and Closed Classes (Questions 5\u20138)",
        "questions": [
            {
                "num": 5,
                "text": "Which of the following word classes is a closed class?",
                "options": ["Verbs", "Nouns", "Pronouns", "Adverbs"],
                "answer": "c",
                "explanations": {
                    "a": "open class", "b": "open class",
                    "c": "closed class (a fixed, limited set: I, you, he, she, it, we, they, etc.)",
                    "d": "open class"
                }
            },
            {
                "num": 6,
                "text": 'In the phrase "a small child in the classroom," which word is the head of the noun phrase?',
                "options": ["a", "small", "child", "classroom"],
                "answer": "c",
                "explanations": {
                    "a": "determiner (modifier)", "b": "adjective (modifier)",
                    "c": "The head noun. Cannot be removed without destroying the phrase.",
                    "d": "head of a nested NP inside the PP, not the head of the main NP."
                }
            },
            {
                "num": 7,
                "text": '"The artist displayed his paintings at the gallery." What is the antecedent of "his"?',
                "options": ["The gallery", "The artist", "paintings", "his paintings"],
                "answer": "b",
                "explanations": {
                    "b": '"his" (third person, singular, masculine) refers back to "The artist."'
                }
            },
            {
                "num": 8,
                "text": "Which of the following words is NOT a preposition?",
                "options": ["under", "between", "always", "beside"],
                "answer": "c",
                "explanations": {
                    "a": "preposition", "b": "preposition",
                    "c": "adverb (modifies verbs and describes frequency; does not introduce a PP)",
                    "d": "preposition"
                }
            },
        ]
    },
    "A3": {
        "title": "Section A3: Parts of Speech in Context (Questions 9\u201313)",
        "preamble": 'Sentence for Questions 9\u201310: "Many eager volunteers cheerfully organized the annual fundraiser."',
        "questions": [
            {
                "num": 9,
                "text": "Which word in the sentence above is a determiner?",
                "options": ["eager", "Many", "cheerfully", "annual"],
                "answer": "b",
                "explanations": {
                    "a": "adjective",
                    "b": 'determiner (indicates amount; cannot be modified by "very"; appears before adjectives)',
                    "c": "adverb", "d": "adjective"
                }
            },
            {
                "num": 10,
                "text": "Which word in the sentence above is an adverb?",
                "options": ["Many", "eager", "cheerfully", "fundraiser"],
                "answer": "c",
                "explanations": {
                    "a": "determiner", "b": "adjective",
                    "c": 'adverb (modifies the verb "organized"; formed with -ly from cheerful)',
                    "d": "noun"
                }
            },
            {
                "num": 11,
                "text": '"His brother quickly repaired our broken fence." Which word is functioning as a verb?',
                "options": ["His", "brother", "quickly", "repaired"],
                "answer": "d",
                "explanations": {
                    "a": "determiner", "b": "noun", "c": "adverb",
                    "d": "verb (the main action; takes past tense -ed)"
                }
            },
            {
                "num": 12,
                "text": '"Two curious puppies chased the red ball." Which word is an adjective?',
                "options": ["Two", "curious", "chased", "ball"],
                "answer": "b",
                "explanations": {
                    "a": "determiner",
                    "b": 'adjective (describes a quality; can be modified by "very": very curious)',
                    "c": "verb", "d": "noun"
                }
            },
            {
                "num": 13,
                "text": 'In the phrase "the remarkably bright student," what word class does "remarkably" belong to?',
                "options": ["Adjective", "Adverb", "Determiner", "Noun"],
                "answer": "b",
                "explanations": {
                    "b": '"Remarkably" modifies the adjective "bright," telling the degree of brightness.'
                }
            },
        ]
    }
}

# Section B sentence data
SECTION_B = [
    {
        "num": 14,
        "sentence": "The dog barked.",
        "words": ["The", "dog", "barked"],
        "pos": ["DET", "N", "V"],
        "phrases": ["NP", "", "VP"],
        "roles": ["Subject", "", "Predicate"],
        "bracket": "[S [NP [DET The] [N dog]] [VP [V barked]]]",
        "diagram": "sg_q14_dog_barked",
        "width": 3.5,
    },
    {
        "num": 15,
        "sentence": "The happy children played quietly in the park.",
        "words": ["The", "happy", "children", "played", "quietly", "in", "the", "park"],
        "pos": ["DET", "ADJ", "N", "V", "ADV", "PREP", "DET", "N"],
        "phrases": ["NP", "", "", "VP", "ADVP", "PP", "", ""],
        "roles": ["Subject", "", "", "Predicate", "", "", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ happy]] [N children]] [VP [V played] [ADVP [ADV quietly]] [PP [PREP in] [NP [DET the] [N park]]]]]",
        "diagram": "sg_q15_children_played",
        "width": 5.5,
    },
    {
        "num": 16,
        "sentence": "The teacher from Texas arrived early.",
        "words": ["The", "teacher", "from", "Texas", "arrived", "early"],
        "pos": ["DET", "N", "PREP", "N", "V", "ADV"],
        "phrases": ["NP", "", "PP", "", "VP", "ADVP"],
        "roles": ["Subject", "", "", "", "Predicate", ""],
        "bracket": "[S [NP [DET The] [N teacher] [PP [PREP from] [NP [N Texas]]]] [VP [V arrived] [ADVP [ADV early]]]]",
        "diagram": "sg_q16_teacher_arrived",
        "width": 5.0,
    },
    {
        "num": 17,
        "sentence": "Many talented musicians performed brilliantly during the annual festival.",
        "words": ["Many", "talented", "musicians", "performed", "brilliantly", "during", "the", "annual", "festival"],
        "pos": ["DET", "ADJ", "N", "V", "ADV", "PREP", "DET", "ADJ", "N"],
        "phrases": ["NP", "", "", "VP", "ADVP", "PP", "", "", ""],
        "roles": ["Subject", "", "", "Predicate", "", "", "", "", ""],
        "bracket": "[S [NP [DET Many] [ADJP [ADJ talented]] [N musicians]] [VP [V performed] [ADVP [ADV brilliantly]] [PP [PREP during] [NP [DET the] [ADJP [ADJ annual]] [N festival]]]]]",
        "diagram": "sg_q17_musicians_performed",
        "width": 6.0,
    },
    {
        "num": 18,
        "sentence": "The young artist with curly hair painted quite skillfully.",
        "words": ["The", "young", "artist", "with", "curly", "hair", "painted", "quite", "skillfully"],
        "pos": ["DET", "ADJ", "N", "PREP", "ADJ", "N", "V", "ADV", "ADV"],
        "phrases": ["NP", "", "", "PP", "", "", "VP", "ADVP", ""],
        "roles": ["Subject", "", "", "", "", "", "Predicate", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ young]] [N artist] [PP [PREP with] [NP [ADJP [ADJ curly]] [N hair]]]] [VP [V painted] [ADVP [ADVP [ADV quite]] [ADV skillfully]]]]",
        "diagram": "sg_q18_artist_painted",
        "width": 6.0,
    },
]

EXAMPLE_SENTENCE = {
    "sentence": "The old man sat quietly.",
    "words": ["The", "old", "man", "sat", "quietly"],
    "pos": ["DET", "ADJ", "N", "V", "ADV"],
    "phrases": ["NP", "", "", "VP", "ADVP"],
    "roles": ["Subject", "", "", "Predicate", ""],
    "bracket": "[S [NP [DET The] [ADJP [ADJ old]] [N man]] [VP [V sat] [ADVP [ADV quietly]]]]",
    "diagram": "example_old_man_sat",
    "width": 5.5,
}

BONUS_QUESTIONS = [
    {
        "num": 1,
        "text": "Write a word that contains at least three morphemes. List each morpheme and label it as free or bound.",
        "example": '"disagreeable" \u2014 dis- (bound) + agree (free) + -able (bound) = 3 morphemes',
        "rubric": "Award 1 pt for a valid word with 3+ morphemes and 1 pt for correct labeling."
    },
    {
        "num": 2,
        "text": "Write a noun phrase that contains a determiner, an adjective, and a noun.",
        "example": '"the cheerful dog"',
        "rubric": "Award 1 pt for including all three elements and 1 pt for correct word class identification (DET + ADJ + N)."
    },
    {
        "num": 3,
        "text": "Write a sentence containing a prepositional phrase that modifies a noun. Underline or circle the prepositional phrase.",
        "example": 'The book on the table belongs to me.',
        "rubric": "Award 1 pt for a valid sentence and 1 pt for correctly identifying a PP that modifies a noun (not a verb)."
    },
]


def create_document(is_answer_key=False):
    """Create either the student study guide or the answer key."""
    doc = Document()

    # Set up default style
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Open Sans'
        h_style.font.bold = True

    # =========================================================================
    # HEADER
    # =========================================================================
    title_text = "ENGL 3110 - Exam One Study Guide - Spring 2026"
    if is_answer_key:
        title_text += " \u2014 ANSWER KEY"

    title = doc.add_heading(title_text, level=1)
    for run in title.runs:
        run.font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=4)

    p = doc.add_paragraph()
    add_run(p, "This study guide mirrors the format and topics of Exam One. Use it to practice before the exam.", italic=True, font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=4)

    # Section overview
    p = doc.add_paragraph()
    add_run(p, "Sections:", bold=True, font_size=11)
    for line in [
        "Section A: Multiple Choice \u2014 13 questions",
        "Section B: Sentence Labeling and Diagramming \u2014 5 sentences",
        "Bonus \u2014 3 questions"
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, line, font_size=10)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "STUDENT NAME: ", bold=True, font_size=11)
        add_run(p, "_" * 60, font_size=11)

    # =========================================================================
    # SECTION A: MULTIPLE CHOICE
    # =========================================================================
    add_section_header(doc, "Section A: Multiple Choice", level=2, font_size=14)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "Instructions: ", bold=True, font_size=11)
        add_run(p, "For each question, select the best answer and write the letter in the blank provided.", font_size=11)

    letters = ["a", "b", "c", "d"]

    for section_key in ["A1", "A2", "A3"]:
        section = MC_QUESTIONS[section_key]
        add_section_header(doc, section["title"], level=3, font_size=12)

        if "preamble" in section:
            p = doc.add_paragraph()
            add_run(p, section["preamble"], italic=True, font_size=11)
            set_paragraph_spacing(p, space_before=4, space_after=4)

        for q in section["questions"]:
            add_question(doc, q["num"], q["text"])

            if is_answer_key:
                add_answer_line(doc, q["answer"])
            else:
                add_answer_line(doc, None)

            for i, option_text in enumerate(q["options"]):
                letter = letters[i]
                is_correct = (letter == q["answer"])
                explanation = q.get("explanations", {}).get(letter, "")

                if is_answer_key:
                    add_mc_option(doc, letter, option_text, is_answer=is_correct, explanation=explanation)
                else:
                    add_mc_option(doc, letter, option_text)

    # =========================================================================
    # SECTION B: SENTENCE LABELING AND DIAGRAMMING
    # =========================================================================
    doc.add_page_break()
    add_section_header(doc, "Section B: Sentence Labeling and Diagramming", level=2, font_size=14)

    # Instructions
    p = doc.add_paragraph()
    add_run(p, "Instructions: ", bold=True, font_size=11)
    add_run(p, "For each sentence below, take two steps.", font_size=11)

    p = doc.add_paragraph()
    add_run(p, "1. ", bold=True, font_size=11)
    add_run(p, "Complete the labeling table ", bold=True, font_size=11)
    add_run(p, "by filling in the Role, Phrase, and Part of Speech (POS) rows.", font_size=11)

    p = doc.add_paragraph()
    add_run(p, "2. ", bold=True, font_size=11)
    add_run(p, "Write the bracket notation ", bold=True, font_size=11)
    add_run(p, "for the sentence's tree diagram.", font_size=11)

    # Abbreviation key table
    p = doc.add_paragraph()
    add_run(p, "Abbreviation Key", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=6, space_after=2)

    abbrev_table = doc.add_table(rows=8, cols=3)
    abbrev_table.style = 'Table Grid'
    headers = ["Parts of Speech", "Phrase Types", "Roles"]
    abbrev_data = [
        ["Noun = N", "Noun Phrase = NP", "Subject = Subj"],
        ["Verb = V", "Verb Phrase = VP", "Predicate = Pred"],
        ["Adjective = ADJ", "", ""],
        ["Adverb = ADV", "Adjective Phrase = ADJP", ""],
        ["Determiner = DET", "Adverb Phrase = ADVP", ""],
        ["Preposition = PREP", "Prepositional Phrase = PP", ""],
        ["", "Sentence = S", ""],
    ]
    for j, h in enumerate(headers):
        cell = abbrev_table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, font_size=9)
        set_cell_shading(cell, "E8E8E8")
    for i, row_data in enumerate(abbrev_data):
        for j, val in enumerate(row_data):
            cell = abbrev_table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_run(p, val, font_size=9)

    # Notes
    p = doc.add_paragraph()
    add_run(p, "Notes:", bold=True, font_size=10)
    for note in [
        "In the Phrase row, label the main verb as VP (Verb Phrase). The VP label in the table applies to just the main verb itself.",
        "In the Role row, label the subject noun phrase as Subject and the verb phrase as Predicate. Leave other role cells empty.",
        "Always include a phrase level for nouns, verbs, adjectives, adverbs, and prepositions."
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, note, font_size=10)

    # Worked example
    p = doc.add_paragraph()
    add_run(p, "Example: ", bold=True, font_size=11)
    add_run(p, EXAMPLE_SENTENCE["sentence"], italic=True, font_size=11)
    set_paragraph_spacing(p, space_before=8, space_after=4)

    add_labeling_table(
        doc,
        EXAMPLE_SENTENCE["words"],
        pos_labels=EXAMPLE_SENTENCE["pos"],
        phrase_labels=EXAMPLE_SENTENCE["phrases"],
        role_labels=EXAMPLE_SENTENCE["roles"],
        font_size=10,
    )

    p = doc.add_paragraph()
    add_run(p, "Bracket notation: ", bold=True, font_size=10)
    add_run(p, EXAMPLE_SENTENCE["bracket"], font_size=9, font_name="Consolas")
    set_paragraph_spacing(p, space_before=4, space_after=2)

    add_diagram_image(doc, EXAMPLE_SENTENCE["diagram"], width_inches=EXAMPLE_SENTENCE["width"])

    # Questions 14-18
    for q in SECTION_B:
        doc.add_page_break()

        p = doc.add_paragraph()
        add_run(p, f"Question {q['num']}: ", bold=True, font_size=12)
        add_run(p, q["sentence"], italic=True, font_size=11)
        set_paragraph_spacing(p, space_before=6, space_after=4)

        if is_answer_key:
            add_labeling_table(
                doc,
                q["words"],
                pos_labels=q["pos"],
                phrase_labels=q["phrases"],
                role_labels=q["roles"],
                font_size=10,
            )
        else:
            add_labeling_table(
                doc,
                q["words"],
                font_size=10,
            )

        p = doc.add_paragraph()
        if is_answer_key:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
            add_run(p, q["bracket"], font_size=9, font_name="Consolas")
        else:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
            add_run(p, "_" * 70, font_size=10)
        set_paragraph_spacing(p, space_before=4, space_after=2)

        if is_answer_key:
            add_diagram_image(doc, q["diagram"], width_inches=q["width"])

    # =========================================================================
    # BONUS
    # =========================================================================
    doc.add_page_break()
    add_section_header(doc, "Bonus", level=2, font_size=14)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "Instructions: ", bold=True, font_size=11)
        add_run(p, "Write original sentences or phrases meeting each requirement below.", font_size=11)
    else:
        p = doc.add_paragraph()
        add_run(p, "Answers will vary. Accept any response that correctly demonstrates the required feature.", italic=True, font_size=11)

    for bq in BONUS_QUESTIONS:
        p = doc.add_paragraph()
        add_run(p, f"Bonus {bq['num']}: ", bold=True, font_size=11)
        add_run(p, f"{bq['text']} (2 pts)", font_size=11)
        set_paragraph_spacing(p, space_before=6, space_after=2)

        if is_answer_key:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_run(p, "Example: ", bold=True, font_size=10)
            add_run(p, bq["example"], italic=True, font_size=10)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_run(p, bq["rubric"], font_size=10)
        else:
            if bq["num"] == 1:
                # Word + morphemes fields
                p = doc.add_paragraph()
                add_run(p, "Word: ", font_size=11)
                add_run(p, "_" * 30, font_size=11)
                p = doc.add_paragraph()
                add_run(p, "Morphemes: ", font_size=11)
                add_run(p, "_" * 60, font_size=11)
            else:
                p = doc.add_paragraph()
                add_run(p, "_" * 75, font_size=11)

    # =========================================================================
    # QUICK ANSWER REFERENCE (answer key only)
    # =========================================================================
    if is_answer_key:
        doc.add_page_break()
        add_section_header(doc, "Quick Answer Reference", level=3, font_size=12)

        answers = {
            1: "c", 2: "b", 3: "c", 4: "a", 5: "c", 6: "c", 7: "b",
            8: "c", 9: "b", 10: "c", 11: "d", 12: "b", 13: "b"
        }

        ref_table = doc.add_table(rows=8, cols=4)
        ref_table.style = 'Table Grid'
        for j, h in enumerate(["Question", "Answer", "Question", "Answer"]):
            cell = ref_table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, h, bold=True, font_size=10)
            set_cell_shading(cell, "E8E8E8")

        for row_idx in range(7):
            q1 = row_idx + 1
            q2 = row_idx + 8

            cell = ref_table.rows[row_idx + 1].cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(q1), font_size=10)

            cell = ref_table.rows[row_idx + 1].cells[1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, answers[q1], bold=True, font_size=10)

            if q2 <= 13:
                cell = ref_table.rows[row_idx + 1].cells[2]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, str(q2), font_size=10)

                cell = ref_table.rows[row_idx + 1].cells[3]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, answers[q2], bold=True, font_size=10)

    return doc


def main():
    print("=" * 60)
    print("Generating Exam One Study Guide .docx files")
    print("=" * 60)

    # Generate student study guide
    sg_path = SCRIPT_DIR / "ENGL 3110 - S26 - Exam One Study Guide.docx"
    print(f"\nGenerating student study guide: {sg_path.name}")
    doc = create_document(is_answer_key=False)
    doc.save(str(sg_path))
    print(f"  Saved: {sg_path}")

    # Generate answer key
    key_path = SCRIPT_DIR / "ENGL 3110 - S26 - Exam One Study Guide - Answer Key.docx"
    print(f"\nGenerating answer key: {key_path.name}")
    doc = create_document(is_answer_key=True)
    doc.save(str(key_path))
    print(f"  Saved: {key_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()

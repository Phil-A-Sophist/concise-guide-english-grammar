#!/usr/bin/env python3
"""
Generate Morphology Bonus Assignment .docx files (student version and answer key).
No diagrams — students divide 20 words into morphemes and label each free or bound.
Answer format: one morpheme per line, e.g. "play — free" / "-ful — bound"
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent


# =============================================================================
# HELPERS
# =============================================================================

def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_run(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_section_header(doc, text, level=2, font_size=14):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(font_size)
    return h


def fill_answer_cell(cell, morphemes, labels, font_size=10):
    """Fill a table cell with one morpheme-label line per morpheme."""
    for j, (morph, label) in enumerate(zip(morphemes, labels)):
        p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, morph, bold=True, font_size=font_size)
        add_run(p, f"  \u2014  {label}", font_size=font_size)


def fill_blank_answer_cell(cell):
    """Leave a tall blank area for student to type their answer."""
    p = cell.paragraphs[0]
    p.text = ""
    set_paragraph_spacing(p, space_before=0, space_after=0)
    # Extra blank lines give visual height so the cell isn't cramped
    for _ in range(3):
        ep = cell.add_paragraph()
        set_paragraph_spacing(ep, space_before=0, space_after=0)


# =============================================================================
# WORD DATA
# Each entry: word, breakdown (list of morphemes), labels (list of free/bound)
# Prefixes shown with trailing hyphen: "un-"; suffixes with leading hyphen: "-ness"
# =============================================================================

WORDS = [
    {
        "word": "playful",
        "morphemes": ["play", "-ful"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "darkness",
        "morphemes": ["dark", "-ness"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "quickly",
        "morphemes": ["quick", "-ly"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "kindness",
        "morphemes": ["kind", "-ness"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "bookcase",
        "morphemes": ["book", "case"],
        "labels":    ["free", "free"],
        "note": "compound noun",
    },
    {
        "word": "truthful",
        "morphemes": ["truth", "-ful"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "powerless",
        "morphemes": ["power", "-less"],
        "labels":    ["free", "bound"],
        "note": "",
    },
    {
        "word": "misunderstand",
        "morphemes": ["mis-", "understand"],
        "labels":    ["bound", "free"],
        "note": "understand is a single morpheme in modern English",
    },
    {
        "word": "rebuilding",
        "morphemes": ["re-", "build", "-ing"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "teachers",
        "morphemes": ["teach", "-er", "-s"],
        "labels":    ["free", "bound", "bound"],
        "note": "",
    },
    {
        "word": "carefully",
        "morphemes": ["care", "-ful", "-ly"],
        "labels":    ["free", "bound", "bound"],
        "note": "",
    },
    {
        "word": "unfriendly",
        "morphemes": ["un-", "friend", "-ly"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "notebooks",
        "morphemes": ["note", "book", "-s"],
        "labels":    ["free", "free", "bound"],
        "note": "compound noun + plural",
    },
    {
        "word": "unlockable",
        "morphemes": ["un-", "lock", "-able"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "songwriter",
        "morphemes": ["song", "write", "-er"],
        "labels":    ["free", "free", "bound"],
        "note": "compound base + derivational suffix",
    },
    {
        "word": "replayable",
        "morphemes": ["re-", "play", "-able"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "unhappiness",
        "morphemes": ["un-", "happy", "-ness"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "unpredictable",
        "morphemes": ["un-", "predict", "-able"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "disrespectful",
        "morphemes": ["dis-", "respect", "-ful"],
        "labels":    ["bound", "free", "bound"],
        "note": "",
    },
    {
        "word": "disagreements",
        "morphemes": ["dis-", "agree", "-ment", "-s"],
        "labels":    ["bound", "free", "bound", "bound"],
        "note": "",
    },
]

EXAMPLE = {
    "word": "unhelpful",
    "morphemes": ["un-", "help", "-ful"],
    "labels":    ["bound", "free", "bound"],
    "note": "",
}


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================

def build_example_table(doc, is_answer_key, font_size=10):
    """Add a small completed-example table showing the expected answer format."""
    num_cols = 4 if is_answer_key else 3
    table = doc.add_table(rows=2, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = (["#", "Word", "Answer", "Notes"] if is_answer_key
               else ["#", "Word", "Answer"])
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, h, bold=True, font_size=font_size)
        set_cell_shading(cell, "E8E8E8")

    row = table.rows[1]

    # Col 0: placeholder number
    cell = row.cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "\u2014", font_size=font_size)

    # Col 1: word
    cell = row.cells[1]
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, EXAMPLE["word"], bold=True, font_size=font_size)

    # Col 2: answer (always filled in — it's the example)
    fill_answer_cell(row.cells[2], EXAMPLE["morphemes"], EXAMPLE["labels"], font_size)

    # Col 3: notes (answer key only)
    if is_answer_key:
        cell = row.cells[3]
        cell.text = ""


def create_document(is_answer_key=False):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Open Sans'
        h_style.font.bold = True

    # -------------------------------------------------------------------------
    # HEADER
    # -------------------------------------------------------------------------
    title_text = "ENGL 3110 \u2014 Morphology Bonus Assignment \u2014 Spring 2026"
    if is_answer_key:
        title_text += " \u2014 ANSWER KEY"

    title = doc.add_heading(title_text, level=1)
    for run in title.runs:
        run.font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=4)

    p = doc.add_paragraph()
    add_run(p, "Total Points: 20  \u2022  1 point per word", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=4)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "STUDENT NAME: ", bold=True, font_size=11)
        add_run(p, "_" * 55, font_size=11)
        set_paragraph_spacing(p, space_before=0, space_after=6)

    # -------------------------------------------------------------------------
    # INSTRUCTIONS
    # -------------------------------------------------------------------------
    add_section_header(doc, "Instructions", level=2, font_size=13)

    p = doc.add_paragraph()
    add_run(p, "For each word below, divide it into its component morphemes and label each morpheme as ", font_size=11)
    add_run(p, "free", bold=True, font_size=11)
    add_run(p, " (can stand alone as a word) or ", font_size=11)
    add_run(p, "bound", bold=True, font_size=11)
    add_run(p, " (must attach to another morpheme).", font_size=11)

    p = doc.add_paragraph()
    add_run(p, "List one morpheme per line in the Answer column. Mark prefixes with a trailing hyphen (", font_size=11)
    add_run(p, "un-", italic=True, font_size=11)
    add_run(p, ") and suffixes with a leading hyphen (", font_size=11)
    add_run(p, "-ness", italic=True, font_size=11)
    add_run(p, ").", font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=6)

    # Completed example
    p = doc.add_paragraph()
    add_run(p, "Example (completed):", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=4, space_after=2)

    build_example_table(doc, is_answer_key, font_size=10)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=6, space_after=4)

    # -------------------------------------------------------------------------
    # MAIN TABLE
    # -------------------------------------------------------------------------
    num_cols = 4 if is_answer_key else 3
    table = doc.add_table(rows=len(WORDS) + 1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = (["#", "Word", "Answer", "Notes"] if is_answer_key
               else ["#", "Word", "Answer"])
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, h, bold=True, font_size=10)
        set_cell_shading(cell, "E8E8E8")

    for i, entry in enumerate(WORDS):
        row = table.rows[i + 1]

        # Col 0: number
        cell = row.cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(i + 1), font_size=10)

        # Col 1: word
        cell = row.cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, entry["word"], bold=True, font_size=10)

        # Col 2: answer
        if is_answer_key:
            fill_answer_cell(row.cells[2], entry["morphemes"], entry["labels"], font_size=10)
        else:
            fill_blank_answer_cell(row.cells[2])

        # Col 3: notes (answer key only)
        if is_answer_key:
            cell = row.cells[3]
            cell.text = ""
            p = cell.paragraphs[0]
            if entry["note"]:
                add_run(p, entry["note"], font_size=9, italic=True)

        # Alternating row shading
        if (i + 1) % 2 == 0:
            for j in range(num_cols):
                set_cell_shading(table.rows[i + 1].cells[j], "F5F5F5")

    # -------------------------------------------------------------------------
    # SCORING NOTE (student version only)
    # -------------------------------------------------------------------------
    if not is_answer_key:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=8, space_after=2)
        add_run(p, "Scoring: ", bold=True, font_size=10)
        add_run(p, "1 point per word. Full credit requires correctly identifying ALL morphemes and labeling each one.", font_size=10)

    return doc


def main():
    print("=" * 60)
    print("Generating Morphology Bonus Assignment .docx files")
    print("=" * 60)

    student_path = SCRIPT_DIR / "ENGL 3110 - S26 - Morphology Bonus Assignment.docx"
    print(f"\nGenerating student version: {student_path.name}")
    doc = create_document(is_answer_key=False)
    doc.save(str(student_path))
    print(f"  Saved: {student_path}")

    key_path = SCRIPT_DIR / "ENGL 3110 - S26 - Morphology Bonus Assignment - Answer Key.docx"
    print(f"\nGenerating answer key: {key_path.name}")
    doc = create_document(is_answer_key=True)
    doc.save(str(key_path))
    print(f"  Saved: {key_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()

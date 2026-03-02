#!/usr/bin/env python3
"""
Generate Bonus Assignment .docx files (student version and answer key).
Run generate_bonus_diagrams.py first to create the PNG files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
DIAGRAM_DIR = SCRIPT_DIR / "diagrams"


# =============================================================================
# HELPERS (same pattern as exam scripts)
# =============================================================================

def set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(spacing)


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_run(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_section_header(doc, text, level=2, font_size=14):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(font_size)
    return h


def compute_spans(labels):
    """Convert flat label list to (label, span_count, start_index) tuples."""
    spans = []
    i = 0
    while i < len(labels):
        label = labels[i]
        if label:
            span = 1
            while i + span < len(labels) and labels[i + span] == "":
                span += 1
            spans.append((label, span, i))
            i += span
        else:
            i += 1
    return spans


def blank_labels(labels):
    """Convert answer labels to blank labels preserving span structure."""
    return ["" if label == "" else " " for label in labels]


def add_labeling_table(doc, words, pos_labels=None, phrase_labels=None, role_labels=None, font_size=10):
    """Add a 4-row sentence labeling table (Role, Phrase, Word, POS)."""
    num_cols = len(words) + 1
    table = doc.add_table(rows=4, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_headers = ["Role", "Phrase", "Word", "POS"]
    row_data = [role_labels, phrase_labels, words, pos_labels]

    for i, (header, data) in enumerate(zip(row_headers, row_data)):
        cell = table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, header, bold=True, font_size=font_size)
        set_cell_shading(cell, "E8E8E8")

        if i <= 1 and data:
            spans = compute_spans(data)
            for label, span, start_idx in spans:
                col_start = start_idx + 1
                col_end = col_start + span - 1
                if span > 1:
                    table.rows[i].cells[col_start].merge(table.rows[i].cells[col_end])
                merged_cell = table.rows[i].cells[col_start]
                for paragraph in merged_cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                p = merged_cell.paragraphs[0]
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, label, font_size=font_size)
        else:
            for j, val in enumerate(data if data else [""] * len(words)):
                cell = table.rows[i].cells[j + 1]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 2:
                    add_run(p, val, font_size=font_size)
                elif data:
                    add_run(p, val, font_size=font_size)

    return table


def add_diagram_image(doc, image_name, width_inches=5.0):
    img_path = DIAGRAM_DIR / f"{image_name}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        set_paragraph_spacing(p, space_before=4, space_after=4)
        return p
    else:
        p = doc.add_paragraph(f"[Diagram not found: {image_name}.png]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p


def add_diagram_box(doc):
    """Add a bordered blank area for students to paste their diagram."""
    p = doc.add_paragraph()
    add_run(p, "Diagram: ", bold=True, font_size=10)
    add_run(p, "Please paste your diagram below.", font_size=10, italic=True)
    set_paragraph_spacing(p, space_before=4, space_after=2)

    # Blank space (4 empty paragraphs)
    for _ in range(4):
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, space_before=0, space_after=0)

    return p


# =============================================================================
# SENTENCE DATA
# =============================================================================

SENTENCES = [
    # --- Sentence 1 ---
    # Long NP (DET+ADJ+ADJ+N) with post-nominal PP modifying the subject noun.
    # Shows: PP → N (subject head)
    {
        "num": 1,
        "sentence": "The tall energetic student near the door arrived late.",
        "words":   ["The", "tall", "energetic", "student", "near", "the", "door", "arrived", "late"],
        "pos":     ["DET", "ADJ", "ADJ",        "N",       "PREP", "DET", "N",    "V",       "ADV"],
        "phrases": ["NP",  "",    "",            "",        "PP",   "",    "",     "VP",      "ADVP"],
        # NP spans cols 1-4 (DET+ADJ+ADJ+N), PP spans cols 5-7 (post-nominal, inside NP)
        "roles":   ["Subject", "", "", "", "", "", "", "Predicate", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ tall]] [ADJP [ADJ energetic]] [N student] [PP [PREP near] [NP [DET the] [N door]]]] [VP [V arrived] [ADVP [ADV late]]]]",
        "diagram": "bonus_q01_student_arrived",
        "width": 5.5,
    },
    # --- Sentence 2 ---
    # Very long NP (DET+ADJ+ADJ+N+PP) where the PP object itself contains an ADJ.
    # Shows: PP → N (subject head); long NP
    {
        "num": 2,
        "sentence": "The brilliant young professor from the northern campus spoke clearly.",
        "words":   ["The", "brilliant", "young", "professor", "from", "the", "northern", "campus", "spoke",  "clearly"],
        "pos":     ["DET", "ADJ",       "ADJ",   "N",         "PREP", "DET", "ADJ",      "N",      "V",      "ADV"],
        "phrases": ["NP",  "",          "",      "",           "PP",   "",    "",          "",      "VP",     "ADVP"],
        # NP spans cols 1-4, PP spans cols 5-8 (post-nominal inside NP)
        "roles":   ["Subject", "", "", "", "", "", "", "", "Predicate", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ brilliant]] [ADJP [ADJ young]] [N professor] [PP [PREP from] [NP [DET the] [ADJP [ADJ northern]] [N campus]]]] [VP [V spoke] [ADVP [ADV clearly]]]]",
        "diagram": "bonus_q02_professor_spoke",
        "width": 6.0,
    },
    # --- Sentence 3 ---
    # Post-nominal ADJP with a PP complement modifying the adjective.
    # Shows: PP → ADJ
    {
        "num": 3,
        "sentence": "The student afraid of failure studied very diligently.",
        "words":   ["The", "student", "afraid", "of",   "failure", "studied", "very", "diligently"],
        "pos":     ["DET", "N",       "ADJ",    "PREP", "N",       "V",       "ADV",  "ADV"],
        "phrases": ["NP",  "",        "ADJP",   "",     "",        "VP",      "ADVP", ""],
        # NP spans cols 1-2; ADJP spans cols 3-5 (post-nominal, inside NP; PP inside ADJP in bracket)
        "roles":   ["Subject", "", "", "", "", "Predicate", "", ""],
        "bracket": "[S [NP [DET The] [N student] [ADJP [ADJ afraid] [PP [PREP of] [NP [N failure]]]]] [VP [V studied] [ADVP [ADVP [ADV very]] [ADV diligently]]]]",
        "diagram": "bonus_q03_student_studied",
        "width": 5.5,
    },
    # --- Sentence 4 ---
    # ADVP whose head adverb "far" is modified by the following PP.
    # Shows: PP → ADV
    {
        "num": 4,
        "sentence": "They lived quite far from the noisy old city.",
        "words":   ["They", "lived", "quite", "far", "from", "the", "noisy", "old", "city"],
        "pos":     ["PRO",  "V",     "ADV",   "ADV", "PREP", "DET", "ADJ",   "ADJ", "N"],
        "phrases": ["NP",   "VP",    "ADVP",  "",    "",     "",    "",      "",    ""],
        # ADVP spans cols 3-9; the PP "from the noisy old city" is inside it, modifying "far"
        "roles":   ["Subject", "Predicate", "", "", "", "", "", "", ""],
        "bracket": "[S [NP [PRO They]] [VP [V lived] [ADVP [ADVP [ADV quite]] [ADV far] [PP [PREP from] [NP [DET the] [ADJP [ADJ noisy]] [ADJP [ADJ old]] [N city]]]]]]",
        "diagram": "bonus_q04_they_lived",
        "width": 6.0,
    },
    # --- Sentence 5 ---
    # Nested PP: the outer PP's NP object contains its own post-nominal PP.
    # Shows: PP → N (as object of a preposition)
    {
        "num": 5,
        "sentence": "He waited quietly at the bench near the old fountain.",
        "words":   ["He",  "waited", "quietly", "at",   "the", "bench", "near", "the", "old", "fountain"],
        "pos":     ["PRO", "V",      "ADV",     "PREP", "DET", "N",     "PREP", "DET", "ADJ", "N"],
        "phrases": ["NP",  "VP",     "ADVP",    "PP",   "",    "",      "PP",   "",    "",    ""],
        # PP1 spans cols 4-6 (head of outer PP: "at the bench")
        # PP2 spans cols 7-10 (inner PP modifying "bench", inside the NP object of "at")
        "roles":   ["Subject", "Predicate", "", "", "", "", "", "", "", ""],
        "bracket": "[S [NP [PRO He]] [VP [V waited] [ADVP [ADV quietly]] [PP [PREP at] [NP [DET the] [N bench] [PP [PREP near] [NP [DET the] [ADJP [ADJ old]] [N fountain]]]]]]]",
        "diagram": "bonus_q05_he_waited",
        "width": 6.0,
    },
    # --- Sentence 6 ---
    # Very long VP: main verb + ADVP + two adverbial PPs.
    # Shows: PP → V (adverbial, two instances); long VP
    {
        "num": 6,
        "sentence": "She sat quite peacefully near the tall fountain in the old courtyard.",
        "words":   ["She", "sat", "quite", "peacefully", "near", "the", "tall",  "fountain", "in",   "the", "old", "courtyard"],
        "pos":     ["PRO", "V",   "ADV",   "ADV",        "PREP", "DET", "ADJ",   "N",        "PREP", "DET", "ADJ", "N"],
        "phrases": ["NP",  "VP",  "ADVP",  "",           "PP",   "",    "",      "",         "PP",   "",    "",    ""],
        # ADVP spans cols 3-4; PP1 spans cols 5-8; PP2 spans cols 9-12
        "roles":   ["Subject", "Predicate", "", "", "", "", "", "", "", "", "", ""],
        "bracket": "[S [NP [PRO She]] [VP [V sat] [ADVP [ADVP [ADV quite]] [ADV peacefully]] [PP [PREP near] [NP [DET the] [ADJP [ADJ tall]] [N fountain]]] [PP [PREP in] [NP [DET the] [ADJP [ADJ old]] [N courtyard]]]]]",
        "diagram": "bonus_q06_she_sat",
        "width": 6.5,
    },
    # --- Sentence 7 ---
    # Long NP (DET+ADJ+ADJ+N+PP) AND long VP (V+ADVP+PP).
    # Shows: PP → N (subject head) and PP → V (adverbial); both NP and VP are complex
    {
        "num": 7,
        "sentence": "The tired old dog from the farm slept quietly near the fire.",
        "words":   ["The", "tired", "old", "dog", "from", "the", "farm", "slept",  "quietly", "near", "the", "fire"],
        "pos":     ["DET", "ADJ",   "ADJ", "N",   "PREP", "DET", "N",    "V",      "ADV",     "PREP", "DET", "N"],
        "phrases": ["NP",  "",      "",    "",     "PP",   "",    "",     "VP",     "ADVP",    "PP",   "",    ""],
        # NP spans cols 1-4; PP1 (subject-NP post-nominal) cols 5-7;
        # VP col 8; ADVP col 9; PP2 (adverbial) cols 10-12
        "roles":   ["Subject", "", "", "", "", "", "", "Predicate", "", "", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ tired]] [ADJP [ADJ old]] [N dog] [PP [PREP from] [NP [DET the] [N farm]]]] [VP [V slept] [ADVP [ADV quietly]] [PP [PREP near] [NP [DET the] [N fire]]]]]",
        "diagram": "bonus_q07_dog_slept",
        "width": 6.5,
    },
    # --- Sentence 8 ---
    # Simple NP; VP has ADVP + PP where the PP object NP contains two ADJs.
    # Shows: PP → V (adverbial); long PP object
    {
        "num": 8,
        "sentence": "The nervous boy sat rigidly in the hard narrow chair.",
        "words":   ["The", "nervous", "boy", "sat", "rigidly", "in",   "the", "hard", "narrow", "chair"],
        "pos":     ["DET", "ADJ",     "N",   "V",   "ADV",     "PREP", "DET", "ADJ",  "ADJ",    "N"],
        "phrases": ["NP",  "",        "",    "VP",  "ADVP",    "PP",   "",    "",     "",       ""],
        "roles":   ["Subject", "", "", "Predicate", "", "", "", "", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ nervous]] [N boy]] [VP [V sat] [ADVP [ADV rigidly]] [PP [PREP in] [NP [DET the] [ADJP [ADJ hard]] [ADJP [ADJ narrow]] [N chair]]]]]",
        "diagram": "bonus_q08_boy_sat",
        "width": 5.5,
    },
    # --- Sentence 9 ---
    # VP with two AVPs: the second ADVP has "deep" as its head with a PP complement.
    # Shows: PP → ADV (a second example); long VP
    {
        "num": 9,
        "sentence": "The small birds sang very loudly deep in the dark forest.",
        "words":   ["The", "small", "birds", "sang", "very", "loudly", "deep", "in",   "the", "dark", "forest"],
        "pos":     ["DET", "ADJ",   "N",     "V",    "ADV",  "ADV",    "ADV",  "PREP", "DET", "ADJ",  "N"],
        "phrases": ["NP",  "",      "",      "VP",   "ADVP", "",       "ADVP", "",     "",    "",     ""],
        # ADVP1 spans cols 5-6 ("very loudly"); ADVP2 spans cols 7-11 ("deep in the dark forest")
        "roles":   ["Subject", "", "", "Predicate", "", "", "", "", "", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ small]] [N birds]] [VP [V sang] [ADVP [ADVP [ADV very]] [ADV loudly]] [ADVP [ADV deep] [PP [PREP in] [NP [DET the] [ADJP [ADJ dark]] [N forest]]]]]]",
        "diagram": "bonus_q09_birds_sang",
        "width": 6.0,
    },
    # --- Sentence 10 ---
    # Very long NP: pre-nominal ADJ + N + post-nominal ADJP, where the ADJP has a PP complement.
    # Shows: PP → ADJ (a second example); very long NP
    {
        "num": 10,
        "sentence": "The old woman confident in her skills spoke quite passionately.",
        "words":   ["The", "old", "woman", "confident", "in",   "her", "skills", "spoke", "quite", "passionately"],
        "pos":     ["DET", "ADJ", "N",     "ADJ",       "PREP", "DET", "N",      "V",     "ADV",   "ADV"],
        "phrases": ["NP",  "",    "",      "ADJP",      "",     "",    "",       "VP",    "ADVP",  ""],
        # NP spans cols 1-3 (DET+ADJ+N); ADJP spans cols 4-7 (post-nominal, inside NP;
        # PP "in her skills" is inside the ADJP in bracket notation)
        "roles":   ["Subject", "", "", "", "", "", "", "Predicate", "", ""],
        "bracket": "[S [NP [DET The] [ADJP [ADJ old]] [N woman] [ADJP [ADJ confident] [PP [PREP in] [NP [DET her] [N skills]]]]] [VP [V spoke] [ADVP [ADVP [ADV quite]] [ADV passionately]]]]",
        "diagram": "bonus_q10_woman_spoke",
        "width": 6.0,
    },
]


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================

def build_abbrev_table(doc):
    """Add the standard abbreviation key table."""
    p = doc.add_paragraph()
    add_run(p, "Abbreviation Key", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=6, space_after=2)

    abbrev_table = doc.add_table(rows=8, cols=3)
    abbrev_table.style = 'Table Grid'

    headers = ["Parts of Speech", "Phrase Types", "Roles"]
    abbrev_data = [
        ["Noun = N",          "Noun Phrase = NP",          "Subject = Subj"],
        ["Verb = V",          "Verb Phrase = VP",          "Predicate = Pred"],
        ["Adjective = ADJ",   "Adjective Phrase = ADJP",   ""],
        ["Adverb = ADV",      "Adverb Phrase = ADVP",      ""],
        ["Pronoun = PRO",     "Prepositional Phrase = PP",  ""],
        ["Determiner = DET",  "Sentence = S",              ""],
        ["Preposition = PREP","",                           ""],
    ]

    for j, h in enumerate(headers):
        cell = abbrev_table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, font_size=9)
        set_cell_shading(cell, "E8E8E8")

    for i, row_vals in enumerate(abbrev_data):
        for j, val in enumerate(row_vals):
            cell = abbrev_table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_run(p, val, font_size=9)


def build_notes(doc):
    """Add the standard notes block."""
    p = doc.add_paragraph()
    add_run(p, "Notes:", bold=True, font_size=10)
    notes = [
        "In the Phrase row, label the main verb as VP. The VP label in the table applies to just the main verb itself.",
        "In the Role row, label the subject noun phrase as Subject and the main verb as Predicate. Leave other cells empty.",
        "Always include a Phrase-level label for every noun, verb, adjective, adverb, and preposition.",
        "Determiners (DET) group under the NP — they do not get their own Phrase label.",
    ]
    for note in notes:
        bp = doc.add_paragraph(style='List Bullet')
        bp.text = ""
        add_run(bp, note, font_size=10)


def create_document(is_answer_key=False):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Open Sans'
        h_style.font.bold = True

    # -------------------------------------------------------------------------
    # HEADER
    # -------------------------------------------------------------------------
    title_text = "ENGL 3110 \u2014 Bonus Assignment \u2014 Spring 2026"
    if is_answer_key:
        title_text += " \u2014 ANSWER KEY"

    title = doc.add_heading(title_text, level=1)
    for run in title.runs:
        run.font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=4)

    p = doc.add_paragraph()
    add_run(p, "Total Points: 20  \u2022  2 points per sentence", bold=True, font_size=11)
    set_paragraph_spacing(p, space_before=0, space_after=4)

    if not is_answer_key:
        p = doc.add_paragraph()
        add_run(p, "STUDENT NAME: ", bold=True, font_size=11)
        add_run(p, "_" * 55, font_size=11)
        set_paragraph_spacing(p, space_before=0, space_after=6)

    # -------------------------------------------------------------------------
    # INSTRUCTIONS
    # -------------------------------------------------------------------------
    add_section_header(doc, "Instructions", level=2, font_size=13)

    p = doc.add_paragraph()
    add_run(p, "For each sentence below, complete two tasks:", font_size=11)

    for item in [
        ("1. ", "Complete the labeling table ", "by filling in the Role, Phrase, and Part of Speech (POS) rows."),
        ("2. ", "Create a syntax tree diagram ", "and paste your diagram below."),
    ]:
        p = doc.add_paragraph()
        add_run(p, item[0], bold=True, font_size=11)
        add_run(p, item[1], bold=True, font_size=11)
        add_run(p, item[2], font_size=11)

    p = doc.add_paragraph()
    add_run(p, "Scoring (per sentence): ", bold=True, font_size=10)
    add_run(p, "1 point for a correct labeling table  \u2022  1 point for a correct diagram", font_size=10)
    set_paragraph_spacing(p, space_before=2, space_after=6)

    build_abbrev_table(doc)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=4, space_after=0)
    build_notes(doc)

    # -------------------------------------------------------------------------
    # SENTENCES
    # -------------------------------------------------------------------------
    for q in SENTENCES:
        doc.add_page_break()

        p = doc.add_paragraph()
        add_run(p, f"Sentence {q['num']}: ", bold=True, font_size=12)
        add_run(p, f"(2 pts) ", font_size=11)
        add_run(p, q["sentence"], italic=True, font_size=11)
        set_paragraph_spacing(p, space_before=6, space_after=4)

        # Labeling table
        if is_answer_key:
            add_labeling_table(
                doc, q["words"],
                pos_labels=q["pos"],
                phrase_labels=q["phrases"],
                role_labels=q["roles"],
                font_size=10,
            )
        else:
            add_labeling_table(
                doc, q["words"],
                phrase_labels=blank_labels(q["phrases"]),
                role_labels=blank_labels(q["roles"]),
                font_size=10,
            )

        # Bracket notation
        p = doc.add_paragraph()
        if is_answer_key:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
            add_run(p, q["bracket"], font_size=9, font_name="Consolas")
        else:
            add_run(p, "Bracket notation: ", bold=True, font_size=10)
        set_paragraph_spacing(p, space_before=4, space_after=2)

        # Diagram
        if is_answer_key:
            add_diagram_image(doc, q["diagram"], width_inches=q["width"])
        else:
            add_diagram_box(doc)

    return doc


def main():
    print("=" * 60)
    print("Generating Bonus Assignment .docx files")
    print("=" * 60)

    student_path = SCRIPT_DIR / "ENGL 3110 - S26 - Bonus Assignment.docx"
    print(f"\nGenerating student version: {student_path.name}")
    doc = create_document(is_answer_key=False)
    doc.save(str(student_path))
    print(f"  Saved: {student_path}")

    key_path = SCRIPT_DIR / "ENGL 3110 - S26 - Bonus Assignment - Answer Key.docx"
    print(f"\nGenerating answer key: {key_path.name}")
    doc = create_document(is_answer_key=True)
    doc.save(str(key_path))
    print(f"  Saved: {key_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()

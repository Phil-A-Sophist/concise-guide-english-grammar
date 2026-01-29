#!/usr/bin/env python3
"""
Convert homework markdown files to Word documents.
Creates separate homework and answer key documents.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_answer_lines(doc, num_lines=3, line_spacing=1.5):
    """Add blank lines for student answers."""
    for _ in range(num_lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run('_' * 70)
        run.font.color.rgb = None  # Use default color

def add_answer_box(doc, height_inches=1.5):
    """Add a bordered box for longer answers."""
    # Add an empty paragraph with borders
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Add multiple lines for the box effect
    for _ in range(int(height_inches * 4)):
        doc.add_paragraph()

def process_markdown_to_word(md_content, doc, is_answer_key=False):
    """Process markdown content and add to Word document."""
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Headers
        if line.startswith('# ') and not line.startswith('# ANSWER'):
            text = line[2:].strip()
            doc.add_heading(text, level=0)
        elif line.startswith('## Part') or (line.startswith('## ') and 'Part' in line):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=1)

        # Horizontal rules
        elif line == '---':
            doc.add_paragraph()

        # Instructions blocks
        elif line.startswith('**Instructions:**') or line.startswith('**Instructions**'):
            para = doc.add_paragraph()
            run = para.add_run('Instructions: ')
            run.bold = True
            text = re.sub(r'\*\*Instructions:?\*\*\s*', '', line)
            text = clean_markdown(text)
            para.add_run(text)

        # Numbered questions
        elif re.match(r'^\*\*\d+\.\*\*', line):
            text = clean_markdown(line)
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.add_run(text)

            # If not answer key, add answer space
            if not is_answer_key:
                # Look ahead for continuation lines (like "Number of morphemes:")
                j = i + 1
                while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith('**') and not lines[j].strip().startswith('#') and not lines[j].strip().startswith('---'):
                    next_line = lines[j].strip()
                    if next_line and not next_line.startswith('-'):
                        para = doc.add_paragraph()
                        para.add_run(clean_markdown(next_line))
                    j += 1
                i = j - 1
                # Add answer space
                add_answer_lines(doc, 2)

        # List items
        elif line.startswith('- ') or line.startswith('* '):
            text = clean_markdown(line[2:])
            doc.add_paragraph(text, style='List Bullet')

        # Example blocks
        elif line.startswith('### Example') or 'Example' in line and '**' in line:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            run = para.add_run('Example: ')
            run.bold = True
            run.italic = True
            text = re.sub(r'###?\s*Example[^:]*:?\s*', '', line)
            text = clean_markdown(text)
            if text:
                para.add_run(text)

        # Regular paragraphs
        else:
            text = clean_markdown(line)
            if text:
                para = doc.add_paragraph()
                para.add_run(text)

        i += 1

def clean_markdown(text):
    """Remove markdown formatting while preserving content."""
    # Remove bold markers but keep content
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic markers but keep content
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove inline code markers
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def split_homework_and_answers(content):
    """Split markdown content into homework questions and answer key."""
    # Find the answer key section (marked by double horizontal rule or "# ANSWER KEY")
    patterns = [
        r'\n---\s*\n---\s*\n\s*#\s*ANSWER\s*KEY',
        r'\n---\s*---\s*\n\s*#\s*ANSWER\s*KEY',
        r'\n#\s*ANSWER\s*KEY'
    ]

    for pattern in patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            homework = content[:match.start()].strip()
            answers = content[match.end():].strip()
            return homework, answers

    # No answer key found
    return content.strip(), None

def create_word_documents(md_path, output_dir):
    """Create Word documents from a homework markdown file."""

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Get chapter info from filename
    filename = os.path.basename(md_path)
    match = re.search(r'ch(\d+)', filename, re.IGNORECASE)
    if match:
        chapter_num = int(match.group(1))
    else:
        chapter_num = 0

    # Get title from first header
    title_match = re.search(r'^#\s*Homework:\s*(.+)$', content, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
    else:
        title = f'Chapter {chapter_num}'

    # Split into homework and answers
    homework_content, answer_content = split_homework_and_answers(content)

    # Create homework document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add title
    title_para = doc.add_heading(f'Homework: {title}', level=0)

    # Add name/date fields
    doc.add_paragraph()
    doc.add_paragraph('Name: ________________________________')
    doc.add_paragraph('Date: ________________________________')
    doc.add_paragraph()

    # Process homework content
    process_markdown_to_word(homework_content, doc, is_answer_key=False)

    # Save homework document
    hw_filename = f'Chapter {chapter_num:02d} Homework.docx'
    hw_path = os.path.join(output_dir, hw_filename)
    doc.save(hw_path)
    print(f'Created: {hw_path}')

    # Create answer key document if answers exist
    if answer_content:
        doc_key = Document()

        style = doc_key.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        title_para = doc_key.add_heading(f'Answer Key: {title}', level=0)
        doc_key.add_paragraph()

        process_markdown_to_word(answer_content, doc_key, is_answer_key=True)

        key_filename = f'Chapter {chapter_num:02d} Answer Key.docx'
        key_path = os.path.join(output_dir, key_filename)
        doc_key.save(key_path)
        print(f'Created: {key_path}')

def main():
    """Process all homework markdown files."""
    homework_dir = r'C:\Users\irphy\Documents\bookmaker\Writing Projects\concise_guide_to_english_grammar\Homework'

    # Find all homework markdown files
    md_files = [f for f in os.listdir(homework_dir)
                if f.endswith('.md') and 'homework' in f.lower()]

    print(f'Found {len(md_files)} homework markdown files')

    for md_file in sorted(md_files):
        md_path = os.path.join(homework_dir, md_file)
        print(f'\nProcessing: {md_file}')
        try:
            create_word_documents(md_path, homework_dir)
        except Exception as e:
            print(f'Error processing {md_file}: {e}')

if __name__ == '__main__':
    main()

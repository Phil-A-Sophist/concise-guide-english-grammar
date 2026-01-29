#!/usr/bin/env python3
"""
Script to create Word documents from homework markdown files.
Creates both homework (questions only) and answer key documents.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_homework_doc(chapter_num, title, homework_content, answer_key_content, output_dir):
    """Create Word documents for homework and answer key."""

    # Create homework document (questions only)
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add title
    title_para = doc.add_heading(f'Chapter {chapter_num} Homework: {title}', level=1)

    # Add name/date fields
    doc.add_paragraph('Name: _________________________________')
    doc.add_paragraph('Date: _________________________________')
    doc.add_paragraph()

    # Process homework content
    add_markdown_to_doc(doc, homework_content)

    # Save homework document
    homework_path = os.path.join(output_dir, f'Chapter {chapter_num:02d} Homework.docx')
    doc.save(homework_path)
    print(f'Created: {homework_path}')

    # Create answer key document
    if answer_key_content:
        doc_key = Document()

        style = doc_key.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        title_para = doc_key.add_heading(f'Chapter {chapter_num} Answer Key: {title}', level=1)

        add_markdown_to_doc(doc_key, answer_key_content)

        key_path = os.path.join(output_dir, f'Chapter {chapter_num:02d} Answer Key.docx')
        doc_key.save(key_path)
        print(f'Created: {key_path}')

def add_markdown_to_doc(doc, content):
    """Convert markdown content to Word document paragraphs."""
    lines = content.split('\n')
    current_para = []
    in_list = False

    for line in lines:
        # Handle headers
        if line.startswith('## Part') or line.startswith('### Part'):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            level = 2 if line.startswith('## ') else 3
            text = re.sub(r'^#+\s*', '', line)
            doc.add_heading(text, level=level)
        elif line.startswith('## ') or line.startswith('### '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            level = 2 if line.startswith('## ') else 3
            text = re.sub(r'^#+\s*', '', line)
            doc.add_heading(text, level=level)
        elif line.startswith('---'):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            # Add a paragraph break for horizontal rules
            doc.add_paragraph()
        elif line.startswith('**Instructions:**') or line.startswith('**Instructions**'):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            # Bold instructions
            para = doc.add_paragraph()
            run = para.add_run('Instructions: ')
            run.bold = True
            text = re.sub(r'\*\*Instructions:?\*\*\s*', '', line)
            text = clean_markdown(text)
            para.add_run(text)
        elif re.match(r'^\*\*\d+\.\*\*', line):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            # Numbered question
            text = clean_markdown(line)
            para = doc.add_paragraph(text)
            # Add answer space
            doc.add_paragraph('_' * 60)
            doc.add_paragraph()
        elif line.startswith('- ') or line.startswith('* '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            text = clean_markdown(line[2:])
            doc.add_paragraph(text, style='List Bullet')
        elif line.strip() == '':
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
        else:
            text = clean_markdown(line)
            if text:
                current_para.append(text)

    if current_para:
        doc.add_paragraph(' '.join(current_para))

def clean_markdown(text):
    """Remove markdown formatting while preserving content."""
    # Remove bold markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def extract_homework_from_chapter(chapter_path):
    """Extract homework section from chapter markdown file."""
    with open(chapter_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Find homework section
    match = re.search(r'## Homework:?\s*([^\n]+)\n(.*?)(?=\n## (?!Part)|$)', content, re.DOTALL)
    if match:
        title = match.group(1).strip()
        homework_content = match.group(2).strip()
        return title, homework_content
    return None, None

if __name__ == '__main__':
    chapters_dir = r'C:\Users\irphy\Documents\bookmaker\Writing Projects\concise_guide_to_english_grammar\chapters'
    homework_dir = r'C:\Users\irphy\Documents\bookmaker\Writing Projects\concise_guide_to_english_grammar\Homework'

    # Process each chapter
    for i in range(1, 22):
        # Find the chapter file
        chapter_files = [f for f in os.listdir(chapters_dir)
                        if f.startswith(f'chapter_{i:02d}') or f.startswith(f'chapter_{i}_')]

        if not chapter_files:
            # Try without leading zero
            chapter_files = [f for f in os.listdir(chapters_dir)
                           if f.startswith(f'chapter_{i}_')]

        if chapter_files:
            chapter_path = os.path.join(chapters_dir, chapter_files[0])
            title, homework = extract_homework_from_chapter(chapter_path)
            if title and homework:
                print(f'Processing Chapter {i}: {title}')
                create_homework_doc(i, title, homework, None, homework_dir)
            else:
                print(f'No homework found in Chapter {i}')
        else:
            print(f'Chapter {i} file not found')
